"""
Churchgate Group HRIS v5.0
Enterprise-Grade AI-Powered Human Resource Information System
"""

import streamlit as st
from streamlit_option_menu import option_menu
import pandas as pd
import numpy as np
import plotly.express as px
import plotly.graph_objects as go
from datetime import datetime, timedelta, date
import hashlib
from pathlib import Path
import sys
import time
import json
import base64
import io
import os
import random
from PIL import Image
import calendar

sys.path.append(str(Path(__file__).parent))

# =============================================
# FIX: READ SECRETS FROM HUGGING FACE ENVIRONMENT
# =============================================
# This wraps st.secrets to also check environment variables
class HuggingFaceSecrets:
    def __getitem__(self, key):
        # Try st.secrets first
        try:
            # Check if st.secrets has the key
            if hasattr(st, 'secrets') and hasattr(st.secrets, '_secrets'):
                if key in st.secrets._secrets:
                    return st.secrets._secrets[key]
        except:
            pass
        
        # Fallback to environment variables
        return os.environ.get(key)
    
    def get(self, key, default=None):
        try:
            val = self.__getitem__(key)
            return val if val is not None else default
        except:
            return default
    
    def __contains__(self, key):
        try:
            return self.__getitem__(key) is not None
        except:
            return False

# Backup the original secrets if it exists
_original_secrets = None
if hasattr(st, 'secrets') and st.secrets:
    _original_secrets = st.secrets

# Replace st.secrets with our wrapper
st.secrets = HuggingFaceSecrets()
# =============================================

from utils.database import DatabaseManager
from utils.ai_agent import AIRecruitmentAgent
from utils.linkedin_parser import LinkedInParser
from utils.email_service import EmailService
from utils.chat_service import ChatService
from utils.training_service import TrainingService
from streamlit_quill import st_quill
# ============================================================
# SECTION 4: REST OF YOUR CODE (logo, page config, etc)
# ============================================================
logo_icon = Path(__file__).parent / "churchgate-logo.jpeg"

logo_icon = Path(__file__).parent / "churchgate-logo.jpeg"
if logo_icon.exists():
    st.set_page_config(page_title="Churchgate Group HRIS", page_icon=str(logo_icon), layout="wide", initial_sidebar_state="expanded")
else:
    st.set_page_config(page_title="Churchgate Group HRIS", page_icon="🏢", layout="wide", initial_sidebar_state="expanded")

# Browser Notification Setup
st.markdown("""
<script>
if ('Notification' in window && Notification.permission === 'default') {
    setTimeout(() => {
        Notification.requestPermission();
    }, 3000);
}
function showNotification(title, body) {
    if ('Notification' in window && Notification.permission === 'granted') {
        new Notification(title, {
            body: body,
            icon: 'https://raw.githubusercontent.com/eetuk-churchgate/churchgate-hris/main/churchgate-logo-192.png',
            badge: 'https://raw.githubusercontent.com/eetuk-churchgate/churchgate-hris/main/churchgate-logo-192.png',
            vibrate: [200, 100, 200]
        });
    }
}
</script>
""", unsafe_allow_html=True)

CHURCHGATE_RED = "#CC0000"
CHURCHGATE_DARK = "#1a1a1a"
CHURCHGATE_GREY = "#4a4a4a"
CHURCHGATE_LIGHT = "#f5f5f5"
CHURCHGATE_WHITE = "#ffffff"

CHURCHGATE_PURPOSE = "To improve the lives of all those we serve."
CHURCHGATE_VISION = "To become the premier property developer in Nigeria, impacting millions, while having fun!"
CHURCHGATE_MISSION = "To provide our customers with innovative and sustainable real estate solutions that enable them to thrive."

CHURCHGATE_VALUES = {
    "Humility": "We pride in selflessness and dignity with a natural ability to improve our services and make our offerings better.",
    "Gratitude": "Our customers are at the crux of our activities, thus we are devoted to ensuring the best experience with us.",
    "Integrity": "We are responsible for our actions, and our results. We keep to our word!",
    "Hunger": "We are driven by our commitment and dedication to make our services better.",
    "Passion": "Your satisfaction is our pleasure. We strive to always leave you with a bigger smile than you came with."
}

CHURCHGATE_PORTFOLIO = [
    "World Trade Center Abuja",
    "Churchgate Tower 1, Lagos",
    "Churchgate Tower 2, Lagos",
    "Churchgate Plaza, Abuja",
    "Warehouses",
    "Ocean Terrace"
]

st.markdown("""
<style>
    @import url('https://fonts.googleapis.com/css2?family=Georgia&family=Inter:wght@300;400;500;600;700;800;900&display=swap');
    
    * { font-family: 'Inter', sans-serif; }
    h1, h2, h3, h4, h5, h6 { font-family: 'Georgia', serif; }
    
    .stApp { background: #f5f0e8; }
    .main > div { background: #f5f0e8; }
    
    section[data-testid="stSidebar"] { background-color: #d5d5d5 !important; }
    section[data-testid="stSidebar"] * { color: #333333 !important; }
    section[data-testid="stSidebar"] .stButton > button { background-color: #c0c0c0 !important; border: 1px solid #a0a0a0 !important; color: #333333 !important; }
    section[data-testid="stSidebar"] .nav-link { color: #333333 !important; }
    section[data-testid="stSidebar"] .nav-link span { color: #333333 !important; font-size: 14px !important; }
    section[data-testid="stSidebar"] .nav-link svg { color: #C8A951 !important; }
    section[data-testid="stSidebar"] .nav-link-selected { background-color: #c0c0c0 !important; border-left: 3px solid #C8A951 !important; }
    section[data-testid="stSidebar"] .nav-link-selected span { color: #C8A951 !important; font-weight: 700 !important; }
    
    .churchgate-header { background: white; padding: 1.5rem 2rem; border-radius: 8px; margin-bottom: 1.5rem; border-left: 4px solid #C8A951; box-shadow: 0 1px 3px rgba(0,0,0,0.1); }
    .churchgate-header h1 { color: #1a1a1a; font-size: 1.8rem; font-weight: 700; margin: 0; font-family: 'Georgia', serif; }
    .churchgate-header p { color: #666; font-size: 0.9rem; margin-top: 0.3rem; }
    
    .metric-card { background: white; padding: 1.2rem; border-radius: 8px; box-shadow: 0 1px 3px rgba(0,0,0,0.08); text-align: center; border: 1px solid #cccccc; transition: all 0.2s ease; }
    .metric-card:hover { transform: translateY(-3px); box-shadow: 0 4px 12px rgba(0,0,0,0.12); }
    .metric-value { font-size: 1.8rem; font-weight: 700; color: #C8A951; font-family: 'Georgia', serif; }
    .metric-label { color: #666; font-size: 0.8rem; text-transform: uppercase; letter-spacing: 0.5px; }
    
    .stButton > button { background: #CC0000 !important; color: white !important; border: none !important; padding: 0.5rem 1rem !important; border-radius: 6px !important; font-weight: 600 !important; }
    .stButton > button:hover { background: #aa0000 !important; }
    
    .mission-banner { background: #d5d5d5; padding: 1.5rem; border-radius: 8px; text-align: center; margin: 1.5rem 0; border: 2px solid #C8A951; }
    .mission-banner h2 { color: #C8A951; font-size: 1.3rem; }
    .mission-banner h3 { color: #1a1a1a; }
    .mission-banner p { color: #333333; }
    
    .value-card { background: white; padding: 0.8rem; border-radius: 6px; text-align: center; border: 1px solid #cccccc; }
    .value-card h4 { color: #C8A951; font-size: 0.85rem; margin: 0 0 0.3rem 0; }
    .value-card p { font-size: 0.7rem; color: #666; margin: 0; line-height: 1.3; }
    
    .tier-1-badge { background: #38a169; color: white; padding: 0.3rem 0.6rem; border-radius: 15px; font-weight: 600; font-size: 0.8rem; }
    .tier-2-badge { background: #C8A951; color: #1a1a1a; padding: 0.3rem 0.6rem; border-radius: 15px; font-weight: 600; font-size: 0.8rem; }
    .tier-3-badge { background: #CC0000; color: white; padding: 0.3rem 0.6rem; border-radius: 15px; font-weight: 600; font-size: 0.8rem; }
    
    .status-active { background: #38a169; color: white; padding: 0.2rem 0.6rem; border-radius: 15px; font-size: 0.8rem; }
    .status-pending { background: #C8A951; color: #1a1a1a; padding: 0.2rem 0.6rem; border-radius: 15px; font-size: 0.8rem; }
    .status-at-risk { background: #CC0000; color: white; padding: 0.2rem 0.6rem; border-radius: 15px; font-size: 0.8rem; }
    
    .stImage { display: flex; justify-content: center; }
    .chat-container { max-height: 400px; overflow-y: auto; padding: 1rem; background: white; border-radius: 8px; margin-bottom: 1rem; }
    
    .greeting-header {
        background: #d5d5d5;
        padding: 1rem 2rem;
        border-radius: 10px;
        margin-bottom: 1rem;
        border-left: 4px solid #C8A951;
        display: flex;
        align-items: center;
        justify-content: space-between;
    }
    .greeting-header h1 { color: #1a1a1a; font-size: 1.5rem; font-weight: 700; margin: 0; font-family: 'Georgia', serif; }
    .greeting-header p { color: #555; font-size: 0.85rem; margin: 0.2rem 0 0 0; }
    
    .mission-banner {
        background: #d5d5d5;
        padding: 1rem 1.5rem;
        border-radius: 8px;
        text-align: center;
        margin: 1rem 0;
        border: 2px solid #C8A951;
    }
    .mission-banner h2 { color: #C8A951; font-size: 1.2rem; margin-bottom: 0.5rem; }
    
    .mission-item { 
        background: white; 
        padding: 0.8rem; 
        border-radius: 6px; 
        text-align: center; 
        height: 100%;
        min-height: 80px;
        display: flex;
        flex-direction: column;
        justify-content: center;
    }
    .mission-item h3 { color: #C8A951; font-size: 0.9rem; margin: 0 0 0.3rem 0; }
    .mission-item p { font-size: 0.75rem; color: #666; margin: 0; line-height: 1.3; }
    
    .value-card {
        background: white;
        padding: 0.8rem;
        border-radius: 6px;
        text-align: center;
        border: 1px solid #e0e0e0;
        height: 100%;
        min-height: 100px;
        display: flex;
        flex-direction: column;
        justify-content: center;
    }
    .value-card h4 { color: #C8A951; font-size: 0.85rem; margin: 0 0 0.3rem 0; }
    .value-card p { font-size: 0.7rem; color: #666; margin: 0; line-height: 1.3; }
    
    [data-testid="stExpander"] details summary {
        font-weight: 600 !important;
        font-size: 0.95rem !important;
        background: white !important;
        border-radius: 8px !important;
        border-left: 4px solid #C8A951 !important;
        padding: 0.5rem 1rem !important;
        margin-bottom: 0.3rem !important;
    }
    [data-testid="stExpander"] details summary:hover {
        background: #f8f8f8 !important;
    }
    [data-testid="stExpander"] details summary svg {
        color: #C8A951 !important;
    }
    
    #MainMenu {{ visibility: hidden; }}
    header {{ visibility: hidden; }}
    footer {{ visibility: hidden; }}
</style>
""", unsafe_allow_html=True)

@st.cache_resource
def init_resources():
    db = DatabaseManager()
    db.create_tables()
    ai_agent = AIRecruitmentAgent()
    linkedin_parser = LinkedInParser()
    email_service = EmailService()
    chat_service = ChatService()
    training_service = TrainingService()
    return db, ai_agent, linkedin_parser, email_service, chat_service, training_service

db, ai_agent, linkedin_parser, email_service, chat_service, training_service = init_resources()

if 'user' not in st.session_state:
    st.session_state.user = None
if 'current_jd' not in st.session_state:
    st.session_state.current_jd = None
if 'candidates_batch' not in st.session_state:
    st.session_state.candidates_batch = []
if 'chat_messages' not in st.session_state:
    st.session_state.chat_messages = []
if 'bot_conversation' not in st.session_state:
    st.session_state.bot_conversation = []
if 'dashboard_metrics' not in st.session_state:
    st.session_state.dashboard_metrics = {
        'total_employees': 48, 'occupancy_rate': 87,
        'revenue_vs_budget': 94, 'tenant_satisfaction': 4.2, 'open_positions': 5
    }

def get_logo():
    logo_paths = [Path(__file__).parent / "churchgate-logo.png", Path(__file__).parent / "churchgate_logo.png"]
    for path in logo_paths:
        if path.exists():
            return Image.open(path)
    return None

def generate_initials(name):
    parts = name.split()
    if len(parts) >= 2:
        return f"{parts[0][0]}{parts[-1][0]}".upper()
    return name[:2].upper()

def save_uploaded_file(uploaded_file):
    if uploaded_file is None:
        return None
    try:
        if uploaded_file.type == "application/pdf":
            import PyPDF2
            pdf_reader = PyPDF2.PdfReader(uploaded_file)
            text = ""
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
            return text
        elif uploaded_file.type == "text/plain":
            return uploaded_file.read().decode('utf-8')
        elif "word" in uploaded_file.type or "docx" in uploaded_file.type:
            import docx
            doc = docx.Document(uploaded_file)
            return "\n".join([p.text for p in doc.paragraphs])
        else:
            return uploaded_file.read().decode('utf-8', errors='ignore')
    except Exception as e:
        return f"[Error: {str(e)}]"

def generate_summary_pdf(dept_name, dept_data, summary):
    """Generate a Fortune 500 Executive Summary PDF"""
    import fpdf
    FPDF = fpdf.FPDF
    
    pdf = FPDF(orientation='P', unit='mm', format='A4')
    pdf.set_auto_page_break(auto=True, margin=12)
    pdf.add_page()
    
    # ===== COVER HEADER =====
    pdf.set_fill_color(26, 26, 26)
    pdf.rect(0, 0, 210, 35, 'F')
    pdf.set_fill_color(204, 0, 0)
    pdf.rect(0, 35, 210, 4, 'F')
    pdf.set_font('Helvetica', 'B', 22)
    pdf.set_text_color(255, 255, 255)
    pdf.cell(0, 18, 'CHURCHGATE GROUP', ln=True, align='C')
    pdf.set_font('Helvetica', 'B', 13)
    pdf.set_text_color(204, 0, 0)
    pdf.cell(0, 10, 'EXECUTIVE PERFORMANCE SUMMARY', ln=True, align='C')
    pdf.ln(8)
    
    # ===== DEPT & DATE BAR =====
    pdf.set_fill_color(245, 245, 245)
    pdf.rect(10, pdf.get_y(), 190, 10, 'F')
    pdf.set_font('Helvetica', 'B', 10)
    pdf.set_text_color(26, 26, 26)
    pdf.cell(95, 10, f'  Department: {dept_name}', 0, 0, 'L')
    pdf.cell(95, 10, f'Generated: {datetime.now().strftime("%B %d, %Y")}  ', 0, 0, 'R')
    pdf.ln(14)
    
    # ===== OVERALL SCORE HERO =====
    total_weighted = sum(p['progress'] * p['weight'] / 100 for p in dept_data.values())
    on_track = sum(1 for p in dept_data.values() if p['status'] in ['On Track', 'Exceeding'])
    at_risk = sum(1 for p in dept_data.values() if p['status'] == 'At Risk')
    completed = sum(1 for p in dept_data.values() if p['status'] == 'Completed')
    
    pdf.set_fill_color(26, 26, 26)
    pdf.rect(10, pdf.get_y(), 190, 18, 'F')
    pdf.set_fill_color(204, 0, 0)
    pdf.rect(10, pdf.get_y()+18, 190, 2, 'F')
    pdf.set_font('Helvetica', 'B', 12)
    pdf.set_text_color(255, 255, 255)
    pdf.cell(190, 18, f'OVERALL WEIGHTED SCORE: {total_weighted:.1f}%', ln=True, align='C')
    pdf.ln(6)
    
    # ===== KPI CARDS ROW =====
    pdf.set_font('Helvetica', 'B', 9)
    col_w = 42
    col_h = 14
    x_start = 14
    y_pos = pdf.get_y()
    
    cards = [
        ('ON TRACK', str(on_track), (56, 161, 105)),
        ('AT RISK', str(at_risk), (204, 0, 0)),
        ('COMPLETED', str(completed), (49, 130, 206)),
        ('TOTAL PILLARS', '4', (100, 100, 100)),
    ]
    
    for i, (label, value, color) in enumerate(cards):
        x = x_start + i * (col_w + 6)
        pdf.set_fill_color(*color)
        pdf.rect(x, y_pos, col_w, col_h, 'F')
        pdf.set_text_color(255, 255, 255)
        pdf.set_xy(x, y_pos + 1)
        pdf.cell(col_w, 6, label, 0, 0, 'C')
        pdf.set_xy(x, y_pos + 7)
        pdf.set_font('Helvetica', 'B', 14)
        pdf.cell(col_w, 7, value, 0, 0, 'C')
        pdf.set_font('Helvetica', 'B', 9)
    
    pdf.set_y(y_pos + col_h + 8)
    
    # ===== PILLAR TABLE =====
    pdf.set_fill_color(26, 26, 26)
    pdf.set_text_color(255, 255, 255)
    pdf.set_font('Helvetica', 'B', 9)
    pdf.cell(68, 8, '  STRATEGIC PILLAR', 1, 0, 'L', True)
    pdf.cell(18, 8, 'WEIGHT', 1, 0, 'C', True)
    pdf.cell(22, 8, 'PROGRESS', 1, 0, 'C', True)
    pdf.cell(30, 8, 'STATUS', 1, 0, 'C', True)
    pdf.cell(52, 8, 'DEADLINE', 1, 0, 'C', True)
    pdf.ln()
    
    pdf.set_font('Helvetica', '', 8)
    for pn, pdata in dept_data.items():
        if pdata['status'] in ['On Track', 'Exceeding']:
            sc = (230, 255, 230)
            tc = (56, 161, 105)
        elif pdata['status'] == 'In Progress':
            sc = (255, 248, 230)
            tc = (214, 158, 46)
        else:
            sc = (255, 230, 230)
            tc = (204, 0, 0)
        
        pdf.set_fill_color(*sc)
        pdf.set_text_color(26, 26, 26)
        pdf.cell(68, 9, f'  {pn[:45]}', 1, 0, 'L', True)
        pdf.cell(18, 9, f'{pdata["weight"]}%', 1, 0, 'C', True)
        pdf.cell(22, 9, f'{pdata["progress"]}%', 1, 0, 'C', True)
        pdf.set_text_color(*tc)
        pdf.set_font('Helvetica', 'B', 8)
        pdf.cell(30, 9, pdata['status'], 1, 0, 'C', True)
        pdf.set_text_color(26, 26, 26)
        pdf.set_font('Helvetica', '', 8)
        pdf.cell(52, 9, pdata['deadline'], 1, 0, 'C', True)
        pdf.ln()
    
    pdf.ln(4)
    
    # ===== PROGRESS BARS =====
    pdf.set_font('Helvetica', 'B', 11)
    pdf.set_text_color(26, 26, 26)
    pdf.cell(0, 8, 'PILLAR PROGRESS', ln=True)
    pdf.ln(2)
    
    for pn, pdata in dept_data.items():
        pdf.set_font('Helvetica', '', 8)
        pdf.set_text_color(80, 80, 80)
        pdf.cell(190, 4, pn[:55], ln=True)
        
        bar_y = pdf.get_y()
        pdf.set_fill_color(230, 230, 230)
        pdf.rect(10, bar_y, 190, 5, 'F')
        
        if pdata['progress'] >= 85:
            bar_color = (56, 161, 105)
        elif pdata['progress'] >= 65:
            bar_color = (214, 158, 46)
        else:
            bar_color = (204, 0, 0)
        
        pdf.set_fill_color(*bar_color)
        pdf.rect(10, bar_y, 190 * pdata['progress'] / 100, 5, 'F')
        pdf.set_y(bar_y + 7)
    
    pdf.ln(4)
    
    # ===== RACI MATRIX TABLE =====
    pdf.set_fill_color(26, 26, 26)
    pdf.set_text_color(255, 255, 255)
    pdf.set_font('Helvetica', 'B', 9)
    pdf.cell(190, 8, '  RACI MATRIX', 1, 0, 'L', True)
    pdf.ln()
    
    pdf.set_fill_color(245, 245, 245)
    pdf.set_text_color(26, 26, 26)
    pdf.set_font('Helvetica', 'B', 8)
    pdf.cell(35, 7, '  Role', 1, 0, 'L', True)
    pdf.cell(155, 7, '  Party', 1, 0, 'L', True)
    pdf.ln()
    
    raci_data = [
        ('RESPONSIBLE', 'Department Heads / HODs'),
        ('ACCOUNTABLE', 'GMD / COO'),
        ('CONSULTED', 'All HODs / Key Stakeholders'),
        ('INFORMED', 'Board of Directors'),
    ]
    
    pdf.set_font('Helvetica', '', 8)
    for role, party in raci_data:
        pdf.set_fill_color(255, 255, 255)
        pdf.cell(35, 7, f'  {role}', 1, 0, 'L', True)
        pdf.cell(155, 7, f'  {party}', 1, 0, 'L', True)
        pdf.ln()
    
    # ===== FOOTER =====
    pdf.set_y(-18)
    pdf.set_fill_color(26, 26, 26)
    pdf.rect(0, pdf.get_y()-2, 210, 20, 'F')
    pdf.set_font('Helvetica', 'I', 7)
    pdf.set_text_color(180, 180, 180)
    pdf.cell(0, 6, 'Churchgate Group - Confidential - Fortune 500 Standard HRIS', ln=True, align='C')
    pdf.cell(0, 5, 'World Trade Center, Abuja | hr@churchgate.com', ln=True, align='C')
    
    return bytes(pdf.output())

def generate_performance_pdf(dept_name, dept_data, report_data):
    """Generate a Fortune 500 Performance Report PDF"""
    import fpdf
    FPDF = fpdf.FPDF
    
    pdf = FPDF(orientation='L', unit='mm', format='A4')
    pdf.set_auto_page_break(auto=True, margin=12)
    pdf.add_page()
    
    pdf.set_fill_color(26, 26, 26)
    pdf.rect(0, 0, 297, 28, 'F')
    pdf.set_fill_color(204, 0, 0)
    pdf.rect(0, 28, 297, 3, 'F')
    pdf.set_font('Helvetica', 'B', 22)
    pdf.set_text_color(255, 255, 255)
    pdf.cell(0, 16, 'CHURCHGATE GROUP', ln=True, align='C')
    pdf.set_font('Helvetica', 'B', 11)
    pdf.set_text_color(204, 0, 0)
    pdf.cell(0, 8, 'DETAILED PERFORMANCE REPORT', ln=True, align='C')
    pdf.ln(6)
    
    total_weighted = sum(p['progress'] * p['weight'] / 100 for p in dept_data.values())
    pdf.set_fill_color(245, 245, 245)
    pdf.rect(10, pdf.get_y(), 277, 8, 'F')
    pdf.set_font('Helvetica', 'B', 9)
    pdf.set_text_color(26, 26, 26)
    pdf.cell(92, 8, f'  Department: {dept_name}', 0, 0, 'L')
    pdf.cell(92, 8, f'Overall Score: {total_weighted:.1f}%', 0, 0, 'C')
    pdf.cell(93, 8, f'Generated: {datetime.now().strftime("%B %d, %Y %I:%M %p")}  ', 0, 0, 'R')
    pdf.ln(12)
    
    pdf.set_fill_color(26, 26, 26)
    pdf.set_text_color(255, 255, 255)
    pdf.set_font('Helvetica', 'B', 8)
    pdf.cell(8, 8, ' #', 1, 0, 'C', True)
    pdf.cell(80, 8, ' STRATEGIC PILLAR / KPI', 1, 0, 'L', True)
    pdf.cell(22, 8, 'TARGET', 1, 0, 'C', True)
    pdf.cell(22, 8, 'CURRENT', 1, 0, 'C', True)
    pdf.cell(24, 8, 'STATUS', 1, 0, 'C', True)
    pdf.cell(28, 8, 'DEADLINE', 1, 0, 'C', True)
    pdf.cell(93, 8, ' REMARKS', 1, 0, 'L', True)
    pdf.ln()
    
    pdf.set_font('Helvetica', '', 7)
    row_num = 0
    for pn, pdata in dept_data.items():
        pdf.set_fill_color(240, 240, 240)
        pdf.set_text_color(26, 26, 26)
        pdf.set_font('Helvetica', 'B', 8)
        pdf.cell(8, 7, '', 1, 0, 'C', True)
        pdf.cell(80, 7, f' {pn}', 1, 0, 'L', True)
        pdf.cell(22, 7, f'Weight: {pdata["weight"]}%', 1, 0, 'C', True)
        pdf.cell(22, 7, f'{pdata["progress"]}%', 1, 0, 'C', True)
        
        if pdata['status'] in ['On Track', 'Exceeding']:
            sc = (56, 161, 105)
        elif pdata['status'] == 'In Progress':
            sc = (214, 158, 46)
        else:
            sc = (204, 0, 0)
        pdf.set_fill_color(*sc)
        pdf.set_text_color(255, 255, 255)
        pdf.cell(24, 7, pdata['status'], 1, 0, 'C', True)
        pdf.set_text_color(26, 26, 26)
        pdf.set_font('Helvetica', '', 7)
        pdf.cell(28, 7, pdata['deadline'], 1, 0, 'C', True)
        pdf.cell(93, 7, f' {len(pdata.get("kpis", []))} KPIs defined', 1, 0, 'L', True)
        pdf.ln()
        
        for kpi in pdata.get('kpis', []):
            row_num += 1
            pdf.set_fill_color(255, 255, 255)
            pdf.set_text_color(80, 80, 80)
            pdf.set_font('Helvetica', '', 7)
            pdf.cell(8, 6, str(row_num), 1, 0, 'C', True)
            pdf.cell(80, 6, f'   {kpi["kpi"][:55]}', 1, 0, 'L', True)
            pdf.cell(22, 6, kpi['target'], 1, 0, 'C', True)
            pdf.cell(22, 6, kpi['current'], 1, 0, 'C', True)
            
            if kpi['status'] in ['On Track', 'Completed']:
                kc = (56, 161, 105)
            elif kpi['status'] in ['Near Target', 'In Progress']:
                kc = (214, 158, 46)
            else:
                kc = (204, 0, 0)
            pdf.set_text_color(*kc)
            pdf.set_font('Helvetica', 'B', 7)
            pdf.cell(24, 6, kpi['status'], 1, 0, 'C', True)
            pdf.set_text_color(80, 80, 80)
            pdf.set_font('Helvetica', '', 7)
            pdf.cell(28, 6, kpi['deadline'], 1, 0, 'C', True)
            pdf.cell(93, 6, '', 1, 0, 'L', True)
            pdf.ln()
    
    pdf.ln(4)
    
    pdf.set_fill_color(26, 26, 26)
    pdf.set_text_color(255, 255, 255)
    pdf.set_font('Helvetica', 'B', 9)
    pdf.cell(277, 8, '  RACI MATRIX', 1, 0, 'L', True)
    pdf.ln()
    
    raci = [
        ('RESPONSIBLE', 'Department Heads'),
        ('ACCOUNTABLE', 'GMD / COO'),
        ('CONSULTED', 'All HODs'),
        ('INFORMED', 'Board'),
    ]
    for role, party in raci:
        pdf.set_fill_color(255, 255, 255)
        pdf.set_text_color(26, 26, 26)
        pdf.set_font('Helvetica', 'B', 7)
        pdf.cell(35, 6, f'  {role}', 1, 0, 'L', True)
        pdf.set_font('Helvetica', '', 7)
        pdf.cell(242, 6, f'  {party}', 1, 0, 'L', True)
        pdf.ln()
    
    pdf.set_y(-15)
    pdf.set_fill_color(26, 26, 26)
    pdf.rect(0, pdf.get_y()-2, 297, 17, 'F')
    pdf.set_font('Helvetica', 'I', 7)
    pdf.set_text_color(180, 180, 180)
    pdf.cell(0, 5, 'Churchgate Group - Confidential - Fortune 500 Standard HRIS | World Trade Center, Abuja', ln=True, align='C')
    
    return bytes(pdf.output())

def generate_ref(prefix):
    return f"{prefix}-{datetime.now().strftime('%Y%m%d')}-{str(time.time())[-6:]}"

def show_churchgate_mission():
    st.markdown("""
    <div class="mission-banner">
        <h2 style="font-size: 2rem; font-weight: 900; letter-spacing: 2px;">CHURCHGATE GROUP</h2>
        <div style="display: flex; justify-content: space-around; margin: 1rem 0; flex-wrap: wrap; gap: 0.5rem;">
            <div style="flex: 1; min-width: 180px;">
                <div class="mission-item">
                <h3>🎯 Our Purpose</h3>
                <p>To improve the lives of all those we serve.</p>
                </div>
            </div>
            <div style="flex: 1; min-width: 180px;">
                <div class="mission-item">
                <h3>🔭 Our Vision</h3>
                <p>To become the premier property developer in Nigeria, impacting millions, while having fun!</p>
                </div>
            </div>
            <div style="flex: 1; min-width: 180px;">
                <div class="mission-item">
                <h3>📋 Our Mission</h3>
                <p>To provide our customers with innovative and sustainable real estate solutions that enable them to thrive.</p>
                </div>
            </div>
        </div>
    </div>
    """, unsafe_allow_html=True)
    st.markdown("### Our Values Are Dear To Us")
    st.markdown("We deliver exceptional properties and infrastructure because we consistently sustain our core values of:")
    cols = st.columns(5)
    for i, (value, desc) in enumerate(CHURCHGATE_VALUES.items()):
        with cols[i]:
            st.markdown(f"""<div class="value-card"><h4>{value}</h4><p>{desc}</p></div>""", unsafe_allow_html=True)

def login_section():
    if 'show_forgot_password' not in st.session_state:
        st.session_state.show_forgot_password = False
    if 'show_reset_form' not in st.session_state:
        st.session_state.show_reset_form = False
    if 'reset_code' not in st.session_state:
        st.session_state.reset_code = None
    if 'reset_email' not in st.session_state:
        st.session_state.reset_email = None
    
    col1, col2, col3 = st.columns([1, 2, 1])
    with col2:
        logo = get_logo()
        if logo:
            ca, cb, cc = st.columns([1, 1, 1])
            with cb:
                st.image(logo, width=300)
        st.markdown("""<div style="text-align: center; padding: 1rem 0;"><h1 style="color: #1a1a1a; font-size: 2rem; font-weight: 700;">HRIS Portal</h1><p style="color: #666666; font-size: 0.9rem;">Human Resource Information System</p></div>""", unsafe_allow_html=True)
        
        # WTC Abuja Image
        wtc_path = Path(__file__).parent / "WTC Abuja 7 (1).jpg"
        if wtc_path.exists():
            st.image(str(wtc_path), use_container_width=True)
            st.markdown("<p style='text-align:center;color:#888;font-size:0.8rem;'>Churchgate Group</p>", unsafe_allow_html=True)
        
        with st.form("login_form", clear_on_submit=False):
            email = st.text_input("📧 Corporate Email", placeholder="Enter your corporate email")
            password = st.text_input("🔒 Password", type="password", placeholder="Enter your password")
            st.markdown("<div style='text-align:right;margin-top:-0.5rem;margin-bottom:0.5rem;'></div>", unsafe_allow_html=True)
            
            col1, col2 = st.columns([1, 1])
            with col1:
                submit = st.form_submit_button("🔐 Sign In", use_container_width=True)
            with col2:
                forgot_clicked = st.form_submit_button("🔑 Forgot Password?", use_container_width=True)
            
            if submit:
                if 'login_attempts' not in st.session_state:
                    st.session_state.login_attempts = 0
                if 'login_locked_until' not in st.session_state:
                    st.session_state.login_locked_until = None
                
                if st.session_state.login_locked_until and datetime.now() < st.session_state.login_locked_until:
                    remaining = int((st.session_state.login_locked_until - datetime.now()).total_seconds() / 60)
                    st.error(f"🔒 Too many failed attempts. Try again in {remaining} minute{'s' if remaining > 1 else ''}.")
                elif email and password:
                    user = db.verify_user(email, password)
                    if user:
                        st.session_state.user = user
                        st.session_state.authenticated = True
                        st.query_params['logged_in'] = user.get('email', 'true')
                        st.rerun()
                    else:
                        st.session_state.login_attempts += 1
                        remaining = 3 - st.session_state.login_attempts
                        if st.session_state.login_attempts >= 3:
                            st.session_state.login_locked_until = datetime.now() + timedelta(minutes=15)
                            st.error("🔒 Account locked for 15 minutes due to too many failed attempts.")
                        else:
                            st.error(f"❌ Invalid credentials. {remaining} attempt{'s' if remaining > 1 else ''} remaining.")
                else:
                    st.warning("⚠️ Please enter your email and password.")
            
            if forgot_clicked:
                st.session_state.show_forgot_password = True
                st.rerun()
        
        # FORGOT PASSWORD FORM
        if st.session_state.show_forgot_password:
            st.markdown("---")
            st.markdown("### 🔑 Reset Your Password")
            st.info("Enter your corporate email. We'll send you a reset code.")
            
            reset_email = st.text_input("📧 Corporate Email", key="reset_email_input")
            
            c1, c2 = st.columns([1, 1])
            with c1:
                if st.button("📤 Send Reset Code", use_container_width=True):
                    if reset_email:
                        try:
                            result = db.supabase.table("users").select("*").eq("email", reset_email).execute()
                            if result.data and len(result.data) > 0:
                                import random
                                reset_code = str(random.randint(100000, 999999))
                                st.session_state.reset_code = reset_code
                                st.session_state.reset_email = reset_email
                                
                                try:
                                    from utils.email_service import EmailService
                                    EmailService().send_email(
                                        reset_email,
                                        "Password Reset - Churchgate Group HRIS",
                                        f"Your password reset code is: {reset_code}\n\nEnter this code below to reset your password.\n\nIf you did not request this, please ignore this email.\n\nChurchgate Group HR"
                                    )
                                    st.success(f"✅ Reset code sent to {reset_email}")
                                except:
                                    st.success(f"✅ Reset code: {reset_code} (check your email)")
                                st.session_state.show_reset_form = True
                                st.rerun()
                            else:
                                st.error("❌ Email not found in our system.")
                        except:
                            st.error("❌ Error checking email.")
                    else:
                        st.warning("⚠️ Please enter your email.")
            
            with c2:
                if st.button("❌ Cancel", use_container_width=True):
                    st.session_state.show_forgot_password = False
                    st.session_state.show_reset_form = False
                    st.rerun()
            
            if st.session_state.get('show_reset_form'):
                st.markdown("---")
                st.markdown("### 🔐 Enter Reset Code & New Password")
                user_code = st.text_input("6-Digit Reset Code", key="user_code_input")
                new_pw = st.text_input("New Password", type="password", key="new_pw_input")
                confirm_pw = st.text_input("Confirm Password", type="password", key="confirm_pw_input")
                
                if st.button("✅ Reset Password", use_container_width=True):
                    if user_code == st.session_state.get('reset_code', ''):
                        if new_pw == confirm_pw:
                            if len(new_pw) >= 6:
                                import bcrypt
                                hashed_pw = bcrypt.hashpw(new_pw.encode('utf-8'), bcrypt.gensalt()).decode('utf-8')
                                db._patch("users", {"password_hash": hashed_pw}, {"email": st.session_state.reset_email})
                                st.success("✅ Password reset successfully! Please login with your new password.")
                                st.balloons()
                                st.session_state.show_forgot_password = False
                                st.session_state.show_reset_form = False
                                st.session_state.reset_code = None
                                st.session_state.reset_email = None
                                time.sleep(2)
                                st.rerun()
                            else:
                                st.warning("⚠️ Password must be at least 6 characters.")
                        else:
                            st.warning("⚠️ Passwords do not match.")
                    else:
                        st.error("❌ Invalid reset code.")
        

def sidebar_navigation():
    with st.sidebar:
        logo = get_logo()
        if logo:
            st.image(logo, width=220)
        st.markdown("""<div style="text-align: center; padding: 0.8rem 0; background: #4a4a4a; border-radius: 6px; margin-bottom: 1rem; border: 1px solid #666666;"><h3 style="color: #ffffff; margin: 0; font-size: 1.1rem; font-weight: 700;">CHURCHGATE GROUP</h3><p style="color: #cccccc; font-size: 0.7rem; margin: 0;">HRIS v5.0</p></div>""", unsafe_allow_html=True)
        if st.session_state.user:
            user = st.session_state.user
            # Load real data from employees table
            emp_id = user.get('employee_id', '')
            try:
                emp_result = db._get("employees", {"employee_id": emp_id})
                if emp_result and len(emp_result) > 0:
                    emp = emp_result[0]
                    user['name'] = f"{emp.get('first_name', '')} {emp.get('last_name', '')}".strip()
                    user['department'] = emp.get('department', user.get('department', ''))
                    user['position'] = emp.get('position', user.get('position', ''))
            except:
                pass
            
            initials = generate_initials(user['name'])
            db_pic = db.get_profile_picture(int(user.get('id', 0))) if user.get('id') else None
            
            if db_pic is not None:
                import base64
                profile_html = f'<img src="data:image/png;base64,{base64.b64encode(db_pic).decode()}" style="width: 40px; height: 40px; border-radius: 50%; object-fit: cover;">'
            else:
                profile_html = f'<div style="width: 40px; height: 40px; border-radius: 50%; background: #CC0000; display: flex; align-items: center; justify-content: center; font-weight: 700; font-size: 1rem; color: white;">{initials}</div>'
            st.markdown(f"""<div style="background: rgba(255,255,255,0.08); padding: 0.8rem; border-radius: 8px; margin-bottom: 1rem; border: 1px solid rgba(204, 0, 0, 0.2);"><div style="display: flex; align-items: center; gap: 0.6rem;">{profile_html}<div><p style="color: #333; margin: 0; font-weight: 600; font-size: 0.85rem;">{user['name']}</p><p style="color: #666; margin: 0; font-size: 0.7rem;">{user['role']} • {user.get('department', '')}</p></div></div></div>""", unsafe_allow_html=True)
        user_role = st.session_state.user['role'] if st.session_state.user else 'Employee'
        if user_role in ['Admin', 'HR Director']:
            menu_options = ["🏠 Employee Dashboard", "📊 Executive Dashboard", "👥 Employee Management", "📈 Performance & OKRs", "🚀 Promotions", "💼 Recruitment Hub", "🤖 AI Recruitment Agent", "📊 Reports & Analytics", "💬 Chat & Communications", "🎓 Training & Development", "🔔 Notifications", "📋 My Documents", "💡 Ideas Box", "📅 Calendar", "🎯 My Goals", "🔄 Requests Hub", "🌐 Directory", "📚 Knowledge Base", "🎉 Wellness & Perks", "🎓 LMS", "📋 Audit Log", "📊 Advanced Analytics", "👤 My Profile"]
        elif user_role == 'Manager':
            menu_options = ["🏠 Employee Dashboard", "💼 Recruitment Hub", "🤖 AI Recruitment Agent", "📈 Performance & OKRs", "💬 Chat & Communications", "🎓 Training & Development", "📋 My Documents", "💡 Ideas Box", "📅 Calendar", "🎯 My Goals", "🔄 Requests Hub", "🌐 Directory", "📚 Knowledge Base", "🎉 Wellness & Perks", "🎓 LMS", "👤 My Profile"]
        else:
            menu_options = ["🏠 Employee Dashboard", "📈 My Performance & OKRs", "💬 Chat & Communications", "🎓 Training & Development", "📋 My Documents", "💡 Ideas Box", "📅 Calendar", "🎯 My Goals", "🔄 Requests Hub", "🌐 Directory", "📚 Knowledge Base", "🎉 Wellness & Perks", "🎓 LMS", "👤 My Profile"]

        all_icons = ["house-fill", "speedometer2", "people-fill", "graph-up-arrow", "trophy-fill", "briefcase-fill", "robot", "file-earmark-bar-graph", "chat-dots-fill", "book-fill", "bell-fill", "folder-fill", "lightbulb-fill", "calendar-fill", "bullseye", "inbox-fill", "person-lines-fill", "book-half", "heart-fill", "mortarboard-fill", "shield-fill", "graph-up", "person-circle"]
        selected = option_menu(menu_title=None, options=menu_options, icons=all_icons[:len(menu_options)], menu_icon="cast", default_index=0, styles={"container": {"padding": "0!important", "background-color": "transparent"}, "icon": {"color": "#CC0000", "font-size": "16px"}, "nav-link": {"font-size": "13px", "text-align": "left", "margin": "3px 0", "color": "#333333", "--hover-color": "rgba(204, 0, 0, 0.1)", "border-radius": "6px"}, "nav-link-selected": {"background-color": "rgba(204, 0, 0, 0.15)", "color": "#CC0000", "border-left": "3px solid #CC0000", "font-weight": "700"}})
        st.markdown("---")
        if user_role in ['Admin', 'HR Director', 'Manager']:
            st.markdown("<p style='color: #CC0000; font-size: 0.75rem; margin-bottom: 0.5rem;'>⚡ QUICK ACTIONS</p>", unsafe_allow_html=True)
            c1, c2 = st.columns(2)
            with c1:
                if st.button("📝 Post Job", use_container_width=True, key="qa_job"):
                    st.session_state['navigate_to'] = "💼 Recruitment Hub"
                    st.rerun()
            with c2:
                if st.button("🤖 AI Screen", use_container_width=True, key="qa_ai"):
                    st.session_state['navigate_to'] = "🤖 AI Recruitment Agent"
                    st.rerun()
        if st.session_state.user:
            if st.button("🚪 Sign Out", use_container_width=True):
                st.session_state.user = None
                st.query_params.clear()
                st.rerun()
        st.markdown("---")
        st.markdown("### 🌐 Quick Links")
        col1, col2 = st.columns(2)
        with col1:
            st.markdown("[🏠 Home](/)", unsafe_allow_html=True)
        with col2:
            st.markdown("[💼 Careers](/Careers)", unsafe_allow_html=True)
        st.markdown("""<div style="text-align: center; padding: 0.5rem; margin-top: 1rem;"><p style="color: #888; font-size: 0.65rem; margin: 0;">© 2026 Churchgate Group</p><p style="color: #888; font-size: 0.65rem; margin: 0;">HRIS v5.0</p></div>""", unsafe_allow_html=True)
        return selected

def employee_dashboard():
    user = st.session_state.user
    user_name = user['name'] if user else 'Staff'
    user_id = user.get('employee_id', '') if user else ''
    user_dept = user.get('department', '') if user else ''
    user_position = user.get('position', '') if user else ''
    user_email = user.get('email', '') if user else ''
    
    # Override with real employee data from database
    try:
        emp_result = db._get("employees", {"employee_id": user_id})
        if emp_result and len(emp_result) > 0:
            emp = emp_result[0]
            user_name = f"{emp.get('first_name', '')} {emp.get('last_name', '')}".strip()
            user_dept = emp.get('department', user_dept)
            user_position = emp.get('position', user_position)
            user_email = emp.get('email', user_email)
            st.session_state.user['name'] = user_name
            st.session_state.user['department'] = user_dept
            st.session_state.user['position'] = user_position
    except:
        pass
    
    from datetime import timezone, timedelta
    wat = timezone(timedelta(hours=1))
    hour = datetime.now(wat).hour
    if hour < 12: greeting = "Good Morning"
    elif hour < 17: greeting = "Good Afternoon"
    else: greeting = "Good Evening"
    
    initials = generate_initials(user_name)
    
    # Load profile picture for greeting
    greeting_pic_html = f'<div style="width:60px;height:60px;border-radius:50%;background:linear-gradient(135deg,#CC0000,#e53e3e);display:flex;align-items:center;justify-content:center;font-size:1.5rem;font-weight:700;color:white;min-width:60px;overflow:hidden;">{initials}</div>'
    db_pic = db.get_profile_picture(int(user.get('id', 0))) if user.get('id') else None
    if db_pic is not None:
        import base64
        greeting_pic_html = f'<img src="data:image/png;base64,{base64.b64encode(db_pic).decode()}" style="width:60px;height:60px;border-radius:50%;object-fit:cover;min-width:60px;">'
    
    st.markdown(f"""
    <div class="greeting-header animate-fade-in">
        <div style="display:flex;align-items:center;gap:1rem;">
            {greeting_pic_html}
            <div>
                <h1>👋 {greeting}, {user_name}!</h1>
                <p>{user_position} • {user_dept} • ID: {user_id}</p>
            </div>
        </div>
        <div style="text-align:right;color:#ccc;font-size:0.85rem;">
            {datetime.now().strftime('%A, %B %d, %Y')}
        </div>
    </div>
    """, unsafe_allow_html=True)
    
    show_churchgate_mission()
    
    # ============ NEWS TICKER ============
    announcements = [
        "📢 Q2 Performance Reviews due by June 30, 2026",
        "🎓 New training courses available in Learning Hub",
        "🏢 Churchgate Tower 2 renovation project starting July",
        "🎉 Employee Appreciation Day coming up - June 15"
    ]
    ticker_text = " • ".join(announcements)
    st.markdown(f"""
    <div style="background:#d5d5d5;color:#333;padding:0.6rem 1.5rem;border-radius:8px;margin-bottom:1rem;overflow:hidden;white-space:nowrap;">
        <marquee behavior="scroll" direction="left" scrollamount="3">{ticker_text}</marquee>
    </div>
    """, unsafe_allow_html=True)
    
    # ============ TOP METRICS ROW ============
    # Real data calculations
    team_count = 0
    try:
        emp_df = db.get_all_employees()
        if not emp_df.empty:
            team_count = len(emp_df[emp_df['department'] == user_dept])
    except:
        team_count = 0
    
    # Performance score from appraisal
    perf_score = 0
    try:
        if 'self_assessments' in st.session_state and user_name in st.session_state.self_assessments:
            assessment = st.session_state.self_assessments[user_name]
            if assessment.get('hod_scores'):
                scores = [s for s in assessment['hod_scores'].values() if isinstance(s, (int, float))]
                if scores:
                    perf_score = sum(scores) / len(scores)
    except:
        pass
    
    # Leave balance (editable)
    if 'leave_balance' not in st.session_state:
        st.session_state.leave_balance = 18
    
    c1, c2, c3, c4, c5 = st.columns(5)
    with c1:
        st.markdown(f"""<div class="metric-card"><div class="metric-label">📊 Performance</div><div class="metric-value">{perf_score:.0f}%</div><small style="color:#38a169;">Latest appraisal</small></div>""", unsafe_allow_html=True)
    with c2:
        st.markdown(f"""<div class="metric-card"><div class="metric-label">👥 Team Members</div><div class="metric-value">{team_count}</div><small>{user_dept}</small></div>""", unsafe_allow_html=True)
    with c3:
        st.markdown(f"""<div class="metric-card"><div class="metric-label">🏖️ Leave Days</div><div class="metric-value">{st.session_state.leave_balance}</div><small style="color:#38a169;">Remaining</small></div>""", unsafe_allow_html=True)
    with c4:
        st.markdown(f"""<div class="metric-card"><div class="metric-label">🎯 KPIs</div><div class="metric-value">4</div><small>Strategic pillars</small></div>""", unsafe_allow_html=True)
    with c5:
        peer_rating = st.session_state.get('peer_rating', 0)
        st.markdown(f"""<div class="metric-card"><div class="metric-label">⭐ Rating</div><div class="metric-value">{peer_rating}</div><small>Peer recognition</small></div>""", unsafe_allow_html=True)
    
    # ============ QUICK ACTIONS ============
    st.markdown("---")
    st.markdown("### ⚡ Quick Actions")
    qc1, qc2, qc3, qc4, qc5 = st.columns(5)
    with qc1:
        if st.button("🎯 Set My KPIs", use_container_width=True):
            st.session_state['navigate_to'] = "📈 Performance & OKRs"
            st.rerun()
    with qc2:
        if st.button("📝 Self-Assessment", use_container_width=True):
            st.session_state['navigate_to'] = "📈 Performance & OKRs"
            st.rerun()
    with qc3:
        if st.button("👤 Edit Profile", use_container_width=True):
            st.session_state['navigate_to'] = "👤 My Profile"
            st.rerun()
    with qc4:
        if st.button("🏖️ Request Leave", use_container_width=True):
            st.success("Leave request feature coming soon!")
    with qc5:
        if st.button("📚 Training", use_container_width=True):
            st.session_state['navigate_to'] = "🎓 Training & Development"
            st.rerun()
    
    # ============ MAIN CONTENT ============
    st.markdown("---")
    col1, col2 = st.columns([2, 1])
    
    with col1:
        # KPI Progress
        st.subheader("🎯 My KPI Progress")
        try:
            perf_data = db.get_performance_data(user_name)
            if not perf_data.empty:
                # Sort by pillar order 1-2-3-4
                pillar_order = ['1. Occupancy & Revenue Growth', '2. Process Simplification', '3. Asset Reliability & Digitalization', '4. People & Culture']
                perf_data['sort_order'] = perf_data['pillar_name'].apply(lambda x: pillar_order.index(x) if x in pillar_order else 99)
                perf_data = perf_data.sort_values('sort_order')
                for _, row in perf_data.iterrows():
                    progress = row.get('progress', 0)
                    color = "#38a169" if progress >= 85 else "#d69e2e" if progress >= 65 else "#CC0000"
                    st.markdown(f"""
                    <div style="background:white;padding:0.7rem 1rem;border-radius:8px;margin-bottom:0.4rem;">
                        <div style="display:flex;justify-content:space-between;">
                            <span style="font-weight:600;font-size:0.9rem;">{row.get('pillar_name', 'Pillar')}</span>
                            <span style="color:{color};font-weight:700;">{int(progress)}%</span>
                        </div>
                        <div style="background:#e0e0e0;height:5px;border-radius:3px;margin-top:0.3rem;">
                            <div style="background:{color};width:{int(progress)}%;height:5px;border-radius:3px;"></div>
                        </div>
                    </div>
                    """, unsafe_allow_html=True)
            else:
                st.info("Set your KPIs in Performance & OKRs to track progress here.")
        except:
            st.info("KPI progress will appear here.")
        
        # Recognition Wall
        st.markdown("---")
        st.subheader("🌟 Recognition Wall")
        st.markdown("""
        <div style="background:white;padding:1rem;border-radius:8px;text-align:center;border:1px dashed #d69e2e;">
            <p style="color:#888;">🎉 Recognitions from colleagues will appear here!</p>
            <small style="color:#d69e2e;">Coming soon: Peer-to-peer recognition</small>
        </div>
        """, unsafe_allow_html=True)
    
    with col2:
        # Profile Card
        initials = generate_initials(user_name)
        st.markdown(f"""
        <div style="background:white;padding:1.5rem;border-radius:10px;text-align:center;box-shadow:0 2px 8px rgba(0,0,0,0.05);margin-bottom:1rem;">
            <div style="width:70px;height:70px;border-radius:50%;background:linear-gradient(135deg,#CC0000,#e53e3e);display:flex;align-items:center;justify-content:center;font-size:1.8rem;font-weight:700;color:white;margin:0 auto;">{initials}</div>
            <h3 style="margin-top:0.8rem;">{user_name}</h3>
            <p style="color:#666;">{user_position}</p>
            <p style="color:#888;font-size:0.85rem;">🏢 {user_dept}</p>
            <p style="color:#888;font-size:0.85rem;">📧 {user_email}</p>
        </div>
        """, unsafe_allow_html=True)
        
        # Team Mood
        st.markdown("---")
        st.subheader("😊 Team Mood Check")
        mood_col1, mood_col2, mood_col3 = st.columns(3)
        with mood_col1:
            if st.button("😊 Great", use_container_width=True, key="mood_great"):
                st.success("Thanks for sharing! 😊")
        with mood_col2:
            if st.button("😐 Okay", use_container_width=True, key="mood_okay"):
                st.success("Thanks for sharing! 😐")
        with mood_col3:
            if st.button("😤 Stressed", use_container_width=True, key="mood_stressed"):
                st.success("We hear you. HR is here to help. 😤")
        
        # Birthdays This Month
        st.markdown("---")
        st.subheader("🎂 Birthdays This Month")
        try:
            emp_df = db.get_all_employees()
            if not emp_df.empty:
                today = datetime.now()
                found_birthday = False
                for _, emp in emp_df.iterrows():
                    dob = emp.get('date_of_birth')
                    if dob and str(dob) != 'None' and str(dob) != 'nan':
                        try:
                            dob_date = pd.to_datetime(dob)
                            if dob_date.month == today.month:
                                found_birthday = True
                                st.markdown(f"🎂 **{emp['first_name']} {emp['last_name']}** — {dob_date.strftime('%B %d')} ({emp.get('department', '')})")
                        except:
                            pass
                if not found_birthday:
                    st.info("No birthdays this month.")
        except:
            pass
        
        # Work Anniversaries
        st.markdown("---")
        st.subheader("⭐ Work Anniversaries This Month")
        try:
            if not emp_df.empty:
                today = datetime.now()
                found_anniversary = False
                for _, emp in emp_df.iterrows():
                    join_date = emp.get('join_date')
                    if join_date and str(join_date) != 'None' and str(join_date) != 'nan':
                        try:
                            jd = pd.to_datetime(join_date)
                            if jd.month == today.month:
                                years = today.year - jd.year
                                if years > 0:
                                    found_anniversary = True
                                    st.markdown(f"⭐ **{emp['first_name']} {emp['last_name']}** — {years} year{'s' if years > 1 else ''} ({emp.get('department', '')})")
                        except:
                            pass
                if not found_anniversary:
                    st.info("No work anniversaries this month.")
        except:
            pass
        
        # Upcoming Holidays
        st.markdown("---")
        st.subheader("🏖️ Upcoming Holidays")
        st.markdown("📅 **Democracy Day** — June 12 (18 days)")
        st.markdown("📅 **Eid al-Adha** — June 17 (23 days)")
        
        # Wellness Tip
        st.markdown("---")
        st.subheader("💪 Wellness Tip")
        tips = [
            "Take a 5-minute stretch break every hour.",
            "Stay hydrated! Aim for 8 glasses of water today.",
            "Practice deep breathing for 2 minutes between meetings.",
            "Step away from your screen for lunch.",
            "A short walk boosts creativity and energy."
        ]
        import random
        st.info(f"💡 {random.choice(tips)}")
    # ============ ORGANIZATIONAL STRUCTURE ============
    st.markdown("---")
    st.subheader("📊 Organizational Structure")
    st.info("GMD → COO (All Depts) / VP Sales (Sales & Mkt) / GEA | Regions: Abuja & Lagos")
    
    # Simple org view for dashboard
    col1, col2, col3, col4 = st.columns(4)
    with col1:
        st.markdown("""
        <div style="background:white;padding:0.8rem;border-radius:8px;text-align:center;border-top:3px solid #CC0000;">
            <strong>Vinay Mahtani</strong><br>
            <small style="color:#CC0000;">GMD/CEO</small><br>
            <small style="color:#888;">Group-wide</small>
        </div>
        """, unsafe_allow_html=True)
    with col2:
        st.markdown("""
        <div style="background:white;padding:0.8rem;border-radius:8px;text-align:center;border-top:3px solid #e53e3e;">
            <strong>Jerome Das</strong><br>
            <small style="color:#e53e3e;">COO</small><br>
            <small style="color:#888;">All Departments</small>
        </div>
        """, unsafe_allow_html=True)
    with col3:
        st.markdown("""
        <div style="background:white;padding:0.8rem;border-radius:8px;text-align:center;border-top:3px solid #dd6b20;">
            <strong>Ahmed Karim</strong><br>
            <small style="color:#dd6b20;">VP Sales</small><br>
            <small style="color:#888;">Sales & Marketing</small>
        </div>
        """, unsafe_allow_html=True)
    with col4:
        st.markdown("""
        <div style="background:white;padding:0.8rem;border-radius:8px;text-align:center;border-top:3px solid #805ad5;">
            <strong>Partab Lalchandani</strong><br>
            <small style="color:#805ad5;">GEA</small><br>
            <small style="color:#888;">Group Advisor</small>
        </div>
        """, unsafe_allow_html=True)
    
    # Department heads from real data
    st.markdown("---")
    st.markdown("**Department Heads**")
    try:
        emp_df = db.get_all_employees()
        if not emp_df.empty:
            hod_df = emp_df[emp_df['grade'].isin(['Manager', 'HOD', 'C-Level'])]
            dept_list = hod_df['department'].unique()
            cols = st.columns(min(len(dept_list), 5))
            for i, dept in enumerate(dept_list[:10]):
                dept_hod = hod_df[hod_df['department'] == dept].head(1)
                if len(dept_hod) > 0:
                    hod = dept_hod.iloc[0]
                    with cols[i % 5]:
                        st.markdown(f"""
                        <div style="background:white;padding:0.6rem;border-radius:6px;text-align:center;margin-bottom:0.5rem;border:1px solid #e0e0e0;">
                            <strong style="font-size:0.85rem;">{hod['first_name']} {hod['last_name']}</strong><br>
                            <small style="color:#888;font-size:0.75rem;">{dept}</small>
                        </div>
                        """, unsafe_allow_html=True)
    except:
        pass
    
    with st.expander("📊 View Full Org Chart", expanded=False):
        st.markdown("### 🔗 Group Reporting Hierarchy")
        st.info("GMD → COO (All Depts) / VP Sales (Sales & Mkt) / GEA | Regions: Abuja & Lagos")
        
        labels = ['GMD', 'COO', 'VP Sales', 'GEA',
            'Technology (Abuja)', 'Technology (Lagos)',
            'Facility Mgmt (Abuja)', 'Facility Mgmt (Lagos)',
            'Engineering/MEP', 'HR', 'Accounts & Finance',
            'Sales & Marketing', 'Procurement', 'Security',
            'Legal', 'Operations']
        
        colors = ['#CC0000', '#e53e3e', '#dd6b20', '#805ad5',
            '#3182ce', '#3182ce', '#38a169', '#38a169',
            '#d53f8c', '#d69e2e', '#805ad5',
            '#dd6b20', '#2b6cb0', '#718096',
            '#e53e3e', '#319795']
        
        sources = [0, 0, 0, 1, 1, 1, 1, 1, 1, 1, 1, 1, 1, 2, 3]
        targets = [1, 2, 3, 4, 5, 6, 7, 8, 9, 10, 12, 13, 14, 11, 15]
        values = [1, 1, 1, 1, 1, 1, 1, 1, 1, 1, 1, 1, 1, 1, 1]
        
        for i in range(4, 16):
            sources.append(i)
            targets.append(16)
            values.append(1)
        
        sources += [16, 16, 17, 17]
        targets += [17, 18, 18, 19]
        values += [10, 5, 12, 30]
        
        labels += ['Heads of Department', 'Sr. Managers', 'Managers', 'Team Leads', 'Team Members']
        colors += ['#FF6B35', '#38a169', '#d69e2e', '#2b6cb0', '#718096']
        
        fig = go.Figure(data=[go.Sankey(
            node=dict(pad=20, thickness=18, label=labels, color=colors),
            link=dict(source=sources, target=targets, value=values,
                color=['rgba(204,0,0,0.2)'] * len(sources))
        )])
        fig.update_layout(height=500)
        st.plotly_chart(fig, use_container_width=True)
        
        st.markdown("---")
        st.markdown("### 👥 Department Heads")
        try:
            emp_df = db.get_all_employees()
            if not emp_df.empty:
                hod_df = emp_df[emp_df['grade'].isin(['Manager', 'HOD', 'C-Level'])]
                hod_data = []
                for _, row in hod_df.iterrows():
                    hod_data.append({
                        'Department': row['department'],
                        'Name': f"{row['first_name']} {row['last_name']}",
                        'Position': row['position']
                    })
                if hod_data:
                    st.dataframe(pd.DataFrame(hod_data), use_container_width=True, hide_index=True)
        except:
            pass
    # ============ PEER RATING ============
    st.markdown("---")
    st.subheader("⭐ Rate a Colleague")
    
    # Get colleagues list
    try:
        emp_df = db.get_all_employees()
        if not emp_df.empty:
            colleague_list = [f"{row['first_name']} {row['last_name']}" for _, row in emp_df.iterrows() if f"{row['first_name']} {row['last_name']}" != user_name]
        else:
            colleague_list = ["No colleagues available"]
    except:
        colleague_list = ["No colleagues available"]
    
    col1, col2, col3 = st.columns([2, 1, 1])
    with col1:
        rate_colleague = st.selectbox("👤 Colleague", ["Select..."] + colleague_list, key="rate_colleague")
    with col2:
        rating = st.selectbox("⭐ Rating", [5, 4, 3, 2, 1], format_func=lambda x: "⭐" * x, key="rating_value")
    with col3:
        is_anonymous = st.checkbox("Anonymous", value=False, key="rate_anonymous")
    
    rating_comment = st.text_area("💬 Comment (Optional)", placeholder="Why are you giving this rating?", key="rate_comment")
    
    if st.button("🌟 Submit Rating", use_container_width=True):
        if rate_colleague != "Select...":
            db._post("peer_ratings", {
                "rated_by": "Anonymous" if is_anonymous else user_name,
                "rated_user": rate_colleague,
                "rating": rating,
                "comment": rating_comment,
                "is_anonymous": is_anonymous,
                "created_at": datetime.now().strftime('%Y-%m-%d %H:%M')
            })
            st.success(f"🌟 Rating submitted for {rate_colleague}!")
            st.balloons()
            st.rerun()
        else:
            st.warning("⚠️ Please select a colleague to rate.")
    
    # Show recent ratings received
    st.markdown("---")
    st.markdown("### 📊 My Recent Ratings")
    try:
        all_ratings = db._get("peer_ratings")
        my_ratings = [r for r in all_ratings if r.get('rated_user') == user_name] if all_ratings else []
        if my_ratings:
            avg_rating = sum(r.get('rating', 0) for r in my_ratings) / len(my_ratings)
            stars = "⭐" * round(avg_rating)
            st.markdown(f"**Average Rating:** {stars} ({avg_rating:.1f}/5) from {len(my_ratings)} review{'s' if len(my_ratings) > 1 else ''}")
            
            # Show latest 3 ratings
            for r in sorted(my_ratings, key=lambda x: x.get('created_at', ''), reverse=True)[:3]:
                rater = r.get('rated_by', 'Unknown')
                stars_display = "⭐" * r.get('rating', 0)
                comment = r.get('comment', '')
                st.markdown(f"{stars_display} — {'Anonymous' if r.get('is_anonymous') else rater}")
                if comment:
                    st.caption(f"\"{comment}\"")
        else:
            st.info("No ratings yet.")
    except:
        pass
    
    # Update the hardcoded peer rating metric
    try:
        all_ratings = db._get("peer_ratings")
        my_ratings = [r for r in all_ratings if r.get('rated_user') == user_name] if all_ratings else []
        if my_ratings:
            st.session_state['peer_rating'] = round(sum(r.get('rating', 0) for r in my_ratings) / len(my_ratings), 1)
        else:
            st.session_state['peer_rating'] = 0
    except:
        st.session_state['peer_rating'] = 0
    
    # ============ UPCOMING TRAINING ============
    st.markdown("---")
    st.subheader("🎓 Upcoming Training & Webinars")
    
    try:
        training_data = db._get("training_courses")
        if training_data:
            upcoming = sorted([t for t in training_data if t.get('start_date', '') >= datetime.now().strftime('%Y-%m-%d')], 
                            key=lambda x: x.get('start_date', ''))[:3]
            
            if upcoming:
                tc1, tc2, tc3 = st.columns(3)
                for i, course in enumerate(upcoming):
                    col = [tc1, tc2, tc3][i]
                    with col:
                        st.markdown(f"""
                        <div class="metric-card">
                            <h4>{course.get('title', 'Training')[:40]}</h4>
                            <p style="font-size:0.8rem;">{course.get('description', '')[:80]}</p>
                            <small style="color:#CC0000;">📅 {course.get('start_date', '')}</small>
                            <br><small style="color:#888;">{course.get('provider', '')} • {course.get('format', '')}</small>
                        </div>
                        """, unsafe_allow_html=True)
            else:
                st.info("No upcoming training sessions. Check the Training & Development tab.")
        else:
            st.info("Training catalog loading...")
    except:
        st.info("Training data will appear here.")
    
    if st.button("📚 View All Training Courses", use_container_width=True):
        st.session_state['navigate_to'] = "🎓 Training & Development"
        st.rerun()




def executive_dashboard():
    show_churchgate_mission()
    st.markdown("""<div class="churchgate-header"><h1>📊 Executive Dashboard</h1><p>Corporate Strategy 2026-2027 | Real-Time Group Performance Intelligence</p></div>""", unsafe_allow_html=True)
    
    # ============ LOAD REAL DATA ============
    total_employees = 0
    active_employees = 0
    departments = 0
    open_positions = 0
    male_count = 0
    female_count = 0
    emp_df = pd.DataFrame()
    
    try:
        emp_df = db.get_all_employees()
        if not emp_df.empty:
            total_employees = len(emp_df)
            active_employees = len(emp_df[emp_df['status'] == 'Active'])
            departments = len(emp_df['department'].unique())
            if 'gender' in emp_df.columns:
                male_count = len(emp_df[emp_df['gender'].str.lower() == 'male'])
                female_count = len(emp_df[emp_df['gender'].str.lower() == 'female'])
    except:
        pass
    
    try:
        all_reqs = db.get_all_job_requisitions()
        if all_reqs:
            open_positions = len([r for r in all_reqs if r.get('status') == 'Approved - Live'])
        candidates_df = db.get_all_candidates()
        if not candidates_df.empty:
            open_positions += len(candidates_df[candidates_df['status'] == 'New'])
    except:
        pass
    
    escalated_count = 0
    try:
        if 'self_assessments' in st.session_state:
            escalated_count = len([v for v in st.session_state.self_assessments.values() if v.get('acceptance') == 'Rejected'])
    except:
        pass
    
    # ============ LOAD/INITIALIZE PORTFOLIO METRICS ============
    if 'portfolio_metrics' not in st.session_state:
        st.session_state.portfolio_metrics = {
            'occupancy': 87, 'revenue': 94, 'rating': 4.2,
            'portfolio_data': {
                'World Trade Center Abuja': {'occupancy': 87, 'revenue': 94},
                'Churchgate Tower 1, Lagos': {'occupancy': 92, 'revenue': 98},
                'Churchgate Tower 2, Lagos': {'occupancy': 85, 'revenue': 88},
                'Churchgate Plaza, Abuja': {'occupancy': 78, 'revenue': 82},
                'Warehouses': {'occupancy': 95, 'revenue': 97},
                'Ocean Terrace': {'occupancy': 90, 'revenue': 91}
            }
        }
        try:
            saved = db.get_portfolio_metrics()
            if saved:
                if 'occupancy' in saved:
                    st.session_state.portfolio_metrics['occupancy'] = int(float(saved['occupancy']))
                if 'revenue' in saved:
                    st.session_state.portfolio_metrics['revenue'] = int(float(saved['revenue']))
                if 'rating' in saved:
                    st.session_state.portfolio_metrics['rating'] = float(saved['rating'])
                for prop in st.session_state.portfolio_metrics['portfolio_data']:
                    occ_key = f"{prop}_occupancy"
                    rev_key = f"{prop}_revenue"
                    if occ_key in saved:
                        st.session_state.portfolio_metrics['portfolio_data'][prop]['occupancy'] = int(float(saved[occ_key]))
                    if rev_key in saved:
                        st.session_state.portfolio_metrics['portfolio_data'][prop]['revenue'] = int(float(saved[rev_key]))
        except:
            pass
    
    metrics = st.session_state.portfolio_metrics
    
    
    # ============ ALERTS BAR ============
    alerts = []
    if escalated_count > 0:
        alerts.append(f"🚨 {escalated_count} escalated appraisal(s) need attention")
    if open_positions > 0:
        alerts.append(f"📢 {open_positions} open positions to fill")
    
    if alerts:
        alert_html = " | ".join(alerts)
        st.markdown(f"""
        <div style="background:#fff3cd;padding:0.8rem 1.5rem;border-radius:8px;margin-bottom:1rem;border-left:4px solid #d69e2e;">
            <strong>⚠️ Executive Alerts:</strong> {alert_html}
        </div>
        """, unsafe_allow_html=True)
    
    # ============ TOP KPI CARDS - ALL REAL DATA ============
    c1, c2, c3, c4, c5, c6 = st.columns(6)
    with c1:
        st.markdown(f"""<div class="metric-card"><div class="metric-label">👥 Total Employees</div><div class="metric-value">{total_employees}</div><small style="color:#38a169;">{active_employees} active</small></div>""", unsafe_allow_html=True)
    with c2:
        st.markdown(f"""<div class="metric-card"><div class="metric-label">🏢 Departments</div><div class="metric-value">{departments}</div><small style="color:#38a169;">2 regions</small></div>""", unsafe_allow_html=True)
    with c3:
        st.markdown(f"""<div class="metric-card"><div class="metric-label">📋 Open Positions</div><div class="metric-value">{open_positions}</div><small style="color:#CC0000;">Active hiring</small></div>""", unsafe_allow_html=True)
    with c4:
        st.markdown(f"""<div class="metric-card"><div class="metric-label">🏠 Occupancy</div><div class="metric-value">{metrics['occupancy']}%</div><small style="color:#38a169;">Portfolio avg</small></div>""", unsafe_allow_html=True)
    with c5:
        st.markdown(f"""<div class="metric-card"><div class="metric-label">💰 Revenue</div><div class="metric-value">{metrics['revenue']}%</div><small style="color:#d69e2e;">vs budget</small></div>""", unsafe_allow_html=True)
    with c6:
        st.markdown(f"""<div class="metric-card"><div class="metric-label">⭐ Rating</div><div class="metric-value">{metrics['rating']}</div><small style="color:#38a169;">Tenant satisfaction</small></div>""", unsafe_allow_html=True)
    
    # ============ ADMIN UPDATE METRICS ============
    is_admin = st.session_state.user['role'] in ['Admin', 'HR Director'] if st.session_state.user else False
    is_sr_mgmt = st.session_state.user.get('department') == 'Senior Management' if st.session_state.user else False
    
    # ============ CELEBRATION EMAIL BLAST ============
    if is_admin:
        today = datetime.now()
        has_celebration = False
        try:
            emp_df = db.get_all_employees()
            if not emp_df.empty:
                for _, emp in emp_df.iterrows():
                    dob = emp.get('date_of_birth')
                    if dob and str(dob) != 'None' and str(dob) != 'nan':
                        try:
                            dob_date = pd.to_datetime(dob)
                            if dob_date.month == today.month and dob_date.day == today.day:
                                has_celebration = True
                                break
                        except:
                            pass
                    jd = emp.get('join_date')
                    if jd and str(jd) != 'None' and str(jd) != 'nan':
                        try:
                            jd_date = pd.to_datetime(jd)
                            if jd_date.month == today.month and jd_date.day == today.day:
                                years = today.year - jd_date.year
                                if years > 0:
                                    has_celebration = True
                                    break
                        except:
                            pass
        except:
            pass
        
        if has_celebration:
            st.markdown("""
            <div style="background: linear-gradient(135deg, #fff5f5, #fffbf0); padding: 0.8rem 1.5rem; border-radius: 8px; margin-bottom: 1rem; border: 2px solid #CC0000;">
                <strong>🎉 Today's Celebrations!</strong> There are birthdays or work anniversaries today. Send celebration emails to all employees.
            </div>
            """, unsafe_allow_html=True)
            
            col_btn1, col_btn2 = st.columns([1, 3])
            with col_btn1:
                if st.button("📧 Send Celebration Emails", use_container_width=True, type="primary"):
                    with st.spinner("📧 Sending celebration emails to all employees..."):
                        bdays, annivs, msg = send_celebration_emails()
                        if bdays > 0 or annivs > 0:
                            st.success(f"✅ Sent! {bdays} birthday(s) and {annivs} anniversary(s). {msg}")
                            st.balloons()
                        else:
                            st.info(msg)
            with col_btn2:
                st.caption("Emails are sent to all employees at once. Use this in the morning (7:30 AM recommended).")
    
    if is_admin or is_sr_mgmt:
        with st.expander("⚙️ Update Portfolio Metrics (Admin)", expanded=False):
            st.markdown("### Update Key Metrics")
            c1, c2, c3 = st.columns(3)
            with c1:
                new_occupancy = st.slider("Occupancy Rate %", 0, 100, metrics['occupancy'])
            with c2:
                new_revenue = st.slider("Revenue vs Budget %", 0, 100, metrics['revenue'])
            with c3:
                new_rating = st.slider("Tenant Rating /5", 1.0, 5.0, metrics['rating'], 0.1)
            
            st.markdown("### Update Portfolio Properties")
            for prop in metrics['portfolio_data']:
                c1, c2 = st.columns(2)
                with c1:
                    metrics['portfolio_data'][prop]['occupancy'] = st.slider(f"{prop} - Occupancy %", 0, 100, metrics['portfolio_data'][prop]['occupancy'], key=f"occ_{prop}")
                with c2:
                    metrics['portfolio_data'][prop]['revenue'] = st.slider(f"{prop} - Revenue %", 0, 100, metrics['portfolio_data'][prop]['revenue'], key=f"rev_{prop}")
            
            if st.button("💾 Update All Metrics", use_container_width=True):
                st.session_state.portfolio_metrics['occupancy'] = new_occupancy
                st.session_state.portfolio_metrics['revenue'] = new_revenue
                st.session_state.portfolio_metrics['rating'] = new_rating
                try:
                    db.save_portfolio_metric('occupancy', new_occupancy)
                    db.save_portfolio_metric('revenue', new_revenue)
                    db.save_portfolio_metric('rating', new_rating)
                    for prop in metrics['portfolio_data']:
                        db.save_portfolio_metric(f"{prop}_occupancy", metrics['portfolio_data'][prop]['occupancy'])
                        db.save_portfolio_metric(f"{prop}_revenue", metrics['portfolio_data'][prop]['revenue'])
                except:
                    pass
                st.success("✅ Metrics updated & saved!")
                st.rerun()
    
    # ============ QUICK ACTIONS ============
    st.markdown("---")
    quick_col1, quick_col2, quick_col3, quick_col4 = st.columns(4)
    with quick_col1:
        if st.button("📝 Post Job", use_container_width=True):
            st.session_state['navigate_to'] = "💼 Recruitment Hub"
            st.rerun()
    with quick_col2:
        if st.button("👥 View Employees", use_container_width=True):
            st.session_state['navigate_to'] = "👥 Employee Management"
            st.rerun()
    with quick_col3:
        if st.button("📈 Performance", use_container_width=True):
            st.session_state['navigate_to'] = "📈 Performance & OKRs"
            st.rerun()
    with quick_col4:
        if st.button("🤖 AI Screening", use_container_width=True):
            st.session_state['navigate_to'] = "🤖 AI Recruitment Agent"
            st.rerun()
    
    # ============ ROW 1: PORTFOLIO + DEPARTMENT DISTRIBUTION ============
    st.markdown("---")
    col1, col2 = st.columns(2)
    
    with col1:
        st.subheader("🏢 Portfolio Performance")
        props = list(metrics['portfolio_data'].keys())
        occ_vals = [metrics['portfolio_data'][p]['occupancy'] for p in props]
        rev_vals = [metrics['portfolio_data'][p]['revenue'] for p in props]
        
        portfolio_data = pd.DataFrame({'Property': props, 'Occupancy %': occ_vals, 'Revenue %': rev_vals})
        fig = go.Figure()
        fig.add_trace(go.Bar(name='Occupancy %', x=portfolio_data['Property'], y=portfolio_data['Occupancy %'], marker_color='#CC0000', text=portfolio_data['Occupancy %'], textposition='outside'))
        fig.add_trace(go.Bar(name='Revenue %', x=portfolio_data['Property'], y=portfolio_data['Revenue %'], marker_color='#4a4a4a', text=portfolio_data['Revenue %'], textposition='outside'))
        fig.update_layout(height=350, barmode='group', margin=dict(t=20))
        st.plotly_chart(fig, use_container_width=True)
    
    with col2:
        st.subheader("👥 Department Distribution")
        if not emp_df.empty:
            dept_counts = emp_df['department'].value_counts()
            fig2 = px.pie(values=dept_counts.values, names=dept_counts.index, hole=0.4,
                         color_discrete_sequence=['#CC0000', '#3182ce', '#38a169', '#d69e2e', '#805ad5', '#dd6b20', '#2b6cb0', '#718096', '#e53e3e', '#319795', '#d53f8c'])
            fig2.update_layout(height=350, margin=dict(t=20))
            st.plotly_chart(fig2, use_container_width=True)
    
    # ============ ROW 2: STRATEGIC PILLARS + GENDER DIVERSITY ============
    st.markdown("---")
    col1, col2 = st.columns(2)
    
    with col1:
        st.subheader("🎯 Strategic Pillars 2026-2027")
        
        # Load real performance data
        pillar_progress = {
            '1. Occupancy & Revenue Growth': {'progress': 0, 'color': '#CC0000'},
            '2. Process Simplification': {'progress': 0, 'color': '#38a169'},
            '3. Asset Reliability & Digitalization': {'progress': 0, 'color': '#3182ce'},
            '4. People & Culture': {'progress': 0, 'color': '#d69e2e'},
        }
        
        try:
            perf_data = db.get_performance_data()
            if not perf_data.empty:
                for _, row in perf_data.iterrows():
                    pillar = row.get('pillar_name', '')
                    progress = row.get('progress', 0)
                    for key in pillar_progress:
                        if key[:2] in pillar or pillar[:2] in key:
                            pillar_progress[key]['progress'] = int(progress)
                            break
        except:
            pass
        
        for name, data in pillar_progress.items():
            st.markdown(f"""
            <div style="background:white;padding:0.8rem 1rem;border-radius:8px;margin-bottom:0.5rem;border-left:4px solid {data['color']};">
                <div style="display:flex;justify-content:space-between;align-items:center;">
                    <span style="font-weight:600;font-size:0.9rem;">{name}</span>
                    <span style="color:{data['color']};font-weight:700;">{data['progress']}%</span>
                </div>
                <div style="background:#e0e0e0;height:6px;border-radius:3px;margin-top:0.4rem;">
                    <div style="background:{data['color']};width:{data['progress']}%;height:6px;border-radius:3px;"></div>
                </div>
            </div>
            """, unsafe_allow_html=True)
    
    with col2:
        st.subheader("🌍 Gender Diversity (Real-Time)")
        if male_count > 0 or female_count > 0:
            gender_data = pd.DataFrame({'Gender': ['Male', 'Female'], 'Count': [male_count, female_count]})
            fig3 = px.pie(gender_data, values='Count', names='Gender', hole=0.6, color_discrete_sequence=['#3182ce', '#CC0000'])
            fig3.update_layout(height=300, margin=dict(t=20))
            st.plotly_chart(fig3, use_container_width=True)
            
            st.markdown(f"""
            <div style="text-align:center;margin-top:-5rem;position:relative;z-index:1;">
                <span style="font-size:2rem;font-weight:900;color:#CC0000;">{total_employees}</span><br>
                <small style="color:#888;">Total Workforce</small>
            </div>
            """, unsafe_allow_html=True)
        else:
            st.info("Gender data will appear as employees are updated with gender information.")
    
    # ============ ROW 3: REVENUE TREND + UPCOMING DEADLINES ============
    st.markdown("---")
    col1, col2 = st.columns(2)
    
    with col1:
        st.subheader("📈 Revenue Trend (₦ Billions)")
        
        if 'revenue_trend' not in st.session_state:
            st.session_state.revenue_trend = {
                'Jan': 1.8, 'Feb': 2.0, 'Mar': 2.1, 'Apr': 2.2, 'May': 2.3, 'Jun': 2.5
            }
        
        months = list(st.session_state.revenue_trend.keys())
        actual = list(st.session_state.revenue_trend.values())
        target = [v * 1.05 for v in actual]
        
        fig4 = go.Figure()
        fig4.add_trace(go.Scatter(x=months, y=actual, mode='lines+markers', name='Actual', line=dict(color='#CC0000', width=3), fill='tozeroy', fillcolor='rgba(204,0,0,0.1)'))
        fig4.add_trace(go.Scatter(x=months, y=target, mode='lines', name='Target', line=dict(color='#4a4a4a', width=2, dash='dash')))
        fig4.update_layout(height=300, margin=dict(t=20))
        st.plotly_chart(fig4, use_container_width=True)
        
        if is_admin or is_sr_mgmt:
            with st.expander("✏️ Edit Revenue Data"):
                new_trend = {}
                c1, c2, c3 = st.columns(3)
                for i, month in enumerate(months):
                    col = [c1, c2, c3][i % 3]
                    with col:
                        new_trend[month] = st.number_input(f"{month} (₦B)", value=float(st.session_state.revenue_trend[month]), step=0.1, key=f"rev_{month}")
                if st.button("💾 Update Revenue Trend", use_container_width=True):
                    st.session_state.revenue_trend = new_trend
                    st.success("✅ Revenue trend updated!")
                    st.rerun()
    
    with col2:
        st.subheader("📅 Upcoming Deadlines")
        
        if 'custom_deadlines' not in st.session_state:
            st.session_state.custom_deadlines = []
        
        real_deadlines = []
        
        try:
            all_reqs = db.get_all_job_requisitions()
            if all_reqs:
                for r in all_reqs:
                    if r.get('status') == 'Approved - Live' and r.get('closing'):
                        real_deadlines.append({
                            "task": f"Applications Close: {r.get('title', 'Job')}",
                            "date": r['closing'],
                            "priority": "High"
                        })
        except:
            pass
        
        if st.session_state.get('appraisal_cycle_active') and st.session_state.get('appraisal_end'):
            real_deadlines.append({
                "task": f"Appraisal Cycle Ends: {st.session_state.get('appraisal_cycle_name', '')}",
                "date": st.session_state['appraisal_end'],
                "priority": "Medium"
            })
        
        all_deadlines = real_deadlines + st.session_state.custom_deadlines
        try:
            all_deadlines.sort(key=lambda x: x['date'])
        except:
            pass
        
        if all_deadlines:
            for d in all_deadlines[:8]:
                try:
                    due = datetime.strptime(d['date'], '%Y-%m-%d')
                    days_left = (due - datetime.now()).days
                    days_str = f"{days_left} days left" if days_left > 0 else "🔴 OVERDUE"
                    urgency_color = "#CC0000" if days_left < 30 else "#d69e2e" if days_left < 60 else "#38a169"
                except:
                    days_str = d['date']
                    urgency_color = "#718096"
                
                st.markdown(f"""
                <div style="background:white;padding:0.7rem 1rem;border-radius:6px;margin-bottom:0.4rem;display:flex;justify-content:space-between;align-items:center;border-left:3px solid {urgency_color};">
                    <div>
                        <strong style="font-size:0.9rem;">{d['task']}</strong>
                        <span style="background:{'#CC0000' if d['priority']=='High' else '#d69e2e'};color:white;padding:0.1rem 0.5rem;border-radius:10px;font-size:0.7rem;margin-left:0.5rem;">{d['priority']}</span>
                    </div>
                    <small style="color:{urgency_color};font-weight:600;">{days_str}</small>
                </div>
                """, unsafe_allow_html=True)
        else:
            st.info("No upcoming deadlines.")
        
        if is_admin or is_sr_mgmt:
            with st.expander("✏️ Add Custom Deadline"):
                with st.form("add_deadline"):
                    c1, c2 = st.columns(2)
                    with c1:
                        new_task = st.text_input("Task Name")
                        new_priority = st.selectbox("Priority", ["High", "Medium", "Low"])
                    with c2:
                        new_date = st.date_input("Due Date")
                    if st.form_submit_button("➕ Add Deadline"):
                        if new_task:
                            st.session_state.custom_deadlines.append({
                                "task": new_task,
                                "date": new_date.strftime('%Y-%m-%d'),
                                "priority": new_priority
                            })
                            st.success("✅ Deadline added!")
                            st.rerun()
                
                if st.session_state.custom_deadlines:
                    st.markdown("**Custom Deadlines:**")
                    for i, cd in enumerate(st.session_state.custom_deadlines):
                        c1, c2 = st.columns([4, 1])
                        with c1:
                            st.markdown(f"- {cd['task']} ({cd['date']})")
                        with c2:
                            if st.button("🗑️", key=f"del_dead_{i}"):
                                st.session_state.custom_deadlines.pop(i)
                                st.rerun()
    
    # ============ ROW 4: RECENT ACTIVITY ============
    st.markdown("---")
    st.subheader("🕐 Recent Activity Feed")
    
    # Pull real activities from audit trail and system state
    activities = []
    
    # From audit trail (last 5 entries)
    if 'audit_trail' in st.session_state and st.session_state.audit_trail:
        for entry in st.session_state.audit_trail[-3:]:
            activities.append({
                "icon": "📋",
                "action": entry.get('action', 'Activity'),
                "detail": entry.get('details', ''),
                "time": entry.get('timestamp', 'Recent')
            })
    
    # Current system state
    activities.append({
        "icon": "👥",
        "action": "Workforce",
        "detail": f"{total_employees} employees across {departments} departments",
        "time": "Current"
    })
    
    activities.append({
        "icon": "📋",
        "action": "Open Positions",
        "detail": f"{open_positions} positions actively hiring",
        "time": "Current"
    })
    
    if st.session_state.get('appraisal_cycle_active'):
        activities.append({
            "icon": "📊",
            "action": "Appraisal Cycle Active",
            "detail": st.session_state.get('appraisal_cycle_name', 'Appraisal in progress'),
            "time": "Active"
        })
    
    for activity in activities[-5:]:
        st.markdown(f"""
        <div style="background:white;padding:0.8rem;border-radius:8px;margin-bottom:0.4rem;display:flex;align-items:center;gap:1rem;border-left:3px solid #CC0000;">
            <span style="font-size:1.5rem;">{activity['icon']}</span>
            <div style="flex:1;"><strong>{activity['action']}</strong><br><small>{activity['detail']}</small></div>
            <small style="color:#888;">{activity['time']}</small>
        </div>
        """, unsafe_allow_html=True)

def employee_management():
    st.markdown("""<div class="churchgate-header"><h1>👥 Employee Management</h1><p>Comprehensive workforce management | Real-time Data | Churchgate Group</p></div>""", unsafe_allow_html=True)
    
    user_role = st.session_state.user['role'] if st.session_state.user else 'Team Member'
    user_dept = st.session_state.user.get('department', '') if st.session_state.user else ''
    is_admin = user_role in ['Admin', 'HR Director'] or user_dept == 'Senior Management'
    
    # Subsidiary options by region
    SUBSIDIARY_OPTIONS = {
        'Abuja': ['World Trade Center(WTC)', 'Agroline Ventures Limited'],
        'Lagos': [
            'First Continental Properties Limited', 'R. B Properties Limited', 
            'Churchgate Nigeria Limited', 'Aba Textile Mills PLC',
            'Associated Textile Manufacturing Company Limited',
            'Food & Confectionery Products (Nig.) Limited',
            'First Spinners PLC', 'HotelInvest & Resorts Limited',
            'International Textile Industries (Nig.) Limited',
            'Intercott Limited', 'Ocean Fisheries (Nig.) Limited',
            'Platinum Travel Limited', 'Reliance Mills Limited',
            'Vineyard Designs Nig. Limited'
        ],
        'Aba': ['Aba Textile Mills PLC']
    }
    
    def load_employees():
        try:
            df = db.get_all_employees()
            if df is None or df.empty:
                df = pd.DataFrame(columns=['employee_id', 'first_name', 'last_name', 'email', 'phone', 'department', 'position', 'grade', 'employment_type', 'join_date', 'status', 'region', 'subsidiary', 'reports_to'])
            return df
        except:
            return pd.DataFrame(columns=['employee_id', 'first_name', 'last_name', 'email', 'phone', 'department', 'position', 'grade', 'employment_type', 'join_date', 'status', 'region', 'subsidiary', 'reports_to'])
    
    employees_df = load_employees()
    
    # Build list of all employee names for Reports To dropdown
    all_employee_names = []
    if not employees_df.empty:
        all_employee_names = sorted(list((employees_df['first_name'].astype(str) + ' ' + employees_df['last_name'].astype(str)).unique()))
    
    dept_colors = {
        'Senior Management': '#CC0000', 'Technology Group': '#3182ce', 'Facility Management': '#38a169',
        'Human Resources': '#d69e2e', 'Accounts & Finance': '#805ad5', 'Sales & Marketing': '#dd6b20',
        'Procurement': '#2b6cb0', 'Security': '#718096', 'Legal': '#e53e3e', 'Operations': '#319795',
        'Engineering': '#d53f8c'
    }
    
    tab1, tab2, tab3, tab4, tab5, tab6, tab7, tab8 = st.tabs([
        "📋 Directory", "➕ Add Employee", "📤 Bulk Upload", 
        "🔑 Generate Logins", "🏢 Departments", "📊 Org Chart", "📈 Demographics", "📥 Export"
    ])
    
    # ============ TAB 1: DIRECTORY ============
    with tab1:
        st.subheader("📋 Employee Directory")
        
        total_emp = len(employees_df)
        active_emp = len(employees_df[employees_df['status'] == 'Active']) if not employees_df.empty else 0
        new_this_month = 0
        try:
            current_month = datetime.now().month
            for _, emp in employees_df.iterrows():
                jd = emp.get('join_date')
                if jd and pd.notna(jd):
                    if hasattr(jd, 'month') and jd.month == current_month:
                        new_this_month += 1
        except:
            pass
        
        c1, c2, c3, c4 = st.columns(4)
        with c1: st.metric("👥 Total", total_emp)
        with c2: st.metric("✅ Active", active_emp)
        with c3: st.metric("🏢 Departments", len(employees_df['department'].unique()) if not employees_df.empty else 0)
        with c4: st.metric("🆕 New This Month", new_this_month)
        
        st.markdown("---")
        st.markdown("### 🎂 Birthdays & Anniversaries This Month")
        bday_col1, bday_col2 = st.columns(2)
        with bday_col1:
            st.markdown("**🎂 Birthdays:** Chika Ikwuegbu (May 13), Francis Asuquo (May 19), Rhoda Ajibola (May 25), Alice Agbo (May 28)")
        with bday_col2:
            st.markdown("**⭐ Anniversaries:** Augustine Oleh (4 yrs), Shem Waziri (3 yrs), Charles Okere (7 yrs), Chika Ikwuegbu (3 yrs)")
        
        st.markdown("---")
        
        # Search & Filters
        c1, c2, c3, c4, c5 = st.columns([2, 1, 1, 1, 1])
        with c1:
            search = st.text_input("🔍 Search", placeholder="Name, ID, email, department, position...")
        with c2:
            all_depts = ['All'] + sorted(list(employees_df['department'].dropna().unique())) if not employees_df.empty else ['All']
            dept_filter = st.selectbox("Department", all_depts)
        with c3:
            all_regions = ['All'] + sorted(list(employees_df['region'].dropna().unique())) if not employees_df.empty and 'region' in employees_df.columns else ['All']
            region_filter = st.selectbox("Region", all_regions)
        with c4:
            all_grades = ['All'] + sorted(list(employees_df['grade'].dropna().unique())) if not employees_df.empty else ['All']
            grade_filter = st.selectbox("Grade", all_grades)
        with c5:
            status_filter = st.selectbox("Status", ["All", "Active", "On Leave", "Probation", "Terminated", "Archived", "Inactive"])
        
        filtered_df = employees_df.copy()
        if not filtered_df.empty:
            if search:
                s = search.lower()
                filtered_df = filtered_df[
                    filtered_df['first_name'].str.lower().str.contains(s, na=False) |
                    filtered_df['last_name'].str.lower().str.contains(s, na=False) |
                    filtered_df['employee_id'].str.lower().str.contains(s, na=False) |
                    filtered_df['email'].str.lower().str.contains(s, na=False) |
                    filtered_df['department'].str.lower().str.contains(s, na=False) |
                    filtered_df['position'].str.lower().str.contains(s, na=False)
                ]
            if dept_filter != 'All':
                filtered_df = filtered_df[filtered_df['department'] == dept_filter]
            if region_filter != 'All' and 'region' in filtered_df.columns:
                filtered_df = filtered_df[filtered_df['region'] == region_filter]
            if grade_filter != 'All':
                filtered_df = filtered_df[filtered_df['grade'] == grade_filter]
            if status_filter != 'All':
                filtered_df = filtered_df[filtered_df['status'] == status_filter]
        
        st.markdown(f"**Showing {len(filtered_df)} of {total_emp} employees**")
        
        if not filtered_df.empty:
            items_per_page = 10
            total_pages = max(1, (len(filtered_df) + items_per_page - 1) // items_per_page)
            
            if 'dir_page' not in st.session_state:
                st.session_state.dir_page = 1
            
            pg_col1, pg_col2, pg_col3 = st.columns([1, 2, 1])
            with pg_col1:
                if st.button("⬅️ Previous", disabled=st.session_state.dir_page <= 1, use_container_width=True):
                    st.session_state.dir_page -= 1; st.rerun()
            with pg_col2:
                st.markdown(f"<p style='text-align:center;color:#666;'>Page <strong>{st.session_state.dir_page}</strong> of <strong>{total_pages}</strong></p>", unsafe_allow_html=True)
            with pg_col3:
                if st.button("Next ➡️", disabled=st.session_state.dir_page >= total_pages, use_container_width=True):
                    st.session_state.dir_page += 1; st.rerun()
            
            start_idx = (st.session_state.dir_page - 1) * items_per_page
            end_idx = min(start_idx + items_per_page, len(filtered_df))
            
            for _, emp in filtered_df.iloc[start_idx:end_idx].iterrows():
                initials = generate_initials(f"{emp['first_name']} {emp['last_name']}")
                status_color = "#38a169" if emp.get('status') == 'Active' else "#d69e2e" if emp.get('status') == 'On Leave' else "#CC0000"
                status_bg = "#e6f9e6" if emp.get('status') == 'Active' else "#fff8e6" if emp.get('status') == 'On Leave' else "#ffe6e6"
                border_color = dept_colors.get(emp.get('department', ''), '#CC0000')
                
                with st.expander(f"👤 {emp['first_name']} {emp['last_name']} • {emp.get('position', 'N/A')}", expanded=False):
                    col1, col2, col3 = st.columns([1, 3, 1])
                    with col1:
                        st.markdown(f"""<div style="width:55px;height:55px;border-radius:50%;background:linear-gradient(135deg,{border_color},#e53e3e);display:flex;align-items:center;justify-content:center;font-weight:700;font-size:1.3rem;color:white;">{initials}</div>""", unsafe_allow_html=True)
                    with col2:
                        region_str = f" • 🌍 {emp.get('region', 'N/A')}" if emp.get('region') else ""
                        subsidiary_str = f" • 🏢 {emp.get('subsidiary', '')}" if emp.get('subsidiary') else ""
                        reports_to_str = f" • 👔 Reports to: {emp.get('reports_to', 'N/A')}" if emp.get('reports_to') else ""
                        st.markdown(f"""<div style="line-height:1.6;"><strong style="font-size:1.1rem;">{emp['first_name']} {emp['last_name']}</strong><br><span style="color:#666;">💼 {emp.get('position', 'N/A')}</span><br><span style="color:#888;font-size:0.85rem;">🏢 {emp.get('department', 'N/A')}{region_str}{subsidiary_str}{reports_to_str} • 🆔 {emp.get('employee_id', 'N/A')}</span></div>""", unsafe_allow_html=True)
                    with col3:
                        st.markdown(f"""<div style="text-align:right;"><span style="background:{status_bg};color:{status_color};padding:0.3rem 0.8rem;border-radius:20px;font-size:0.8rem;font-weight:600;border:1px solid {status_color};">{emp.get('status', 'Active')}</span><br><small style="color:#888;">📅 {emp.get('join_date', 'N/A')}</small></div>""", unsafe_allow_html=True)
                    
                    st.markdown("---")
                    c1, c2, c3, c4 = st.columns(4)
                    with c1: st.markdown(f"📧 <small>{emp.get('email', 'N/A')}</small>", unsafe_allow_html=True)
                    with c2: st.markdown(f"📱 <small>{emp.get('phone', 'N/A')}</small>", unsafe_allow_html=True)
                    with c3: st.markdown(f"📊 <small>Grade: {emp.get('grade', 'N/A')}</small>", unsafe_allow_html=True)
                    with c4: st.markdown(f"💼 <small>{emp.get('employment_type', 'N/A')}</small>", unsafe_allow_html=True)
                    
                    if is_admin:
                        st.markdown("---")
                        
                        # GET CURRENT VALUES
                        current_region = str(emp.get('region', 'Lagos'))
                        current_subsidiary = str(emp.get('subsidiary', ''))
                        current_reports_to = str(emp.get('reports_to', ''))
                        
                        # Auto-detect region from subsidiary
                        for reg, subs in SUBSIDIARY_OPTIONS.items():
                            if current_subsidiary in subs:
                                current_region = reg
                                break
                        
                        # REGION & SUBSIDIARY & REPORTS TO OUTSIDE FORM
                        region_options = ['Abuja', 'Lagos', 'Aba']
                        region_idx = region_options.index(current_region) if current_region in region_options else 1
                        
                        st.markdown("#### ✏️ Quick Edit")
                        pre_col1, pre_col2, pre_col3 = st.columns(3)
                        with pre_col1:
                            new_region = st.selectbox("Region", region_options, index=region_idx, key=f"reg_out_{emp['employee_id']}_{st.session_state.dir_page}")
                        with pre_col2:
                            sub_opts = SUBSIDIARY_OPTIONS.get(new_region, SUBSIDIARY_OPTIONS.get('Lagos', []))
                            sub_idx = sub_opts.index(current_subsidiary) if current_subsidiary in sub_opts else 0
                            new_subsidiary = st.selectbox("Subsidiary", sub_opts, index=sub_idx, key=f"sub_out_{emp['employee_id']}_{st.session_state.dir_page}")
                        with pre_col3:
                            report_options = ['None'] + all_employee_names
                            report_idx = report_options.index(current_reports_to) if current_reports_to in report_options else 0
                            new_reports_to = st.selectbox("Reports To", report_options, index=report_idx, key=f"rpt_out_{emp['employee_id']}_{st.session_state.dir_page}")
                        
                        st.markdown("---")
                        
                        with st.form(f"edit_{emp['employee_id']}"):
                            ec1, ec2, ec3 = st.columns(3)
                            with ec1:
                                current_dept = str(emp.get('department', 'Technology Group'))
                                dept_options = ['Senior Management', 'Technology Group', 'Facility Management', 'Human Resources', 'Accounts & Finance', 'Sales & Marketing', 'Procurement', 'Security', 'Legal', 'Operations', 'Engineering']
                                dept_idx = dept_options.index(current_dept) if current_dept in dept_options else 1
                                new_dept = st.selectbox("Department", dept_options, index=dept_idx, key=f"dept_{emp['employee_id']}_{st.session_state.dir_page}")
                                
                                current_grade = str(emp.get('grade', 'Junior'))
                                grade_options = ['Junior', 'Senior', 'Manager', 'HOD', 'C-Level']
                                grade_idx = grade_options.index(current_grade) if current_grade in grade_options else 0
                                new_grade = st.selectbox("Grade", grade_options, index=grade_idx, key=f"grd_{emp['employee_id']}_{st.session_state.dir_page}")
                            with ec2:
                                new_position = st.text_input("Position", value=str(emp.get('position', '')), key=f"pos_{emp['employee_id']}_{st.session_state.dir_page}")
                                
                                current_status = str(emp.get('status', 'Active'))
                                status_options = ['Active', 'On Leave', 'Probation', 'Terminated', 'Archived', 'Inactive']
                                status_idx = status_options.index(current_status) if current_status in status_options else 0
                                new_status = st.selectbox("Status", status_options, index=status_idx, key=f"sts_{emp['employee_id']}_{st.session_state.dir_page}")
                                
                                current_gender = str(emp.get('gender', 'Male'))
                                gender_options = ['Male', 'Female']
                                gender_idx = 0 if current_gender == 'Male' else 1
                                new_gender = st.selectbox("Gender", gender_options, index=gender_idx, key=f"gen_{emp['employee_id']}_{st.session_state.dir_page}")
                            with ec3:
                                new_role = st.selectbox("System Role", ['Admin', 'HOD', 'Manager', 'Team Lead', 'Team Member'], key=f"role_{emp['employee_id']}_{st.session_state.dir_page}")
                                new_email = st.text_input("Email", value=str(emp.get('email', '')), key=f"eml_{emp['employee_id']}_{st.session_state.dir_page}")
                                new_leave = st.number_input("Leave Days", value=int(emp.get('leave_balance', 20) or 20), min_value=0, max_value=365, key=f"leave_{emp['employee_id']}_{st.session_state.dir_page}")
                                
                                current_dob = emp.get('date_of_birth', '')
                                if current_dob and str(current_dob) != 'None' and str(current_dob) != 'nan':
                                    try:
                                        current_dob_date = pd.to_datetime(current_dob).date()
                                        if current_dob_date < date(1920, 1, 1) or current_dob_date > date(2020, 12, 31):
                                            current_dob_date = date(1990, 1, 1)
                                    except: current_dob_date = date(1990, 1, 1)
                                else: current_dob_date = date(1990, 1, 1)
                                new_dob = st.date_input("Date of Birth", value=current_dob_date, min_value=date(1920, 1, 1), max_value=date(2026, 12, 31), key=f"dob_{emp['employee_id']}_{st.session_state.dir_page}")
                            
                            if st.form_submit_button("💾 Save Changes", use_container_width=True):
                                try:
                                    db._patch("employees", {
                                        "department": new_dept, "grade": new_grade,
                                        "position": new_position, "status": new_status,
                                        "email": new_email, "gender": new_gender,
                                        "leave_balance": new_leave,
                                        "date_of_birth": new_dob.strftime('%Y-%m-%d'),
                                        "region": new_region, "subsidiary": new_subsidiary,
                                        "reports_to": new_reports_to if new_reports_to != 'None' else ''
                                    }, {"employee_id": emp['employee_id']})
                                    db._patch("users", {"role": new_role, "department": new_dept, "name": f"{emp['first_name']} {emp['last_name']}"}, {"email": new_email})
                                    st.success(f"✅ {emp['first_name']} {emp['last_name']} updated!")
                                    st.cache_data.clear(); time.sleep(1); st.rerun()
                                except Exception as e:
                                    st.error(f"Update failed: {str(e)}")
                        
                        st.markdown("---")
                        st.markdown("#### 🗄️ Employee Actions")
                        action_col1, action_col2, action_col3 = st.columns([2, 1, 1])
                        with action_col1:
                            current_status = str(emp.get('status', 'Active'))
                            if current_status == 'Archived': st.info("📦 This employee is archived")
                            else: st.caption(f"Status: {current_status}")
                        with action_col2:
                            if current_status == 'Archived':
                                if st.button("🔄 Restore", key=f"restore_{emp['employee_id']}_{st.session_state.dir_page}", use_container_width=True):
                                    db._patch("employees", {"status": "Active"}, {"employee_id": emp['employee_id']})
                                    st.success(f"✅ Restored!"); st.cache_data.clear(); time.sleep(1); st.rerun()
                            else:
                                if st.button("📦 Archive", key=f"archive_{emp['employee_id']}_{st.session_state.dir_page}", use_container_width=True):
                                    db._patch("employees", {"status": "Archived"}, {"employee_id": emp['employee_id']})
                                    st.success(f"📦 Archived!"); st.cache_data.clear(); time.sleep(1); st.rerun()
                        with action_col3:
                            if st.button("🗑️ Delete", key=f"del_{emp['employee_id']}_{st.session_state.dir_page}", use_container_width=True):
                                st.error("Permanently delete?")
                                if st.button("⚠️ Yes, Delete", key=f"confirm_del_{emp['employee_id']}_{st.session_state.dir_page}"):
                                    try:
                                        db._delete("employees", {"employee_id": emp['employee_id']})
                                        st.success("🗑️ Deleted!"); st.cache_data.clear(); time.sleep(1); st.rerun()
                                    except Exception as e: st.error(f"Failed: {str(e)}")
        else:
            st.info("No employees match your search criteria.")
    
    # ============ TAB 2: ADD EMPLOYEE ============
    with tab2:
        st.subheader("➕ Add New Employee")
        
        # Region/Subsidiary OUTSIDE form
        st.markdown("### Location & Reporting")
        loc_col1, loc_col2, loc_col3 = st.columns(3)
        with loc_col1:
            if 'add_region' not in st.session_state: st.session_state.add_region = 'Abuja'
            add_region = st.selectbox("Region *", ['Abuja', 'Lagos', 'Aba'], key="add_region_selector")
            st.session_state.add_region = add_region
        with loc_col2:
            add_sub_opts = SUBSIDIARY_OPTIONS.get(add_region, SUBSIDIARY_OPTIONS.get('Lagos', []))
            add_subsidiary = st.selectbox("Subsidiary *", add_sub_opts, key="add_subsidiary_selector")
        with loc_col3:
            report_options = ['None'] + all_employee_names
            add_reports_to = st.selectbox("Reports To *", report_options, key="add_reports_to_selector")
        
        st.markdown("---")
        
        with st.form("add_employee_form"):
            st.markdown("### Personal Information")
            c1, c2, c3 = st.columns(3)
            with c1:
                first_name = st.text_input("First Name *")
                last_name = st.text_input("Last Name *")
                email = st.text_input("Email *")
                phone = st.text_input("Phone")
            with c2:
                employee_id = st.text_input("Employee ID *", placeholder="e.g., AN00001")
                department = st.selectbox("Department *", ['Senior Management', 'Technology Group', 'Facility Management', 'Human Resources', 'Accounts & Finance', 'Sales & Marketing', 'Procurement', 'Security', 'Legal', 'Operations', 'Engineering'])
                position = st.text_input("Position *")
                grade = st.selectbox("Grade", ['Junior', 'Senior', 'Manager', 'HOD', 'C-Level'])
            with c3:
                employment_type = st.selectbox("Employment Type", ['Full-time', 'Contract', 'Part-time', 'Intern'])
                join_date = st.date_input("Join Date")
                date_of_birth = st.date_input("Date of Birth *", min_value=date(1920, 1, 1), max_value=date(2026, 12, 31), value=date(1990, 1, 1))
                system_role = st.selectbox("System Role", ['Admin', 'HOD', 'Manager', 'Team Lead', 'Team Member'])
                status = st.selectbox("Status", ['Active', 'Probation'])
            
            if st.form_submit_button("✅ Add Employee", use_container_width=True):
                if first_name and last_name and employee_id and department and position:
                    try:
                        result = db._post("employees", {
                            "employee_id": employee_id, "first_name": first_name, "last_name": last_name,
                            "email": email, "phone": phone, "department": department,
                            "region": add_region, "subsidiary": add_subsidiary,
                            "reports_to": add_reports_to if add_reports_to != 'None' else '',
                            "position": position, "grade": grade, "employment_type": employment_type,
                            "join_date": join_date.strftime('%Y-%m-%d'), 
                            "date_of_birth": date_of_birth.strftime('%Y-%m-%d'),
                            "status": status
                        })
                        if result:
                            st.success(f"✅ {first_name} {last_name} added!"); st.balloons()
                            st.cache_data.clear(); time.sleep(0.5); st.rerun()
                        else: st.error("❌ Insert failed - check Supabase")
                    except Exception as e: st.error(f"❌ Error: {str(e)}")
                else: st.error("❌ Required fields missing!")
    
    # ============ TAB 3: BULK UPLOAD ============
    with tab3:
        st.subheader("📤 Bulk Employee Upload")
        st.info("Upload CSV with: employee_id, first_name, last_name, email, phone, department, position, grade, employment_type, join_date, region, subsidiary, reports_to")
        template_df = pd.DataFrame(columns=['employee_id', 'first_name', 'last_name', 'email', 'phone', 'department', 'position', 'grade', 'employment_type', 'join_date', 'region', 'subsidiary', 'reports_to'])
        st.download_button("📥 Download Template", template_df.to_csv(index=False), "employee_template.csv", "text/csv")
        uploaded_file = st.file_uploader("Upload CSV", type=['csv'])
        if uploaded_file:
            df = pd.read_csv(uploaded_file)
            st.write(f"**{len(df)} employees in file**")
            st.dataframe(df.head(), use_container_width=True)
            if st.button("📤 Upload All", use_container_width=True):
                success, fail = 0, 0
                for _, row in df.iterrows():
                    try:
                        db._post("employees", {
                            "employee_id": str(row.get('employee_id', '')), "first_name": str(row.get('first_name', '')),
                            "last_name": str(row.get('last_name', '')), "email": str(row.get('email', '')),
                            "phone": str(row.get('phone', '')), "department": str(row.get('department', '')),
                            "position": str(row.get('position', '')), "grade": str(row.get('grade', 'Junior')),
                            "employment_type": str(row.get('employment_type', 'Full-time')),
                            "join_date": str(row.get('join_date', '')), "status": "Active",
                            "region": str(row.get('region', 'Lagos')),
                            "subsidiary": str(row.get('subsidiary', '')),
                            "reports_to": str(row.get('reports_to', ''))
                        })
                        success += 1
                    except: fail += 1
                st.success(f"✅ {success} uploaded! ({fail} skipped)"); st.balloons(); st.cache_data.clear()
    
    # ============ TAB 4: GENERATE LOGINS ============
    with tab4:
        st.subheader("🔑 Generate Employee Login Credentials")
        st.markdown("### ⚡ Quick Single Employee")
        with st.form("single_login_form"):
            c1, c2 = st.columns(2)
            with c1:
                single_email = st.text_input("Employee Email *", placeholder="e.g., employee@churchgate.com")
                single_name = st.text_input("Full Name *")
                single_pw = st.text_input("Password", value="churchgate2026")
            with c2:
                single_dept = st.selectbox("Department", ['Senior Management', 'Technology Group', 'Facility Management', 'Human Resources', 'Accounts & Finance', 'Sales & Marketing', 'Procurement', 'Security', 'Legal', 'Operations', 'Engineering'], key="single_dept")
                single_role = st.selectbox("Role", ['Admin', 'HOD', 'Manager', 'Team Lead', 'Team Member'], key="single_role")
                single_id = st.text_input("Employee ID", placeholder="e.g., AN00001")
            if st.form_submit_button("🔑 Create Single Login", use_container_width=True):
                if single_email and single_name:
                    try:
                        db.create_user(single_id, single_name, single_email, single_pw, single_role, single_dept, 'Staff')
                        st.success(f"✅ Login created!"); st.balloons()
                    except: st.warning("User may already exist.")
                else: st.error("❌ Email and Name required!")
        
        st.markdown("---")
        st.markdown("### 👥 Bulk Generate for All Employees")
        if not employees_df.empty:
            default_pw = st.text_input("Default Password for Bulk", value="churchgate2026")
            emp_list = []
            for _, emp in employees_df.iterrows():
                emp_list.append({'Name': f"{emp['first_name']} {emp['last_name']}", 'ID': emp['employee_id'], 'Email': emp.get('email', 'N/A'), 'Department': emp.get('department', ''), 'Role': 'Team Member'})
            st.dataframe(pd.DataFrame(emp_list), use_container_width=True, hide_index=True)
            if st.button("🔑 Generate Logins for All", use_container_width=True):
                count = 0
                for emp in emp_list:
                    if emp['Email'] and emp['Email'] != 'N/A':
                        try:
                            db.create_user(emp['ID'], emp['Name'], emp['Email'], default_pw, 'Team Member', emp['Department'], 'Staff')
                            count += 1
                        except: pass
                st.success(f"✅ {count} logins generated!")
                st.info(f"Default password: **{default_pw}**")
                st.download_button("📥 Download Login List", pd.DataFrame(emp_list).to_csv(index=False), "logins.csv", "text/csv")
        else: st.info("No employees found.")
    
    # ============ TAB 5: DEPARTMENTS ============
    with tab5:
        st.subheader("🏢 Department Analytics")
        if not employees_df.empty:
            dept_counts = employees_df['department'].value_counts()
            c1, c2 = st.columns(2)
            for i, (dept, count) in enumerate(dept_counts.items()):
                color = dept_colors.get(dept, '#CC0000')
                with (c1 if i % 2 == 0 else c2):
                    st.markdown(f"""<div style="background:white;padding:1.2rem;border-radius:10px;margin-bottom:0.8rem;border-left:4px solid {color};box-shadow:0 2px 8px rgba(0,0,0,0.05);"><strong>{dept}</strong><span style="float:right;font-size:1.5rem;font-weight:700;color:{color};">{count}</span><br><small style="color:#888;">staff members</small></div>""", unsafe_allow_html=True)
    
    # ============ TAB 6: ORG CHART ============
    with tab6:
        st.subheader("📊 Organizational Structure — Churchgate Group")
        st.markdown("### 🌟 Key Leadership")
        c1, c2, c3, c4 = st.columns(4)
        with c1:
            st.markdown("""<div style="background:white;padding:1rem;border-radius:10px;text-align:center;border-top:3px solid #CC0000;box-shadow:0 2px 8px rgba(0,0,0,0.05);"><div style="width:50px;height:50px;border-radius:50%;background:#CC0000;margin:0 auto;display:flex;align-items:center;justify-content:center;font-weight:700;color:white;">VM</div><strong style="display:block;margin-top:0.5rem;">Vinay Mahtani</strong><small style="color:#888;">GMD/CEO</small><br><small style="color:#CC0000;">👥 Group-wide</small></div>""", unsafe_allow_html=True)
        with c2:
            st.markdown("""<div style="background:white;padding:1rem;border-radius:10px;text-align:center;border-top:3px solid #e53e3e;box-shadow:0 2px 8px rgba(0,0,0,0.05);"><div style="width:50px;height:50px;border-radius:50%;background:#e53e3e;margin:0 auto;display:flex;align-items:center;justify-content:center;font-weight:700;color:white;">JD</div><strong style="display:block;margin-top:0.5rem;">Jerome Das</strong><small style="color:#888;">COO</small><br><small style="color:#e53e3e;">👥 All Departments</small></div>""", unsafe_allow_html=True)
        with c3:
            st.markdown("""<div style="background:white;padding:1rem;border-radius:10px;text-align:center;border-top:3px solid #dd6b20;box-shadow:0 2px 8px rgba(0,0,0,0.05);"><div style="width:50px;height:50px;border-radius:50%;background:#dd6b20;margin:0 auto;display:flex;align-items:center;justify-content:center;font-weight:700;color:white;">AK</div><strong style="display:block;margin-top:0.5rem;">Ahmed Karim</strong><small style="color:#888;">VP Sales</small><br><small style="color:#dd6b20;">👥 Sales & Marketing</small></div>""", unsafe_allow_html=True)
        with c4:
            st.markdown("""<div style="background:white;padding:1rem;border-radius:10px;text-align:center;border-top:3px solid #805ad5;box-shadow:0 2px 8px rgba(0,0,0,0.05);"><div style="width:50px;height:50px;border-radius:50%;background:#805ad5;margin:0 auto;display:flex;align-items:center;justify-content:center;font-weight:700;color:white;">PL</div><strong style="display:block;margin-top:0.5rem;">Partab Lalchandani</strong><small style="color:#888;">GEA</small><br><small style="color:#805ad5;">👥 Group Advisor</small></div>""", unsafe_allow_html=True)
        
        st.markdown("---")
        st.markdown("### 🔗 Group Reporting Hierarchy")
        st.info("GMD → COO (All Depts) / VP Sales (Sales & Mkt) / GEA | Regions: Abuja, Lagos & Aba")
        
        labels = ['GMD', 'COO', 'VP Sales', 'GEA', 'Technology (Abuja)', 'Technology (Lagos)', 'Facility Mgmt (Abuja)', 'Facility Mgmt (Lagos)', 'Engineering/MEP', 'HR', 'Accounts & Finance', 'Sales & Marketing', 'Procurement', 'Security', 'Legal', 'Operations', 'Heads of Department', 'Sr. Managers', 'Managers', 'Team Leads', 'Team Members']
        colors = ['#CC0000', '#e53e3e', '#dd6b20', '#805ad5', '#3182ce', '#3182ce', '#38a169', '#38a169', '#d53f8c', '#d69e2e', '#805ad5', '#dd6b20', '#2b6cb0', '#718096', '#e53e3e', '#319795', '#FF6B35', '#38a169', '#d69e2e', '#2b6cb0', '#718096']
        sources = [0, 0, 0, 1, 1, 1, 1, 1, 1, 1, 1, 1, 1, 2, 3]
        targets = [1, 2, 3, 4, 5, 6, 7, 8, 9, 10, 12, 13, 14, 11, 15]
        values = [1]*15
        for i in range(4, 16): sources.append(i); targets.append(16); values.append(1)
        sources.append(16); targets.append(17); values.append(11)
        sources += [17, 17, 18, 18]; targets += [18, 19, 19, 20]; values += [10, 5, 12, 30]
        
        fig = go.Figure(data=[go.Sankey(node=dict(pad=20, thickness=18, label=labels, color=colors), link=dict(source=sources, target=targets, value=values, color=['rgba(204,0,0,0.2)']*len(sources)))])
        fig.update_layout(height=600)
        st.plotly_chart(fig, use_container_width=True)
        
        st.markdown("---")
        col1, col2 = st.columns(2)
        with col1:
            st.markdown("### 🏢 Abuja Region — Department Heads")
            st.dataframe(pd.DataFrame({'Department': ['Technology Group', 'Facility Management', 'Engineering (MEP)', 'Human Resources', 'Accounts & Finance', 'Sales & Marketing', 'Procurement', 'Security', 'Legal', 'Operations'], 'HOD': ['Emmanuel Etuk', 'David Effiong', 'Sanjeev Purwar', 'Adebayo Sakote', 'Jeff Arikawe', 'Ahmed Karim (VP)', 'Anand Bora', 'Usman Sani', 'David Aiyedun', 'Ibukun Adeogun'], 'Team': [12, 20, 8, 6, 8, 12, 6, 15, 3, 10]}), use_container_width=True, hide_index=True)
        with col2:
            st.markdown("### 🏢 Lagos Region — Department Heads")
            st.dataframe(pd.DataFrame({'Department': ['Technology Group', 'Facility Management', 'Engineering (MEP)', 'Human Resources', 'Accounts & Finance', 'Sales & Marketing', 'Procurement', 'Security', 'Legal', 'Operations'], 'HOD': ['Lawal Mohammed', 'TBD', 'TBD', 'TBD', 'TBD', 'TBD', 'TBD', 'TBD', 'TBD', 'TBD'], 'Team': ['TBD']*10}), use_container_width=True, hide_index=True)
        
        st.markdown("---")
        st.markdown("### 👥 Span of Control")
        span_data = pd.DataFrame({'Leader': ['Vinay Mahtani', 'Jerome Das', 'Ahmed Karim', 'Emmanuel Etuk', 'David Effiong', 'Sanjeev Purwar', 'All HODs (Avg)'], 'Role': ['GMD', 'COO', 'VP Sales', 'HOD Tech (Abuja)', 'HOD FM (Abuja)', 'HOD Engr (MEP)', 'Heads of Dept'], 'Region': ['Group', 'Group', 'Group', 'Abuja', 'Abuja', 'Abuja', 'Group'], 'Direct Reports': [5, 12, 6, 12, 20, 8, 10]})
        fig2 = px.bar(span_data, x='Leader', y='Direct Reports', color='Region', text='Direct Reports', color_discrete_sequence=['#CC0000', '#3182ce', '#38a169'])
        fig2.update_layout(height=350); fig2.update_traces(textposition='outside')
        st.plotly_chart(fig2, use_container_width=True)
        
        st.markdown("---")
        st.markdown("### 🔍 Find Reporting Chain")
        chain_search = st.text_input("Enter employee name", placeholder="e.g., Francis Asuquo", key="chain_search")
        if chain_search:
            found = False; chain = ""
            try:
                for _, emp in employees_df.iterrows():
                    full_name = f"{emp['first_name']} {emp['last_name']}".lower()
                    if chain_search.lower() in full_name:
                        found = True
                        role = emp.get('position', ''); dept = emp.get('department', ''); region = emp.get('region', 'Abuja')
                        reports_to = emp.get('reports_to', '')
                        if reports_to: chain = f"📋 **{emp['first_name']} {emp['last_name']}** → Reports to: {reports_to}"
                        elif 'GMD' in role or 'CEO' in role: chain = f"📋 **{emp['first_name']} {emp['last_name']}** → Reports to Board"
                        elif 'COO' in role: chain = f"📋 **{emp['first_name']} {emp['last_name']}** → COO → GMD"
                        elif 'HOD' in role: chain = f"📋 **{emp['first_name']} {emp['last_name']}** → HOD ({dept}, {region}) → COO → GMD"
                        elif 'Manager' in role: chain = f"📋 **{emp['first_name']} {emp['last_name']}** → Manager ({dept}) → HOD → COO → GMD"
                        elif 'Team Lead' in role: chain = f"📋 **{emp['first_name']} {emp['last_name']}** → Team Lead ({dept}) → Manager → HOD → COO → GMD"
                        else: chain = f"📋 **{emp['first_name']} {emp['last_name']}** → Team Member ({dept}) → Team Lead → Manager → HOD → COO → GMD"
                        break
                if not found: chain = f"📋 **{chain_search}** not found."
            except: chain = "📋 Lookup unavailable."
            st.info(chain)
    
    # ============ TAB 7: DEMOGRAPHICS ============
    with tab7:
        st.subheader("📈 Demographics & Inclusion")
        st.markdown("### 👥 Gender Distribution")
        gender_data = pd.DataFrame({'Gender': ['Male', 'Female'], 'Count': [38, 18]})
        fig = px.pie(gender_data, values='Count', names='Gender', hole=0.5, color_discrete_sequence=['#3182ce', '#CC0000'])
        fig.update_layout(height=350); st.plotly_chart(fig, use_container_width=True)
        
        st.markdown("---"); st.markdown("### 🏢 Department Gender Split")
        dept_gender = pd.DataFrame({'Department': ['Technology Group', 'Facility Management', 'Human Resources', 'Sales & Marketing', 'Accounts & Finance', 'Procurement', 'Security', 'Legal', 'Operations', 'Engineering'], 'Male': [10, 10, 3, 6, 4, 3, 10, 1, 4, 3], 'Female': [4, 3, 5, 4, 2, 2, 2, 1, 1, 1]})
        fig2 = px.bar(dept_gender, x='Department', y=['Male', 'Female'], barmode='group', color_discrete_sequence=['#3182ce', '#CC0000'])
        fig2.update_layout(height=400); st.plotly_chart(fig2, use_container_width=True)
        
        st.markdown("---"); st.markdown("### 📅 Tenure Distribution")
        try:
            tenure_data = []
            for _, emp in employees_df.iterrows():
                jd = emp.get('join_date')
                if jd and pd.notna(jd):
                    try:
                        years = (datetime.now() - pd.to_datetime(jd)).days / 365
                        if years < 2: tenure_data.append('0-2 years')
                        elif years < 5: tenure_data.append('3-5 years')
                        elif years < 10: tenure_data.append('6-10 years')
                        else: tenure_data.append('10+ years')
                    except: pass
            if tenure_data:
                tenure_df = pd.DataFrame(pd.Series(tenure_data).value_counts()).reset_index()
                tenure_df.columns = ['Tenure', 'Count']
                fig3 = px.bar(tenure_df, x='Tenure', y='Count', color='Tenure', color_discrete_sequence=['#CC0000', '#3182ce', '#38a169', '#d69e2e'])
                fig3.update_layout(height=350, showlegend=False); st.plotly_chart(fig3, use_container_width=True)
        except: pass
        
        st.markdown("---"); st.markdown("### 📊 Grade Distribution")
        if not employees_df.empty:
            grade_counts = employees_df['grade'].value_counts()
            grade_df = pd.DataFrame({'Grade': grade_counts.index, 'Count': grade_counts.values})
            fig4 = px.pie(grade_df, values='Count', names='Grade', hole=0.4, color_discrete_sequence=['#CC0000', '#3182ce', '#38a169', '#d69e2e', '#805ad5'])
            fig4.update_layout(height=350); st.plotly_chart(fig4, use_container_width=True)
    
    # ============ TAB 8: EXPORT ============
    with tab8:
        st.subheader("📥 Export Employee Data")
        if not employees_df.empty:
            st.download_button("📥 Download Full Directory (CSV)", employees_df.to_csv(index=False), "churchgate_employees.csv", "text/csv")
            st.markdown("---"); st.markdown("### 📊 Export by Department")
            selected_export_dept = st.selectbox("Select Department", ['All'] + list(employees_df['department'].unique()) if not employees_df.empty else ['All'])
            if selected_export_dept != 'All':
                dept_df = employees_df[employees_df['department'] == selected_export_dept]
                st.download_button(f"📥 Download {selected_export_dept} (CSV)", dept_df.to_csv(index=False), f"{selected_export_dept}_employees.csv", "text/csv")
            st.markdown("---"); st.markdown("### 📊 Quick Stats")
            st.dataframe(employees_df.describe(), use_container_width=True)


def performance_okrs():
    """
    Churchgate Group HRIS - Performance & OKRs Module v8.0
    Fortune 500 Grade: Appraisal Committee | Certificates | Full Admin Control | Advanced Analytics
    """
    
    # ============================================================
    # CSS INJECTION
    # ============================================================
    st.markdown("""
    <style>
    .glass-card {
        background: rgba(255, 255, 255, 0.85);
        backdrop-filter: blur(12px);
        -webkit-backdrop-filter: blur(12px);
        border-radius: 16px;
        border: 1px solid rgba(204, 0, 0, 0.1);
        box-shadow: 0 8px 32px rgba(0, 0, 0, 0.06);
        padding: 1.5rem;
        margin-bottom: 1rem;
        transition: all 0.3s cubic-bezier(0.4, 0, 0.2, 1);
    }
    .glass-card:hover { box-shadow: 0 12px 40px rgba(204, 0, 0, 0.12); transform: translateY(-2px); }
    .kpi-card {
        background: white; border-radius: 12px; padding: 1rem; margin: 0.4rem 0;
        border-left: 5px solid #CC0000; box-shadow: 0 2px 8px rgba(0,0,0,0.04); transition: all 0.25s ease;
    }
    .kpi-card:hover { box-shadow: 0 6px 20px rgba(204, 0, 0, 0.15); transform: translateX(4px); }
    .region-header {
        background: linear-gradient(135deg, #1a1a1a, #2d2d2d); color: white; padding: 1rem 1.5rem;
        border-radius: 12px; margin: 1rem 0 0.5rem 0; font-weight: 700; font-size: 1.1rem; border-left: 5px solid #CC0000;
    }
    .subsidiary-header {
        background: linear-gradient(135deg, #2d2d2d, #3d3d3d); color: white; padding: 0.7rem 1.2rem;
        border-radius: 10px; margin: 0.4rem 0 0.4rem 1.5rem; font-weight: 600; border-left: 4px solid #CC0000;
    }
    .metric-mini {
        background: rgba(255,255,255,0.9); border-radius: 10px; padding: 0.8rem 1rem;
        text-align: center; box-shadow: 0 2px 8px rgba(0,0,0,0.04); backdrop-filter: blur(8px);
    }
    .metric-mini .value { font-size: 1.5rem; font-weight: 700; color: #CC0000; }
    .metric-mini .label { font-size: 0.7rem; color: #888; text-transform: uppercase; }
    .badge { display: inline-block; padding: 3px 10px; border-radius: 12px; font-size: 0.7rem; font-weight: 600; margin: 2px; }
    .badge-green { background: #c6f6d5; color: #22543d; }
    .badge-yellow { background: #fefcbf; color: #744210; }
    .badge-gray { background: #e2e8f0; color: #2d3748; }
    .badge-red { background: #fed7d7; color: #742a2a; }
    .badge-escalated { background: #fbb6ce; color: #742a2a; }
    .certificate-card {
        background: linear-gradient(135deg, #fffef5, #fff8e1);
        border: 2px solid #d69e2e;
        border-radius: 16px;
        padding: 2rem;
        text-align: center;
        box-shadow: 0 8px 32px rgba(214, 158, 46, 0.2);
    }
    @keyframes pulse { 0%, 100% { opacity: 1; } 50% { opacity: 0.5; } }
    .pulse-dot { display: inline-block; width: 8px; height: 8px; border-radius: 50%; background: #38a169; animation: pulse 2s infinite; margin-right: 6px; }
    input[type="number"] { border: 2px solid #e0e0e0; border-radius: 8px; padding: 8px 12px; font-size: 1rem; transition: all 0.2s; }
    input[type="number"]:focus { border-color: #CC0000; box-shadow: 0 0 0 3px rgba(204, 0, 0, 0.1); outline: none; }
    ::-webkit-scrollbar { width: 6px; }
    ::-webkit-scrollbar-track { background: #f1f1f1; border-radius: 3px; }
    ::-webkit-scrollbar-thumb { background: #CC0000; border-radius: 3px; }
    </style>
    """, unsafe_allow_html=True)
    
    # ============================================================
    # INITIALIZATION
    # ============================================================
    from datetime import timezone, timedelta
    import re, os
    from collections import defaultdict
    
    wat = timezone(timedelta(hours=1))
    now_wat = datetime.now(wat)
    
    user_role = st.session_state.user['role'] if st.session_state.user else 'Employee'
    user_dept = st.session_state.user.get('department', '') if st.session_state.user else ''
    user_name = st.session_state.user['name'] if st.session_state.user else 'Staff'
    user_email = st.session_state.user['email'] if st.session_state.user else ''
    is_admin = user_role in ['Admin', 'HR Director'] or user_dept == 'Senior Management'
    is_sr_mgmt = user_dept == 'Senior Management'
    is_hod = is_admin or user_role in ['Manager', 'HOD']
    is_team_lead_or_manager = user_role in ['Manager', 'Team Lead', 'HOD', 'Admin', 'HR Director']
    is_super_admin = user_email == 'admin@churchgate.com' or user_role in ['Admin', 'HR Director']
    
    all_depts = ['Senior Management', 'Technology Group', 'Facility Management', 'Human Resources', 
                 'Accounts & Finance', 'Sales & Marketing', 'Procurement', 'Security', 'Legal', 
                 'Operations', 'Engineering', 'Central Stores', 'Project Development', 'Trade Services']
    
    SUBSIDIARY_REGIONS = {
        'World Trade Center(WTC)': 'Abuja', 'World Trade Center': 'Abuja', 'WTC': 'Abuja',
        'Agroline Ventures Limited': 'Abuja', 'Agroline Ventures': 'Abuja',
        'Aba Textile Mills PLC': 'Aba', 'Aba Textile Mills': 'Aba',
    }
    
    def get_region(subsidiary):
        if not subsidiary: return 'Lagos'
        sub_lower = str(subsidiary).lower().strip()
        for key, region in SUBSIDIARY_REGIONS.items():
            if key.lower() in sub_lower or sub_lower in key.lower(): return region
        return 'Lagos'
    
    # ============================================================
    # DATABASE-BACKED APPRAISAL CYCLE STATE
    # ============================================================
    def load_appraisal_cycle_from_db():
        try:
            cycles = db._get("appraisal_cycles", {"is_active": "true"})
            if cycles and len(cycles) > 0:
                cycle = cycles[0]
                st.session_state.appraisal_cycle_active = True
                st.session_state.appraisal_cycle_name = cycle.get('cycle_name', '2026 Half-Year Appraisal')
                st.session_state.appraisal_start = cycle.get('start_date', '2026-06-01')
                st.session_state.appraisal_end = cycle.get('end_date', '2026-12-31')
                st.session_state.appraisal_locked = cycle.get('is_locked', False)
                return True
            else:
                st.session_state.appraisal_cycle_active = False
                return False
        except:
            return False
    
    def save_appraisal_cycle_to_db():
        try:
            existing = db._get("appraisal_cycles")
            if existing:
                for c in existing: db._patch("appraisal_cycles", {"is_active": False}, {"id": c['id']})
            db._post("appraisal_cycles", {
                "cycle_name": st.session_state.appraisal_cycle_name,
                "start_date": st.session_state.appraisal_start,
                "end_date": st.session_state.appraisal_end,
                "is_active": True,
                "is_locked": st.session_state.appraisal_locked,
                "activated_by": user_name,
                "activated_at": now_wat.strftime('%Y-%m-%d %H:%M WAT')
            })
            return True
        except:
            return False
    
    # Initialize session state
    if 'appraisal_cycle_active' not in st.session_state: st.session_state.appraisal_cycle_active = False
    if 'appraisal_cycle_name' not in st.session_state: st.session_state.appraisal_cycle_name = "2026 Half-Year Appraisal"
    if 'appraisal_start' not in st.session_state: st.session_state.appraisal_start = "2026-06-01"
    if 'appraisal_end' not in st.session_state: st.session_state.appraisal_end = "2026-12-31"
    if 'appraisal_locked' not in st.session_state: st.session_state.appraisal_locked = False
    
    if 'cycle_loaded_from_db' not in st.session_state:
        load_appraisal_cycle_from_db()
        st.session_state.cycle_loaded_from_db = True
    
    if 'self_assessments' not in st.session_state: st.session_state.self_assessments = {}
    if 'audit_trail' not in st.session_state: st.session_state.audit_trail = []
    if 'exceptional_achievements' not in st.session_state: st.session_state.exceptional_achievements = {}
    if 'editing_kpi' not in st.session_state: st.session_state.editing_kpi = None
    if 'dept_expanded' not in st.session_state: st.session_state.dept_expanded = {}
    if 'subsidiary_expanded' not in st.session_state: st.session_state.subsidiary_expanded = {}
    if 'team_lead_reviews' not in st.session_state: st.session_state.team_lead_reviews = {}
    
    # Load appraisals from DB
    try:
        all_appraisals = db.get_all_appraisals()
        for a in all_appraisals:
            st.session_state.self_assessments[a['user_name']] = {
                'scores': a.get('scores', {}), 'comments': a.get('comments', ''),
                'pillar_comments': a.get('pillar_comments', {}), 'date': a.get('submitted_date', ''),
                'status': a.get('status', 'Submitted'), 'department': a.get('department', ''),
                'email': a.get('user_email', ''), 'hod_scores': a.get('hod_scores'),
                'hod_comments': a.get('hod_comments'), 'hod_pillar_comments': a.get('hod_pillar_comments'),
                'acceptance': a.get('acceptance'), 'sr_decision': a.get('sr_decision'),
                'rejection_comment': a.get('rejection_comment', ''), 'rejection_docs': a.get('rejection_docs', '[]'),
                'reviewer_type': a.get('reviewer_type', 'HOD'),
                'tl_scores': a.get('tl_scores'), 'tl_comments': a.get('tl_comments'),
            }
    except: pass
    
    # Load appraisal history for accepted/completed
    try:
        appraisal_history = db.get_appraisal_history()
        if appraisal_history:
            for h in appraisal_history:
                uname = h.get('user_name', '')
                if uname and uname not in st.session_state.self_assessments:
                    st.session_state.self_assessments[uname] = {
                        'scores': h.get('scores', {}), 'comments': h.get('comments', ''),
                        'status': 'Completed', 'acceptance': 'Accepted',
                        'hod_scores': h.get('hod_scores'), 'hod_comments': h.get('hod_comments', ''),
                        'date': h.get('completed_date', ''),
                    }
    except: pass
    
    if not st.session_state.audit_trail:
        try:
            for a in (db.get_audit_trail() or []):
                st.session_state.audit_trail.append({'action': a.get('action', ''), 'details': a.get('details', ''), 'user': a.get('user_name', ''), 'timestamp': a.get('timestamp_text', '')})
        except: pass
    
    # ============================================================
    # HELPER FUNCTIONS
    # ============================================================
    def log_audit(action, details):
        entry = {'action': action, 'details': details, 'user': user_name, 'timestamp': now_wat.strftime('%Y-%m-%d %H:%M WAT')}
        st.session_state.audit_trail.append(entry)
        try: db.save_audit(action, details, user_name, now_wat.strftime('%Y-%m-%d %H:%M WAT'))
        except: pass
    
    def get_kpi_status(progress):
        if progress >= 85: return 'On Track', "#38a169"
        elif progress >= 65: return 'Near Target', "#d69e2e"
        else: return 'At Risk', "#CC0000"
    
    def natural_sort_key(item):
        key = item[0] if isinstance(item, tuple) else str(item)
        parts = re.split(r'(\d+)', key)
        return [int(p) if p.isdigit() else p for p in parts]
    
    @st.cache_data(ttl=120)
    def get_all_perf_cached(): 
        try: return db.get_performance_data()
        except: return pd.DataFrame()
    
    @st.cache_data(ttl=300)
    def get_all_emp_cached(): 
        try: return db.get_all_employees()
        except: return pd.DataFrame()
    
    def get_user_perf(): 
        try: return db.get_performance_data(user_name)
        except: return pd.DataFrame()
    
    def load_user_pillar_data():
        user_perf = get_user_perf()
        pillar_data = {}
        for pillar in ['1. Occupancy & Revenue Growth', '2. Process Simplification', '3. Asset Reliability & Digitalization', '4. People & Culture']:
            pillar_data[pillar] = {'weight': 0, 'progress': 0, 'status': 'Not Started', 'deadline': '2026-12-31', 'kpis': [], 'submission_status': 'Draft'}
        if not user_perf.empty:
            for _, row in user_perf.iterrows():
                p_name = row.get('pillar_name', '')
                if p_name in pillar_data:
                    kpi_list = json.loads(row.get('kpi_data', '[]')) if row.get('kpi_data') else []
                    pillar_data[p_name] = {'weight': row.get('weight', 25), 'progress': row.get('progress', 0), 'status': row.get('status', 'Not Started'), 'deadline': row.get('deadline', '2026-12-31'), 'kpis': kpi_list, 'submission_status': row.get('submission_status', 'Draft')}
        return pillar_data
    
    # Employee lookup maps
    employees_df = get_all_emp_cached()
    emp_dept_map, emp_subsidiary_map, emp_email_map, emp_role_map = {}, {}, {}, {}
    
    if not employees_df.empty:
        for _, emp in employees_df.iterrows():
            first = str(emp.get('first_name', '')).strip()
            last = str(emp.get('last_name', '')).strip()
            full_name = f"{first} {last}".strip()
            emp_id = str(emp.get('employee_id', '')).strip()
            dept = emp.get('department', 'General')
            sub = emp.get('subsidiary', '')
            email_addr = emp.get('email', '')
            role = str(emp.get('role', 'Team Member'))
            
            for key in [full_name, first, emp_id]:
                if key:
                    emp_dept_map[key] = dept
                    emp_subsidiary_map[key] = sub
                    emp_email_map[key] = email_addr
                    emp_role_map[key] = role
    
    def get_employee_dept(name):
        clean = ' '.join(str(name).split())
        return emp_dept_map.get(clean, emp_dept_map.get(clean.split()[0] if ' ' in clean else clean, 'General'))
    
    def get_employee_subsidiary(name):
        clean = ' '.join(str(name).split())
        return emp_subsidiary_map.get(clean, emp_subsidiary_map.get(clean.split()[0] if ' ' in clean else clean, ''))
    
    def get_employee_email(name):
        clean = ' '.join(str(name).split())
        return emp_email_map.get(clean, emp_email_map.get(clean.split()[0] if ' ' in clean else clean, ''))
    
    def find_hod_email_for_dept(dept):
        if not employees_df.empty:
            hod_rows = employees_df[(employees_df['department'] == dept) & (employees_df['role'].isin(['HOD', 'Admin', 'HR Director']))]
            if not hod_rows.empty: return hod_rows.iloc[0].get('email', '')
        return ''
    
    # ============================================================
    # HEADER & CYCLE STATUS
    # ============================================================
    st.markdown("""<div class="glass-card" style="text-align:center;margin-bottom:1.5rem;"><h1 style="margin:0;color:#1a1a1a;">📈 Performance & Appraisal Engine</h1><p style="color:#888;margin:0.5rem 0 0 0;">Fortune 500 Grade | KPI Management | Self-Assessment | Team Lead Review | HOD Review | Appraisal Committee | Certificates</p></div>""", unsafe_allow_html=True)
    
    if st.session_state.appraisal_cycle_active:
        try:
            end_date = datetime.strptime(st.session_state.appraisal_end, '%Y-%m-%d')
            days_left = (end_date - datetime.now()).days
            user_perf = get_user_perf()
            has_approved = any(row.get('submission_status') == 'Approved' for _, row in user_perf.iterrows()) if not user_perf.empty else False
            if has_approved:
                color = "#38a169" if days_left > 14 else "#d69e2e" if days_left > 7 else "#CC0000"
                st.markdown(f"""<div class="glass-card" style="border-left:4px solid {color};padding:0.8rem 1rem;"><strong>📊 Appraisal Active: {st.session_state.appraisal_cycle_name}</strong><span style="float:right;color:{color};font-weight:700;">⏰ {days_left} day{'s' if days_left > 1 else ''} remaining</span><br><small style="color:#38a169;">✅ Your KPIs are approved — ready for self-assessment!</small></div>""", unsafe_allow_html=True)
            else:
                st.warning("⚠️ Appraisal cycle active, but your KPIs haven't been approved yet.")
        except: pass
    else:
        st.info("📊 No active appraisal cycle. HR will activate it when ready.")
    
    # ============================================================
    # 9 TABS (Added Appraisal Committee)
    # ============================================================
    tab1, tab2, tab3, tab4, tab5, tab6, tab7, tab8, tab9 = st.tabs([
        "🎯 Strategic Pillars", "✏️ My KPIs", "📝 Self-Assessment", 
        "👔 HOD Review", "👥 Team Lead Review", "🌟 Exceptional Achievements",
        "⚙️ Appraisal Settings", "📊 Dashboard", "🏛️ Appraisal Committee"
    ])
    
    # ============================================================
    # TAB 1: STRATEGIC PILLARS
    # ============================================================
    with tab1:
        st.markdown('<div class="glass-card"><h3>🎯 My Strategic Pillars</h3></div>', unsafe_allow_html=True)
        pillar_data = load_user_pillar_data()
        overall_status = 'Draft'
        statuses = [p['submission_status'] for p in pillar_data.values()]
        if 'Submitted' in statuses: overall_status = 'Submitted'
        if 'Approved' in statuses and 'Submitted' not in statuses: overall_status = 'Approved'
        
        total_weighted = sum(p['progress'] * p['weight'] / 100 for p in pillar_data.values())
        c1, c2, c3, c4 = st.columns(4)
        with c1: st.markdown(f'<div class="metric-mini"><div class="label">Overall Score</div><div class="value">{total_weighted:.1f}%</div></div>', unsafe_allow_html=True)
        with c2: st.markdown(f'<div class="metric-mini"><div class="label">On Track</div><div class="value" style="color:#38a169;">{sum(1 for p in pillar_data.values() if p["status"] in ["On Track","Exceeding"])}</div></div>', unsafe_allow_html=True)
        with c3: st.markdown(f'<div class="metric-mini"><div class="label">At Risk</div><div class="value" style="color:#CC0000;">{sum(1 for p in pillar_data.values() if p["status"]=="At Risk")}</div></div>', unsafe_allow_html=True)
        with c4: st.markdown(f'<div class="metric-mini"><div class="label">KPI Status</div><div class="value" style="font-size:1rem;">{overall_status.upper()}</div></div>', unsafe_allow_html=True)
        
        if overall_status == 'Draft' and any(len(p['kpis']) > 0 for p in pillar_data.values()):
            if st.button("🚀 Final Submit All KPIs", use_container_width=True, type="primary"):
                all_rows = db._get("performance_data", {"user_name": user_name})
                for row in (all_rows or []): db._patch("performance_data", {"submission_status": "Submitted"}, {"id": row['id']})
                send_kpi_notification('submitted_to_employee', user_name, user_email)
                hod_email = find_hod_email_for_dept(user_dept)
                if hod_email: send_kpi_notification('submitted_to_hod', user_name, '', hod_email)
                log_audit("KPIs Submitted", f"All KPIs submitted by {user_name}")
                st.cache_data.clear()
                st.success("✅ All KPIs submitted!"); st.balloons(); time.sleep(1.5); st.rerun()
        
        for pillar_name in ['1. Occupancy & Revenue Growth', '2. Process Simplification', '3. Asset Reliability & Digitalization', '4. People & Culture']:
            pd_data = pillar_data[pillar_name]
            status_text, color = get_kpi_status(pd_data['progress'])
            is_locked = pd_data['submission_status'] != 'Draft'
            with st.expander(f"📌 {pillar_name} | {pd_data['progress']}% | {pd_data['status']}", expanded=not is_locked):
                st.progress(pd_data['progress'] / 100)
                if pd_data['kpis']:
                    for i, kpi in enumerate(pd_data['kpis']):
                        try: kpi_prog = int(float(str(kpi.get('current', '0')).replace('%', '')))
                        except: kpi_prog = 0
                        kpi_stat, kpi_col = get_kpi_status(kpi_prog)
                        st.markdown(f"""<div class="kpi-card" style="border-left-color:{kpi_col};"><strong>{kpi.get('kpi', 'Untitled')}</strong><br><small>🎯 Target: {kpi.get('target', 'N/A')} | 📊 Current: {kpi.get('current', '0')} | ⚖️ Weight: {kpi.get('weight', 0)}%</small><br><small style="color:{kpi_col};">● {kpi_stat}</small></div>""", unsafe_allow_html=True)
                        if not is_locked:
                            c1, c2 = st.columns([1, 1])
                            with c1:
                                if st.button("✏️ Edit", key=f"qedit_{pillar_name}_{i}"): st.session_state.editing_kpi = {'pillar': pillar_name, 'index': i, 'data': kpi}; st.rerun()
                            with c2:
                                if st.button("🗑️ Delete", key=f"qdel_{pillar_name}_{i}"):
                                    pd_data['kpis'].pop(i)
                                    db.save_performance_data(user_name, pillar_name, pd_data['weight'], pd_data['progress'], pd_data['status'], pd_data['deadline'], pd_data['kpis'])
                                    st.cache_data.clear(); st.success("Deleted!"); st.rerun()
                else: st.info("No KPIs in this pillar yet.")
    
    # ============================================================
    # TAB 2: MY KPIs
    # ============================================================
    with tab2:
        st.markdown('<div class="glass-card"><h3>✏️ My KPIs & Objectives</h3></div>', unsafe_allow_html=True)
        st.info("Set unlimited KPIs aligned to the 4 strategic pillars.")
        pillar_data = load_user_pillar_data()
        overall_status = 'Draft'
        statuses = [p['submission_status'] for p in pillar_data.values()]
        if 'Submitted' in statuses: overall_status = 'Submitted'
        if 'Approved' in statuses and 'Submitted' not in statuses: overall_status = 'Approved'
        
        if overall_status != 'Draft':
            st.warning(f"🔒 KPIs are locked ({overall_status}). Contact your HOD if changes are needed.")
        else:
            editing = st.session_state.get('editing_kpi')
            if editing: st.info(f"✏️ Editing KPI: {editing['data'].get('kpi', '')[:60]}")
            with st.form("kpi_add_form", clear_on_submit=not editing):
                c1, c2 = st.columns(2)
                with c1:
                    pillar_list = ['1. Occupancy & Revenue Growth', '2. Process Simplification', '3. Asset Reliability & Digitalization', '4. People & Culture']
                    pillar_choice = st.selectbox("Strategic Pillar *", pillar_list, index=pillar_list.index(editing['pillar']) if editing else 0)
                    kpi_title = st.text_input("KPI Title *", value=editing['data'].get('kpi', '') if editing else "", placeholder="What will you achieve?")
                    kpi_target = st.text_input("Target *", value=editing['data'].get('target', '') if editing else "", placeholder="e.g., 15% increase")
                with c2:
                    kpi_weight = st.number_input("Weight (%)", 0, 100, value=int(editing['data'].get('weight', 25)) if editing else 25)
                    kpi_deadline = st.date_input("Target Deadline *", value=datetime.strptime(editing['data'].get('deadline', '2026-12-31'), '%Y-%m-%d') if editing and editing['data'].get('deadline') else datetime.now())
                    kpi_current = st.text_input("Current Progress", value=editing['data'].get('current', '0') if editing else "0", placeholder="e.g., 10%")
                kpi_cycle = st.selectbox("Appraisal Cycle *", ['Half-Year Appraisal', 'Full-Year Appraisal', 'HOD Mock Appraisal', 'Team Mock Appraisal'])
                c1, c2 = st.columns(2)
                with c1: save_add = st.form_submit_button("💾 Save & Add Another", use_container_width=True)
                with c2: save_done = st.form_submit_button("✅ Save & Finish", use_container_width=True)
                
                if save_add or save_done:
                    if not kpi_title or not kpi_target: st.error("❌ Title and Target are required!")
                    else:
                        new_kpi = {'kpi': kpi_title, 'target': kpi_target, 'current': kpi_current, 'weight': kpi_weight, 'deadline': kpi_deadline.strftime('%Y-%m-%d'), 'cycle': kpi_cycle, 'owner': user_name, 'status': 'In Progress'}
                        if editing: pillar_data[pillar_choice]['kpis'][editing['index']] = new_kpi; st.session_state.editing_kpi = None; log_audit("KPI Updated", f"KPI '{kpi_title}' updated")
                        else: pillar_data[pillar_choice]['kpis'].append(new_kpi); log_audit("KPI Added", f"KPI '{kpi_title}' added")
                        total_weight = sum(k.get('weight', 0) for k in pillar_data[pillar_choice]['kpis'])
                        if total_weight > 0: pillar_data[pillar_choice]['weight'] = total_weight
                        db.save_performance_data(user_name, pillar_choice, pillar_data[pillar_choice]['weight'], pillar_data[pillar_choice]['progress'], pillar_data[pillar_choice]['status'], pillar_data[pillar_choice]['deadline'], pillar_data[pillar_choice]['kpis'])
                        st.cache_data.clear(); st.success("✅ KPI saved!")
                        if save_done: st.rerun()
                        else: time.sleep(0.5); st.rerun()
    
    # ============================================================
    # TAB 3: SELF-ASSESSMENT
    # ============================================================
    with tab3:
        st.markdown('<div class="glass-card"><h3>📝 Self-Assessment</h3></div>', unsafe_allow_html=True)
        
        if st.session_state.appraisal_cycle_active:
            user_perf = get_user_perf()
            has_approved = not user_perf.empty and any(row.get('submission_status') == 'Approved' for _, row in user_perf.iterrows())
            
            if not has_approved: st.warning("⚠️ Your KPIs must be approved before self-assessment.")
            elif st.session_state.appraisal_locked: st.warning("🔒 Scores are locked.")
            elif st.session_state.self_assessments.get(user_name, {}).get('status') in ['Submitted', 'Approved', 'Awaiting HOD Re-review', 'Awaiting TL Re-review']: pass
            else:
                st.success(f"🔓 Ready for Self-Assessment — {st.session_state.appraisal_cycle_name}")
                with st.form("self_assessment_form"):
                    scores, pillar_comments, evidence_files = {}, {}, {}
                    for pillar_name in ['1. Occupancy & Revenue Growth', '2. Process Simplification', '3. Asset Reliability & Digitalization', '4. People & Culture']:
                        pillar_rows = user_perf[user_perf['pillar_name'] == pillar_name]
                        if not pillar_rows.empty:
                            kpi_list = json.loads(pillar_rows.iloc[0].get('kpi_data', '[]')) if pillar_rows.iloc[0].get('kpi_data') else []
                            if kpi_list:
                                st.markdown(f"### {pillar_name}")
                                st.caption("📎 Evidence (Optional — up to 5 files)")
                                ev_cols = st.columns(5); pillar_files = []
                                for j in range(5):
                                    with ev_cols[j]:
                                        ev = st.file_uploader(f"File {j+1}", type=['pdf','docx','jpg','png','xlsx'], key=f"ev_{pillar_name}_{j}", label_visibility="collapsed")
                                        if ev: pillar_files.append(ev)
                                evidence_files[pillar_name] = pillar_files if pillar_files else None
                                for i, kpi in enumerate(kpi_list):
                                    score_key = f"{pillar_name}_{i}"
                                    col1, col2 = st.columns([3, 1])
                                    with col1: st.markdown(f"**{kpi.get('kpi', 'KPI')[:80]}**"); st.caption(f"Target: {kpi.get('target', 'N/A')}")
                                    with col2: scores[score_key] = st.number_input("Score %", 0, 100, 50, 1, key=f"score_{pillar_name}_{i}")
                                pillar_comments[pillar_name] = st.text_area(f"Justification *", key=f"just_{pillar_name}")
                                st.markdown("---")
                    overall_comments = st.text_area("Overall Comments *")
                    if st.form_submit_button("📤 Submit Self-Assessment", use_container_width=True, type="primary"):
                        if not scores: st.error("❌ Please score at least one KPI!")
                        elif not overall_comments: st.error("❌ Overall comments required!")
                        else:
                            empty_just = [k for k, v in pillar_comments.items() if v is not None and not v]
                            if empty_just: st.error(f"❌ Justification required for: {', '.join(empty_just)}")
                            else:
                                evidence_urls = {}
                                for p_name, files in evidence_files.items():
                                    if files:
                                        urls = []
                                        for f in files:
                                            try:
                                                url = db.upload_file("evidence", f"{user_name}_{p_name}_{f.name}", f.read(), f.type)
                                                if url: urls.append(url)
                                            except: pass
                                        if urls: evidence_urls[p_name] = urls
                                try: db.save_appraisal(user_name, user_email, user_dept, st.session_state.appraisal_cycle_name, 'Submitted', scores, overall_comments, pillar_comments, None, None, None, None, None, now_wat.strftime('%Y-%m-%d %H:%M WAT'))
                                except: pass
                                st.session_state.self_assessments[user_name] = {'scores': scores, 'comments': overall_comments, 'pillar_comments': pillar_comments, 'evidence_files': json.dumps(evidence_urls), 'date': now_wat.strftime('%Y-%m-%d %H:%M WAT'), 'status': 'Submitted', 'department': user_dept, 'email': user_email, 'hod_scores': None, 'hod_comments': None, 'acceptance': None}
                                log_audit('Self-Assessment Submitted', f'Submitted by {user_name}')
                                st.success("✅ Submitted!"); st.balloons(); time.sleep(1.5); st.rerun()
        else: st.info("⏳ No active appraisal cycle.")
        
        # Acceptance/Rejection
        if user_name in st.session_state.self_assessments:
            a = st.session_state.self_assessments[user_name]
            reviewer_type = a.get('reviewer_type', 'HOD')
            reviewer_scores = a.get('hod_scores') or a.get('tl_scores')
            
            if a.get('status') == 'Submitted': st.info("📝 Your self-assessment has been submitted and is awaiting review.")
            elif reviewer_scores and not a.get('acceptance'):
                st.markdown("---"); st.success(f"✅ {reviewer_type} review complete — awaiting your acceptance")
                with st.expander("📊 Score Comparison", expanded=True):
                    for score_key, staff_score in sorted(a['scores'].items(), key=natural_sort_key):
                        r_score = reviewer_scores.get(score_key, 'N/A') if reviewer_scores else 'N/A'
                        c1, c2 = st.columns(2)
                        c1.metric("Your Score", f"{staff_score}%"); c2.metric(f"{reviewer_type} Score", f"{r_score}%" if r_score != 'N/A' else 'N/A')
                
                st.markdown("---"); st.markdown(f"### 🔍 Accept or Reject {reviewer_type} Review")
                with st.form(f"accept_reject_form_{user_name}", clear_on_submit=True):
                    rejection_comment = st.text_area("Comments/Remarks *", placeholder="Required if rejecting...", key=f"rejection_comment_{user_name}")
                    st.caption("📎 Attach Documents (Optional — up to 2 files)")
                    d1, d2 = st.columns(2); rejection_docs = []
                    with d1:
                        doc1 = st.file_uploader("Document 1", type=['pdf','docx','jpg','png','xlsx'], key=f"rej_doc1_{user_name}")
                        if doc1: rejection_docs.append(doc1)
                    with d2:
                        doc2 = st.file_uploader("Document 2", type=['pdf','docx','jpg','png','xlsx'], key=f"rej_doc2_{user_name}")
                        if doc2: rejection_docs.append(doc2)
                    col1, col2 = st.columns(2)
                    with col1: 
                        accept_btn = st.form_submit_button("✅ Accept Review", use_container_width=True, type="primary")
                    with col2: 
                        reject_btn = st.form_submit_button("❌ Reject - Request Re-review", use_container_width=True)
                    
                    if accept_btn:
                        st.session_state.self_assessments[user_name]['acceptance'] = 'Accepted'
                        st.session_state.self_assessments[user_name]['status'] = 'Completed'
                        try: 
                            db.archive_appraisal(user_name, user_email, user_dept, st.session_state.appraisal_cycle_name, 'Accepted', a['scores'], reviewer_scores or {}, a.get('comments', ''), a.get('hod_comments', ''), now_wat.strftime('%Y-%m-%d %H:%M WAT'))
                            # Also update appraisals table
                            db._patch("appraisals", {"status": "Completed", "acceptance": "Accepted"}, {"user_name": user_name, "cycle_name": st.session_state.appraisal_cycle_name})
                        except Exception as e: 
                            st.error(f"Save error: {str(e)}")
                        log_audit('Appraisal Accepted', f'{user_name} accepted {reviewer_type} review')
                        st.success("✅ Appraisal Accepted! Congratulations!"); st.balloons(); time.sleep(2); st.rerun()
                    
                    if reject_btn:
                        if not rejection_comment.strip(): 
                            st.error("❌ You MUST provide comments when rejecting!")
                        else:
                            rej_urls = []
                            for doc in rejection_docs:
                                try:
                                    url = db.upload_file("rejection_evidence", f"reject_{user_name}_{doc.name}", doc.read(), doc.type)
                                    if url: rej_urls.append(url)
                                except: pass
                            st.session_state.self_assessments[user_name]['acceptance'] = 'Rejected'
                            st.session_state.self_assessments[user_name]['status'] = f'Awaiting {reviewer_type} Re-review'
                            st.session_state.self_assessments[user_name]['rejection_comment'] = rejection_comment
                            st.session_state.self_assessments[user_name]['rejection_docs'] = json.dumps(rej_urls)
                            st.session_state.self_assessments[user_name]['reject_count'] = a.get('reject_count', 0) + 1
                            
                            # Save to DB
                            try:
                                db.save_appraisal(user_name, user_email, user_dept, st.session_state.appraisal_cycle_name, f'Awaiting {reviewer_type} Re-review', a['scores'], a.get('comments', ''), a.get('pillar_comments', {}), a.get('hod_scores'), a.get('hod_comments', ''), a.get('hod_pillar_comments', {}), 'Rejected', None, a.get('date', ''))
                            except: pass
                            
                            # Notify reviewer
                            reviewer_email = find_hod_email_for_dept(user_dept) if reviewer_type == 'HOD' else get_employee_email(user_name)
                            if reviewer_email:
                                try: 
                                    EmailService().send_email(reviewer_email, f"🔄 Appraisal Rejected by {user_name}", f"Dear {reviewer_type},\n\n{user_name} has rejected your review.\n\nReason: {rejection_comment}\n\nPlease log in to review.\n\nhttps://churchgate-churchgate-hris.hf.space\n\nChurchgate Group HR")
                                except: pass
                            
                            log_audit('Appraisal Rejected', f'{user_name} rejected {reviewer_type} review')
                            st.warning(f"⚠️ Rejected! {reviewer_type} notified."); time.sleep(2); st.rerun()
            elif a.get('acceptance') == 'Accepted': st.success("🎉 Appraisal Complete!")
            elif a.get('acceptance') == 'Rejected': st.warning(f"🔄 Awaiting {reviewer_type} re-review")
    
    # ============================================================
    # TAB 4: HOD REVIEW (FIXED - Stand Firm + Admin sees ALL)
    # ============================================================
    with tab4:
        st.markdown('<div class="glass-card"><h3>👔 HOD Review Hub</h3></div>', unsafe_allow_html=True)
        if not is_hod: st.info("This section is for Managers, HODs, and Admins only.")
        else:
            # ===== SECTION 1: KPI APPROVAL =====
            st.markdown("### 📊 Team KPI Submissions")
            try:
                all_perf = db._get("performance_data"); team_submissions = {}
                for row in (all_perf or []):
                    if row.get('submission_status') == 'Submitted':
                        clean_name = ' '.join(str(row.get('user_name', '')).split())
                        # Admin sees ALL, HOD sees their department
                        if is_admin or get_employee_dept(clean_name) == user_dept:
                            if clean_name not in team_submissions: team_submissions[clean_name] = []
                            team_submissions[clean_name].append({'pillar': row.get('pillar_name', ''), 'kpis': json.loads(row.get('kpi_data', '[]')) if row.get('kpi_data') else [], 'row_id': row.get('id')})
                if team_submissions:
                    st.success(f"📋 {len(team_submissions)} team member(s)")
                    for emp_name, submissions in team_submissions.items():
                        with st.expander(f"👤 {emp_name}", expanded=False):
                            for sub in submissions:
                                st.markdown(f"**{sub['pillar']}**")
                                for kpi in sub['kpis']: st.markdown(f"• {kpi.get('kpi', 'N/A')}")
                            hod_comment = st.text_area(f"💬 Comment", key=f"hod_comment_{emp_name}")
                            c1, c2 = st.columns(2)
                            with c1:
                                if st.button(f"✅ Approve", key=f"app_{emp_name}", type="primary"):
                                    for sub in submissions: db._patch("performance_data", {"submission_status": "Approved"}, {"id": sub['row_id']})
                                    emp_email_addr = get_employee_email(emp_name)
                                    if emp_email_addr: send_kpi_notification('approved', emp_name, emp_email_addr)
                                    log_audit("KPIs Approved", f"HOD approved KPIs for {emp_name}")
                                    st.success("✅ Approved!"); st.balloons(); time.sleep(1); st.rerun()
                            with c2:
                                if st.button(f"🔄 Revise", key=f"rev_{emp_name}"):
                                    if hod_comment:
                                        for sub in submissions:
                                            db._patch("performance_data", {"submission_status": "Draft"}, {"id": sub['row_id']})
                                        emp_email_addr = get_employee_email(emp_name)
                                        if emp_email_addr:
                                            send_kpi_notification('revision_requested', emp_name, emp_email_addr)
                                        st.warning("🔄 Revision requested")
                                        time.sleep(1)
                                        st.rerun()
                                    else:
                                        st.error("❌ Please provide a comment!")
                else:
                    st.info("No pending KPI submissions.")
            except Exception as e:
                st.error(f"Error: {str(e)}")
            
            # ===== SECTION 1B: APPROVED KPIs PREVIEW =====
            st.markdown("---")
            st.markdown("### ✅ Team Approved KPIs")
            try:
                all_perf = db._get("performance_data")
                team_approved = {}
                for row in all_perf:
                    if row.get('submission_status') == 'Approved':
                        clean_name = ' '.join(str(row.get('user_name', '')).split())
                        if is_admin or get_employee_dept(clean_name) == user_dept:
                            if clean_name not in team_approved:
                                team_approved[clean_name] = []
                            kpi_list = json.loads(row.get('kpi_data', '[]')) if row.get('kpi_data') else []
                            team_approved[clean_name].append({
                                'pillar': row.get('pillar_name', ''),
                                'kpis': kpi_list,
                                'weight': row.get('weight', 0)
                            })
                if team_approved:
                    st.success(f"✅ {len(team_approved)} team member(s) with approved KPIs")
                    for emp_name, kpi_data in team_approved.items():
                        with st.expander(f"✅ {emp_name} — {len(kpi_data)} pillar(s) approved", expanded=False):
                            for entry in kpi_data:
                                st.markdown(f"**{entry['pillar']}** (Weight: {entry['weight']}%)")
                                for kpi in entry['kpis']:
                                    st.markdown(f"• {kpi.get('kpi', 'N/A')} — Target: {kpi.get('target', 'N/A')} — Weight: {kpi.get('weight', 'N/A')}%")
                                st.markdown("")
                else:
                    st.info("No team members have approved KPIs yet.")
            except:
                pass
            
            # ===== SECTION 2: APPRAISAL REVIEW (FIXED) =====
            st.markdown("---")
            st.markdown("### 📝 Appraisal Review")
            
            if is_admin:
                submitted_appraisals = {k: v for k, v in st.session_state.self_assessments.items() 
                                       if v['status'] in ['Submitted', 'Awaiting HOD Re-review', 'Escalated from TL']}
            else:
                submitted_appraisals = {k: v for k, v in st.session_state.self_assessments.items() 
                                       if get_employee_dept(k) == user_dept and v['status'] in ['Submitted', 'Awaiting HOD Re-review', 'Escalated from TL']}
            
            if submitted_appraisals:
                st.success(f"📋 {len(submitted_appraisals)} appraisal(s) for review")
                
                for staff_name, assessment in submitted_appraisals.items():
                    is_escalated = assessment.get('status') == 'Escalated from TL'
                    is_re_review = assessment.get('status') == 'Awaiting HOD Re-review'
                    
                    expander_title = f"{'🚨 ESCALATED: ' if is_escalated else '🔄 RE-REVIEW: ' if is_re_review else '📋 '}{staff_name} | {get_employee_dept(staff_name)}"
                    
                    with st.expander(expander_title, expanded=True):
                        if is_escalated: 
                            st.warning(f"🚨 Escalated from Team Lead. TL scores available for reference.")
                            if assessment.get('tl_scores'):
                                st.markdown("**Team Lead Scores:**")
                                for k, v in assessment['tl_scores'].items():
                                    st.markdown(f"- {k[:60]}: {v}%")
                        
                        if is_re_review: 
                            st.warning(f"⚠️ Staff rejected your review (Rejection #{assessment.get('reject_count', 1)})")
                            st.markdown(f"**Staff Rejection Reason:** {assessment.get('rejection_comment', 'No comment provided')}")
                        
                        st.markdown(f"**Staff Overall Comments:** {assessment.get('comments', 'N/A')}")
                        
                        st.markdown("---")
                        st.markdown("### 📊 Your Review Scores")
                        
                        hod_scores = {}
                        for score_key, staff_score in sorted(assessment['scores'].items(), key=natural_sort_key):
                            c1, c2 = st.columns(2)
                            with c1: 
                                st.markdown(f"**Staff Score:** {staff_score}%")
                                st.caption(score_key[:80])
                            with c2:
                                # Show previous HOD score for re-reviews
                                prev_hod = assessment.get('hod_scores', {}).get(score_key, staff_score) if is_re_review else staff_score
                                hod_scores[score_key] = st.number_input(
                                    "Your Score",
                                    min_value=0, max_value=100,
                                    value=int(prev_hod) if prev_hod else int(staff_score),
                                    step=1,
                                    key=f"hod_{staff_name}_{score_key}"
                                )
                        
                        st.markdown("---")
                        hod_overall = st.text_area(
                            f"Your Overall Comments for {staff_name} *",
                            value=assessment.get('hod_comments', '') if is_re_review else '',
                            key=f"hod_app_{staff_name}",
                            placeholder="Provide detailed feedback..."
                        )
                        
                        # BUTTONS: Submit + Stand Firm (for re-reviews)
                        if is_re_review or is_escalated:
                            c1, c2, c3 = st.columns(3)
                            with c1:
                                if st.button(f"✅ Submit Revised Review", key=f"submit_{staff_name}", type="primary"):
                                    if not hod_overall: 
                                        st.error("❌ Comments required!")
                                    else:
                                        st.session_state.self_assessments[staff_name].update({
                                            'status': 'Approved', 'hod_scores': hod_scores, 
                                            'hod_comments': hod_overall, 'acceptance': None, 'reviewer_type': 'HOD'
                                        })
                                        try:
                                            db.save_appraisal(staff_name, assessment.get('email', ''), 
                                                get_employee_dept(staff_name), st.session_state.appraisal_cycle_name, 
                                                'Approved', assessment['scores'], assessment.get('comments', ''),
                                                assessment.get('pillar_comments', {}), hod_scores, hod_overall, 
                                                {}, None, None, assessment.get('date', ''))
                                        except: pass
                                        # Notify employee
                                        emp_email = get_employee_email(staff_name)
                                        if emp_email:
                                            try:
                                                EmailService().send_email(emp_email, 
                                                    f"📝 Updated HOD Review - {st.session_state.appraisal_cycle_name}",
                                                    f"Dear {staff_name},\n\nYour HOD has submitted an updated review.\n\nPlease log in to accept or reject.\n\nChurchgate Group HR")
                                            except: pass
                                        log_audit('HOD Revised Review', f'{staff_name} review revised by HOD {user_name}')
                                        st.success("✅ Revised review submitted!"); st.balloons(); time.sleep(1.5); st.rerun()
                            
                            with c2:
                                if st.button(f"✋ Stand Firm - Send to Committee", key=f"standfirm_{staff_name}"):
                                    if not hod_overall:
                                        # Use existing HOD comments if no new ones
                                        hod_overall = assessment.get('hod_comments', 'Standing firm on original review.')
                                    
                                    st.session_state.self_assessments[staff_name].update({
                                        'status': 'Escalated from TL' if is_escalated else 'Approved',
                                        'acceptance': 'Rejected',
                                        'hod_scores': hod_scores if hod_scores else assessment.get('hod_scores', {}),
                                        'hod_comments': hod_overall,
                                        'sr_decision': 'Pending Committee'
                                    })
                                    try:
                                        db.save_appraisal(staff_name, assessment.get('email', ''), 
                                            get_employee_dept(staff_name), st.session_state.appraisal_cycle_name,
                                            'Escalated from TL' if is_escalated else 'Approved',
                                            assessment['scores'], assessment.get('comments', ''),
                                            assessment.get('pillar_comments', {}),
                                            st.session_state.self_assessments[staff_name].get('hod_scores', {}),
                                            hod_overall, {}, 'Rejected', 'Pending Committee',
                                            assessment.get('date', ''))
                                    except: pass
                                    
                                    # Notify Senior Management
                                    try:
                                        sr_emails = employees_df[employees_df['department'] == 'Senior Management']['email'].dropna().tolist() if not employees_df.empty else []
                                        for sr_email in sr_emails:
                                            if sr_email and '@' in str(sr_email):
                                                EmailService().send_email(sr_email,
                                                    f"🚨 Appraisal Escalated to Committee: {staff_name}",
                                                    f"Dear Committee Member,\n\n{staff_name}'s appraisal has been escalated.\n\nHOD: {user_name}\nHOD Comments: {hod_overall}\n\nPlease review in the Appraisal Committee Board.\n\nChurchgate Group HR")
                                    except: pass
                                    
                                    log_audit('HOD Escalated to Committee', f'{staff_name} sent to committee by HOD {user_name}')
                                    st.warning("✋ Escalated to Appraisal Committee!"); time.sleep(1.5); st.rerun()
                            
                            with c3:
                                if st.button(f"💬 Request Staff Revision", key=f"request_rev_{staff_name}"):
                                    st.session_state.self_assessments[staff_name]['status'] = 'Revision Requested by HOD'
                                    emp_email = get_employee_email(staff_name)
                                    if emp_email:
                                        try:
                                            EmailService().send_email(emp_email,
                                                f"🔄 Revision Requested - {st.session_state.appraisal_cycle_name}",
                                                f"Dear {staff_name},\n\nYour HOD has requested additional information.\n\nComments: {hod_overall}\n\nPlease update your self-assessment.\n\nChurchgate Group HR")
                                        except: pass
                                    log_audit('HOD Requested Revision', f'{staff_name} revision requested by HOD')
                                    st.info("💬 Revision requested from staff"); time.sleep(1.5); st.rerun()
                        else:
                            # First-time review
                            if st.button(f"✅ Submit HOD Review for {staff_name}", key=f"submit_{staff_name}", type="primary"):
                                if not hod_overall: 
                                    st.error("❌ Comments required!")
                                else:
                                    st.session_state.self_assessments[staff_name].update({
                                        'status': 'Approved', 'hod_scores': hod_scores, 
                                        'hod_comments': hod_overall, 'acceptance': None, 'reviewer_type': 'HOD'
                                    })
                                    try:
                                        db.save_appraisal(staff_name, assessment.get('email', ''), 
                                            get_employee_dept(staff_name), st.session_state.appraisal_cycle_name, 
                                            'Approved', assessment['scores'], assessment.get('comments', ''),
                                            assessment.get('pillar_comments', {}), hod_scores, hod_overall, 
                                            {}, None, None, assessment.get('date', ''))
                                    except: pass
                                    # Notify employee
                                    emp_email = get_employee_email(staff_name)
                                    if emp_email:
                                        try:
                                            EmailService().send_email(emp_email,
                                                f"📝 HOD Review Complete - {st.session_state.appraisal_cycle_name}",
                                                f"Dear {staff_name},\n\nYour HOD has completed your performance review.\n\nPlease log in to accept or reject.\n\nChurchgate Group HR")
                                        except: pass
                                    log_audit('HOD Review Submitted', f'{staff_name} reviewed by HOD {user_name}')
                                    st.success("✅ Review submitted!"); st.balloons(); time.sleep(1.5); st.rerun()
            else:
                st.info("No pending appraisals for review.")
    
    # ============================================================
    # TAB 5: TEAM LEAD/MANAGER REVIEW (STRICT REPORTS-TO FILTER)
    # ============================================================
    with tab5:
        st.markdown('<div class="glass-card"><h3>👥 Team Lead / Manager Review</h3><p style="color:#888;">Review ONLY team members who report directly to you</p></div>', unsafe_allow_html=True)
        
        if not is_team_lead_or_manager:
            st.info("This section is for Team Leads and Managers only.")
        else:
            # STRICT FILTER: Find employees whose reports_to matches current user's name
            team_members = []
            if not employees_df.empty:
                for _, emp in employees_df.iterrows():
                    emp_name = f"{emp['first_name']} {emp['last_name']}".strip()
                    reports_to = str(emp.get('reports_to', '')).strip()
                    
                    # Check if this employee reports to the current user
                    if reports_to.lower() == user_name.lower():
                        team_members.append(emp_name)
            
            if not team_members:
                st.warning("⚠️ No team members report directly to you. You can only review employees who have you listed as their 'Reports To' in the employee directory.")
                st.info("If team members should report to you, ask HR/Admin to update their 'Reports To' field in Employee Management.")
            else:
                st.success(f"👥 {len(team_members)} team member(s) reporting directly to you")
                
                for member_name in team_members:
                    assessment = st.session_state.self_assessments.get(member_name, {})
                    member_status = assessment.get('status', 'Not Submitted')
                    member_dept = get_employee_dept(member_name)
                    
                    with st.expander(f"👤 {member_name} | {member_dept} | Status: {member_status}", expanded=(member_status in ['Submitted', 'Awaiting TL Re-review'])):
                        if member_status == 'Not Submitted' or not assessment:
                            st.info(f"⏳ {member_name} has not submitted their self-assessment yet.")
                        elif member_status == 'Submitted' or member_status == 'Awaiting TL Re-review':
                            st.markdown(f"**Staff Comments:** {assessment.get('comments', 'N/A')}")
                            
                            if member_status == 'Awaiting TL Re-review':
                                st.warning(f"⚠️ Staff rejected your previous review (Rejection #{assessment.get('reject_count', 1)})")
                                st.markdown(f"**Rejection Reason:** {assessment.get('rejection_comment', 'No comment provided')}")
                            
                            if assessment.get('scores'):
                                st.markdown("---")
                                st.markdown("### 📊 Score Review")
                                tl_scores = {}
                                for score_key, staff_score in sorted(assessment['scores'].items(), key=natural_sort_key):
                                    c1, c2 = st.columns(2)
                                    with c1:
                                        st.markdown(f"**Staff Score:** {staff_score}%")
                                        st.caption(score_key[:80])
                                    with c2:
                                        existing_tl_score = assessment.get('tl_scores', {}).get(score_key, staff_score) if member_status == 'Awaiting TL Re-review' else staff_score
                                        tl_scores[score_key] = st.number_input(
                                            "Your Score",
                                            min_value=0, max_value=100,
                                            value=int(existing_tl_score) if existing_tl_score else int(staff_score),
                                            step=1,
                                            key=f"tl_{member_name}_{score_key}"
                                        )
                                
                                st.markdown("---")
                                tl_comments = st.text_area(
                                    f"Your Overall Comments for {member_name} *",
                                    value=assessment.get('tl_comments', '') if member_status == 'Awaiting TL Re-review' else '',
                                    key=f"tl_comment_{member_name}",
                                    placeholder="Provide detailed feedback on their performance..."
                                )
                                
                                c1, c2, c3 = st.columns(3)
                                with c1:
                                    if st.button(f"✅ Submit Review", key=f"tl_submit_{member_name}", type="primary"):
                                        if not tl_comments:
                                            st.error("❌ Comments are required!")
                                        else:
                                            # Save to session state
                                            st.session_state.self_assessments[member_name].update({
                                                'status': 'Approved',
                                                'tl_scores': tl_scores,
                                                'tl_comments': tl_comments,
                                                'hod_scores': tl_scores,
                                                'hod_comments': tl_comments,
                                                'acceptance': None,
                                                'reviewer_type': 'Team Lead'
                                            })
                                            
                                            # Save to database
                                            try:
                                                db.save_appraisal(
                                                    member_name,
                                                    assessment.get('email', ''),
                                                    member_dept,
                                                    st.session_state.appraisal_cycle_name,
                                                    'Approved',
                                                    assessment['scores'],
                                                    assessment.get('comments', ''),
                                                    assessment.get('pillar_comments', {}),
                                                    tl_scores,
                                                    tl_comments,
                                                    {},
                                                    None,
                                                    None,
                                                    assessment.get('date', '')
                                                )
                                            except Exception as e:
                                                st.error(f"Save error: {str(e)}")
                                            
                                            # Send email to team member
                                            member_email = get_employee_email(member_name)
                                            if member_email:
                                                try:
                                                    EmailService().send_email(
                                                        member_email,
                                                        f"📝 Performance Review Complete - {st.session_state.appraisal_cycle_name}",
                                                        f"Dear {member_name},\n\n"
                                                        f"Your Team Lead ({user_name}) has completed your performance review.\n\n"
                                                        f"Please log in to the HRIS to view your scores and accept or reject the review.\n\n"
                                                        f"https://churchgate-churchgate-hris.hf.space\n\n"
                                                        f"Churchgate Group HR"
                                                    )
                                                except:
                                                    pass
                                            
                                            log_audit('TL Review Submitted', f'{member_name} reviewed by Team Lead {user_name}')
                                            st.success(f"✅ Review submitted for {member_name}!")
                                            st.balloons()
                                            time.sleep(1.5)
                                            st.rerun()
                                
                                with c2:
                                    if member_status == 'Awaiting TL Re-review' and st.button(f"🚨 Escalate to HOD", key=f"tl_escalate_{member_name}"):
                                        st.session_state.self_assessments[member_name].update({
                                            'status': 'Escalated from TL',
                                            'tl_scores': tl_scores if tl_scores else assessment.get('tl_scores', {}),
                                            'tl_comments': tl_comments if tl_comments else assessment.get('tl_comments', ''),
                                            'acceptance': 'Rejected'
                                        })
                                        
                                        # Save to database
                                        try:
                                            db.save_appraisal(
                                                member_name,
                                                assessment.get('email', ''),
                                                member_dept,
                                                st.session_state.appraisal_cycle_name,
                                                'Escalated from TL',
                                                assessment['scores'],
                                                assessment.get('comments', ''),
                                                assessment.get('pillar_comments', {}),
                                                st.session_state.self_assessments[member_name].get('tl_scores', {}),
                                                st.session_state.self_assessments[member_name].get('tl_comments', ''),
                                                {},
                                                'Rejected',
                                                None,
                                                assessment.get('date', '')
                                            )
                                        except:
                                            pass
                                        
                                        # Notify HOD
                                        hod_email = find_hod_email_for_dept(member_dept)
                                        if hod_email:
                                            try:
                                                EmailService().send_email(
                                                    hod_email,
                                                    f"🚨 Escalated Appraisal: {member_name}",
                                                    f"Dear HOD,\n\n"
                                                    f"The appraisal for {member_name} ({member_dept}) has been escalated by Team Lead {user_name}.\n\n"
                                                    f"The team member rejected the Team Lead's review.\n"
                                                    f"Rejection reason: {assessment.get('rejection_comment', 'No comment')}\n\n"
                                                    f"Please review and make the final decision.\n\n"
                                                    f"https://churchgate-churchgate-hris.hf.space\n\n"
                                                    f"Churchgate Group HR"
                                                )
                                            except:
                                                pass
                                        
                                        log_audit('TL Escalated to HOD', f'{member_name} escalated to HOD by Team Lead {user_name}')
                                        st.warning(f"🚨 Escalated to HOD for final decision!")
                                        time.sleep(1.5)
                                        st.rerun()
                                
                                with c3:
                                    if st.button(f"💬 Request Revision", key=f"tl_request_revision_{member_name}"):
                                        st.session_state.self_assessments[member_name]['status'] = 'Revision Requested by TL'
                                        member_email = get_employee_email(member_name)
                                        if member_email:
                                            try:
                                                EmailService().send_email(
                                                    member_email,
                                                    f"🔄 Revision Requested - {st.session_state.appraisal_cycle_name}",
                                                    f"Dear {member_name},\n\n"
                                                    f"Your Team Lead ({user_name}) has requested a revision to your self-assessment.\n\n"
                                                    f"Please review and update your submission.\n\n"
                                                    f"https://churchgate-churchgate-hris.hf.space\n\n"
                                                    f"Churchgate Group HR"
                                                )
                                            except:
                                                pass
                                        log_audit('TL Requested Revision', f'{member_name} revision requested by {user_name}')
                                        st.info(f"💬 Revision requested from {member_name}")
                                        time.sleep(1.5)
                                        st.rerun()
                            else:
                                st.info("No scores available for review.")
                        elif member_status == 'Approved' and not assessment.get('acceptance'):
                            st.info(f"✅ Review submitted. Waiting for {member_name} to accept or reject.")
                        elif assessment.get('acceptance') == 'Accepted':
                            st.success(f"🎉 Appraisal complete for {member_name}!")
                        elif assessment.get('acceptance') == 'Rejected' and member_status == 'Escalated from TL':
                            st.warning(f"🚨 Escalated to HOD for final decision.")
                        else:
                            st.info(f"Status: {member_status}")
    
    # ============================================================
    # TAB 6: EXCEPTIONAL ACHIEVEMENTS
    # ============================================================
    with tab6:
        st.markdown('<div class="glass-card"><h3>🌟 Exceptional Achievements</h3></div>', unsafe_allow_html=True)
        st.info("Document accomplishments outside your formal KPIs.")
        categories = {"💡 Innovation": "New ideas", "👑 Leadership": "Leading teams", "😊 Customer Impact": "Going above and beyond", "💰 Cost Savings": "Efficiency gains", "🚨 Crisis Management": "Handling emergencies", "🤝 Teamwork": "Collaboration"}
        my_achievements = st.session_state.exceptional_achievements.get(user_name, [])
        c1, c2, c3 = st.columns(3)
        c1.metric("🏆 Achievements", len(my_achievements))
        c2.metric("⭐ Avg Impact", f"{sum(3 if a.get('impact')=='Organization' else 2 if a.get('impact')=='Department' else 1 for a in my_achievements) / len(my_achievements):.1f}" if my_achievements else "N/A")
        c3.metric("🎖️ Badges", len(set(a.get('category', '') for a in my_achievements)))
        if my_achievements:
            for i, ach in enumerate(sorted(my_achievements, key=lambda x: x.get('date', ''), reverse=True)):
                with st.expander(f"{ach.get('category', '💡')} {ach.get('title', '')}", expanded=i==0):
                    st.markdown(f"**Description:** {ach.get('description', '')}")
                    st.markdown(f"**Outcome:** {ach.get('outcome', '')}")
        with st.form("add_achievement"):
            c1, c2 = st.columns(2)
            with c1: ach_title = st.text_input("Title *"); ach_category = st.selectbox("Category", list(categories.keys())); ach_impact = st.selectbox("Impact", ["Individual", "Team", "Department", "Organization"]); ach_date = st.date_input("Date")
            with c2: ach_description = st.text_area("Description *", height=100); ach_outcome = st.text_area("Outcome *", height=100)
            if st.form_submit_button("💾 Save", use_container_width=True):
                if ach_title and ach_description and ach_outcome:
                    if user_name not in st.session_state.exceptional_achievements: st.session_state.exceptional_achievements[user_name] = []
                    st.session_state.exceptional_achievements[user_name].append({'title': ach_title, 'category': ach_category, 'description': ach_description, 'impact': ach_impact, 'outcome': ach_outcome, 'date': ach_date.strftime('%Y-%m-%d')})
                    st.success("✅ Saved!"); st.balloons(); time.sleep(1); st.rerun()
                else: st.error("❌ All fields required!")
    
    # ============================================================
    # TAB 7: APPRAISAL SETTINGS (ADMIN ONLY)
    # ============================================================
    with tab7:
        st.markdown('<div class="glass-card"><h3>⚙️ Appraisal Cycle Settings</h3><p style="color:#888;">Admin & HR Only — Changes apply to ALL users</p></div>', unsafe_allow_html=True)
        
        if not is_admin:
            st.info("⛔ This section is restricted to Admin and HR Director only.")
        else:
            all_perf_data = get_all_perf_cached()
            approved_kpi_count = len(all_perf_data[all_perf_data['submission_status'] == 'Approved']) if not all_perf_data.empty else 0
            submitted_appraisal_count = len([v for v in st.session_state.self_assessments.values() if v['status'] == 'Submitted'])
            completed_count = len([v for v in st.session_state.self_assessments.values() if v.get('acceptance') == 'Accepted'])
            escalated_count = len([v for v in st.session_state.self_assessments.values() if v.get('status') == 'Escalated from TL'])
            
            c1, c2, c3, c4 = st.columns(4)
            c1.metric("✅ KPIs Approved", approved_kpi_count)
            c2.metric("📝 Appraisals Submitted", submitted_appraisal_count)
            c3.metric("🎉 Completed", completed_count)
            c4.metric("🚨 Escalated", escalated_count)
            
            st.markdown("---")
            if st.session_state.appraisal_cycle_active:
                st.success(f"🟢 Appraisal Cycle ACTIVE: **{st.session_state.appraisal_cycle_name}**")
                st.markdown(f"Period: {st.session_state.appraisal_start} to {st.session_state.appraisal_end}")
            else:
                st.info("⚪ No active appraisal cycle.")
            
            st.markdown("---")
            st.markdown("### ⚙️ Cycle Configuration")
            st.session_state.appraisal_cycle_active = st.checkbox("Activate Appraisal Cycle", value=st.session_state.appraisal_cycle_active)
            cycle_options = ['Half-Year Appraisal', 'Full-Year Appraisal', 'HOD Mock Appraisal', 'Team Mock Appraisal']
            st.session_state.appraisal_cycle_name = st.selectbox("Select Appraisal Cycle", cycle_options, index=0 if 'Half-Year' in st.session_state.appraisal_cycle_name else 0)
            c1, c2 = st.columns(2)
            with c1:
                start_val = st.session_state.appraisal_start
                if isinstance(start_val, str):
                    try: start_val = datetime.strptime(start_val, '%Y-%m-%d')
                    except: start_val = datetime.now()
                st.session_state.appraisal_start = st.date_input("Start Date", value=start_val if hasattr(start_val, 'strftime') else datetime.now())
            with c2:
                end_val = st.session_state.appraisal_end
                if isinstance(end_val, str):
                    try: end_val = datetime.strptime(end_val, '%Y-%m-%d')
                    except: end_val = datetime.now()
                st.session_state.appraisal_end = st.date_input("End Date", value=end_val if hasattr(end_val, 'strftime') else datetime.now())
            st.session_state.appraisal_locked = st.checkbox("Lock Scores", value=st.session_state.appraisal_locked)
            
            col1, col2 = st.columns(2)
            with col1:
                if st.button("💾 Activate & Notify All Eligible", use_container_width=True, type="primary"):
                    st.session_state.appraisal_start = st.session_state.appraisal_start.strftime('%Y-%m-%d') if hasattr(st.session_state.appraisal_start, 'strftime') else str(st.session_state.appraisal_start)
                    st.session_state.appraisal_end = st.session_state.appraisal_end.strftime('%Y-%m-%d') if hasattr(st.session_state.appraisal_end, 'strftime') else str(st.session_state.appraisal_end)
                    save_appraisal_cycle_to_db()
                    try:
                        es = EmailService(); sent_count = 0
                        for _, emp in employees_df.iterrows():
                            emp_name = f"{emp['first_name']} {emp['last_name']}"; emp_email_addr = emp.get('email', '')
                            if emp_email_addr and '@' in str(emp_email_addr):
                                emp_perf = db.get_performance_data(emp_name)
                                if not emp_perf.empty and any(row.get('submission_status') == 'Approved' for _, row in emp_perf.iterrows()):
                                    es.send_email(emp_email_addr, f"📊 Appraisal Cycle Open: {st.session_state.appraisal_cycle_name}", f"Dear {emp_name},\n\nThe appraisal cycle is now active.\n\nPeriod: {st.session_state.appraisal_start} to {st.session_state.appraisal_end}\n\nPlease complete your self-assessment.\n\nhttps://churchgate-churchgate-hris.hf.space\n\nChurchgate Group HR")
                                    sent_count += 1
                        st.success(f"📧 Emails sent to {sent_count} employees.")
                    except: pass
                    send_browser_notification("📊 Appraisal Cycle Active!", f"{st.session_state.appraisal_cycle_name} is now open.")
                    st.balloons(); st.rerun()
            with col2:
                if st.button("🛑 Deactivate Cycle", use_container_width=True):
                    st.session_state.appraisal_cycle_active = False
                    existing = db._get("appraisal_cycles")
                    if existing:
                        for c in existing: db._patch("appraisal_cycles", {"is_active": False}, {"id": c['id']})
                    st.warning("🛑 Cycle deactivated."); st.rerun()
    
    # ============================================================
    # TAB 8: ADVANCED DASHBOARD - NO st.rerun() - OPTIMIZED
    # ============================================================
    with tab8:
        st.markdown('<div class="glass-card"><h3>📊 Advanced Dashboard</h3><p style="color:#888;">Full Group-Wide Performance Disclosure & Analytics</p></div>', unsafe_allow_html=True)
        
        if is_admin or user_role in ['HR Director']:
            admin_dash_tabs = st.tabs(["🏢 Group Dashboard", "👤 My Performance"])
            
            # ============================================================
            # ADMIN SUB-TAB 0: FULL GROUP DASHBOARD
            # ============================================================
            with admin_dash_tabs[0]:
                
                SUBSIDIARY_OPTIONS_DASH = {
                    'Abuja': ['World Trade Center(WTC)', 'Agroline Ventures Limited'],
                    'Lagos': ['First Continental Properties Limited', 'R. B Properties Limited', 'Churchgate Nigeria Limited', 'Aba Textile Mills PLC', 'Associated Textile Manufacturing Company Limited', 'Food & Confectionery Products (Nig.) Limited', 'First Spinners PLC', 'HotelInvest & Resorts Limited', 'International Textile Industries (Nig.) Limited', 'Intercott Limited', 'Ocean Fisheries (Nig.) Limited', 'Platinum Travel Limited', 'Reliance Mills Limited', 'Vineyard Designs Nig. Limited'],
                    'Aba': ['Aba Textile Mills PLC']
                }
                
                st.markdown("### 🔍 Filter Dashboard Data")
                filter_col1, filter_col2, filter_col3 = st.columns(3)
                with filter_col1:
                    dash_region_filter = st.selectbox("🌍 Region", ['All Regions', 'Abuja', 'Lagos', 'Aba'], key="dash_region_filter")
                with filter_col2:
                    if dash_region_filter == 'All Regions':
                        all_subs = []
                        for subs in SUBSIDIARY_OPTIONS_DASH.values():
                            all_subs.extend(subs)
                    else:
                        all_subs = SUBSIDIARY_OPTIONS_DASH.get(dash_region_filter, [])
                    dash_sub_filter = st.selectbox("🏢 Subsidiary", ['All Subsidiaries'] + sorted(all_subs), key="dash_sub_filter")
                with filter_col3:
                    dash_dept_filter = st.selectbox("🏭 Department", ['All Departments'] + all_depts, key="dash_dept_filter")
                
                def apply_dash_filters(region, subsidiary, department):
                    if dash_region_filter != 'All Regions' and region != dash_region_filter:
                        return False
                    if dash_sub_filter != 'All Subsidiaries' and subsidiary != dash_sub_filter:
                        return False
                    if dash_dept_filter != 'All Departments' and department != dash_dept_filter:
                        return False
                    return True
                
                st.markdown("---")
                
                dash_tabs = st.tabs(["📈 Group Metrics", "👥 Reviewer Transparency", "📊 Score Matrix", "🔍 Rejection Analysis", "🏢 Hierarchy View"])
                all_perf = get_all_perf_cached()
                
                # ============================================================
                # DASH TAB 0: GROUP METRICS
                # ============================================================
                with dash_tabs[0]:
                    st.subheader("📈 Group-Wide Performance Metrics")
                    total_approved = len(all_perf[all_perf['submission_status'] == 'Approved']) if not all_perf.empty else 0
                    total_submitted = len(all_perf[all_perf['submission_status'] == 'Submitted']) if not all_perf.empty else 0
                    total_appraisals_in = len([v for v in st.session_state.self_assessments.values() if v['status'] in ['Submitted', 'Approved', 'Completed']])
                    total_completed = len([v for v in st.session_state.self_assessments.values() if v.get('acceptance') == 'Accepted'])
                    total_escalated = len([v for v in st.session_state.self_assessments.values() if v.get('status') == 'Escalated from TL'])
                    total_rejected = len([v for v in st.session_state.self_assessments.values() if v.get('acceptance') == 'Rejected'])
                    
                    c1, c2, c3, c4, c5, c6 = st.columns(6)
                    c1.metric("✅ KPIs Approved", total_approved)
                    c2.metric("📝 Submitted", total_submitted)
                    c3.metric("📊 Appraisals In", total_appraisals_in)
                    c4.metric("🎉 Completed", total_completed)
                    c5.metric("🚨 Escalated", total_escalated)
                    c6.metric("❌ Rejected", total_rejected)
                    
                    st.markdown("---")
                    st.subheader("🌍 Performance by Region → Subsidiary → Department")
                    
                    region_breakdown = defaultdict(lambda: defaultdict(lambda: defaultdict(lambda: {'total_employees': 0, 'approved_kpis': 0, 'appraisals_completed': 0, 'appraisals_rejected': 0, 'appraisals_escalated': 0})))
                    
                    if not employees_df.empty:
                        for _, emp in employees_df.iterrows():
                            emp_name = f"{emp['first_name']} {emp['last_name']}".strip()
                            if not emp_name: continue
                            dept = emp.get('department', 'General'); sub = emp.get('subsidiary', 'Unassigned'); region = get_region(sub)
                            if not apply_dash_filters(region, sub or 'Unassigned', dept): continue
                            stats = region_breakdown[region][sub or 'Unassigned'][dept]; stats['total_employees'] += 1
                            emp_perf = all_perf[all_perf['user_name'] == emp_name] if not all_perf.empty else pd.DataFrame()
                            if not emp_perf.empty and any(emp_perf['submission_status'] == 'Approved'): stats['approved_kpis'] += 1
                            assessment = st.session_state.self_assessments.get(emp_name, {})
                            if assessment.get('acceptance') == 'Accepted': stats['appraisals_completed'] += 1
                            if assessment.get('acceptance') == 'Rejected': stats['appraisals_rejected'] += 1
                            if assessment.get('status') == 'Escalated from TL': stats['appraisals_escalated'] += 1
                    
                    total_filtered = 0
                    for region in ['Abuja', 'Lagos', 'Aba']:
                        if region not in region_breakdown: continue
                        if dash_region_filter != 'All Regions' and region != dash_region_filter: continue
                        region_color = {'Abuja': '#CC0000', 'Lagos': '#1a1a1a', 'Aba': '#38a169'}.get(region, '#1a1a1a')
                        region_total = sum(d['total_employees'] for s in region_breakdown[region].values() for d in s.values())
                        region_approved = sum(d['approved_kpis'] for s in region_breakdown[region].values() for d in s.values())
                        region_completed = sum(d['appraisals_completed'] for s in region_breakdown[region].values() for d in s.values())
                        if region_total == 0: continue
                        total_filtered += region_total
                        st.markdown(f"""<div class="region-header" style="background:linear-gradient(135deg, {region_color}, #2d2d2d);display:flex;justify-content:space-between;align-items:center;"><span>🌍 {region} Region</span><span style="font-size:0.85rem;opacity:0.85;">👥 {region_total} | ✅ {region_approved} KPIs | 🎉 {region_completed} Completed</span></div>""", unsafe_allow_html=True)
                        
                        for subsidiary in sorted(region_breakdown[region].keys()):
                            sub_data = region_breakdown[region][subsidiary]
                            sub_emps = sum(d['total_employees'] for d in sub_data.values()); sub_approved = sum(d['approved_kpis'] for d in sub_data.values()); sub_completed = sum(d['appraisals_completed'] for d in sub_data.values())
                            if sub_emps == 0: continue
                            sub_key = f"metric_sub_{region}_{subsidiary}".replace(' ', '_').replace('(', '').replace(')', '').replace('.', '')
                            if sub_key not in st.session_state: st.session_state[sub_key] = False
                            expand_icon = "▼" if st.session_state[sub_key] else "▶"
                            col1, col2 = st.columns([1, 20])
                            with col1:
                                if st.button(expand_icon, key=f"btn_{sub_key}"):
                                    st.session_state[sub_key] = not st.session_state[sub_key]
                            with col2:
                                st.markdown(f"""<div class="subsidiary-header" style="margin:0.3rem 0;">🏢 {subsidiary} — 👥 {sub_emps} | ✅ {sub_approved} KPIs | 🎉 {sub_completed} Completed</div>""", unsafe_allow_html=True)
                            if st.session_state[sub_key]:
                                for department in sorted(sub_data.keys()):
                                    dept_stats = sub_data[department]
                                    st.markdown(f"""<div style="background:#f8f9fa;padding:0.5rem 1rem;border-radius:6px;margin:0.2rem 0 0.2rem 3rem;border-left:3px solid #CC0000;display:flex;justify-content:space-between;align-items:center;"><span>🏭 {department}</span><span style="font-size:0.8rem;">👥 {dept_stats['total_employees']} | ✅ {dept_stats['approved_kpis']} KPIs | 🎉 {dept_stats['appraisals_completed']} Done | ❌ {dept_stats['appraisals_rejected']} Rejected | 🚨 {dept_stats['appraisals_escalated']} Escalated</span></div>""", unsafe_allow_html=True)
                    
                    if total_filtered == 0: st.info("No data matches the selected filters.")
                    
                    st.markdown("---")
                    region_chart_data = []
                    for region in ['Abuja', 'Lagos', 'Aba']:
                        if region in region_breakdown and (dash_region_filter == 'All Regions' or region == dash_region_filter):
                            total_c = sum(d['appraisals_completed'] for s in region_breakdown[region].values() for d in s.values())
                            total_r = sum(d['appraisals_rejected'] for s in region_breakdown[region].values() for d in s.values())
                            region_chart_data.append({'Region': region, 'Completed': total_c, 'Rejected': total_r})
                    if region_chart_data:
                        chart_df = pd.DataFrame(region_chart_data)
                        fig = px.bar(chart_df, x='Region', y=['Completed', 'Rejected'], barmode='group', color_discrete_sequence=['#38a169', '#CC0000']); fig.update_layout(height=350)
                        st.plotly_chart(fig, use_container_width=True)
                
                # ============================================================
                # DASH TAB 1: REVIEWER TRANSPARENCY
                # ============================================================
                with dash_tabs[1]:
                    st.subheader("👥 Reviewer Scoring Transparency")
                    reviewer_breakdown = defaultdict(lambda: defaultdict(lambda: defaultdict(list)))
                    for emp_name, assessment in st.session_state.self_assessments.items():
                        scores = assessment.get('hod_scores') or assessment.get('tl_scores') or {}
                        if scores:
                            avg_score = sum(int(v) for v in scores.values() if v) / len(scores) if scores else 0
                            dept = get_employee_dept(emp_name); sub = get_employee_subsidiary(emp_name); region = get_region(sub)
                            if not apply_dash_filters(region, sub or 'Unassigned', dept): continue
                            reviewer_breakdown[region][sub or 'Unassigned'][dept].append({'reviewer': assessment.get('reviewer_type', 'N/A'), 'employee': emp_name, 'avg_score': round(avg_score, 1), 'status': assessment.get('acceptance', assessment.get('status', 'Pending'))})
                    
                    total_reviewers = sum(len(v) for r in reviewer_breakdown.values() for s in r.values() for v in s.values())
                    if total_reviewers == 0: st.info("No review data matches the selected filters.")
                    else:
                        for region in ['Abuja', 'Lagos', 'Aba']:
                            if region not in reviewer_breakdown: continue
                            region_color = {'Abuja': '#CC0000', 'Lagos': '#1a1a1a', 'Aba': '#38a169'}.get(region, '#1a1a1a')
                            region_count = sum(len(v) for s in reviewer_breakdown[region].values() for v in s.values())
                            st.markdown(f"""<div class="region-header" style="background:linear-gradient(135deg, {region_color}, #2d2d2d);">🌍 {region} — {region_count} reviews</div>""", unsafe_allow_html=True)
                            for subsidiary in sorted(reviewer_breakdown[region].keys()):
                                sub_key = f"rev_sub_{region}_{subsidiary}".replace(' ', '_').replace('(', '').replace(')', '').replace('.', '')
                                if sub_key not in st.session_state: st.session_state[sub_key] = False
                                expand_icon = "▼" if st.session_state[sub_key] else "▶"
                                col1, col2 = st.columns([1, 20])
                                with col1:
                                    if st.button(expand_icon, key=f"revbtn_{sub_key}"):
                                        st.session_state[sub_key] = not st.session_state[sub_key]
                                with col2: st.markdown(f"""<div class="subsidiary-header" style="margin:0.3rem 0;">🏢 {subsidiary}</div>""", unsafe_allow_html=True)
                                if st.session_state[sub_key]:
                                    for department in sorted(reviewer_breakdown[region][subsidiary].keys()):
                                        reviews = reviewer_breakdown[region][subsidiary][department]
                                        dept_avg = sum(r['avg_score'] for r in reviews) / len(reviews) if reviews else 0
                                        with st.expander(f"🏭 {department} — {len(reviews)} reviews | Avg: {dept_avg:.0f}%", expanded=False):
                                            for r in reviews:
                                                badge_class = {'Accepted': 'badge-green', 'Rejected': 'badge-red'}.get(r['status'], 'badge-yellow')
                                                st.markdown(f"""<div style="padding:0.4rem;margin:0.2rem 0;border-left:3px solid #CC0000;"><strong>{r['employee']}</strong> | Reviewer: {r['reviewer']} | Score: <strong>{r['avg_score']}%</strong> | <span class="badge {badge_class}">{r['status']}</span></div>""", unsafe_allow_html=True)
                        
                        dept_avg_scores = {}
                        for region_data in reviewer_breakdown.values():
                            for sub_data in region_data.values():
                                for dept, reviews in sub_data.items():
                                    if reviews: dept_avg_scores[dept] = sum(r['avg_score'] for r in reviews) / len(reviews)
                        if dept_avg_scores:
                            dept_avg_df = pd.DataFrame({'Department': list(dept_avg_scores.keys()), 'Avg Reviewer Score': list(dept_avg_scores.values())})
                            fig = px.bar(dept_avg_df, x='Department', y='Avg Reviewer Score', color='Avg Reviewer Score', color_continuous_scale=['#38a169', '#d69e2e', '#CC0000']); fig.update_layout(height=350)
                            st.plotly_chart(fig, use_container_width=True)
                
                # ============================================================
                # DASH TAB 2: SCORE MATRIX
                # ============================================================
                with dash_tabs[2]:
                    st.subheader("📊 Score Comparison Matrix")
                    score_breakdown = defaultdict(lambda: defaultdict(lambda: defaultdict(list)))
                    for emp_name, assessment in st.session_state.self_assessments.items():
                        self_scores = assessment.get('scores', {}); reviewer_scores = assessment.get('hod_scores') or assessment.get('tl_scores') or {}
                        if not self_scores: continue
                        dept = get_employee_dept(emp_name); sub = get_employee_subsidiary(emp_name); region = get_region(sub)
                        if not apply_dash_filters(region, sub or 'Unassigned', dept): continue
                        self_avg = sum(int(v) for v in self_scores.values() if v) / len(self_scores) if self_scores else 0
                        rev_avg = sum(int(v) for v in reviewer_scores.values() if v) / len(reviewer_scores) if reviewer_scores else 0
                        score_breakdown[region][sub or 'Unassigned'][dept].append({'employee': emp_name, 'self_avg': round(self_avg, 1), 'reviewer_avg': round(rev_avg, 1), 'diff': round(self_avg - rev_avg, 1), 'status': assessment.get('acceptance', 'Pending')})
                    
                    total_scores = sum(len(v) for r in score_breakdown.values() for s in r.values() for v in s.values())
                    if total_scores == 0: st.info("No score data matches the selected filters.")
                    else:
                        for region in ['Abuja', 'Lagos', 'Aba']:
                            if region not in score_breakdown: continue
                            region_color = {'Abuja': '#CC0000', 'Lagos': '#1a1a1a', 'Aba': '#38a169'}.get(region, '#1a1a1a')
                            st.markdown(f"""<div class="region-header" style="background:linear-gradient(135deg, {region_color}, #2d2d2d);">🌍 {region}</div>""", unsafe_allow_html=True)
                            for subsidiary in sorted(score_breakdown[region].keys()):
                                sub_key = f"score_sub_{region}_{subsidiary}".replace(' ', '_').replace('(', '').replace(')', '').replace('.', '')
                                if sub_key not in st.session_state: st.session_state[sub_key] = False
                                expand_icon = "▼" if st.session_state[sub_key] else "▶"
                                col1, col2 = st.columns([1, 20])
                                with col1:
                                    if st.button(expand_icon, key=f"scorebtn_{sub_key}"):
                                        st.session_state[sub_key] = not st.session_state[sub_key]
                                with col2: st.markdown(f'<div class="subsidiary-header" style="margin:0.3rem 0;">🏢 {subsidiary}</div>', unsafe_allow_html=True)
                                if st.session_state[sub_key]:
                                    for department in sorted(score_breakdown[region][subsidiary].keys()):
                                        entries = score_breakdown[region][subsidiary][department]
                                        with st.expander(f"🏭 {department} — {len(entries)} employees", expanded=False):
                                            for e in entries:
                                                diff_color = '#38a169' if e['diff'] >= 0 else '#CC0000'; status_badge = {'Accepted': 'badge-green', 'Rejected': 'badge-red'}.get(e['status'], 'badge-yellow')
                                                st.markdown(f"""<div style="padding:0.4rem;margin:0.2rem 0;border-left:3px solid {diff_color};display:flex;justify-content:space-between;align-items:center;"><strong>{e['employee']}</strong><span>Self: <strong>{e['self_avg']}%</strong> | Rev: <strong>{e['reviewer_avg']}%</strong> | Diff: <strong style="color:{diff_color};">{e['diff']:+.1f}%</strong> | <span class="badge {status_badge}">{e['status']}</span></span></div>""", unsafe_allow_html=True)
                
                # ============================================================
                # DASH TAB 3: REJECTION ANALYSIS
                # ============================================================
                with dash_tabs[3]:
                    st.subheader("🔍 Rejection Analysis")
                    rejections = []
                    for emp_name, assessment in st.session_state.self_assessments.items():
                        if assessment.get('acceptance') == 'Rejected':
                            dept = get_employee_dept(emp_name); sub = get_employee_subsidiary(emp_name); region = get_region(sub)
                            if not apply_dash_filters(region, sub or 'N/A', dept): continue
                            rejections.append({'Employee': emp_name, 'Department': dept, 'Subsidiary': sub or 'N/A', 'Region': region, 'Reviewer': assessment.get('reviewer_type', 'HOD'), 'Rejection Reason': assessment.get('rejection_comment', 'No comment'), 'Reject Count': assessment.get('reject_count', 1), 'Date': assessment.get('date', '')})
                    
                    if rejections:
                        rej_df = pd.DataFrame(rejections)
                        c1, c2, c3 = st.columns(3)
                        c1.metric("❌ Total Rejections", len(rejections)); c2.metric("🔄 Avg Reject Count", f"{sum(r['Reject Count'] for r in rejections) / len(rejections):.1f}"); c3.metric("🏢 Departments", len(set(r['Department'] for r in rejections)))
                        st.markdown("---")
                        rej_breakdown = defaultdict(lambda: defaultdict(lambda: defaultdict(list)))
                        for r in rejections: rej_breakdown[r['Region']][r['Subsidiary']][r['Department']].append(r)
                        for region in ['Abuja', 'Lagos', 'Aba']:
                            if region not in rej_breakdown: continue
                            region_color = {'Abuja': '#CC0000', 'Lagos': '#1a1a1a', 'Aba': '#38a169'}.get(region, '#1a1a1a')
                            st.markdown(f'<div class="region-header" style="background:linear-gradient(135deg, {region_color}, #2d2d2d);">🌍 {region} — {sum(len(v) for s in rej_breakdown[region].values() for v in s.values())} rejections</div>', unsafe_allow_html=True)
                            for sub in sorted(rej_breakdown[region].keys()):
                                sub_key = f"rej_sub_{region}_{sub}".replace(' ', '_').replace('(', '').replace(')', '').replace('.', '')
                                if sub_key not in st.session_state: st.session_state[sub_key] = False
                                expand_icon = "▼" if st.session_state[sub_key] else "▶"
                                col1, col2 = st.columns([1, 20])
                                with col1:
                                    if st.button(expand_icon, key=f"rejbtn_{sub_key}"):
                                        st.session_state[sub_key] = not st.session_state[sub_key]
                                with col2: st.markdown(f'<div class="subsidiary-header" style="margin:0.3rem 0;">🏢 {sub}</div>', unsafe_allow_html=True)
                                if st.session_state[sub_key]:
                                    for dept in sorted(rej_breakdown[region][sub].keys()):
                                        dept_rejs = rej_breakdown[region][sub][dept]
                                        with st.expander(f"🏭 {dept} — {len(dept_rejs)} rejection(s)", expanded=False):
                                            for r in dept_rejs:
                                                st.markdown(f"""<div style="padding:0.5rem;margin:0.3rem 0;border-left:4px solid #CC0000;background:#fff5f5;border-radius:6px;"><strong>👤 {r['Employee']}</strong> | Reviewer: {r['Reviewer']} | #{r['Reject Count']}<br><small style="color:#CC0000;">💬 {r['Rejection Reason'][:200]}</small></div>""", unsafe_allow_html=True)
                        
                        dept_rej_counts = rej_df['Department'].value_counts()
                        fig3 = px.pie(values=dept_rej_counts.values, names=dept_rej_counts.index, hole=0.5, title="Rejections by Department", color_discrete_sequence=['#CC0000', '#d69e2e', '#3182ce', '#38a169']); st.plotly_chart(fig3, use_container_width=True)
                        
                        if st.button("🤖 AI Analyze Rejection Patterns", use_container_width=True, type="primary"):
                            with st.spinner("🧠 Analyzing..."):
                                try:
                                    import openai
                                    openai_key = os.environ.get("OPENAI_API_KEY", st.secrets.get("OPENAI_API_KEY", ""))
                                    if openai_key:
                                        client = openai.OpenAI(api_key=openai_key)
                                        reasons_text = "\n".join([f"- {r['Employee']} ({r['Department']}, {r['Region']}): {r['Rejection Reason']}" for r in rejections])
                                        response = client.chat.completions.create(model="gpt-3.5-turbo", messages=[{"role": "system", "content": "Analyze these appraisal rejections. Provide: 1) Top 3 themes 2) Departments with issues 3) Recommendations."}, {"role": "user", "content": reasons_text}], temperature=0.5, max_tokens=400)
                                        st.markdown("### 🤖 AI Analysis"); st.success("Analysis complete!"); st.markdown(response.choices[0].message.content)
                                    else: st.info("OpenAI API key not configured.")
                                except Exception as e:
                                    st.warning(f"AI analysis unavailable: {str(e)}")
                                    st.markdown("### 📊 Manual Analysis\n**Common Themes:** Score disagreements, insufficient evidence review, communication gaps.\n**Recommendations:** Pre-review calibration meetings, ensure evidence acknowledgment, add discussion step.")
                    else: st.info("✅ No rejections recorded matching the filters.")
                
                # ============================================================
                # DASH TAB 4: HIERARCHY VIEW
                # ============================================================
                with dash_tabs[4]:
                    st.subheader("🏢 Full Hierarchy View")
                    hierarchy = defaultdict(lambda: defaultdict(lambda: defaultdict(dict)))
                    if not employees_df.empty:
                        for _, emp in employees_df.iterrows():
                            emp_name = f"{emp['first_name']} {emp['last_name']}".strip()
                            if not emp_name: continue
                            dept = emp.get('department', 'General'); sub = emp.get('subsidiary', ''); region = get_region(sub)
                            if not apply_dash_filters(region, sub or 'Unassigned', dept): continue
                            assessment = st.session_state.self_assessments.get(emp_name, {})
                            scores = assessment.get('scores', {}); reviewer_scores = assessment.get('hod_scores') or assessment.get('tl_scores') or {}
                            hierarchy[region][sub or 'Unassigned'][dept][emp_name] = {'self_avg': round(sum(int(v) for v in scores.values() if v) / len(scores), 1) if scores else 0, 'reviewer_avg': round(sum(int(v) for v in reviewer_scores.values() if v) / len(reviewer_scores), 1) if reviewer_scores else 0, 'status': assessment.get('status', 'Not Started'), 'acceptance': assessment.get('acceptance', ''), 'reviewer_type': assessment.get('reviewer_type', 'N/A'), 'rejection_comment': assessment.get('rejection_comment', ''), 'reject_count': assessment.get('reject_count', 0)}
                    
                    if not all_perf.empty:
                        for _, row in all_perf.iterrows():
                            uname = ' '.join(str(row.get('user_name', '')).split())
                            for region in hierarchy:
                                for sub in hierarchy[region]:
                                    for dept in hierarchy[region][sub]:
                                        if uname in hierarchy[region][sub][dept]: hierarchy[region][sub][dept][uname]['kpi_status'] = row.get('submission_status', 'Draft')
                    
                    total_emps_shown = 0
                    for region in ['Abuja', 'Lagos', 'Aba']:
                        if region not in hierarchy: continue
                        region_color = {'Abuja': '#CC0000', 'Lagos': '#1a1a1a', 'Aba': '#38a169'}.get(region, '#1a1a1a')
                        region_emp_count = sum(len(d) for s in hierarchy[region].values() for d in s.values())
                        if region_emp_count == 0: continue
                        st.markdown(f"""<div class="region-header" style="background:linear-gradient(135deg, {region_color}, #2d2d2d);display:flex;justify-content:space-between;align-items:center;"><span>🌍 {region} Region</span><span style="font-size:0.85rem;opacity:0.85;">👥 {region_emp_count} employees</span></div>""", unsafe_allow_html=True)
                        for subsidiary in sorted(hierarchy[region].keys()):
                            departments = hierarchy[region][subsidiary]; sub_emp_count = sum(len(e) for e in departments.values())
                            if sub_emp_count == 0: continue
                            sub_key = f"hier_sub_{region}_{subsidiary}".replace(' ', '_').replace('(', '').replace(')', '').replace('.', '')
                            if sub_key not in st.session_state: st.session_state[sub_key] = False
                            expand_icon = "▼" if st.session_state[sub_key] else "▶"
                            col1, col2 = st.columns([1, 20])
                            with col1:
                                if st.button(expand_icon, key=f"hierbtn_{sub_key}"):
                                    st.session_state[sub_key] = not st.session_state[sub_key]
                            with col2: st.markdown(f"""<div class="subsidiary-header" style="margin:0.3rem 0;">🏢 {subsidiary} — {len(departments)} depts | 👥 {sub_emp_count}</div>""", unsafe_allow_html=True)
                            if st.session_state[sub_key]:
                                for department in sorted(departments.keys()):
                                    employees = departments[department]
                                    if len(employees) == 0: continue
                                    dept_key = f"hier_dept_{region}_{subsidiary}_{department}".replace(' ', '_').replace('(', '').replace(')', '').replace('.', '')
                                    if dept_key not in st.session_state: st.session_state[dept_key] = False
                                    dept_icon = "▼" if st.session_state[dept_key] else "▶"
                                    col1d, col2d = st.columns([1, 20])
                                    with col1d:
                                        if st.button(dept_icon, key=f"hierdeptbtn_{dept_key}"):
                                            st.session_state[dept_key] = not st.session_state[dept_key]
                                    with col2d:
                                        dept_completed = sum(1 for e in employees.values() if e['acceptance'] == 'Accepted'); dept_rejected = sum(1 for e in employees.values() if e['acceptance'] == 'Rejected')
                                        st.markdown(f"""<div style="background:#f8f9fa;padding:0.4rem 1rem;border-radius:6px;margin:0.2rem 0;border-left:3px solid #CC0000;">🏭 {department} — 👥 {len(employees)} | 🎉 {dept_completed} done | ❌ {dept_rejected} rejected</div>""", unsafe_allow_html=True)
                                    if st.session_state[dept_key]:
                                        for emp_name, data in sorted(employees.items()):
                                            total_emps_shown += 1
                                            accept_class = {'Accepted': 'badge-green', 'Rejected': 'badge-red', '': 'badge-gray'}; accept_badge = accept_class.get(data['acceptance'], 'badge-gray')
                                            status_text = data['acceptance'] or data['status'] or 'Not Started'
                                            kpi_badge_class = {'Approved': 'badge-green', 'Submitted': 'badge-yellow', 'Draft': 'badge-gray'}.get(data.get('kpi_status', ''), 'badge-gray')
                                            border_color = '#38a169' if data['acceptance'] == 'Accepted' else '#CC0000' if data['acceptance'] == 'Rejected' else '#d69e2e'
                                            st.markdown(f"""<div style="padding:0.6rem;margin:0.3rem 0 0.3rem 2rem;border-left:4px solid {border_color};background:white;border-radius:6px;box-shadow:0 1px 3px rgba(0,0,0,0.05);"><div style="display:flex;justify-content:space-between;align-items:center;flex-wrap:wrap;gap:8px;"><strong>👤 {emp_name}</strong><span style="font-size:0.8rem;">📊 Self: <strong>{data['self_avg']}%</strong> | Rev: <strong>{data['reviewer_avg']}%</strong> | <span class="badge {kpi_badge_class}">KPI: {data.get('kpi_status', 'N/A')}</span> | <span class="badge {accept_badge}">{status_text}</span> | {data['reviewer_type']}</span></div>{f'<small style="color:#CC0000;">💬 {data["rejection_comment"][:100]}...</small>' if data['rejection_comment'] else ''}</div>""", unsafe_allow_html=True)
                    if total_emps_shown == 0: st.info("No employees match the selected filters.")
            
            # ============================================================
            # ADMIN SUB-TAB 1: MY PERFORMANCE
            # ============================================================
            with admin_dash_tabs[1]:
                st.subheader("👤 My Personal Performance")
                user_assessment = st.session_state.self_assessments.get(user_name, {})
                
                if user_assessment.get('acceptance') == 'Accepted':
                    st.success("🎉 Congratulations! Your appraisal has been accepted!")
                    final_scores = user_assessment.get('hod_scores') or user_assessment.get('tl_scores') or user_assessment.get('scores', {})
                    avg_score = sum(int(v) for v in final_scores.values() if v) / len(final_scores) if final_scores else 0
                    
                    if avg_score >= 90: classification, class_color = 'PLATINUM', '#E5E4E2'; emoji = '🥇'
                    elif avg_score >= 80: classification, class_color = 'GOLD', '#FFD700'; emoji = '🥇'
                    elif avg_score >= 70: classification, class_color = 'SILVER', '#C0C0C0'; emoji = '🥈'
                    elif avg_score >= 60: classification, class_color = 'BRONZE', '#CD7F32'; emoji = '🥉'
                    elif avg_score >= 50: classification, class_color = 'STEEL', '#71797E'; emoji = '🔵'
                    else: classification, class_color = 'DEVELOPMENT', '#A0AEC0'; emoji = '⚪'
                    
                    st.markdown(f"""
                    <div style="background:linear-gradient(135deg, #fffef5, #fff8e1);border:3px solid {class_color};border-radius:16px;padding:2rem;text-align:center;box-shadow:0 8px 32px rgba(0,0,0,0.1);margin:1rem 0;">
                        <h2 style="color:#1a1a1a;margin:0;">🏆 APPRAISAL CERTIFICATE</h2>
                        <h3 style="color:#CC0000;margin:0.5rem 0;">Churchgate Group</h3>
                        <p style="font-size:1.2rem;margin:0.5rem 0;"><strong>{user_name}</strong></p>
                        <p style="margin:0.3rem 0;">Has successfully completed the</p>
                        <p style="font-size:1.1rem;color:#d69e2e;margin:0.3rem 0;"><strong>{st.session_state.appraisal_cycle_name}</strong></p>
                        <p style="margin:0.3rem 0;">with an overall score of <strong style="color:#38a169;font-size:1.3rem;">{avg_score:.1f}%</strong></p>
                        <p style="font-size:1.4rem;font-weight:700;color:{class_color};margin:0.5rem 0;">Classification: {emoji} {classification}</p>
                        <p style="color:#888;font-size:0.85rem;margin:0.3rem 0;">Date: {user_assessment.get('date', now_wat.strftime('%Y-%m-%d %H:%M WAT'))}</p>
                        <p style="color:#888;font-size:0.8rem;">Verified by Churchgate Group HRIS</p>
                    </div>
                    """, unsafe_allow_html=True)
                    
                    if st.button("📜 Download Certificate (PDF)", use_container_width=True, type="primary", key="cert_dl_admin_t8"):
                        try:
                            from fpdf import FPDF
                            safe_name = user_name.encode('ascii', 'ignore').decode('ascii')
                            dept = st.session_state.user.get('department', 'N/A')
                            emp_id = st.session_state.user.get('employee_id', 'N/A')
                            reviewer = user_assessment.get('reviewer_type', 'HOD')
                            
                            if avg_score >= 90: pdf_class_color = (229, 228, 226)
                            elif avg_score >= 80: pdf_class_color = (255, 215, 0)
                            elif avg_score >= 70: pdf_class_color = (192, 192, 192)
                            elif avg_score >= 60: pdf_class_color = (205, 127, 50)
                            elif avg_score >= 50: pdf_class_color = (113, 121, 126)
                            else: pdf_class_color = (160, 174, 192)
                            
                            if avg_score >= 90: grade = "A+"
                            elif avg_score >= 80: grade = "A"
                            elif avg_score >= 70: grade = "B"
                            elif avg_score >= 60: grade = "C"
                            elif avg_score >= 50: grade = "D"
                            else: grade = "F"
                            
                            if avg_score >= 80: rating = "EXCEEDS EXPECTATIONS"
                            elif avg_score >= 60: rating = "MEETS EXPECTATIONS"
                            else: rating = "BELOW EXPECTATIONS"
                            
                            pillar_scores = {}
                            for key, val in final_scores.items():
                                pillar_name = key.split('_')[0] if '_' in key else key[:35]
                                if pillar_name not in pillar_scores: pillar_scores[pillar_name] = []
                                pillar_scores[pillar_name].append(int(val) if val else 0)
                            pillar_avgs = {k: sum(v)/len(v) for k, v in pillar_scores.items()}
                            
                            pdf = FPDF(orientation='L', unit='mm', format='A4')
                            pdf.add_page()
                            pdf.set_draw_color(*pdf_class_color); pdf.set_line_width(2)
                            pdf.rect(6, 6, 285, 198); pdf.set_line_width(0.5); pdf.rect(8, 8, 281, 194)
                            pdf.set_fill_color(26, 26, 26); pdf.rect(10, 10, 277, 40, 'F')
                            pdf.set_fill_color(204, 0, 0); pdf.rect(10, 50, 277, 3, 'F')
                            pdf.set_fill_color(*pdf_class_color); pdf.rect(10, 53, 277, 1.5, 'F')
                            pdf.set_font('Helvetica', 'B', 28); pdf.set_text_color(255, 255, 255)
                            pdf.cell(0, 22, 'CHURCHGATE GROUP', ln=True, align='C')
                            pdf.set_font('Helvetica', 'B', 13); pdf.set_text_color(*pdf_class_color)
                            pdf.cell(0, 8, 'PERFORMANCE APPRAISAL CERTIFICATE OF ACHIEVEMENT', ln=True, align='C')
                            pdf.ln(12); pdf.set_text_color(26, 26, 26); pdf.set_font('Helvetica', '', 13)
                            pdf.cell(0, 8, 'This is to certify that', ln=True, align='C'); pdf.ln(4)
                            pdf.set_font('Helvetica', 'B', 24); pdf.set_text_color(204, 0, 0)
                            pdf.cell(0, 14, safe_name.upper(), ln=True, align='C'); pdf.ln(4)
                            pdf.set_text_color(26, 26, 26); pdf.set_font('Helvetica', '', 12)
                            pdf.cell(0, 8, f'has successfully completed the {st.session_state.appraisal_cycle_name}', ln=True, align='C')
                            pdf.cell(0, 8, f'from {st.session_state.appraisal_start} to {st.session_state.appraisal_end}', ln=True, align='C')
                            pdf.ln(6)
                            box_y = pdf.get_y(); pdf.set_fill_color(245, 245, 245); pdf.rect(25, box_y, 247, 28, 'F')
                            col_w = 247 / 4
                            metrics = [('OVERALL SCORE', f'{avg_score:.1f}%', (56, 161, 105)), ('GRADE', grade, (204, 0, 0)), ('RATING', rating, (*pdf_class_color,)), ('CLASSIFICATION', classification, (*pdf_class_color,))]
                            for i, (label, value, color) in enumerate(metrics):
                                pdf.set_xy(25 + i*col_w, box_y + 4); pdf.set_font('Helvetica', '', 8); pdf.set_text_color(128, 128, 128); pdf.cell(col_w, 6, label, align='C')
                                pdf.set_xy(25 + i*col_w, box_y + 12); pdf.set_font('Helvetica', 'B', 13); pdf.set_text_color(*color[:3]); pdf.cell(col_w, 8, value, align='C')
                            pdf.set_y(box_y + 32); pdf.ln(4)
                            pdf.set_font('Helvetica', 'B', 11); pdf.set_text_color(26, 26, 26)
                            pdf.cell(0, 8, 'PERFORMANCE PILLAR BREAKDOWN', ln=True, align='C'); pdf.ln(3)
                            table_x = 40
                            pdf.set_fill_color(26, 26, 26); pdf.set_text_color(255, 255, 255); pdf.set_font('Helvetica', 'B', 9)
                            pdf.set_xy(table_x, pdf.get_y())
                            pdf.cell(110, 8, ' Strategic Pillar', 1, 0, 'L', True); pdf.cell(40, 8, 'Score', 1, 0, 'C', True); pdf.cell(40, 8, 'Rating', 1, 0, 'C', True); pdf.cell(40, 8, 'Grade', 1, 0, 'C', True); pdf.ln()
                            pdf.set_text_color(26, 26, 26); pdf.set_font('Helvetica', '', 8)
                            for pillar, avg in pillar_avgs.items():
                                if avg >= 90: p_grade = "A+"
                                elif avg >= 80: p_grade = "A"
                                elif avg >= 70: p_grade = "B"
                                elif avg >= 60: p_grade = "C"
                                elif avg >= 50: p_grade = "D"
                                else: p_grade = "F"
                                if avg >= 80: p_rating = 'Exceeds'
                                elif avg >= 60: p_rating = 'Meets'
                                else: p_rating = 'Below'
                                pdf.set_xy(table_x, pdf.get_y()); pdf.cell(110, 7, f' {pillar[:45]}', 1, 0, 'L'); pdf.cell(40, 7, f'{avg:.1f}%', 1, 0, 'C'); pdf.cell(40, 7, p_rating, 1, 0, 'C'); pdf.cell(40, 7, p_grade, 1, 0, 'C'); pdf.ln()
                            pdf.ln(8); pdf.set_font('Helvetica', '', 9); pdf.set_text_color(80, 80, 80)
                            pdf.cell(0, 6, f'Department: {dept}    |    Reviewer: {reviewer}    |    Employee ID: {emp_id}', ln=True, align='C')
                            pdf.cell(0, 6, f'Date of Issue: {now_wat.strftime("%B %d, %Y")}    |    Time: {now_wat.strftime("%H:%M WAT")}', ln=True, align='C')
                            pdf.ln(6); sig_w = 75
                            pdf.set_font('Helvetica', 'B', 9); pdf.set_text_color(26, 26, 26)
                            pdf.cell(sig_w, 7, '_______________________', align='C'); pdf.cell(25, 7, ''); pdf.cell(sig_w, 7, '_______________________', align='C'); pdf.cell(25, 7, ''); pdf.cell(sig_w, 7, '_______________________', align='C'); pdf.ln()
                            pdf.set_font('Helvetica', '', 8); pdf.set_text_color(128, 128, 128)
                            pdf.cell(sig_w, 5, 'Reviewer / HOD', align='C'); pdf.cell(25, 5, ''); pdf.cell(sig_w, 5, 'HR Director', align='C'); pdf.cell(25, 5, ''); pdf.cell(sig_w, 5, 'GMD / CEO', align='C')
                            pdf.ln(12); pdf.set_font('Helvetica', 'I', 7); pdf.set_text_color(160, 160, 160)
                            pdf.cell(0, 5, 'This certificate is electronically generated by Churchgate Group HRIS and is valid without physical signature.', ln=True, align='C')
                            pdf.cell(0, 5, 'Churchgate Group - Nigeria | hris@churchgate.com', ln=True, align='C')
                            
                            pdf_output = pdf.output(dest='S').encode('latin-1')
                            st.download_button("📥 Download PDF Certificate", pdf_output, f"{safe_name.replace(' ', '_')}_Certificate.pdf", "application/pdf", key="cert_dl_btn_admin_t8")
                            st.success("✅ Fortune 500 Certificate ready!")
                        except Exception as e:
                            st.error(f"PDF error: {str(e)}")
                
                if st.session_state.appraisal_cycle_active:
                    current_status = user_assessment.get('status', 'Not Started')
                    acceptance = user_assessment.get('acceptance', '')
                    if acceptance == 'Accepted' or current_status == 'Completed':
                        current_status = 'Completed'
                    steps = [("Set KPIs", 1), ("Self-Assessment", 2), ("Review", 3), ("Complete", 4)]
                    status_map = {'Not Started': 0, 'KPIs Set': 1, 'Submitted': 2, 'Approved': 3, 'Awaiting HOD Re-review': 3, 'Awaiting TL Re-review': 3, 'Accepted': 4, 'Completed': 4}
                    current_step = status_map.get(current_status, 0)
                    cols = st.columns(4)
                    for i, (step_name, step_num) in enumerate(steps):
                        with cols[i]:
                            if step_num < current_step: st.success(f"✅ {step_name}")
                            elif step_num == current_step: st.success(f"✅ {step_name}")
                            else: st.markdown(f"⏳ {step_name}")
                
                pillar_data = load_user_pillar_data()
                for pillar_name in ['1. Occupancy & Revenue Growth', '2. Process Simplification', '3. Asset Reliability & Digitalization', '4. People & Culture']:
                    pd_data = pillar_data[pillar_name]; status_text, color = get_kpi_status(pd_data['progress'])
                    st.markdown(f"""<div class="glass-card" style="border-left:4px solid {color};padding:0.8rem;"><strong>{pillar_name}</strong> ({pd_data['weight']}%)<br><small>Progress: {pd_data['progress']}% | {pd_data['status']}</small><div style="background:#e0e0e0;height:6px;border-radius:3px;margin-top:0.4rem;"><div style="background:{color};width:{pd_data['progress']}%;height:6px;border-radius:3px;"></div></div></div>""", unsafe_allow_html=True)
                
                total_prog = sum(p['progress'] * p['weight'] / 100 for p in pillar_data.values())
                c1, c2, c3 = st.columns(3)
                c1.metric("Weighted Score", f"{total_prog:.1f}%"); c2.metric("On Track", sum(1 for p in pillar_data.values() if p['status'] in ['On Track', 'Exceeding'])); c3.metric("At Risk", sum(1 for p in pillar_data.values() if p['status'] == 'At Risk'))
                
                st.markdown("---"); st.subheader("📜 My Appraisal History")
                try:
                    history = db.get_appraisal_history(user_name)
                    if history:
                        for h in history[-5:]:
                            sc = '#38a169' if 'Accepted' in str(h.get('final_status', '')) else '#CC0000'
                            st.markdown(f"""<div class="kpi-card" style="border-left-color:{sc};"><strong>{h.get('cycle_name', 'N/A')}</strong>: {h.get('final_status', 'N/A')} — {h.get('completed_date', 'N/A')}</div>""", unsafe_allow_html=True)
                    else: st.info("No completed appraisals yet.")
                except: pass
        
        else:
            # ===== NON-ADMIN EMPLOYEE VIEW =====
            st.subheader("📈 My Performance Summary")
            user_assessment = st.session_state.self_assessments.get(user_name, {})
            
            if user_assessment.get('acceptance') == 'Accepted':
                st.success("🎉 Congratulations! Your appraisal has been accepted!")
                final_scores = user_assessment.get('hod_scores') or user_assessment.get('tl_scores') or user_assessment.get('scores', {})
                avg_score = sum(int(v) for v in final_scores.values() if v) / len(final_scores) if final_scores else 0
                
                if avg_score >= 90: classification, class_color = 'PLATINUM', '#E5E4E2'; emoji = '🥇'
                elif avg_score >= 80: classification, class_color = 'GOLD', '#FFD700'; emoji = '🥇'
                elif avg_score >= 70: classification, class_color = 'SILVER', '#C0C0C0'; emoji = '🥈'
                elif avg_score >= 60: classification, class_color = 'BRONZE', '#CD7F32'; emoji = '🥉'
                elif avg_score >= 50: classification, class_color = 'STEEL', '#71797E'; emoji = '🔵'
                else: classification, class_color = 'DEVELOPMENT', '#A0AEC0'; emoji = '⚪'
                
                st.markdown(f"""
                <div style="background:linear-gradient(135deg, #fffef5, #fff8e1);border:3px solid {class_color};border-radius:16px;padding:2rem;text-align:center;box-shadow:0 8px 32px rgba(0,0,0,0.1);margin:1rem 0;">
                    <h2 style="color:#1a1a1a;margin:0;">🏆 APPRAISAL CERTIFICATE</h2>
                    <h3 style="color:#CC0000;margin:0.5rem 0;">Churchgate Group</h3>
                    <p style="font-size:1.2rem;margin:0.5rem 0;"><strong>{user_name}</strong></p>
                    <p style="margin:0.3rem 0;">Has successfully completed the</p>
                    <p style="font-size:1.1rem;color:#d69e2e;margin:0.3rem 0;"><strong>{st.session_state.appraisal_cycle_name}</strong></p>
                    <p style="margin:0.3rem 0;">with an overall score of <strong style="color:#38a169;font-size:1.3rem;">{avg_score:.1f}%</strong></p>
                    <p style="font-size:1.4rem;font-weight:700;color:{class_color};margin:0.5rem 0;">Classification: {emoji} {classification}</p>
                    <p style="color:#888;font-size:0.85rem;margin:0.3rem 0;">Date: {user_assessment.get('date', now_wat.strftime('%Y-%m-%d %H:%M WAT'))}</p>
                    <p style="color:#888;font-size:0.8rem;">Verified by Churchgate Group HRIS</p>
                </div>
                """, unsafe_allow_html=True)
                
                if st.button("📜 Download Certificate (PDF)", use_container_width=True, type="primary", key="cert_dl_emp_t8"):
                    try:
                        from fpdf import FPDF
                        safe_name = user_name.encode('ascii', 'ignore').decode('ascii')
                        dept = st.session_state.user.get('department', 'N/A')
                        emp_id = st.session_state.user.get('employee_id', 'N/A')
                        reviewer = user_assessment.get('reviewer_type', 'HOD')
                        
                        if avg_score >= 90: pdf_class_color = (229, 228, 226)
                        elif avg_score >= 80: pdf_class_color = (255, 215, 0)
                        elif avg_score >= 70: pdf_class_color = (192, 192, 192)
                        elif avg_score >= 60: pdf_class_color = (205, 127, 50)
                        elif avg_score >= 50: pdf_class_color = (113, 121, 126)
                        else: pdf_class_color = (160, 174, 192)
                        
                        if avg_score >= 90: grade = "A+"
                        elif avg_score >= 80: grade = "A"
                        elif avg_score >= 70: grade = "B"
                        elif avg_score >= 60: grade = "C"
                        elif avg_score >= 50: grade = "D"
                        else: grade = "F"
                        
                        if avg_score >= 80: rating = "EXCEEDS EXPECTATIONS"
                        elif avg_score >= 60: rating = "MEETS EXPECTATIONS"
                        else: rating = "BELOW EXPECTATIONS"
                        
                        pillar_scores = {}
                        for key, val in final_scores.items():
                            pillar_name = key.split('_')[0] if '_' in key else key[:35]
                            if pillar_name not in pillar_scores: pillar_scores[pillar_name] = []
                            pillar_scores[pillar_name].append(int(val) if val else 0)
                        pillar_avgs = {k: sum(v)/len(v) for k, v in pillar_scores.items()}
                        
                        pdf = FPDF(orientation='L', unit='mm', format='A4')
                        pdf.add_page()
                        pdf.set_draw_color(*pdf_class_color); pdf.set_line_width(2)
                        pdf.rect(6, 6, 285, 198); pdf.set_line_width(0.5); pdf.rect(8, 8, 281, 194)
                        pdf.set_fill_color(26, 26, 26); pdf.rect(10, 10, 277, 40, 'F')
                        pdf.set_fill_color(204, 0, 0); pdf.rect(10, 50, 277, 3, 'F')
                        pdf.set_fill_color(*pdf_class_color); pdf.rect(10, 53, 277, 1.5, 'F')
                        pdf.set_font('Helvetica', 'B', 28); pdf.set_text_color(255, 255, 255)
                        pdf.cell(0, 22, 'CHURCHGATE GROUP', ln=True, align='C')
                        pdf.set_font('Helvetica', 'B', 13); pdf.set_text_color(*pdf_class_color)
                        pdf.cell(0, 8, 'PERFORMANCE APPRAISAL CERTIFICATE OF ACHIEVEMENT', ln=True, align='C')
                        pdf.ln(12); pdf.set_text_color(26, 26, 26); pdf.set_font('Helvetica', '', 13)
                        pdf.cell(0, 8, 'This is to certify that', ln=True, align='C'); pdf.ln(4)
                        pdf.set_font('Helvetica', 'B', 24); pdf.set_text_color(204, 0, 0)
                        pdf.cell(0, 14, safe_name.upper(), ln=True, align='C'); pdf.ln(4)
                        pdf.set_text_color(26, 26, 26); pdf.set_font('Helvetica', '', 12)
                        pdf.cell(0, 8, f'has successfully completed the {st.session_state.appraisal_cycle_name}', ln=True, align='C')
                        pdf.cell(0, 8, f'from {st.session_state.appraisal_start} to {st.session_state.appraisal_end}', ln=True, align='C')
                        pdf.ln(6)
                        box_y = pdf.get_y(); pdf.set_fill_color(245, 245, 245); pdf.rect(25, box_y, 247, 28, 'F')
                        col_w = 247 / 4
                        metrics = [('OVERALL SCORE', f'{avg_score:.1f}%', (56, 161, 105)), ('GRADE', grade, (204, 0, 0)), ('RATING', rating, (*pdf_class_color,)), ('CLASSIFICATION', classification, (*pdf_class_color,))]
                        for i, (label, value, color) in enumerate(metrics):
                            pdf.set_xy(25 + i*col_w, box_y + 4); pdf.set_font('Helvetica', '', 8); pdf.set_text_color(128, 128, 128); pdf.cell(col_w, 6, label, align='C')
                            pdf.set_xy(25 + i*col_w, box_y + 12); pdf.set_font('Helvetica', 'B', 13); pdf.set_text_color(*color[:3]); pdf.cell(col_w, 8, value, align='C')
                        pdf.set_y(box_y + 32); pdf.ln(4)
                        pdf.set_font('Helvetica', 'B', 11); pdf.set_text_color(26, 26, 26)
                        pdf.cell(0, 8, 'PERFORMANCE PILLAR BREAKDOWN', ln=True, align='C'); pdf.ln(3)
                        table_x = 40
                        pdf.set_fill_color(26, 26, 26); pdf.set_text_color(255, 255, 255); pdf.set_font('Helvetica', 'B', 9)
                        pdf.set_xy(table_x, pdf.get_y())
                        pdf.cell(110, 8, ' Strategic Pillar', 1, 0, 'L', True); pdf.cell(40, 8, 'Score', 1, 0, 'C', True); pdf.cell(40, 8, 'Rating', 1, 0, 'C', True); pdf.cell(40, 8, 'Grade', 1, 0, 'C', True); pdf.ln()
                        pdf.set_text_color(26, 26, 26); pdf.set_font('Helvetica', '', 8)
                        for pillar, avg in pillar_avgs.items():
                            if avg >= 90: p_grade = "A+"
                            elif avg >= 80: p_grade = "A"
                            elif avg >= 70: p_grade = "B"
                            elif avg >= 60: p_grade = "C"
                            elif avg >= 50: p_grade = "D"
                            else: p_grade = "F"
                            if avg >= 80: p_rating = 'Exceeds'
                            elif avg >= 60: p_rating = 'Meets'
                            else: p_rating = 'Below'
                            pdf.set_xy(table_x, pdf.get_y()); pdf.cell(110, 7, f' {pillar[:45]}', 1, 0, 'L'); pdf.cell(40, 7, f'{avg:.1f}%', 1, 0, 'C'); pdf.cell(40, 7, p_rating, 1, 0, 'C'); pdf.cell(40, 7, p_grade, 1, 0, 'C'); pdf.ln()
                        pdf.ln(8); pdf.set_font('Helvetica', '', 9); pdf.set_text_color(80, 80, 80)
                        pdf.cell(0, 6, f'Department: {dept}    |    Reviewer: {reviewer}    |    Employee ID: {emp_id}', ln=True, align='C')
                        pdf.cell(0, 6, f'Date of Issue: {now_wat.strftime("%B %d, %Y")}    |    Time: {now_wat.strftime("%H:%M WAT")}', ln=True, align='C')
                        pdf.ln(6); sig_w = 75
                        pdf.set_font('Helvetica', 'B', 9); pdf.set_text_color(26, 26, 26)
                        pdf.cell(sig_w, 7, '_______________________', align='C'); pdf.cell(25, 7, ''); pdf.cell(sig_w, 7, '_______________________', align='C'); pdf.cell(25, 7, ''); pdf.cell(sig_w, 7, '_______________________', align='C'); pdf.ln()
                        pdf.set_font('Helvetica', '', 8); pdf.set_text_color(128, 128, 128)
                        pdf.cell(sig_w, 5, 'Reviewer / HOD', align='C'); pdf.cell(25, 5, ''); pdf.cell(sig_w, 5, 'HR Director', align='C'); pdf.cell(25, 5, ''); pdf.cell(sig_w, 5, 'GMD / CEO', align='C')
                        pdf.ln(12); pdf.set_font('Helvetica', 'I', 7); pdf.set_text_color(160, 160, 160)
                        pdf.cell(0, 5, 'This certificate is electronically generated by Churchgate Group HRIS and is valid without physical signature.', ln=True, align='C')
                        pdf.cell(0, 5, 'Churchgate Group - Nigeria | hris@churchgate.com', ln=True, align='C')
                        
                        pdf_output = pdf.output(dest='S').encode('latin-1')
                        st.download_button("📥 Download PDF Certificate", pdf_output, f"{safe_name.replace(' ', '_')}_Certificate.pdf", "application/pdf", key="cert_dl_btn_emp_t8")
                        st.success("✅ Fortune 500 Certificate ready!")
                    except Exception as e:
                        st.error(f"PDF error: {str(e)}")
            
            if st.session_state.appraisal_cycle_active:
                current_status = user_assessment.get('status', 'Not Started')
                acceptance = user_assessment.get('acceptance', '')
                if acceptance == 'Accepted' or current_status == 'Completed':
                    current_status = 'Completed'
                steps = [("Set KPIs", 1), ("Self-Assessment", 2), ("Review", 3), ("Complete", 4)]
                status_map = {'Not Started': 0, 'KPIs Set': 1, 'Submitted': 2, 'Approved': 3, 'Awaiting HOD Re-review': 3, 'Awaiting TL Re-review': 3, 'Accepted': 4, 'Completed': 4}
                current_step = status_map.get(current_status, 0)
                cols = st.columns(4)
                for i, (step_name, step_num) in enumerate(steps):
                    with cols[i]:
                        if step_num < current_step: st.success(f"✅ {step_name}")
                        elif step_num == current_step: st.success(f"✅ {step_name}")
                        else: st.markdown(f"⏳ {step_name}")
            
            pillar_data = load_user_pillar_data()
            for pillar_name in ['1. Occupancy & Revenue Growth', '2. Process Simplification', '3. Asset Reliability & Digitalization', '4. People & Culture']:
                pd_data = pillar_data[pillar_name]; status_text, color = get_kpi_status(pd_data['progress'])
                st.markdown(f"""<div class="glass-card" style="border-left:4px solid {color};padding:0.8rem;"><strong>{pillar_name}</strong> ({pd_data['weight']}%)<br><small>Progress: {pd_data['progress']}% | {pd_data['status']}</small><div style="background:#e0e0e0;height:6px;border-radius:3px;margin-top:0.4rem;"><div style="background:{color};width:{pd_data['progress']}%;height:6px;border-radius:3px;"></div></div></div>""", unsafe_allow_html=True)
            
            total_prog = sum(p['progress'] * p['weight'] / 100 for p in pillar_data.values())
            c1, c2, c3 = st.columns(3)
            c1.metric("Weighted Score", f"{total_prog:.1f}%"); c2.metric("On Track", sum(1 for p in pillar_data.values() if p['status'] in ['On Track', 'Exceeding'])); c3.metric("At Risk", sum(1 for p in pillar_data.values() if p['status'] == 'At Risk'))
            
            if st.button("📊 Benchmark Report", use_container_width=True):
                all_perf_data = get_all_perf_cached()
                dept_avg = all_perf_data[all_perf_data['department'] == user_dept]['progress'].mean() if not all_perf_data.empty and 'progress' in all_perf_data.columns else 0
                group_avg = all_perf_data['progress'].mean() if not all_perf_data.empty and 'progress' in all_perf_data.columns else 0
                bench_data = pd.DataFrame({'Metric': ['My Score', 'Dept Average', 'Group Average', 'Target'], 'Score': [total_prog, dept_avg, group_avg, 85]})
                fig = px.bar(bench_data, x='Metric', y='Score', color='Metric', color_discrete_sequence=['#CC0000', '#4a4a4a', '#888888', '#38a169']); fig.add_hline(y=85, line_dash="dash", line_color="#38a169"); fig.update_layout(height=350)
                st.plotly_chart(fig, use_container_width=True)
            
            st.markdown("---"); st.subheader("📜 My Appraisal History")
            try:
                history = db.get_appraisal_history(user_name)
                if history:
                    for h in history[-5:]:
                        sc = '#38a169' if 'Accepted' in str(h.get('final_status', '')) else '#CC0000'
                        st.markdown(f"""<div class="kpi-card" style="border-left-color:{sc};"><strong>{h.get('cycle_name', 'N/A')}</strong>: {h.get('final_status', 'N/A')} — {h.get('completed_date', 'N/A')}</div>""", unsafe_allow_html=True)
                else: st.info("No completed appraisals yet.")
            except: pass
    
    # ============================================================
    # TAB 9: APPRAISAL COMMITTEE - COMPLETE WITH ALL FIXES
    # ============================================================
    with tab9:
        st.markdown('<div class="glass-card"><h3>🏛️ Appraisal Committee Board</h3><p style="color:#888;">Senior Management & Admin — Full Appraisal Oversight & Talent Analytics</p></div>', unsafe_allow_html=True)
        
        if not (is_sr_mgmt or is_super_admin):
            st.info("⛔ Restricted to Senior Management and Admin only.")
        else:
            
            def get_classification(score):
                if score >= 90: return 'PLATINUM', '#E5E4E2'
                elif score >= 80: return 'GOLD', '#FFD700'
                elif score >= 70: return 'SILVER', '#C0C0C0'
                elif score >= 60: return 'BRONZE', '#CD7F32'
                elif score >= 50: return 'STEEL', '#71797E'
                else: return 'DEVELOPMENT', '#A0AEC0'
            
            def get_performance_level(score):
                if score >= 80: return 'High'
                elif score >= 60: return 'Average'
                else: return 'Low'
            
            def get_emp_score(emp_name):
                a = st.session_state.self_assessments.get(emp_name, {})
                scores = a.get('hod_scores') or a.get('tl_scores') or a.get('scores', {})
                return sum(int(v) for v in scores.values() if v) / len(scores) if scores else 0
            
            def get_potential_level(emp_name):
                assessment = st.session_state.self_assessments.get(emp_name, {})
                scores = assessment.get('scores', {})
                reviewer_scores = assessment.get('hod_scores') or assessment.get('tl_scores') or {}
                potential_score = 0
                if scores and reviewer_scores:
                    self_avg = sum(int(v) for v in scores.values() if v) / len(scores)
                    rev_avg = sum(int(v) for v in reviewer_scores.values() if v) / len(reviewer_scores)
                    gap = abs(self_avg - rev_avg)
                    if gap <= 5: potential_score += 30
                    elif gap <= 15: potential_score += 20
                    else: potential_score += 5
                elif scores: potential_score += 15
                all_perf_data = get_all_perf_cached()
                emp_perf = all_perf_data[all_perf_data['user_name'] == emp_name] if not all_perf_data.empty else pd.DataFrame()
                total_kpis = 0
                if not emp_perf.empty:
                    for _, row in emp_perf.iterrows():
                        kpi_list = json.loads(row.get('kpi_data', '[]')) if row.get('kpi_data') else []
                        total_kpis += len(kpi_list)
                if total_kpis >= 16: potential_score += 25
                elif total_kpis >= 10: potential_score += 20
                elif total_kpis >= 5: potential_score += 15
                else: potential_score += 5
                if not emp_perf.empty:
                    if any(emp_perf['submission_status'] == 'Approved'): potential_score += 25
                    elif any(emp_perf['submission_status'] == 'Submitted'): potential_score += 15
                    else: potential_score += 5
                else: potential_score += 5
                if scores:
                    pillar_scores = defaultdict(list)
                    for key, val in scores.items():
                        pillar = key.split('_')[0] if '_' in key else key[:20]
                        pillar_scores[pillar].append(int(val) if val else 0)
                    pillar_avgs = [sum(v)/len(v) for v in pillar_scores.values()]
                    if pillar_avgs and len(pillar_avgs) > 1:
                        consistency = 100 - (max(pillar_avgs) - min(pillar_avgs))
                        if consistency >= 80: potential_score += 20
                        elif consistency >= 60: potential_score += 15
                        else: potential_score += 5
                    else: potential_score += 10
                else: potential_score += 10
                if potential_score >= 75: return 'High'
                elif potential_score >= 50: return 'Moderate'
                else: return 'Low'
            
            def get_9box_position(score, emp_name):
                perf = get_performance_level(score)
                pot = get_potential_level(emp_name)
                positions = {
                    ('High', 'High'): ('⭐ CONSISTENT STAR', '#6bcb77', 'Top Talent', 'Reward, promote'),
                    ('High', 'Average'): ('💎 POTENTIAL GEM', '#ffd93d', 'Rising Talent', 'Develop, coach'),
                    ('High', 'Low'): ('🚀 RISING STAR', '#ff6b6b', 'Diamond in Rough', 'Mentorship'),
                    ('Moderate', 'High'): ('🏆 CURRENT STAR', '#98fb98', 'Strong Performer', 'Recognize, retain'),
                    ('Moderate', 'Average'): ('👔 SOLID PROFESSIONAL', '#87ceeb', 'Core Contributor', 'Engage, upskill'),
                    ('Moderate', 'Low'): ('🔄 INCONSISTENT PLAYER', '#ffa07a', 'Mixed Results', 'Performance plan'),
                    ('Low', 'High'): ('🔧 TECHNOCRAT', '#90ee90', 'Technical Expert', 'Specialist role'),
                    ('Low', 'Average'): ('📋 STABILIZER', '#b0c4de', 'Steady Hand', 'Maintain role'),
                    ('Low', 'Low'): ('⚠️ TALENT AT BAY', '#ff4444', 'At Risk', 'Urgent intervention')
                }
                return positions.get((pot, perf), ('N/A', '#ccc', 'Unknown', 'Review'))
            
            # Build all employee score data
            all_emps_scored = []
            if not employees_df.empty:
                for _, emp in employees_df.iterrows():
                    en = f"{emp['first_name']} {emp['last_name']}".strip()
                    if not en: continue
                    sc = get_emp_score(en)
                    if sc > 0:
                        cl, co = get_classification(sc)
                        pot = get_potential_level(en)
                        perf = get_performance_level(sc)
                        pos_name, pos_color, pos_desc, pos_action = get_9box_position(sc, en)
                        all_emps_scored.append({
                            'name': en, 'score': sc, 'class': cl, 'color': co,
                            'dept': emp.get('department', 'General'),
                            'subsidiary': emp.get('subsidiary', ''),
                            'region': get_region(emp.get('subsidiary', '')),
                            'initials': generate_initials(en),
                            'potential': pot, 'performance': perf,
                            'position_name': pos_name, 'position_color': pos_color,
                            'position_desc': pos_desc, 'position_action': pos_action
                        })
            
            all_assessments = st.session_state.self_assessments
            
            committee_tabs = st.tabs([
                "📊 9-Box Matrix", "📈 Analytics", "📝 Submissions", 
                "🚨 Escalated", "🎉 Completed", "🏆 Rankings", "📋 Recommendations"
            ])
            
            # ============================================================
            # TAB 0: 9-BOX MATRIX
            # ============================================================
            with committee_tabs[0]:
                st.subheader("📊 9-Box Talent Matrix")
                st.markdown("*Performance vs Potential — Data-driven talent assessment*")
                
                box_data = defaultdict(list)
                for e in all_emps_scored:
                    box_data[f"{e['potential']}_{e['performance']}"].append(e)
                
                total_in_matrix = len(all_emps_scored)
                stars = len(box_data.get('High_High', [])) + len(box_data.get('Moderate_High', []))
                at_risk = len(box_data.get('Low_Low', [])) + len(box_data.get('Moderate_Low', []))
                
                m1, m2, m3, m4 = st.columns(4)
                m1.metric("👥 Total", total_in_matrix)
                m2.metric("⭐ Top Talent", stars)
                m3.metric("⚠️ At Risk", at_risk)
                m4.metric("📊 Avg Score", f"{sum(e['score'] for e in all_emps_scored) / total_in_matrix:.1f}%" if total_in_matrix > 0 else "N/A")
                
                st.markdown("---")
                
                st.markdown("""
                <div style="display:flex;justify-content:space-around;margin-bottom:0.5rem;">
                    <div style="text-align:center;width:30%;background:#fff5f5;padding:0.5rem;border-radius:8px;border:2px solid #ff4444;">
                        <strong style="color:#CC0000;">📉 LOW PERFORMANCE</strong><br><small>Below 60%</small>
                    </div>
                    <div style="text-align:center;width:30%;background:#fffef5;padding:0.5rem;border-radius:8px;border:2px solid #ffd93d;">
                        <strong style="color:#d69e2e;">📊 AVERAGE PERFORMANCE</strong><br><small>60-79%</small>
                    </div>
                    <div style="text-align:center;width:30%;background:#f0fff4;padding:0.5rem;border-radius:8px;border:2px solid #38a169;">
                        <strong style="color:#38a169;">📈 HIGH PERFORMANCE</strong><br><small>80-100%</small>
                    </div>
                </div>
                """, unsafe_allow_html=True)
                
                for pot_level, pot_label, pot_emoji in [('High', 'HIGH POTENTIAL', '⭐'), ('Moderate', 'MODERATE POTENTIAL', '👔'), ('Low', 'LOW POTENTIAL', '👤')]:
                    st.markdown(f'<p style="font-weight:700;font-size:1.1rem;margin-top:1rem;">{pot_emoji} {pot_label}</p>', unsafe_allow_html=True)
                    c1, c2, c3 = st.columns(3)
                    for i, perf_level in enumerate(['Low', 'Average', 'High']):
                        key = f"{pot_level}_{perf_level}"
                        emps = box_data.get(key, [])
                        if emps:
                            pos_name = emps[0]['position_name']
                            color = emps[0]['position_color']
                            desc = emps[0]['position_desc']
                            action = emps[0]['position_action']
                        else:
                            pos_name, color, desc, action = 'Empty', '#e0e0e0', 'No employees', 'N/A'
                        initials_html = ''
                        for e in emps[:12]:
                            initials_html += f'<span style="display:inline-block;width:26px;height:26px;border-radius:50%;background:{e["position_color"]};color:white;text-align:center;line-height:26px;font-size:0.6rem;font-weight:700;margin:2px;cursor:pointer;" title="{e["name"]} — {e["score"]:.0f}% — {e["dept"]}">{e["initials"]}</span>'
                        with [c1, c2, c3][i]:
                            st.markdown(f"""
                            <div style="border:2px solid {color};border-radius:12px;padding:0.8rem;text-align:center;min-height:130px;background:linear-gradient(135deg, {color}08, {color}15);">
                                <div style="font-weight:700;font-size:0.85rem;color:{color};">{pos_name}</div>
                                <div style="font-size:0.6rem;color:#888;margin:0.2rem 0;">{desc}</div>
                                <div style="margin:0.4rem 0;">{initials_html if initials_html else '<small style="color:#aaa;">—</small>'}</div>
                                <div style="font-size:0.7rem;font-weight:600;">👥 {len(emps)}</div>
                                <div style="font-size:0.55rem;color:#888;margin-top:0.2rem;">💡 {action}</div>
                            </div>
                            """, unsafe_allow_html=True)
                
                st.markdown("---")
                st.subheader("🏆 Employee Classifications")
                class_counts = defaultdict(int)
                for e in all_emps_scored:
                    class_counts[e['class']] += 1
                
                class_cols = st.columns(6)
                classifications = ['PLATINUM', 'GOLD', 'SILVER', 'BRONZE', 'STEEL', 'DEVELOPMENT']
                class_colors_list = ['#E5E4E2', '#FFD700', '#C0C0C0', '#CD7F32', '#71797E', '#A0AEC0']
                class_ranges = ['90-100%', '80-89%', '70-79%', '60-69%', '50-59%', 'Below 50%']
                class_emojis = ['🥇', '🥇', '🥈', '🥉', '🔵', '⚪']
                
                for i, (cls, color, crange, emoji) in enumerate(zip(classifications, class_colors_list, class_ranges, class_emojis)):
                    with class_cols[i]:
                        count = class_counts.get(cls, 0)
                        st.markdown(f"""
                        <div style="text-align:center;padding:0.6rem;border-radius:10px;background:{color}22;border:2px solid {color};">
                            <div style="font-size:1.5rem;">{emoji}</div>
                            <div style="font-weight:700;font-size:0.7rem;">{cls}</div>
                            <div style="font-size:1.2rem;font-weight:700;color:{color};">{count}</div>
                            <small style="color:#888;font-size:0.6rem;">{crange}</small>
                        </div>
                        """, unsafe_allow_html=True)
            
            # ============================================================
            # COMMITTEE TAB 1: MASSIVE AI-POWERED ANALYTICS
            # ============================================================
            with committee_tabs[1]:
                st.subheader("📈 Advanced Appraisal Analytics Dashboard")
                st.markdown("*Fortune 500 Grade — Comprehensive Performance Intelligence*")
                
                if not all_emps_scored:
                    st.info("No appraisal data available yet.")
                else:
                    # ============================================================
                    # TOP METRICS ROW
                    # ============================================================
                    avg_all = sum(e['score'] for e in all_emps_scored) / len(all_emps_scored)
                    completed = len([v for v in all_assessments.values() if v.get('acceptance') == 'Accepted'])
                    pending_review = len([v for v in all_assessments.values() if v.get('status') in ['Submitted', 'Approved']])
                    pending_acceptance = len([v for v in all_assessments.values() if v.get('status') == 'Approved' and not v.get('acceptance')])
                    escalated = len([v for v in all_assessments.values() if v.get('status') == 'Escalated from TL'])
                    rejected = len([v for v in all_assessments.values() if v.get('acceptance') == 'Rejected'])
                    not_started = len([v for v in all_assessments.values() if v.get('status') == 'Not Started' or not v.get('status')])
                    in_progress = len([v for v in all_assessments.values() if v.get('status') == 'Submitted'])
                    
                    total_participants = len(all_emps_scored)
                    
                    st.markdown("### 📊 Key Performance Indicators")
                    m1, m2, m3, m4, m5, m6 = st.columns(6)
                    m1.metric("👥 Total Participants", total_participants)
                    m2.metric("📊 Avg Score", f"{avg_all:.1f}%")
                    m3.metric("🎉 Completed", completed)
                    m4.metric("⏳ Pending Review", pending_review)
                    m5.metric("🚨 Escalated", escalated)
                    m6.metric("❌ Rejected", rejected)
                    
                    st.markdown("---")
                    
                    # ============================================================
                    # KPI OBJECTIVE STATUS
                    # ============================================================
                    col1, col2 = st.columns(2)
                    
                    with col1:
                        st.subheader("🎯 KPI Objective Status")
                        all_perf_data = get_all_perf_cached()
                        kpi_approved = len(all_perf_data[all_perf_data['submission_status'] == 'Approved']) if not all_perf_data.empty else 0
                        kpi_submitted = len(all_perf_data[all_perf_data['submission_status'] == 'Submitted']) if not all_perf_data.empty else 0
                        kpi_draft = len(all_perf_data[all_perf_data['submission_status'] == 'Draft']) if not all_perf_data.empty else 0
                        kpi_no = total_participants - kpi_approved - kpi_submitted - kpi_draft
                        
                        kpi_data = pd.DataFrame({
                            'Status': ['Approved', 'Pending Approval', 'Draft', 'No Objectives'],
                            'Count': [kpi_approved, kpi_submitted, kpi_draft, max(0, kpi_no)]
                        })
                        kpi_data = kpi_data[kpi_data['Count'] > 0]
                        
                        fig_kpi = px.pie(kpi_data, values='Count', names='Status', hole=0.6,
                                        color_discrete_sequence=['#38a169', '#d69e2e', '#a0aec0', '#CC0000'])
                        fig_kpi.update_layout(height=350, title="KPI Objective Status")
                        st.plotly_chart(fig_kpi, use_container_width=True)
                        
                        # KPI stats
                        for _, row in kpi_data.iterrows():
                            pct = (row['Count'] / total_participants * 100) if total_participants > 0 else 0
                            st.markdown(f"""
                            <div style="display:flex;justify-content:space-between;padding:0.3rem 0;border-bottom:1px solid #eee;">
                                <span>{row['Status']}</span>
                                <strong>{row['Count']} ({pct:.1f}%)</strong>
                            </div>
                            """, unsafe_allow_html=True)
                    
                    with col2:
                        st.subheader("📋 Objective Performance")
                        # Calculate from actual KPI data
                        at_risk_count = 0
                        behind_count = 0
                        on_track_count = 0
                        completed_kpi_count = 0
                        
                        if not all_perf_data.empty:
                            for _, row in all_perf_data.iterrows():
                                status = row.get('status', '')
                                progress = row.get('progress', 0)
                                if status == 'At Risk' or (progress < 40 and row.get('submission_status') == 'Approved'):
                                    at_risk_count += 1
                                elif status == 'Near Target' or (40 <= progress < 65):
                                    behind_count += 1
                                elif status == 'On Track' or (65 <= progress < 100):
                                    on_track_count += 1
                                elif status == 'Completed' or progress >= 100:
                                    completed_kpi_count += 1
                        
                        total_kpis = at_risk_count + behind_count + on_track_count + completed_kpi_count
                        if total_kpis == 0:
                            total_kpis = 1  # Avoid division by zero
                        
                        obj_perf = pd.DataFrame({
                            'Performance': ['At Risk', 'Behind Schedule', 'On Track', 'Completed'],
                            'Count': [at_risk_count, behind_count, on_track_count, completed_kpi_count],
                            'Pct': [
                                at_risk_count/total_kpis*100,
                                behind_count/total_kpis*100,
                                on_track_count/total_kpis*100,
                                completed_kpi_count/total_kpis*100
                            ]
                        })
                        obj_perf = obj_perf[obj_perf['Count'] > 0]
                        
                        fig_obj = px.bar(obj_perf, x='Performance', y='Count', color='Performance',
                                        color_discrete_sequence=['#CC0000', '#d69e2e', '#3182ce', '#38a169'],
                                        text='Pct')
                        fig_obj.update_traces(texttemplate='%{text:.1f}%', textposition='outside')
                        fig_obj.update_layout(height=350, title="Objective Performance Distribution", showlegend=False)
                        st.plotly_chart(fig_obj, use_container_width=True)
                        
                        for _, row in obj_perf.iterrows():
                            st.markdown(f"""
                            <div style="display:flex;justify-content:space-between;padding:0.3rem 0;border-bottom:1px solid #eee;">
                                <span>{row['Performance']}</span>
                                <strong>{row['Count']} ({row['Pct']:.1f}%)</strong>
                            </div>
                            """, unsafe_allow_html=True)
                    
                    st.markdown("---")
                    
                    # ============================================================
                    # APPRAISAL COMPLETION STATUS
                    # ============================================================
                    st.subheader("📋 Appraisal Completion Status")
                    
                    completion_data = {
                        'Not Started': max(0, total_participants - len(all_assessments)),
                        'In Progress (Self)': in_progress,
                        'Pending Reviewer': pending_review,
                        'Pending Acceptance': pending_acceptance,
                        'Pending Committee': escalated,
                        'Completed': completed,
                        'Rejected': rejected
                    }
                    
                    comp_df = pd.DataFrame({
                        'Stage': list(completion_data.keys()),
                        'Count': list(completion_data.values())
                    })
                    comp_df = comp_df[comp_df['Count'] > 0]
                    comp_df['Pct'] = comp_df['Count'] / total_participants * 100
                    
                    col1, col2 = st.columns(2)
                    with col1:
                        fig_comp = px.pie(comp_df, values='Count', names='Stage', hole=0.5,
                                         color_discrete_sequence=['#a0aec0', '#3182ce', '#d69e2e', '#FFD700', '#CC0000', '#38a169', '#ff4444'])
                        fig_comp.update_layout(height=400, title="Appraisal Completion Status")
                        st.plotly_chart(fig_comp, use_container_width=True)
                    
                    with col2:
                        for _, row in comp_df.iterrows():
                            st.markdown(f"""
                            <div style="display:flex;justify-content:space-between;align-items:center;padding:0.5rem;margin:0.3rem 0;background:#f8f9fa;border-radius:6px;">
                                <span>{row['Stage']}</span>
                                <div style="text-align:right;">
                                    <strong>{row['Count']}</strong>
                                    <small style="color:#888;">({row['Pct']:.1f}%)</small>
                                </div>
                            </div>
                            """, unsafe_allow_html=True)
                    
                    st.markdown("---")
                    
                    # ============================================================
                    # AVERAGE SCORE BY DEPARTMENT
                    # ============================================================
                    st.subheader("📊 Average Score by Department")
                    
                    dept_scores = defaultdict(list)
                    for e in all_emps_scored:
                        dept_scores[e['dept']].append(e['score'])
                    
                    dept_avg_data = []
                    for dept, scores in dept_scores.items():
                        dept_avg_data.append({
                            'Department': dept,
                            'Avg Score': round(sum(scores)/len(scores), 1),
                            'Employees': len(scores),
                            'Min': min(scores),
                            'Max': max(scores)
                        })
                    
                    dept_avg_df = pd.DataFrame(dept_avg_data).sort_values('Avg Score', ascending=False)
                    
                    col1, col2 = st.columns(2)
                    with col1:
                        fig_dept = px.bar(dept_avg_df, x='Department', y='Avg Score', color='Avg Score',
                                         color_continuous_scale=['#CC0000', '#d69e2e', '#38a169'],
                                         text='Avg Score')
                        fig_dept.update_traces(texttemplate='%{text:.1f}%', textposition='outside')
                        fig_dept.update_layout(height=400, title="Average Score by Department")
                        st.plotly_chart(fig_dept, use_container_width=True)
                    
                    with col2:
                        st.dataframe(dept_avg_df, use_container_width=True, hide_index=True,
                                    column_config={
                                        'Department': 'Department',
                                        'Avg Score': st.column_config.NumberColumn('Avg Score', format='%.1f%%'),
                                        'Employees': 'Employees',
                                        'Min': st.column_config.NumberColumn('Min', format='%.0f%%'),
                                        'Max': st.column_config.NumberColumn('Max', format='%.0f%%')
                                    })
                    
                    st.markdown("---")
                    
                    # ============================================================
                    # REVIEWER & COMMITTEE RECOMMENDATIONS BREAKDOWN
                    # ============================================================
                    st.subheader("📋 Recommendations Breakdown")
                    
                    # Collect reviewer recommendations
                    reviewer_recs = defaultdict(int)
                    committee_recs = defaultdict(int)
                    
                    for emp_name, assessment in all_assessments.items():
                        sc = get_emp_score(emp_name)
                        if sc == 0: continue
                        
                        # Reviewer recommendation
                        comments = assessment.get('hod_comments', '') or assessment.get('tl_comments', '')
                        if assessment.get('acceptance') == 'Accepted':
                            if 'promot' in comments.lower():
                                reviewer_recs['Promote'] += 1
                            elif 'salary' in comments.lower() or 'increment' in comments.lower():
                                reviewer_recs['Salary Review'] += 1
                            elif 'train' in comments.lower() or 'develop' in comments.lower():
                                reviewer_recs['Training'] += 1
                            elif 'pip' in comments.lower() or 'improve' in comments.lower():
                                reviewer_recs['PIP'] += 1
                            elif 'status quo' in comments.lower() or 'maintain' in comments.lower():
                                reviewer_recs['Status Quo'] += 1
                            else:
                                reviewer_recs['Completed - No Specific Rec'] += 1
                        
                        # Committee recommendation
                        sr_decision = assessment.get('sr_decision', '')
                        if 'Overturned' in str(sr_decision):
                            committee_recs['Overturned'] += 1
                        elif 'Upheld' in str(sr_decision):
                            committee_recs['Upheld'] += 1
                        elif assessment.get('status') == 'Escalated from TL':
                            committee_recs['Pending Decision'] += 1
                    
                    col1, col2, col3 = st.columns(3)
                    
                    with col1:
                        st.subheader("Reviewer Recommendation")
                        if reviewer_recs:
                            rec_df = pd.DataFrame({'Recommendation': list(reviewer_recs.keys()), 'Count': list(reviewer_recs.values())})
                            fig_rec = px.pie(rec_df, values='Count', names='Recommendation', hole=0.5,
                                            color_discrete_sequence=['#38a169', '#3182ce', '#d69e2e', '#CC0000', '#FFD700', '#a0aec0'])
                            fig_rec.update_layout(height=300)
                            st.plotly_chart(fig_rec, use_container_width=True)
                            st.caption(f"Total Participants: {sum(reviewer_recs.values())}")
                        else:
                            st.info("No recommendations yet.")
                    
                    with col2:
                        st.subheader("Committee Decision")
                        if committee_recs:
                            com_df = pd.DataFrame({'Decision': list(committee_recs.keys()), 'Count': list(committee_recs.values())})
                            fig_com = px.pie(com_df, values='Count', names='Decision', hole=0.5,
                                            color_discrete_sequence=['#38a169', '#CC0000', '#d69e2e'])
                            fig_com.update_layout(height=300)
                            st.plotly_chart(fig_com, use_container_width=True)
                            st.caption(f"Total Decisions: {sum(committee_recs.values())}")
                        else:
                            st.info("No committee decisions yet.")
                    
                    with col3:
                        st.subheader("Classification Summary")
                        class_counts = defaultdict(int)
                        for e in all_emps_scored:
                            class_counts[e['class']] += 1
                        class_df = pd.DataFrame({'Class': list(class_counts.keys()), 'Count': list(class_counts.values())})
                        fig_cls = px.pie(class_df, values='Count', names='Class', hole=0.5,
                                        color_discrete_sequence=['#E5E4E2', '#FFD700', '#C0C0C0', '#CD7F32', '#71797E', '#A0AEC0'])
                        fig_cls.update_layout(height=300)
                        st.plotly_chart(fig_cls, use_container_width=True)
                    
                    st.markdown("---")
                    
                    # ============================================================
                    # AVERAGE PERFORMANCE SCORE & REGION BREAKDOWN
                    # ============================================================
                    st.subheader("🌍 Performance Overview")
                    
                    col1, col2, col3 = st.columns(3)
                    
                    with col1:
                        st.markdown(f"""
                        <div style="text-align:center;padding:2rem;background:linear-gradient(135deg, #1a1a1a, #2d2d2d);border-radius:16px;color:white;">
                            <div style="font-size:0.9rem;opacity:0.8;">AVERAGE PERFORMANCE SCORE</div>
                            <div style="font-size:3rem;font-weight:700;color:#38a169;">{avg_all:.1f}%</div>
                            <div style="font-size:0.8rem;opacity:0.7;">Total Employees: {total_participants}</div>
                        </div>
                        """, unsafe_allow_html=True)
                    
                    with col2:
                        # Region breakdown
                        region_scores = defaultdict(list)
                        for e in all_emps_scored:
                            region_scores[e['region']].append(e['score'])
                        
                        region_avg = {r: sum(s)/len(s) for r, s in region_scores.items()}
                        for region, avg in sorted(region_avg.items()):
                            rc = {'Abuja': '#CC0000', 'Lagos': '#1a1a1a', 'Aba': '#38a169'}.get(region, '#CC0000')
                            st.markdown(f"""
                            <div style="text-align:center;padding:1rem;margin:0.3rem 0;background:white;border-radius:10px;border-left:4px solid {rc};">
                                <strong>🌍 {region}</strong><br>
                                <span style="font-size:1.5rem;font-weight:700;color:{rc};">{avg:.1f}%</span><br>
                                <small style="color:#888;">{len(region_scores[region])} employees</small>
                            </div>
                            """, unsafe_allow_html=True)
                    
                    with col3:
                        # 9-Box distribution
                        box_dist = defaultdict(int)
                        for e in all_emps_scored:
                            box_dist[e['position_name']] += 1
                        
                        st.markdown("**9-Box Matrix Distribution**")
                        for pos, count in sorted(box_dist.items(), key=lambda x: x[1], reverse=True)[:9]:
                            st.markdown(f"""
                            <div style="display:flex;justify-content:space-between;padding:0.3rem 0;border-bottom:1px solid #eee;">
                                <small>{pos}</small>
                                <strong>{count}</strong>
                            </div>
                            """, unsafe_allow_html=True)
                    
                    st.markdown("---")
                    
                    # ============================================================
                    # AI-POWERED INSIGHTS
                    # ============================================================
                    st.subheader("🤖 AI-Powered Performance Insights")
                    
                    if st.button("🧠 Generate AI Insights", use_container_width=True, type="primary"):
                        with st.spinner("Analyzing appraisal data with AI..."):
                            try:
                                import openai
                                openai_key = os.environ.get("OPENAI_API_KEY", st.secrets.get("OPENAI_API_KEY", ""))
                                if openai_key:
                                    client = openai.OpenAI(api_key=openai_key)
                                    
                                    # Build comprehensive data summary
                                    insights_text = f"""
                                    Appraisal Cycle: {st.session_state.appraisal_cycle_name}
                                    Total Participants: {total_participants}
                                    Average Score: {avg_all:.1f}%
                                    Completed: {completed}
                                    Escalated: {escalated}
                                    Rejected: {rejected}
                                    
                                    Department Scores:
                                    {dept_avg_df.to_string()}
                                    
                                    Classification Distribution:
                                    {class_df.to_string()}
                                    
                                    9-Box Distribution:
                                    {dict(box_dist)}
                                    
                                    Recommendations:
                                    Reviewer: {dict(reviewer_recs)}
                                    Committee: {dict(committee_recs)}
                                    """
                                    
                                    response = client.chat.completions.create(
                                        model="gpt-3.5-turbo",
                                        messages=[{
                                            "role": "system",
                                            "content": "You are a Fortune 500 HR Analytics Director. Analyze this appraisal data and provide: 1) Top 3 key findings 2) Department performance comparison 3) Talent risk areas 4) 3 strategic recommendations for leadership 5) Predicted trends for next cycle. Be specific and data-driven."
                                        }, {
                                            "role": "user",
                                            "content": insights_text
                                        }],
                                        temperature=0.5,
                                        max_tokens=600
                                    )
                                    
                                    st.markdown("### 🤖 AI-Generated Insights")
                                    st.success("Analysis complete!")
                                    st.markdown(response.choices[0].message.content)
                                else:
                                    st.info("OpenAI API key not configured.")
                            except Exception as e:
                                st.warning(f"AI insights unavailable: {str(e)}")
                                st.markdown("""
                                ### 📊 Manual Analysis
                                
                                **Key Findings:**
                                1. Review completion rates and identify bottlenecks
                                2. Compare department performance for best practices
                                3. Address escalated cases promptly
                                
                                **Recommendations:**
                                1. Implement pre-review calibration sessions
                                2. Provide additional training for reviewers with high rejection rates
                                3. Fast-track high-performing employees for leadership development
                                """)
                    
                    # ============================================================
                    # EXPORT
                    # ============================================================
                    st.markdown("---")
                    st.subheader("📥 Export Analytics")
                    
                    col1, col2 = st.columns(2)
                    with col1:
                        if st.button("📊 Export Full Report (CSV)", use_container_width=True):
                            export_data = []
                            for e in all_emps_scored:
                                assessment = all_assessments.get(e['name'], {})
                                export_data.append({
                                    'Employee': e['name'],
                                    'Department': e['dept'],
                                    'Region': e['region'],
                                    'Score': e['score'],
                                    'Classification': e['class'],
                                    '9-Box Position': e['position_name'],
                                    'Potential': e['potential'],
                                    'Status': assessment.get('acceptance', assessment.get('status', 'N/A')),
                                    'Reviewer': assessment.get('reviewer_type', 'N/A')
                                })
                            export_df = pd.DataFrame(export_data)
                            st.download_button("📥 Download CSV", export_df.to_csv(index=False), "appraisal_analytics.csv", "text/csv")
                    
                    with col2:
                        if st.button("📊 Export Department Summary (CSV)", use_container_width=True):
                            st.download_button("📥 Download CSV", dept_avg_df.to_csv(index=False), "department_summary.csv", "text/csv")
            
            # ============================================================
            # TAB 2: SUBMISSIONS
            # ============================================================
            with committee_tabs[2]:
                st.subheader("📝 All Submitted Appraisals")
                sub = {k: v for k, v in all_assessments.items() if v.get('status') in ['Submitted', 'Approved', 'Awaiting HOD Re-review', 'Awaiting TL Re-review']}
                if sub:
                    st.success(f"📋 {len(sub)} pending appraisal(s)")
                    for en, a in sub.items():
                        sc = get_emp_score(en)
                        cl, co = get_classification(sc) if sc > 0 else ('N/A', '#a0aec0')
                        with st.expander(f"📋 {en} | {get_employee_dept(en)} | {get_region(get_employee_subsidiary(en))} | {cl}", expanded=False):
                            c1, c2, c3 = st.columns(3)
                            c1.metric("Score", f"{sc:.1f}%" if sc else "N/A")
                            c2.metric("Status", a.get('status', 'N/A'))
                            c3.metric("Reviewer", a.get('reviewer_type', 'N/A'))
                            if a.get('rejection_comment'):
                                st.warning(f"💬 Rejection: {a['rejection_comment'][:200]}")
                else:
                    st.info("No pending submissions.")
            
            # ============================================================
            # TAB 3: ESCALATED (FIXED - UPHOLD/OVERTURN WITH PROPER STATE UPDATE)
            # ============================================================
            with committee_tabs[3]:
                st.subheader("🚨 Escalated Appraisals")
                esc = {k: v for k, v in all_assessments.items() if v.get('status') == 'Escalated from TL' or (v.get('acceptance') == 'Rejected' and v.get('status') != 'Awaiting HOD Re-review')}
                if esc:
                    for en, a in esc.items():
                        with st.expander(f"🚨 {en} | {get_employee_dept(en)} | Rejections: {a.get('reject_count', 1)}", expanded=True):
                            st.markdown(f"**Staff Comments:** {a.get('comments', 'N/A')}")
                            st.markdown(f"**Reviewer Comments:** {a.get('hod_comments', a.get('tl_comments', 'N/A'))}")
                            st.markdown(f"**Rejection Reason:** {a.get('rejection_comment', 'N/A')}")
                            
                            c1, c2 = st.columns(2)
                            with c1:
                                if st.button(f"✅ Uphold Reviewer", key=f"cmt_up_{en}"):
                                    # Update employee's session state so certificate shows
                                    st.session_state.self_assessments[en].update({
                                        'acceptance': 'Accepted',
                                        'status': 'Completed',
                                        'sr_decision': 'Upheld by Committee'
                                    })
                                    try:
                                        db.archive_appraisal(en, a.get('email', ''), a.get('department', ''),
                                            st.session_state.appraisal_cycle_name, 'Accepted - Committee Upheld',
                                            a['scores'], a.get('hod_scores', a.get('tl_scores', {})),
                                            a.get('comments', ''), a.get('hod_comments', ''),
                                            now_wat.strftime('%Y-%m-%d %H:%M WAT'))
                                        db._patch("appraisals", 
                                            {"status": "Completed", "acceptance": "Accepted", "sr_decision": "Upheld by Committee"},
                                            {"user_name": en, "cycle_name": st.session_state.appraisal_cycle_name})
                                    except: pass
                                    emp_email = get_employee_email(en)
                                    if emp_email:
                                        try:
                                            EmailService().send_email(emp_email,
                                                f"📋 Appraisal Decision: {st.session_state.appraisal_cycle_name}",
                                                f"Dear {en},\n\nThe Appraisal Committee has upheld the reviewer's decision.\n\n"
                                                f"Your appraisal is now complete. Log in to download your certificate.\n\n"
                                                f"https://churchgate-churchgate-hris.hf.space\n\nChurchgate Group HR")
                                        except: pass
                                    log_audit('Committee Upheld', f'{en} upheld by committee')
                                    st.success("✅ Upheld! Staff notified and certificate available."); st.balloons(); time.sleep(1.5); st.rerun()
                            
                            with c2:
                                if st.button(f"🔄 Overturn - Favor Staff", key=f"cmt_ov_{en}"):
                                    # Update employee's session state so certificate shows immediately
                                    st.session_state.self_assessments[en].update({
                                        'acceptance': 'Accepted',
                                        'hod_scores': a['scores'],  # Use staff scores as final
                                        'status': 'Completed',
                                        'sr_decision': 'Overturned by Committee',
                                        'reviewer_type': 'Committee (Overturned)'
                                    })
                                    try:
                                        db.archive_appraisal(en, a.get('email', ''), a.get('department', ''),
                                            st.session_state.appraisal_cycle_name, 'Accepted - Overturned',
                                            a['scores'], a['scores'],  # Staff scores become final
                                            a.get('comments', ''), a.get('hod_comments', ''),
                                            now_wat.strftime('%Y-%m-%d %H:%M WAT'))
                                        db._patch("appraisals", 
                                            {"status": "Completed", "acceptance": "Accepted", "sr_decision": "Overturned by Committee"},
                                            {"user_name": en, "cycle_name": st.session_state.appraisal_cycle_name})
                                    except: pass
                                    emp_email = get_employee_email(en)
                                    if emp_email:
                                        try:
                                            EmailService().send_email(emp_email,
                                                f"🎉 Appraisal Overturned: {st.session_state.appraisal_cycle_name}",
                                                f"Dear {en},\n\nGreat news! The Appraisal Committee has overturned the decision in your favor.\n\n"
                                                f"Your appraisal is now complete. Log in to download your certificate.\n\n"
                                                f"https://churchgate-churchgate-hris.hf.space\n\nChurchgate Group HR")
                                        except: pass
                                    log_audit('Committee Overturned', f'{en} overturned by committee')
                                    st.success("🔄 Overturned! Staff notified and certificate available."); st.balloons(); time.sleep(1.5); st.rerun()
                else:
                    st.info("No escalated appraisals.")
            
            # ============================================================
            # TAB 4: COMPLETED
            # ============================================================
            with committee_tabs[4]:
                st.subheader("🎉 Completed Appraisals")
                comp = {k: v for k, v in all_assessments.items() if v.get('acceptance') == 'Accepted'}
                if comp:
                    st.success(f"🎉 {len(comp)} completed appraisal(s)")
                    for en, a in comp.items():
                        sc = get_emp_score(en)
                        cl, co = get_classification(sc)
                        st.markdown(f"""
                        <div class="kpi-card" style="border-left-color:{co};">
                            <div style="display:flex;justify-content:space-between;align-items:center;">
                                <div>
                                    <strong>👤 {en}</strong><br>
                                    <small>{get_employee_dept(en)} | {get_region(get_employee_subsidiary(en))}</small>
                                </div>
                                <div style="text-align:right;">
                                    <span style="font-size:1.2rem;font-weight:700;color:{co};">{sc:.1f}%</span><br>
                                    <span style="color:{co};font-weight:600;">{cl}</span>
                                </div>
                            </div>
                        </div>
                        """, unsafe_allow_html=True)
                else:
                    st.info("No completed appraisals yet.")
            
            # ============================================================
            # TAB 5: RANKINGS
            # ============================================================
            with committee_tabs[5]:
                st.subheader("🏆 Performance Rankings")
                sorted_emps = sorted(all_emps_scored, key=lambda x: x['score'], reverse=True)
                c1, c2 = st.columns(2)
                with c1:
                    st.markdown("### 🏆 Top 10 Performers")
                    if sorted_emps:
                        for i, e in enumerate(sorted_emps[:10]):
                            st.markdown(f"""
                            <div class="kpi-card" style="border-left-color:#38a169;">
                                <strong>#{i+1} {e['name']}</strong> | {e['dept']} | 
                                <span style="font-weight:700;color:#38a169;">{e['score']:.1f}%</span> | {e['class']}
                            </div>
                            """, unsafe_allow_html=True)
                    else: st.info("No data.")
                with c2:
                    st.markdown("### ⚠️ Bottom 10 Performers")
                    if sorted_emps:
                        for i, e in enumerate(sorted_emps[-10:]):
                            st.markdown(f"""
                            <div class="kpi-card" style="border-left-color:#CC0000;">
                                <strong>#{len(sorted_emps)-9+i} {e['name']}</strong> | {e['dept']} | 
                                <span style="font-weight:700;color:#CC0000;">{e['score']:.1f}%</span> | {e['class']}
                            </div>
                            """, unsafe_allow_html=True)
                    else: st.info("No data.")
            
            # ============================================================
            # COMMITTEE TAB 6: RECOMMENDATIONS (FROM ACTUAL REVIEWS)
            # ============================================================
            with committee_tabs[6]:
                st.subheader("📋 Appraisal Recommendations")
                st.markdown("*Based on actual reviewer feedback and committee decisions*")
                
                # Collect real recommendations from assessments
                recommendations = []
                
                for emp_name, assessment in all_assessments.items():
                    sc = get_emp_score(emp_name)
                    if sc == 0:
                        continue
                    
                    cl, co = get_classification(sc)
                    dept = get_employee_dept(emp_name)
                    region = get_region(get_employee_subsidiary(emp_name))
                    
                    # Get the actual reviewer recommendation from comments
                    hod_comments = assessment.get('hod_comments', '')
                    tl_comments = assessment.get('tl_comments', '')
                    sr_decision = assessment.get('sr_decision', '')
                    acceptance = assessment.get('acceptance', '')
                    reviewer_type = assessment.get('reviewer_type', 'N/A')
                    
                    # Determine real recommendation
                    recommendation = ""
                    recommendation_source = ""
                    
                    if sr_decision and 'Overturned' in str(sr_decision):
                        recommendation = "OVERTURNED BY COMMITTEE - Staff scores accepted"
                        recommendation_source = "Appraisal Committee"
                    elif sr_decision and 'Upheld' in str(sr_decision):
                        recommendation = "UPHELD BY COMMITTEE - Reviewer decision stands"
                        recommendation_source = "Appraisal Committee"
                    elif acceptance == 'Accepted':
                        # Check reviewer comments for recommendation
                        comments = hod_comments or tl_comments or ''
                        if 'promot' in comments.lower():
                            recommendation = "PROMOTE - " + comments[:100]
                        elif 'salary' in comments.lower() or 'increment' in comments.lower():
                            recommendation = "SALARY REVIEW - " + comments[:100]
                        elif 'train' in comments.lower() or 'develop' in comments.lower():
                            recommendation = "TRAINING & DEVELOPMENT - " + comments[:100]
                        elif 'pip' in comments.lower() or 'improve' in comments.lower():
                            recommendation = "PERFORMANCE IMPROVEMENT PLAN - " + comments[:100]
                        elif 'consistent' in comments.lower() or 'good' in comments.lower() or 'solid' in comments.lower():
                            recommendation = "MAINTAIN - Consistent Performer"
                        else:
                            recommendation = "REVIEW COMPLETED - " + (comments[:80] if comments else "No specific recommendation")
                        recommendation_source = f"{reviewer_type} Review"
                    elif acceptance == 'Rejected':
                        recommendation = "DISPUTED - Awaiting resolution"
                        recommendation_source = "Pending"
                    elif assessment.get('status') == 'Escalated from TL':
                        recommendation = "ESCALATED - Pending Committee Decision"
                        recommendation_source = "Pending Committee"
                    else:
                        recommendation = "IN PROGRESS - Review not yet complete"
                        recommendation_source = "Pending"
                    
                    recommendations.append({
                        'Employee': emp_name,
                        'Department': dept,
                        'Region': region,
                        'Score': sc,
                        'Classification': cl,
                        'Reviewer': reviewer_type,
                        'Recommendation': recommendation,
                        'Source': recommendation_source,
                        'Status': acceptance or assessment.get('status', 'Pending')
                    })
                
                if recommendations:
                    # Summary by recommendation type
                    st.subheader("📊 Recommendation Summary")
                    
                    rec_categories = defaultdict(list)
                    for r in recommendations:
                        cat = r['Recommendation'].split(' - ')[0] if ' - ' in r['Recommendation'] else r['Recommendation'][:50]
                        rec_categories[cat].append(r)
                    
                    c1, c2 = st.columns(2)
                    with c1:
                        rec_summary = {k: len(v) for k, v in rec_categories.items()}
                        if rec_summary:
                            rec_df = pd.DataFrame({'Recommendation': list(rec_summary.keys()), 'Count': list(rec_summary.values())})
                            fig = px.pie(rec_df, values='Count', names='Recommendation', hole=0.5,
                                        color_discrete_sequence=['#38a169', '#3182ce', '#d69e2e', '#CC0000', '#FFD700', '#a0aec0'])
                            fig.update_layout(height=350, title="Recommendations from Reviews")
                            st.plotly_chart(fig, use_container_width=True)
                    
                    with c2:
                        st.subheader("👥 By Status")
                        status_counts = defaultdict(int)
                        for r in recommendations:
                            status_counts[r['Source']] += 1
                        for source, count in status_counts.items():
                            st.markdown(f"""
                            <div class="kpi-card">
                                <strong>{source}</strong>
                                <span style="float:right;font-size:1.3rem;font-weight:700;">{count}</span>
                                <br><small style="color:#888;">employees</small>
                            </div>
                            """, unsafe_allow_html=True)
                    
                    st.markdown("---")
                    st.subheader("📋 Detailed Recommendations by Employee")
                    
                    for r in sorted(recommendations, key=lambda x: x['Score'], reverse=True):
                        border_color = '#38a169' if r['Status'] == 'Accepted' else '#CC0000' if r['Status'] == 'Rejected' else '#d69e2e'
                        source_badge = {
                            'Appraisal Committee': 'badge-green',
                            'HOD Review': 'badge-green',
                            'Team Lead Review': 'badge-green',
                            'Pending': 'badge-yellow',
                            'Pending Committee': 'badge-red'
                        }.get(r['Source'], 'badge-gray')
                        
                        with st.expander(f"{r['Classification']} {r['Employee']} — {r['Score']:.1f}% — {r['Source']}", expanded=False):
                            st.markdown(f"""
                            <div style="padding:0.5rem;border-left:4px solid {border_color};margin:0.3rem 0;">
                                <strong>👤 {r['Employee']}</strong><br>
                                <small>🏢 {r['Department']} | 🌍 {r['Region']} | 🏅 {r['Classification']}</small><br>
                                <small>📊 Score: {r['Score']:.1f}% | Reviewer: {r['Reviewer']}</small><br>
                                <small>Status: {r['Status']} | Source: <span class="badge {source_badge}">{r['Source']}</span></small>
                                <div style="margin-top:0.5rem;padding:0.5rem;background:#f8f9fa;border-radius:6px;">
                                    <strong>💡 Recommendation:</strong><br>
                                    <span style="color:#CC0000;">{r['Recommendation']}</span>
                                </div>
                            </div>
                            """, unsafe_allow_html=True)
                else:
                    st.info("No appraisal data available for recommendations.")



def promotions():
    st.markdown("""<div class="churchgate-header"><h1>🚀 Promotions & Career Progression Console</h1><p>360° A-Player Assessment | Succession Planning | Talent Pipeline Management</p></div>""", unsafe_allow_html=True)
    
    user_role = st.session_state.user['role'] if st.session_state.user else 'Employee'
    user_dept = st.session_state.user.get('department', '') if st.session_state.user else ''
    is_admin = user_role in ['Admin', 'HR Director'] or user_dept == 'Senior Management'
    
    # ALWAYS load from database
    aplayers_data = {}
    try:
        db_players = db.get_all_aplayers()
        if not db_players.empty:
            for _, row in db_players.iterrows():
                dept = row['department']
                if dept not in aplayers_data:
                    aplayers_data[dept] = []
                aplayers_data[dept].append({
                    'name': row['name'], 'position': row['position'],
                    'nominated_by': row['nominated_by'],
                    'perf_score': row['perf_score'], 'leadership': row['leadership'],
                    'strategic': row['strategic'], 'peer_review': row['peer_review'],
                    'junior_review': row['junior_review'], 'independent_review': row['independent_review'],
                    'overall': row['overall'], 'readiness': row['readiness'],
                    'gap': row['gap'], 'risk': row['risk'], 'photo': None
                })
    except:
        pass
    
    all_depts = ['Senior Management', 'Technology Group', 'Facility Management', 'Human Resources', 'Accounts & Finance', 
                 'Sales & Marketing', 'Procurement', 'Security', 'Legal', 'Operations', 
                 'Engineering', 'Central Stores', 'Project Development', 'Trade Services']
    for dept in all_depts:
        if dept not in aplayers_data:
            aplayers_data[dept] = []
    
    try:
        nom_data = db.get_all_nominations()
        nominations_list = nom_data.to_dict('records') if not nom_data.empty else []
    except:
        nominations_list = []
    
    if 'aplayer_nominations' not in st.session_state:
        st.session_state.aplayer_nominations = nominations_list
    
    tab1, tab2, tab3, tab4, tab5 = st.tabs(["🏆 A-Players Board", "📝 Nominate A-Player", "📊 360° Assessment", "📋 Succession Pipeline", "📥 Reports"])
    
    with tab1:
        st.subheader("🏆 Group A-Players — Talent Pipeline")
        
        if st.session_state.aplayer_nominations:
            pending_count = len(st.session_state.aplayer_nominations)
            st.markdown(f"""<div style="background: #fff3cd; padding: 0.8rem 1rem; border-radius: 8px; margin-bottom: 1rem; border-left: 4px solid #d69e2e;"><strong>🔔 {pending_count} Pending Nomination{'s' if pending_count > 1 else ''}</strong></div>""", unsafe_allow_html=True)
        
        all_players = []
        for dept, players in aplayers_data.items():
            for p in players:
                p['department'] = dept
                all_players.append(p)
        
        ready_now = [p for p in all_players if p['readiness'] == 'Ready Now']
        
        c1, c2, c3, c4 = st.columns(4)
        c1.metric("Total A-Players", len(all_players))
        c2.metric("Ready Now", len(ready_now), "🚀")
        c3.metric("Avg Score", f"{sum(p['overall'] for p in all_players) / len(all_players):.1f}%" if all_players else "N/A")
        c4.metric("Departments", len([d for d, p in aplayers_data.items() if p]))
        
        st.markdown("---")
        
        if is_admin:
            view_dept = st.selectbox("🏢 Filter by Department", ["All Departments"] + all_depts)
        else:
            view_dept = "All Departments"
        
        if view_dept == "All Departments":
            display_players = all_players
        else:
            display_players = [p for p in all_players if p['department'] == view_dept]
        
        if display_players:
            for player in display_players:
                player_key = f"{player['name']}_{player['department']}"
                color = "#38a169" if player['readiness'] == 'Ready Now' else "#d69e2e" if 'Q3' in str(player['readiness']) else "#3182ce" if 'Q4' in str(player['readiness']) else "#CC0000"
                initials = generate_initials(player['name'])
                st.markdown(f"""
                <div style="background: white; padding: 1.2rem; border-radius: 10px; margin-bottom: 0.8rem; border-left: 5px solid {color}; box-shadow: 0 2px 8px rgba(0,0,0,0.06);">
                    <div style="display: flex; align-items: center; gap: 1.2rem;">
                        <div style="width: 55px; height: 55px; border-radius: 50%; background: #CC0000; display: flex; align-items: center; justify-content: center; font-weight: 700; font-size: 1.3rem; color: white; min-width: 55px;">{initials}</div>
                        <div style="flex: 1;">
                            <div style="display: flex; justify-content: space-between; align-items: center;">
                                <div>
                                    <strong style="font-size: 1.1rem;">{player['name']}</strong>
                                    <span style="color: #666; font-size: 0.85rem; margin-left: 0.5rem;">{player['department']}</span>
                                    <br><small>{player['position']} → <strong>Nominated by: {player['nominated_by']}</strong></small>
                                </div>
                                <div style="text-align: right;">
                                    <span style="background: {color}; color: white; padding: 0.3rem 0.8rem; border-radius: 15px; font-weight: 700; font-size: 0.85rem;">{player['readiness']}</span>
                                    <br><small style="font-size: 0.8rem;">Overall: {player['overall']}% | Risk: {player['risk']}</small>
                                </div>
                            </div>
                            <div style="margin-top: 0.5rem; display: flex; gap: 0.5rem; flex-wrap: wrap;">
                                <span style="background: #f0f0f0; padding: 0.2rem 0.6rem; border-radius: 10px; font-size: 0.75rem;">Perf: {player['perf_score']}%</span>
                                <span style="background: #f0f0f0; padding: 0.2rem 0.6rem; border-radius: 10px; font-size: 0.75rem;">Leadership: {player['leadership']}%</span>
                                <span style="background: #f0f0f0; padding: 0.2rem 0.6rem; border-radius: 10px; font-size: 0.75rem;">Strategic: {player['strategic']}%</span>
                                <span style="background: #f0f0f0; padding: 0.2rem 0.6rem; border-radius: 10px; font-size: 0.75rem;">Peer: {player['peer_review']}%</span>
                                <span style="background: #f0f0f0; padding: 0.2rem 0.6rem; border-radius: 10px; font-size: 0.75rem;">Junior: {player['junior_review']}%</span>
                                <span style="background: #f0f0f0; padding: 0.2rem 0.6rem; border-radius: 10px; font-size: 0.75rem;">Indep: {player['independent_review']}%</span>
                            </div>
                            <div style="margin-top: 0.3rem;"><small style="color: #CC0000;">Gap: {player['gap']}</small></div>
                        </div>
                    </div>
                </div>""", unsafe_allow_html=True)
                
        else:
            st.info("No A-Players found. Nominate using the 'Nominate A-Player' tab.")
    
    with tab2:
        st.subheader("📝 Nominate A-Player")
        st.info("HODs and Line Managers can nominate top performers as A-Players for the 360-degree assessment process.")
        
        with st.form("nominate_form"):
            c1, c2 = st.columns(2)
            with c1:
                nominee_name = st.text_input("Nominee Full Name *")
                nominee_position = st.text_input("Current Position *")
                nominee_dept = st.selectbox("Department *", all_depts)
            with c2:
                nominated_by = st.text_input("Nominated By (HOD/Line Manager) *")
                nomination_reason = st.text_area("Reason for Nomination")
            if st.form_submit_button("✅ Submit Nomination", use_container_width=True):
                if nominee_name and nominee_dept and nominated_by:
                    nom_entry = {
                        'name': nominee_name, 'position': nominee_position, 'department': nominee_dept,
                        'nominated_by': nominated_by, 'reason': nomination_reason,
                        'date': datetime.now().strftime('%Y-%m-%d'),
                        'submitted_by': st.session_state.user['name'] if st.session_state.user else 'Unknown'
                    }
                    st.session_state.aplayer_nominations.append(nom_entry)
                    db.save_nomination(nominee_name, nominee_position, nominee_dept, nominated_by, nomination_reason,
                                     st.session_state.user['name'] if st.session_state.user else 'Unknown',
                                     st.session_state.user['email'] if st.session_state.user else '')
                    st.success(f"✅ {nominee_name} nominated successfully!")
                    st.balloons()
                    time.sleep(1)
                    st.rerun()
        
        if st.session_state.aplayer_nominations:
            st.markdown("---")
            st.markdown("### 📋 Pending Nominations — Admin Review")
            for i, nom in enumerate(st.session_state.aplayer_nominations):
                c1, c2 = st.columns([3, 1])
                with c1:
                    st.markdown(f"""<div style="background: white; padding: 0.8rem; border-radius: 6px; margin-bottom: 0.3rem; border-left: 3px solid #d69e2e;"><strong>{nom.get('name', '')}</strong> - {nom.get('position', '')}<br><small>{nom.get('department', '')} | By: {nom.get('nominated_by', '')} | {str(nom.get('date', nom.get('created_at', '')))[:10]}</small><br><small style="color: #666;">Reason: {nom.get('reason', 'N/A')}</small></div>""", unsafe_allow_html=True)
                with c2:
                    if is_admin:
                        action_key = f"action_{i}"
                        if action_key not in st.session_state:
                            st.session_state[action_key] = None
                        
                        col_btn1, col_btn2, col_btn3 = st.columns(3)
                        with col_btn1:
                            if st.button(f"✅", key=f"approve_{i}", use_container_width=True, help="Approve"):
                                st.session_state[action_key] = 'approve'
                        with col_btn2:
                            if st.button(f"⏸️", key=f"hold_{i}", use_container_width=True, help="On Hold"):
                                st.session_state[action_key] = 'hold'
                        with col_btn3:
                            if st.button(f"❌", key=f"reject_{i}", use_container_width=True, help="Reject"):
                                st.session_state[action_key] = 'reject'
                        
                        if st.session_state[action_key]:
                            reason = st.text_area(f"Reason for {st.session_state[action_key]}: *", key=f"reason_{i}", placeholder="Please provide a reason...")
                            if st.button(f"✅ Confirm {st.session_state[action_key].title()}", key=f"confirm_{i}", use_container_width=True):
                                if reason:
                                    if st.session_state[action_key] == 'approve':
                                        db.save_aplayer(nom['name'], nom['position'], nom['department'], nom['nominated_by'], 0, 0, 0, 0, 0, 0, 0, 'Pending Assessment', 'TBD', 'TBD')
                                        st.session_state.aplayer_nominations.pop(i)
                                        st.success(f"✅ {nom['name']} approved!")
                                    elif st.session_state[action_key] == 'reject':
                                        st.session_state.aplayer_nominations.pop(i)
                                        st.error(f"❌ {nom['name']} rejected.")
                                    st.session_state[action_key] = None
                                    st.rerun()
    
    with tab3:
        st.subheader("📊 360° Assessment Scorecard")
        st.info("Complete 360-degree assessment. Weights: Performance (35%), Leadership (25%), Strategic (20%), Peer (10%), Junior (5%), Independent (5%).")
        
        if is_admin:
            depts_with_players = all_depts
            if depts_with_players:
                assess_dept = st.selectbox("Select Department", all_depts, key="assess_dept")
                if assess_dept:
                    player_names = [p['name'] for p in aplayers_data.get(assess_dept, [])]
                    if player_names:
                        assess_player_name = st.selectbox("Select A-Player", player_names, key="assess_player")
                        
                        player_data = None
                        for p in aplayers_data.get(assess_dept, []):
                            if p['name'] == assess_player_name:
                                player_data = p
                                break
                        
                        if player_data:
                            st.markdown(f"### Assessing: {player_data['name']} - {player_data['position']}")
                            with st.form("assessment_form"):
                                c1, c2 = st.columns(2)
                                with c1:
                                    perf = st.slider("Performance Score (35%)", 0, 100, int(player_data['perf_score']))
                                    leadership = st.slider("Leadership Competency (25%)", 0, 100, int(player_data['leadership']))
                                    strategic = st.slider("Strategic Impact (20%)", 0, 100, int(player_data['strategic']))
                                with c2:
                                    peer = st.slider("Peer Review (10%)", 0, 100, int(player_data['peer_review']))
                                    junior = st.slider("Junior Review (5%)", 0, 100, int(player_data['junior_review']))
                                    independent = st.slider("Independent Review (5%)", 0, 100, int(player_data['independent_review']))
                                
                                readiness_options = ['Ready Now', 'Q3 2026', 'Q4 2026', 'Q1 2027', 'Pending Assessment']
                                current_readiness = player_data['readiness'] if player_data['readiness'] in readiness_options else 'Pending Assessment'
                                readiness = st.selectbox("Readiness", readiness_options, index=readiness_options.index(current_readiness))
                                gap = st.text_input("Competency Gap", player_data['gap'])
                                risk = st.selectbox("Risk", ['Low', 'Medium', 'High'], index=['Low', 'Medium', 'High'].index(player_data['risk']) if player_data['risk'] in ['Low', 'Medium', 'High'] else 0)
                                
                                if st.form_submit_button("💾 Save Assessment", use_container_width=True):
                                    overall = int(perf * 0.35 + leadership * 0.25 + strategic * 0.20 + peer * 0.10 + junior * 0.05 + independent * 0.05)
                                    db.save_aplayer(player_data['name'], player_data['position'], assess_dept, player_data['nominated_by'], perf, leadership, strategic, peer, junior, independent, overall, readiness, gap, risk)
                                    st.success(f"✅ Assessment saved! Overall: {overall}%")
                                    st.balloons()
                                    time.sleep(1)
                                    st.rerun()
                    else:
                        st.info(f"No A-Players found in {assess_dept}. Nominate first.")
            else:
                st.info("No A-Players available. Nominate first.")
    
    with tab4:
        st.subheader("📋 Succession Pipeline")
        pipeline_data = []
        for dept, players in aplayers_data.items():
            for p in players:
                pipeline_data.append({'Department': dept, 'A-Player': p['name'], 'Position': p['position'], 'Readiness': p['readiness'], 'Score': p['overall'], 'Risk': p['risk'], 'Gap': p['gap']})
        if pipeline_data:
            df = pd.DataFrame(pipeline_data)
            st.dataframe(df, use_container_width=True, hide_index=True)
            
            readiness_counts = df['Readiness'].value_counts()
            fig = px.pie(values=readiness_counts.values, names=readiness_counts.index, hole=0.5, color_discrete_sequence=['#38a169', '#d69e2e', '#3182ce', '#CC0000'])
            fig.update_layout(height=350)
            st.plotly_chart(fig, use_container_width=True)
    
    with tab5:
        st.subheader("📥 Promotion Reports")
        c1, c2 = st.columns(2)
        with c1:
            if st.button("📊 Generate CSV Report", use_container_width=True):
                csv_data = []
                for dept, players in aplayers_data.items():
                    for p in players:
                        csv_data.append({'Department': dept, 'Name': p['name'], 'Position': p['position'], 'Perf': p['perf_score'], 'Leadership': p['leadership'], 'Strategic': p['strategic'], 'Peer': p['peer_review'], 'Junior': p['junior_review'], 'Independent': p['independent_review'], 'Overall': p['overall'], 'Readiness': p['readiness'], 'Gap': p['gap'], 'Risk': p['risk']})
                if csv_data:
                    st.download_button("📥 Download CSV", pd.DataFrame(csv_data).to_csv(index=False), "aplayers_report.csv", "text/csv")
        with c2:
            if st.button("📕 Generate Executive PDF", use_container_width=True):
                try:
                    import fpdf
                    FPDF = fpdf.FPDF
                    pdf = FPDF(orientation='L', unit='mm', format='A4')
                    pdf.set_auto_page_break(auto=True, margin=10)
                    pdf.add_page()
                    
                    pdf.set_fill_color(55, 55, 55)
                    pdf.rect(0, 0, 297, 35, 'F')
                    pdf.set_fill_color(204, 0, 0)
                    pdf.rect(0, 35, 297, 3, 'F')
                    pdf.set_font('Helvetica', 'B', 22)
                    pdf.set_text_color(255, 255, 255)
                    pdf.cell(0, 18, 'CHURCHGATE GROUP', ln=True, align='C')
                    pdf.set_font('Helvetica', 'B', 12)
                    pdf.set_text_color(255, 255, 255)
                    pdf.cell(0, 8, 'A-PLAYERS & SUCCESSION PIPELINE - EXECUTIVE REPORT', ln=True, align='C')
                    pdf.ln(10)
                    
                    total = len(all_players)
                    ready = len(ready_now)
                    avg = sum(p['overall'] for p in all_players) / total if total > 0 else 0
                    
                    pdf.set_fill_color(245, 245, 245)
                    pdf.rect(10, pdf.get_y(), 277, 10, 'F')
                    pdf.set_font('Helvetica', 'B', 9)
                    pdf.set_text_color(26, 26, 26)
                    pdf.cell(92, 10, f'  Total: {total}', 0, 0, 'L')
                    pdf.cell(92, 10, f'Ready Now: {ready}', 0, 0, 'C')
                    pdf.cell(93, 10, f'Avg Score: {avg:.1f}%  ', 0, 0, 'R')
                    pdf.ln(14)
                    
                    cards_data = [('TOTAL', str(total), (26, 26, 26)), ('READY NOW', str(ready), (56, 161, 105)), ('AVG SCORE', f'{avg:.1f}%', (204, 0, 0)), ('DEPTS', str(len([d for d, p in aplayers_data.items() if p])), (49, 130, 206))]
                    col_w = 65
                    col_h = 16
                    x_start = 14
                    y_pos = pdf.get_y()
                    for j, (label, value, color) in enumerate(cards_data):
                        x = x_start + j * (col_w + 5)
                        pdf.set_fill_color(*color)
                        pdf.rect(x, y_pos, col_w, col_h, 'F')
                        pdf.set_text_color(255, 255, 255)
                        pdf.set_xy(x, y_pos + 1)
                        pdf.set_font('Helvetica', 'B', 7)
                        pdf.cell(col_w, 6, label, 0, 0, 'C')
                        pdf.set_xy(x, y_pos + 7)
                        pdf.set_font('Helvetica', 'B', 14)
                        pdf.cell(col_w, 8, value, 0, 0, 'C')
                    pdf.set_y(y_pos + col_h + 8)
                    
                    pdf.set_fill_color(26, 26, 26)
                    pdf.set_text_color(255, 255, 255)
                    pdf.set_font('Helvetica', 'B', 8)
                    pdf.cell(6, 7, ' #', 1, 0, 'C', True)
                    pdf.cell(42, 7, ' NAME', 1, 0, 'L', True)
                    pdf.cell(42, 7, ' DEPARTMENT', 1, 0, 'L', True)
                    pdf.cell(50, 7, ' POSITION', 1, 0, 'L', True)
                    pdf.cell(18, 7, 'SCORE', 1, 0, 'C', True)
                    pdf.cell(32, 7, 'READINESS', 1, 0, 'C', True)
                    pdf.cell(22, 7, 'RISK', 1, 0, 'C', True)
                    pdf.cell(65, 7, ' GAP', 1, 0, 'L', True)
                    pdf.ln()
                    
                    row_num = 0
                    pdf.set_font('Helvetica', '', 7)
                    for dept, players in aplayers_data.items():
                        for p in players:
                            row_num += 1
                            if p['overall'] >= 85:
                                tc = (56, 161, 105)
                            elif p['overall'] >= 70:
                                tc = (214, 158, 46)
                            else:
                                tc = (204, 0, 0)
                            pdf.set_text_color(26, 26, 26)
                            pdf.cell(6, 6, str(row_num), 1, 0, 'C')
                            pdf.cell(42, 6, f' {p["name"][:25]}', 1, 0, 'L')
                            pdf.cell(42, 6, f' {dept[:25]}', 1, 0, 'L')
                            pdf.cell(50, 6, f' {p["position"][:30]}', 1, 0, 'L')
                            pdf.set_text_color(*tc)
                            pdf.set_font('Helvetica', 'B', 7)
                            pdf.cell(18, 6, f'{p["overall"]}%', 1, 0, 'C')
                            pdf.set_text_color(26, 26, 26)
                            pdf.set_font('Helvetica', '', 7)
                            pdf.cell(32, 6, p['readiness'][:18], 1, 0, 'C')
                            pdf.cell(22, 6, p['risk'], 1, 0, 'C')
                            pdf.cell(65, 6, f' {p["gap"][:40]}', 1, 0, 'L')
                            pdf.ln()
                    
                    pdf.set_y(-15)
                    pdf.set_fill_color(26, 26, 26)
                    pdf.rect(0, pdf.get_y()-2, 297, 17, 'F')
                    pdf.set_font('Helvetica', 'I', 7)
                    pdf.set_text_color(180, 180, 180)
                    pdf.cell(0, 5, 'Churchgate Group - Confidential | hr@churchgate.com', ln=True, align='C')
                    
                    st.download_button("📥 Download PDF", bytes(pdf.output()), "aplayers_report.pdf", "application/pdf")
                    st.success("✅ PDF generated!")
                except Exception as e:
                    st.warning(f"PDF Error: {str(e)}")

def recruitment_hub():
    st.markdown("""<div class="churchgate-header"><h1>💼 Recruitment Hub</h1><p>Job Requisition | Auto-Posting | AI Screening | Interview Scheduler | Offer Letters | Background Checks | Onboarding</p></div>""", unsafe_allow_html=True)
    
    # ============================================================
    # AUTO-POST FUNCTION
    # ============================================================
    def auto_post_job(job_title, job_department, job_location, job_type, job_salary, job_jd, job_ref, public_url):
        """Auto-post job to LinkedIn, Indeed, and Glassdoor"""
        import re
        import urllib.parse
        
        clean_jd = re.sub(r'<[^>]+>', '', job_jd)
        results = {}
        
        # LINKEDIN - Auto-post to personal profile + Company page link
        import requests as req_linkedin
        
        linkedin_token = os.environ.get("LINKEDIN_ACCESS_TOKEN", st.secrets.get("LINKEDIN_ACCESS_TOKEN", ""))
        
        linkedin_post = f"""🚀 WE'RE HIRING: {job_title}

📍 {job_location} | 💼 {job_type} | 🏢 {job_department}

{clean_jd[:500]}...

APPLY NOW: {public_url}

#Hiring #Jobs #ChurchgateGroup #Careers #NigeriaJobs"""
        
        messages = []
        
        # 1. AUTO-POST TO PERSONAL PROFILE
        if linkedin_token:
            try:
                user_url = "https://api.linkedin.com/v2/userinfo"
                user_response = req_linkedin.get(user_url, headers={"Authorization": f"Bearer {linkedin_token}"})
                
                if user_response.status_code == 200:
                    user_data = user_response.json()
                    person_urn = f"urn:li:person:{user_data.get('sub')}"
                    
                    post_response = req_linkedin.post(
                        "https://api.linkedin.com/v2/ugcPosts",
                        headers={
                            "Authorization": f"Bearer {linkedin_token}",
                            "Content-Type": "application/json",
                            "X-Restli-Protocol-Version": "2.0.0"
                        },
                        json={
                            "author": person_urn,
                            "lifecycleState": "PUBLISHED",
                            "specificContent": {
                                "com.linkedin.ugc.ShareContent": {
                                    "shareCommentary": {"text": linkedin_post},
                                    "shareMediaCategory": "NONE"
                                }
                            },
                            "visibility": {"com.linkedin.ugc.MemberNetworkVisibility": "PUBLIC"}
                        }
                    )
                    
                    if post_response.status_code in [200, 201]:
                        messages.append('✅ Posted to personal LinkedIn')
                    else:
                        messages.append('⚠️ Personal LinkedIn: Click share link below')
                else:
                    messages.append('⚠️ LinkedIn auth issue')
            except:
                messages.append('⚠️ LinkedIn API error')
        else:
            messages.append('⚠️ No LinkedIn token configured')
        
        # 2. WTC ABUJA PAGE - One-click share link
        wtc_url = f"https://www.linkedin.com/shareArticle?mini=true&url={urllib.parse.quote(public_url)}&title={urllib.parse.quote('WE ARE HIRING: ' + job_title)}&summary={urllib.parse.quote(clean_jd[:200])}"
        messages.append(f'🏢 [Click to post on WTC Abuja Page]({wtc_url})')
        
        # 3. CHURCHGATE GROUP PAGE - One-click share link
        cg_url = f"https://www.linkedin.com/shareArticle?mini=true&url={urllib.parse.quote(public_url)}&title={urllib.parse.quote('WE ARE HIRING: ' + job_title)}&summary={urllib.parse.quote(clean_jd[:200])}"
        messages.append(f'🏢 [Click to post on Churchgate Group Page]({cg_url})')
        
        results['linkedin'] = {
            'status': 'success',
            'post_text': linkedin_post,
            'message': ' | '.join(messages)
        }
        
        # INDEED
        try:
            db_temp = DatabaseManager()
            db_temp.upload_file("job-feeds", f"indeed_{job_ref}.txt", clean_jd[:2500].encode('utf-8'), "text/plain")
            results['indeed'] = {'status': 'posted', 'message': '✅ Indeed feed saved! Upload at indeed.com/hire'}
        except:
            results['indeed'] = {'status': 'ready', 'message': '📋 Indeed content ready'}
        
        # GLASSDOOR
        glassdoor_url = f"https://www.glassdoor.com/employers/post-job/?title={urllib.parse.quote(job_title)}&location={urllib.parse.quote(job_location)}"
        results['glassdoor'] = {'status': 'ready', 'post_url': glassdoor_url, 'message': '🏢 Glassdoor post ready!'}
        
        return results
    
    # ============================================================
    
    user_role = st.session_state.user['role'] if st.session_state.user else 'Employee'
    
    # Auto-expire jobs past closing date
    try:
        today = datetime.now().strftime('%Y-%m-%d')
        for i, req in enumerate(st.session_state.job_requisitions):
            if req.get('status') == 'Approved - Live' and req.get('closing', '') and req.get('closing', '') < today:
                st.session_state.job_requisitions[i]['status'] = 'Expired'
                st.session_state.active_jobs = [j for j in st.session_state.active_jobs if j.get('ref') != req.get('id')]
        st.session_state.active_jobs = [j for j in st.session_state.active_jobs if j.get('closing', '9999-12-31') >= today]
    except:
        pass
    
    user_dept = st.session_state.user.get('department', '') if st.session_state.user else ''
    user_name = st.session_state.user['name'] if st.session_state.user else 'Staff'
    is_admin = user_role in ['Admin', 'HR Director'] or user_dept == 'Senior Management'
    is_manager = is_admin or user_role in ['Manager', 'HOD']
    
    try:
        employees_df = db.get_all_employees()
    except:
        employees_df = pd.DataFrame()
    
    STREAMLIT_URL = "https://churchgate-churchgate-hris.hf.space"
    
    # CHANGE 1: Load from database
    if 'job_requisitions' not in st.session_state:
        st.session_state.job_requisitions = []
        try:
            db_reqs = db.get_all_job_requisitions()
            for r in db_reqs:
                st.session_state.job_requisitions.append({
                    'id': r.get('req_id', ''), 'title': r.get('title', ''),
                    'department': r.get('department', ''), 'location': r.get('location', ''),
                    'type': r.get('job_type', ''), 'salary': r.get('salary', ''),
                    'level': r.get('level', ''), 'positions': r.get('positions', 1),
                    'closing': r.get('closing', ''), 'jd': r.get('jd', ''),
                    'screening': json.loads(r.get('screening', '[]')),
                    'posts': json.loads(r.get('posts', '{}')),
                    'status': r.get('status', ''), 'submitted_by': r.get('submitted_by', ''),
                    'date': r.get('date', ''), 'lm_comment': r.get('lm_comment', ''),
                    'admin_comment': r.get('admin_comment', ''), 'coo_comment': r.get('coo_comment', '')
                })
        except:
            pass
    
    if 'active_jobs' not in st.session_state:
        st.session_state.active_jobs = []
        try:
            all_reqs = db.get_all_job_requisitions()
            for r in all_reqs:
                if r.get('status') == 'Approved - Live':
                    req_id = r.get('req_id', '')
                    public_url = f"{STREAMLIT_URL}/Careers?job={req_id}"
                    st.session_state.active_jobs.append({
                        'ref': req_id,
                        'title': r.get('title', ''),
                        'department': r.get('department', ''),
                        'location': r.get('location', ''),
                        'type': r.get('job_type', ''),
                        'salary': r.get('salary', ''),
                        'jd': r.get('jd', ''),
                        'closing': r.get('closing', ''),
                        'screening': json.loads(r.get('screening', '[]')),
                        'posts': json.loads(r.get('posts', '{}')),
                        'date': r.get('date', ''),
                        'applications': 0,
                        'public_url': public_url
                    })
        except:
            pass
    if 'onboarding_list' not in st.session_state:
        st.session_state.onboarding_list = []
    if 'interviews_scheduled' not in st.session_state:
        st.session_state.interviews_scheduled = []
    if 'offer_letters' not in st.session_state:
        st.session_state.offer_letters = []
    if 'referrals' not in st.session_state:
        st.session_state.referrals = []
    
    tab1, tab2, tab3, tab4, tab5, tab6, tab7, tab8, tab9 = st.tabs([
        "📋 Job Requisition", "📢 Active Jobs", "🌐 Candidate Portal", 
        "🤖 AI Screening", "📅 Interviews", "📝 Offer Letters",
        "🎯 Onboarding", "🔍 Background Checks", "📊 Analytics"
    ])
    
    # ============ TAB 1: JOB REQUISITION ============
    with tab1:
        st.subheader("📋 Job Requisition & Approval Workflow")
        st.info("Workflow: Line Manager → Super Admin → COO → Job Goes LIVE on Careers Page")
        
        

        
        with st.form("job_requisition_form"):
            st.markdown("### New Job Requisition")
            c1, c2 = st.columns(2)
            with c1:
                job_title = st.text_input("Job Title *", placeholder="e.g., Senior Network Engineer")
                department = st.selectbox("Department *", ['Technology Group', 'Facility Management', 'Human Resources', 'Accounts & Finance', 'Sales & Marketing', 'Procurement', 'Security', 'Legal', 'Operations', 'Engineering', 'Central Stores', 'Project Development', 'Trade Services'])
                location = st.selectbox("Location", ["World Trade Center Abuja", "Churchgate Tower 1 Lagos", "Churchgate Tower 2 Lagos", "Churchgate Plaza Abuja", "Remote/Hybrid"])
                employment_type = st.selectbox("Employment Type", ["Full-time", "Contract", "Part-time", "Intern"])
            with c2:
                salary_range = st.text_input("Salary Range", placeholder="e.g., ₦5,000,000 - ₦8,000,000")
                experience_level = st.selectbox("Experience Level", ["Entry Level (0-2 yrs)", "Junior (2-4 yrs)", "Mid-Level (4-7 yrs)", "Senior (7-10 yrs)", "Executive (10+ yrs)"])
                positions = st.number_input("Number of Positions", min_value=1, value=1)
                closing_date = st.date_input("Application Deadline")
            
            st.markdown("---")
            st.markdown("### 📋 Full Job Description *")
            
            if 'jd_html_content' not in st.session_state:
                st.session_state.jd_html_content = ""
            
            # Rich text editor
            jd_text_for_submission = st_quill(
                value=st.session_state.jd_html_content,
                placeholder="Start typing your job description here...",
                html=True,
                key="jd_quill_editor"
            )
            
            if jd_text_for_submission:
                st.session_state.jd_html_content = jd_text_for_submission
                with st.expander("👁️ Live Preview", expanded=True):
                    st.markdown(jd_text_for_submission, unsafe_allow_html=True)
            else:
                st.info("👆 Use the rich text editor above to create your job description. Click a template to get started.")
            
            st.markdown("### Screening Questions (Optional)")
            st.markdown("*Leave blank if not needed*")
            screening_q1 = st.text_input("Screening Question 1", placeholder="e.g., Cisco Certified (CCNP minimum)")
            screening_q2 = st.text_input("Screening Question 2", placeholder="e.g., 5+ years experience in similar role")
            screening_q3 = st.text_input("Screening Question 3", placeholder="e.g., Experience with security systems")
            screening_q4 = st.text_input("Screening Question 4", placeholder="e.g., B.Sc. Computer Science or related field")
            
            st.markdown("### Auto-Post Settings")
            c1, c2, c3 = st.columns(3)
            with c1:
                post_linkedin = st.checkbox("Post to LinkedIn", value=True)
            with c2:
                post_indeed = st.checkbox("Post to Indeed", value=True)
            with c3:
                post_glassdoor = st.checkbox("Post to Glassdoor", value=True)
            
            submitted = st.form_submit_button("📤 Submit for Approval", use_container_width=True)
            
            if submitted:
                if job_title and department and jd_text_for_submission:
                    req = {
                        'id': f"REQ-{datetime.now().strftime('%Y%m%d%H%M')}",
                        'title': job_title, 'department': department, 'location': location,
                        'type': employment_type, 'salary': salary_range, 'level': experience_level,
                        'positions': positions, 'closing': closing_date.strftime('%Y-%m-%d'),
                        'jd': jd_text_for_submission, 'screening': [screening_q1, screening_q2, screening_q3, screening_q4],
                        'posts': {'linkedin': post_linkedin, 'indeed': post_indeed, 'glassdoor': post_glassdoor},
                        'status': 'Pending LM Approval',
                        'submitted_by': user_name, 'date': datetime.now().strftime('%Y-%m-%d %H:%M'),
                        'lm_comment': '', 'admin_comment': '', 'coo_comment': ''
                    }
                    st.session_state.job_requisitions.append(req)
                    # CHANGE 2: Save to database
                    try:
                        db.save_job_requisition(req['id'], req['title'], req['department'], req['location'],
                            req['type'], req['salary'], req['level'], req['positions'], req['closing'],
                            req['jd'], req['screening'], req['posts'], req['status'], req['submitted_by'], req['date'],
                            req['lm_comment'], req['admin_comment'], req['coo_comment'])
                    except:
                        pass
                    # Send email to Line Manager
                    lm_emails = {
                        'Accounts & Finance': 'jeff@churchgate.com',
                        'Facility Management': 'deffiong@churchgate.com',
                        'Human Resources': 'asakote@churchgate.com',
                        'Legal': 'daiyedun@churchgate.com',
                        'Operations': 'adeogun@churchgate.com',
                        'Procurement': 'abora@churchgate.com',
                        'Sales & Marketing': 'akarim@churchgate.com',
                        'Security': 'usabdullahi@churchgate.com',
                        'Technology Group': 'eetuk@churchgate.com',
                        'Engineering': 'purwar@churchgate.com',
                        'Central Stores': 'abora@churchgate.com',
                        'Project Development': 'deffiong@churchgate.com',
                        'Trade Services': 'akarim@churchgate.com'
                    }
                    lm_email = lm_emails.get(department, 'asakote@churchgate.com')
                    try:
                        from utils.email_service import EmailService
                        EmailService().send_email(
                            lm_email,
                            f"🔔 New Job Requisition Awaiting Authorization: {job_title}",
                            f"A new job requisition for '{job_title}' ({department}) has been submitted by {user_name}.\n\nPlease review and authorize in the HRIS: https://churchgate-churchgate-hris.hf.space\n\nRequisition ID: {req['id']}"
                        )
                        st.info(f"📧 Authorization request sent to Line Manager")
                    except:
                        pass
                    
                    st.success(f"✅ Job requisition {req['id']} submitted! Awaiting authorization.")
                    st.balloons()
                else:
                    st.error("❌ Required fields missing!")
        
        if st.session_state.job_requisitions:
            st.markdown("---")
            st.markdown("### 📋 Requisition Approval Dashboard")
            
            # Load email service for notifications
            try:
                from utils.email_service import EmailService
                email_svc = EmailService()
            except:
                email_svc = None
            
            for i, req in enumerate(st.session_state.job_requisitions):
                with st.expander(f"{req['id']} - {req['title']} | {req['department']} | {req['status']}", expanded=True):
                    
                    # ===== FULL DETAIL VIEW =====
                    st.markdown(f"**Submitted By:** {req['submitted_by']} | **Date:** {req['date']}")
                    st.markdown(f"**Department:** {req['department']} | **Location:** {req['location']}")
                    st.markdown(f"**Type:** {req['type']} | **Level:** {req['level']} | **Positions:** {req.get('positions', 1)}")
                    st.markdown(f"**Salary Range:** {req.get('salary', 'Not specified')}")
                    st.markdown(f"**Closing Date:** {req.get('closing', 'Not set')}")
                    
                    st.markdown("---")
                    with st.expander("📋 View Full Job Description", expanded=False):
                        jd_content = req.get('jd', 'No JD provided')
                        # Check if JD is HTML or plain text
                        if '<' in jd_content and '>' in jd_content:
                            st.markdown(jd_content, unsafe_allow_html=True)
                        else:
                            # Format plain text with proper paragraphs
                            formatted = jd_content.replace('\n\n', '<br><br>').replace('\n', '<br>')
                            st.markdown(formatted, unsafe_allow_html=True)
                    
                    if req.get('screening'):
                        st.markdown("---")
                        st.markdown("**❓ Screening Questions:**")
                        for q in req.get('screening', []):
                            if q:
                                st.markdown(f"- {q}")
                    
                    st.markdown("---")
                    st.markdown(f"**Platform Posts:** LinkedIn: {'✅' if req.get('posts', {}).get('linkedin') else '❌'} | Indeed: {'✅' if req.get('posts', {}).get('indeed') else '❌'} | Glassdoor: {'✅' if req.get('posts', {}).get('glassdoor') else '❌'}")
                    
                    if req.get('lm_comment'):
                        st.markdown(f"**LM Comment:** {req['lm_comment']}")
                    if req.get('admin_comment'):
                        st.markdown(f"**Admin Comment:** {req['admin_comment']}")
                    if req.get('coo_comment'):
                        st.markdown(f"**COO Comment:** {req['coo_comment']}")
                    
                    # ===== APPROVAL WORKFLOW WITH EDIT ACCESS =====
                    if is_admin or is_manager:
                        st.markdown("---")
                        
                        # LM APPROVAL
                        if req['status'] == 'Pending LM Approval':
                            st.markdown("#### 👔 Line Manager Review")
                            with st.form(key=f"lm_form_{i}"):
                                st.markdown("**✏️ Edit Job Description:**")
                                edit_jd = st_quill(value=jd_content, html=True, key=f"edit_jd_quill_{i}")
                                edit_screening = st.text_area("Edit Screening Questions (one per line)", value='\n'.join([q for q in req.get('screening', []) if q]), height=80, key=f"edit_screening_{i}")
                                lm_comment = st.text_area("Line Manager Comment *", key=f"lm_comment_{i}", placeholder="Reason for authorization or any notes...")
                                
                                if st.form_submit_button("✅ Authorize & Send to HR", use_container_width=True):
                                    if lm_comment:
                                        st.session_state.job_requisitions[i]['status'] = 'Pending Admin Approval'
                                        st.session_state.job_requisitions[i]['lm_comment'] = lm_comment
                                        st.session_state.job_requisitions[i]['jd'] = edit_jd
                                        st.session_state.job_requisitions[i]['screening'] = [s.strip() for s in edit_screening.split('\n') if s.strip()]
                                        
                                        try:
                                            r = st.session_state.job_requisitions[i]
                                            db.save_job_requisition(r['id'], r['title'], r['department'], r['location'],
                                                r['type'], r['salary'], r['level'], r['positions'], r['closing'],
                                                r['jd'], r['screening'], r['posts'], r['status'], r['submitted_by'], r['date'],
                                                r['lm_comment'], r['admin_comment'], r['coo_comment'])
                                        except:
                                            pass
                                        
                                        if email_svc:
                                            try:
                                                hr_emails = ["asakote@churchgate.com", "gbalogun@churchgate.com", "ichukwunonye@churchgate.com"]
                                                for hr_email in hr_emails:
                                                    email_svc.send_email(
                                                        hr_email,
                                                        f"🔔 Job Requisition Authorized by LM: {req['title']}",
                                                        f"Line Manager has authorized '{req['title']}' ({req['department']}).\n\nValidate at: https://churchgate-churchgate-hris.hf.space\n\nRequisition ID: {req['id']}"
                                                    )
                                                st.info("📧 HR team notified")
                                            except:
                                                pass
                                        
                                        st.success("✅ Authorized!")
                                        st.rerun()
                                    else:
                                        st.error("❌ Comment required!")
                        
                        # ===== ADMIN VALIDATION =====
                        if req['status'] == 'Pending Admin Approval':
                            st.markdown("#### 🔍 HR Admin Validation")
                            
                            # JD Preview
                            st.markdown("**📋 Job Description Preview:**")
                            jd_content = req.get('jd', 'No JD provided')
                            if '<' in jd_content and '>' in jd_content:
                                st.markdown(jd_content, unsafe_allow_html=True)
                            else:
                                st.markdown(jd_content)
                            st.markdown("---")
                            
                            with st.form(key=f"admin_form_{i}"):
                                st.markdown("**✏️ Edit Job Description:**")
                                edit_jd = st_quill(value=req.get('jd', ''), html=True, key=f"edit_jd_admin_quill_{i}")
                                edit_salary = st.text_input("Edit Salary Range", value=req.get('salary', ''), key=f"edit_salary_{i}")
                                admin_comment = st.text_area("Admin Validation Comment *", key=f"admin_comment_{i}", placeholder="Confirm JD quality, budget alignment, grade fit...")
                                
                                if st.form_submit_button("✅ Validate & Send to COO", use_container_width=True):
                                    if admin_comment:
                                        st.session_state.job_requisitions[i]['status'] = 'Pending COO Approval'
                                        st.session_state.job_requisitions[i]['admin_comment'] = admin_comment
                                        st.session_state.job_requisitions[i]['jd'] = edit_jd
                                        st.session_state.job_requisitions[i]['salary'] = edit_salary
                                        
                                        try:
                                            r = st.session_state.job_requisitions[i]
                                            db.save_job_requisition(r['id'], r['title'], r['department'], r['location'],
                                                r['type'], r['salary'], r['level'], r['positions'], r['closing'],
                                                r['jd'], r['screening'], r['posts'], r['status'], r['submitted_by'], r['date'],
                                                r['lm_comment'], r['admin_comment'], r['coo_comment'])
                                        except:
                                            pass
                                        
                                        if email_svc:
                                            try:
                                                email_svc.send_email(
                                                    "jeromedas@churchgate.com",
                                                    f"🔔 Job Requisition Ready for Final Approval: {req['title']}",
                                                    f"HR has validated '{req['title']}' ({req['department']}).\n\nApprove at: https://churchgate-churchgate-hris.hf.space\n\nRequisition ID: {req['id']}\nAdmin Comment: {admin_comment}"
                                                )
                                                st.info("📧 Email sent to COO (Jerome Das)")
                                            except:
                                                pass
                                        
                                        st.success("✅ Validated! Sent to COO for final approval.")
                                        st.rerun()
                                    else:
                                        st.error("❌ Comment required!")
                        
                        # COO APPROVAL
                        if req['status'] == 'Pending COO Approval':
                            st.markdown("#### 🏢 COO Final Approval")
                            with st.form(key=f"coo_form_{i}"):
                                st.markdown("**Review:** Full requisition details are shown above.")
                                coo_comment = st.text_area("COO Comment *", key=f"coo_comment_{i}", placeholder="Final approval notes, rejection reason, or revision request...")
                                
                                col_coo1, col_coo2, col_coo3 = st.columns(3)
                                with col_coo1:
                                    approve_btn = st.form_submit_button("✅ Approve & Activate", use_container_width=True)
                                with col_coo2:
                                    reject_btn = st.form_submit_button("❌ Reject", use_container_width=True)
                                with col_coo3:
                                    revise_btn = st.form_submit_button("🔄 Request Revision", use_container_width=True)
                                
                                if approve_btn:
                                    if coo_comment:
                                        st.session_state.job_requisitions[i]['status'] = 'Approved - Live'
                                        st.session_state.job_requisitions[i]['coo_comment'] = coo_comment
                                        job_ref = req['id']
                                        STREAMLIT_URL = "https://churchgate-churchgate-hris.hf.space"
                                        public_url = f"{STREAMLIT_URL}/Careers?job={job_ref}"
                                        
                                        st.session_state.active_jobs.append({
                                            'ref': job_ref, 'title': req['title'], 'department': req['department'],
                                            'location': req['location'], 'type': req['type'], 'salary': req['salary'],
                                            'jd': req['jd'], 'screening': req['screening'], 'closing': req['closing'],
                                            'posts': req['posts'], 'date': datetime.now().strftime('%Y-%m-%d'),
                                            'applications': 0, 'public_url': public_url
                                        })
                                        
                                        try:
                                            r = st.session_state.job_requisitions[i]
                                            db.save_job_requisition(r['id'], r['title'], r['department'], r['location'],
                                                r['type'], r['salary'], r['level'], r['positions'], r['closing'],
                                                r['jd'], r['screening'], r['posts'], r['status'], r['submitted_by'], r['date'],
                                                r['lm_comment'], r['admin_comment'], r['coo_comment'])
                                        except:
                                            pass
                                        
                                        if email_svc:
                                            try:
                                                submitter_email = req.get('submitted_by', '')
                                                if '@' in str(submitter_email):
                                                    email_svc.send_email(
                                                        submitter_email,
                                                        f"✅ Job Posting LIVE: {req['title']}",
                                                        f"The job requisition for '{req['title']}' has been fully approved and is now LIVE on the Careers Page.\n\nPublic URL: {public_url}\n\nShare this link with candidates!"
                                                    )
                                            except:
                                                pass
                                        
                                        # ===== AUTO-POST TO JOB BOARDS =====
                                        post_settings = req.get('posts', {})
                                        st.write(f"DEBUG posts: {post_settings}, type: {type(post_settings)}")
                                        posting_results = []
                                        
                                        if post_settings.get('linkedin'):
                                            result = auto_post_job(req['title'], req['department'], req['location'], req['type'], req.get('salary', ''), req['jd'], job_ref, public_url)
                                            posting_results.append(result['linkedin']['message'])
                                        
                                        if post_settings.get('indeed'):
                                            result = auto_post_job(req['title'], req['department'], req['location'], req['type'], req.get('salary', ''), req['jd'], job_ref, public_url)
                                            posting_results.append(result['indeed']['message'])
                                        
                                        if post_settings.get('glassdoor'):
                                            result = auto_post_job(req['title'], req['department'], req['location'], req['type'], req.get('salary', ''), req['jd'], job_ref, public_url)
                                            posting_results.append(result['glassdoor']['message'])
                                        
                                        if posting_results:
                                            for msg in posting_results:
                                                st.success(msg)
                                        # ===================================
                                        
                                        st.success(f"✅ Job is LIVE on Careers Page!")
                                        st.code(public_url, language=None)
                                        st.balloons()
                                        st.rerun()
                                    else:
                                        st.error("❌ Comment required!")
                                
                                if reject_btn:
                                    if coo_comment:
                                        st.session_state.job_requisitions[i]['status'] = 'Rejected by COO'
                                        st.session_state.job_requisitions[i]['coo_comment'] = coo_comment
                                        try:
                                            r = st.session_state.job_requisitions[i]
                                            db.save_job_requisition(r['id'], r['title'], r['department'], r['location'],
                                                r['type'], r['salary'], r['level'], r['positions'], r['closing'],
                                                r['jd'], r['screening'], r['posts'], r['status'], r['submitted_by'], r['date'],
                                                r['lm_comment'], r['admin_comment'], r['coo_comment'])
                                        except:
                                            pass
                                        if email_svc:
                                            try:
                                                submitter_email = req.get('submitted_by', '')
                                                if '@' in str(submitter_email):
                                                    email_svc.send_email(
                                                        submitter_email,
                                                        f"❌ Job Requisition Rejected: {req['title']}",
                                                        f"The job requisition for '{req['title']}' has been rejected by the COO.\n\nReason: {coo_comment}\n\nPlease review and resubmit if needed."
                                                    )
                                            except:
                                                pass
                                        st.error(f"❌ Requisition rejected. Reason: {coo_comment}")
                                        st.rerun()
                                    else:
                                        st.error("❌ Comment required for rejection!")
                                
                                if revise_btn:
                                    if coo_comment:
                                        st.session_state.job_requisitions[i]['status'] = 'Revision Requested by COO'
                                        st.session_state.job_requisitions[i]['coo_comment'] = coo_comment
                                        try:
                                            r = st.session_state.job_requisitions[i]
                                            db.save_job_requisition(r['id'], r['title'], r['department'], r['location'],
                                                r['type'], r['salary'], r['level'], r['positions'], r['closing'],
                                                r['jd'], r['screening'], r['posts'], r['status'], r['submitted_by'], r['date'],
                                                r['lm_comment'], r['admin_comment'], r['coo_comment'])
                                        except:
                                            pass
                                        if email_svc:
                                            try:
                                                submitter_email = req.get('submitted_by', '')
                                                if '@' in str(submitter_email):
                                                    email_svc.send_email(
                                                        submitter_email,
                                                        f"🔄 Revision Requested: {req['title']}",
                                                        f"The COO has requested revisions for '{req['title']}'.\n\nRevision Notes: {coo_comment}\n\nPlease update and resubmit for approval."
                                                    )
                                            except:
                                                pass
                                        st.warning(f"🔄 Revision requested. Reason: {coo_comment}")
                                        st.rerun()
                                    else:
                                        st.error("❌ Comment required for revision request!")
                        
                        # Show LIVE status
                        if req['status'] == 'Approved - Live':
                            st.success("🟢 LIVE — Accepting applications on Careers Page")
                            for job in st.session_state.active_jobs:
                                if job['title'] == req['title'] and job['department'] == req['department']:
                                    st.code(job['public_url'], language=None)
                                    break
    
    # ============ TAB 2: ACTIVE JOBS ============
    with tab2:
        st.subheader("📢 Active Job Postings")
        
        # Get expired jobs
        expired_jobs = [req for req in st.session_state.job_requisitions if req.get('status') == 'Expired']
        
        # Stats row
        c1, c2, c3, c4 = st.columns(4)
        with c1:
            st.metric("🟢 Live Jobs", len(st.session_state.active_jobs))
        with c2:
            st.metric("🔴 Expired", len(expired_jobs))
        with c3:
            total_apps = sum(job.get('applications', 0) for job in st.session_state.active_jobs)
            st.metric("👥 Total Applicants", total_apps)
        with c4:
            # Jobs closing within 7 days
            urgent = 0
            for job in st.session_state.active_jobs:
                try:
                    days_left = (datetime.strptime(job['closing'], '%Y-%m-%d') - datetime.now()).days
                    if days_left <= 7:
                        urgent += 1
                except:
                    pass
            st.metric("⏰ Closing Soon", urgent)
        
        st.markdown("---")
        
        # ===== LIVE JOBS =====
        if st.session_state.active_jobs:
            st.markdown("### 🟢 Live Jobs")
            for job in st.session_state.active_jobs:
                try:
                    days_left = (datetime.strptime(job['closing'], '%Y-%m-%d') - datetime.now()).days
                    if days_left <= 3:
                        status_icon = "🔴"
                        status_text = f"Closing in {days_left} days!"
                    elif days_left <= 7:
                        status_icon = "🟡"
                        status_text = f"{days_left} days remaining"
                    else:
                        status_icon = "🟢"
                        status_text = f"{days_left} days remaining"
                except:
                    status_icon = "🟢"
                    status_text = ""
                
                with st.expander(f"{status_icon} {job['ref']} - {job['title']} | {job['department']} | {job.get('applications', 0)} applicants | {status_text}", expanded=False):
                    col1, col2 = st.columns(2)
                    with col1:
                        st.markdown(f"**📍 Location:** {job['location']}")
                        st.markdown(f"**💼 Type:** {job['type']}")
                        st.markdown(f"**💰 Salary:** {job.get('salary', 'Not specified')}")
                        st.markdown(f"**📅 Closes:** {job['closing']} ({status_text})")
                    with col2:
                        st.markdown(f"**📎 Public URL:**")
                        st.code(job['public_url'], language=None)
                        st.markdown(f"**Platforms:** LinkedIn: {'✅' if job['posts'].get('linkedin') else '❌'} | Indeed: {'✅' if job['posts'].get('indeed') else '❌'} | Glassdoor: {'✅' if job['posts'].get('glassdoor') else '❌'}")
                    
                    st.markdown("---")
                    
                    # Quick actions
                    c1, c2, c3 = st.columns(3)
                    with c1:
                        share_url = job['public_url']
                        st.markdown(f"[🔗 LinkedIn Share](https://www.linkedin.com/sharing/share-offsite/?url={share_url})")
                        st.markdown(f"[💬 WhatsApp Share](https://wa.me/?text=Job:{job['title']}%20at%20Churchgate%20Group%20{share_url})")
                    with c2:
                        st.markdown(f"[📋 View JD]({share_url})")
                        st.markdown(f"📊 **Applications:** {job.get('applications', 0)}")
                    with c3:
                        if st.button(f"⏹️ Close Job Early", key=f"close_{job['ref']}"):
                            for i, req in enumerate(st.session_state.job_requisitions):
                                if req.get('id') == job['ref']:
                                    st.session_state.job_requisitions[i]['status'] = 'Expired'
                                    break
                            st.session_state.active_jobs = [j for j in st.session_state.active_jobs if j.get('ref') != job['ref']]
                            st.warning(f"⏹️ {job['title']} closed early!")
                            st.rerun()
        
        # ===== EXPIRED JOBS =====
        if expired_jobs:
            st.markdown("---")
            st.markdown("### 🔴 Expired Jobs")
            for req in expired_jobs:
                with st.expander(f"🔴 {req['id']} - {req['title']} | {req['department']} | Expired: {req.get('closing', 'N/A')}", expanded=False):
                    col1, col2 = st.columns(2)
                    with col1:
                        st.markdown(f"**📍 Location:** {req.get('location', 'N/A')}")
                        st.markdown(f"**💼 Type:** {req.get('type', 'N/A')}")
                        st.markdown(f"**📅 Closed:** {req.get('closing', 'N/A')}")
                    with col2:
                        st.markdown(f"**👤 Submitted By:** {req.get('submitted_by', 'N/A')}")
                        st.markdown(f"**📊 Positions:** {req.get('positions', 1)}")
                    
                    st.markdown("---")
                    c1, c2, c3 = st.columns(3)
                    with c1:
                        if st.button(f"📊 View Report", key=f"report_{req['id']}"):
                            try:
                                candidates = db.get_all_candidates()
                                job_candidates = candidates[candidates['job_id'] == req['id']] if not candidates.empty else []
                                st.metric("Total Applicants", len(job_candidates))
                                if len(job_candidates) > 0:
                                    st.dataframe(job_candidates[['first_name', 'last_name', 'email', 'status']], use_container_width=True)
                            except:
                                st.info("No application data available.")
                    with c2:
                        if st.button(f"🔄 Repost", key=f"repost_{req['id']}"):
                            st.session_state.job_requisitions[i]['status'] = 'Approved - Live'
                            st.session_state.job_requisitions[i]['closing'] = (datetime.now() + timedelta(days=30)).strftime('%Y-%m-%d')
                            st.success(f"🔄 {req['title']} reposted for 30 days!")
                            st.rerun()
                    with c3:
                        if st.button(f"🗑️ Delete", key=f"del_{req['id']}"):
                            st.session_state.job_requisitions = [r for r in st.session_state.job_requisitions if r['id'] != req['id']]
                            st.success("🗑️ Deleted permanently!")
                            st.rerun()
        
        if not st.session_state.active_jobs and not expired_jobs:
            st.info("No active or expired jobs. Submit a requisition in Tab 1.")
    
    # ============ TAB 3: CANDIDATE PORTAL ============
    with tab3:
        st.subheader("🌐 Candidate Management Portal")
        
        # Stats row
        try:
            candidates = db.get_all_candidates()
            total_candidates = len(candidates) if not candidates.empty else 0
            new_today = len(candidates[candidates['created_at'].str.contains(datetime.now().strftime('%Y-%m-%d'))]) if not candidates.empty and 'created_at' in candidates.columns else 0
            shortlisted = len(candidates[candidates['status'] == 'Shortlisted']) if not candidates.empty and 'status' in candidates.columns else 0
            interviewed = len(candidates[candidates['status'].str.contains('Interview', na=False)]) if not candidates.empty and 'status' in candidates.columns else 0
        except:
            total_candidates = 0
            new_today = 0
            shortlisted = 0
            interviewed = 0
        
        c1, c2, c3, c4 = st.columns(4)
        with c1:
            st.metric("👥 Total Candidates", total_candidates)
        with c2:
            st.metric("🆕 New Today", new_today)
        with c3:
            st.metric("⭐ Shortlisted", shortlisted)
        with c4:
            st.metric("🎯 Interviewed", interviewed)
        
        st.markdown("---")
        
        # Sub-tabs for candidate management
        cand_tab1, cand_tab2, cand_tab3 = st.tabs(["📋 All Candidates", "➕ Quick Add", "📱 Preview Career Page"])
        
        # ===== SUB-TAB 1: ALL CANDIDATES =====
        with cand_tab1:
            st.subheader("📋 Candidate Database")
            
            try:
                candidates = db.get_all_candidates()
                if not candidates.empty:
                    # Filters
                    col1, col2, col3, col4 = st.columns(4)
                    with col1:
                        status_filter = st.selectbox("Status", ["All", "New", "Shortlisted", "Interview Scheduled", "Offered", "Hired", "Rejected"], key="cand_status")
                    with col2:
                        job_filter = st.selectbox("Job", ["All Jobs"] + list(candidates['job_id'].dropna().unique()) if 'job_id' in candidates.columns else ["All Jobs"], key="cand_job")
                    with col3:
                        sort_by = st.selectbox("Sort", ["Newest First", "Name A-Z", "Status"], key="cand_sort")
                    with col4:
                        search_cand = st.text_input("🔍 Search", placeholder="Name or email...", key="cand_search")
                    
                    # Apply filters
                    display = candidates.copy()
                    if status_filter != "All":
                        display = display[display['status'] == status_filter]
                    if job_filter != "All Jobs" and 'job_id' in display.columns:
                        display = display[display['job_id'] == job_filter]
                    if search_cand:
                        s = search_cand.lower()
                        display = display[display['first_name'].str.lower().str.contains(s, na=False) | display['last_name'].str.lower().str.contains(s, na=False) | display['email'].str.lower().str.contains(s, na=False)]
                    
                    st.markdown(f"**Showing {len(display)} candidates**")
                    
                    for _, cand in display.iterrows():
                        initials = (str(cand.get('first_name', ''))[:1] + str(cand.get('last_name', ''))[:1]).upper()
                        status = cand.get('status', 'New')
                        status_color = {
                            'New': '#3182ce', 'Shortlisted': '#d69e2e', 'Interview Scheduled': '#38a169',
                            'Offered': '#805ad5', 'Hired': '#38a169', 'Rejected': '#CC0000'
                        }.get(status, '#a0aec0')
                        
                        with st.expander(f"👤 {cand.get('first_name', '')} {cand.get('last_name', '')} | {cand.get('email', '')} | {status}", expanded=False):
                            col1, col2 = st.columns([1, 3])
                            with col1:
                                st.markdown(f"""<div style="width:45px;height:45px;border-radius:50%;background:{status_color};display:flex;align-items:center;justify-content:center;font-weight:700;color:white;">{initials}</div>""", unsafe_allow_html=True)
                                st.markdown(f"<span style='background:{status_color};color:white;padding:0.1rem 0.5rem;border-radius:10px;font-size:0.7rem;'>{status}</span>", unsafe_allow_html=True)
                            with col2:
                                st.markdown(f"**📧** {cand.get('email', 'N/A')} | **📱** {cand.get('phone', 'N/A')}")
                                st.markdown(f"**💼** {cand.get('current_position', 'N/A')} | **📅** {cand.get('years_of_experience', 'N/A')} yrs")
                                st.markdown(f"**🔗** {cand.get('linkedin_url', 'N/A')[:40] if cand.get('linkedin_url') else 'N/A'}")
                            
                            # Quick actions
                            c1, c2, c3, c4 = st.columns(4)
                            with c1:
                                new_status = st.selectbox("Update Status", ["New", "Shortlisted", "Interview Scheduled", "Offered", "Hired", "Rejected"], key=f"status_{cand.get('candidate_ref')}")
                            with c2:
                                if st.button("💾 Update", key=f"upd_{cand.get('candidate_ref')}"):
                                    db._patch("candidates", {"status": new_status}, {"candidate_ref": cand.get('candidate_ref')})
                                    st.success("✅ Updated!")
                                    st.rerun()
                            with c3:
                                if cand.get('resume_text') and len(str(cand.get('resume_text'))) > 10:
                                    with st.expander("📄 View CV"):
                                        st.text_area("CV Content", str(cand.get('resume_text'))[:2000], height=200, key=f"cv_{cand.get('candidate_ref')}")
                            with c4:
                                if st.button("📧 Email", key=f"email_{cand.get('candidate_ref')}"):
                                    st.info(f"Email queued to {cand.get('email')}")
            except:
                st.info("No candidates yet. Share the Careers Page link to start receiving applications.")
        
        # ===== SUB-TAB 2: QUICK ADD CANDIDATE =====
        with cand_tab2:
            st.subheader("➕ Quick Add Candidate")
            st.info("Manually add a candidate from walk-in, referral, or external source.")
            
            with st.form("quick_add_candidate", clear_on_submit=True):
                c1, c2 = st.columns(2)
                with c1:
                    q_fn = st.text_input("First Name *")
                    q_ln = st.text_input("Last Name *")
                    q_em = st.text_input("Email *")
                    q_ph = st.text_input("Phone")
                with c2:
                    q_pos = st.text_input("Position Applied For")
                    q_src = st.selectbox("Source", ["Walk-in", "Employee Referral", "Career Page", "LinkedIn", "Indeed", "Glassdoor", "Agency", "Other"])
                    q_job = st.text_input("Job Reference (if any)")
                    q_linkedin = st.text_input("LinkedIn URL")
                
                st.markdown("---")
                q_cv = st.file_uploader("Upload CV *", type=['pdf', 'docx'])
                q_other = st.file_uploader("Upload Other Documents (Optional)", type=['pdf', 'docx', 'jpg', 'png'], accept_multiple_files=True)
                q_notes = st.text_area("Notes", height=80)
                
                if st.form_submit_button("➕ Add Candidate", use_container_width=True):
                    if q_fn and q_ln and q_em and q_cv:
                        tracking_id = f"CG-{datetime.now().strftime('%Y%m%d')}-{random.randint(1000,9999)}"
                        try:
                            cv_text = ""
                            file_ext = "pdf"
                            if q_cv.type == "application/pdf":
                                import PyPDF2
                                for page in PyPDF2.PdfReader(q_cv).pages:
                                    cv_text += page.extract_text() + "\n"
                            elif "word" in q_cv.type:
                                import docx
                                cv_text = "\n".join([p.text for p in docx.Document(q_cv).paragraphs])
                                file_ext = "docx"
                            
                            # Upload CV
                            cv_url = ""
                            try:
                                q_cv.seek(0)
                                cv_url = db.upload_file("cvs", f"{tracking_id}_{q_fn}_{q_ln}.{file_ext}", q_cv.read(), q_cv.type)
                            except:
                                pass
                            
                            db._post("candidates", {
                                "candidate_ref": tracking_id,
                                "first_name": q_fn,
                                "last_name": q_ln,
                                "email": q_em,
                                "phone": q_ph,
                                "linkedin_url": q_linkedin,
                                "current_position": q_pos,
                                "resume_filename": f"CV_{q_fn}_{q_ln}.{file_ext}",
                                "resume_text": cv_text[:10000],
                                "cv_url": cv_url,
                                "job_id": q_job,
                                "source": q_src,
                                "status": "New"
                            })
                            st.success(f"✅ Candidate {q_fn} {q_ln} added! (Ref: {tracking_id})")
                            st.balloons()
                        except Exception as e:
                            st.error(f"Error: {str(e)}")
                    else:
                        st.error("❌ Name, email, and CV are required!")
        
        # ===== SUB-TAB 3: PREVIEW CAREER PAGE =====
        with cand_tab3:
            st.subheader("📱 Career Page Preview")
            st.info("This is what candidates see on the live Careers Page.")
            
            # Quick preview of the actual career page
            st.markdown("""
            <div style="background:linear-gradient(135deg, #1a1a1a, #2d2a1f);padding:1.5rem;text-align:center;border-radius:8px;margin-bottom:1rem;">
                <h2 style="color:#F5E6CC;margin:0;">🚀 Build Your Career at Churchgate Group</h2>
                <p style="color:#c4b998;">Join a team of innovators, leaders, and changemakers.</p>
            </div>
            """, unsafe_allow_html=True)
            
            # Show active jobs preview
            if st.session_state.active_jobs:
                st.markdown(f"**🟢 {len(st.session_state.active_jobs)} Live Job(s) on Career Page**")
                for job in st.session_state.active_jobs[:3]:
                    st.markdown(f"""
                    <div style="background:white;padding:0.8rem;border-radius:6px;margin:0.3rem 0;border-left:4px solid #D4AF37;">
                        <strong>{job['title']}</strong><br>
                        <small>{job['department']} | {job['location']} | {job['type']}</small>
                    </div>
                    """, unsafe_allow_html=True)
            
            st.markdown("---")
            st.markdown(f"**🔗 Live Career Page:** [hris.churchgate.com/Careers](https://hris.churchgate.com/Careers)")
            st.markdown(f"**📊 Application Form Preview:**")
            
            with st.expander("Preview Application Form", expanded=False):
                st.text_input("First Name *", disabled=True, placeholder="Candidate fills this")
                st.text_input("Last Name *", disabled=True)
                st.text_input("Email *", disabled=True)
                st.text_input("Phone *", disabled=True)
                st.file_uploader("Upload CV *", disabled=True)
                st.file_uploader("Upload Other Documents (Optional)", type=['pdf', 'docx', 'jpg', 'png'], disabled=True, accept_multiple_files=True)
                st.text_area("Screening Questions", disabled=True, placeholder="Auto-generated per job role")
    
     # ============ TAB 4: AI SCREENING ============
    with tab4:
        st.subheader("🤖 AI-Powered Candidate Screening & Talent Intelligence")
        st.caption(f"🧠 AI Engine: {'OpenAI (95%+ confidence)' if ai_agent.use_openai else 'Enhanced Keyword (85%+ confidence)'}")
        
        try:
            candidates = db.get_all_candidates()
        except:
            candidates = pd.DataFrame()
        
        if not candidates.empty:
            total = len(candidates)
            screened = len(candidates[candidates['ai_score'] > 0]) if 'ai_score' in candidates.columns else 0
            tier1 = len(candidates[candidates['ai_tier'].str.contains('Tier 1', na=False)]) if 'ai_tier' in candidates.columns else 0
            tier2 = len(candidates[candidates['ai_tier'].str.contains('Tier 2', na=False)]) if 'ai_tier' in candidates.columns else 0
            avg_score = int(candidates[candidates['ai_score'] > 0]['ai_score'].mean()) if screened > 0 else 0
            
            c1, c2, c3, c4, c5, c6 = st.columns(6)
            c1.metric("📋 Total", total)
            c2.metric("🤖 Screened", screened)
            c3.metric("🌟 Tier 1", tier1)
            c4.metric("👍 Tier 2", tier2)
            c5.metric("📊 Avg Score", f"{avg_score}%")
            c6.metric("⏳ Pending", total - screened)
            
            st.markdown("---")
            st.markdown("### 🎯 Screening Controls")
            
            # Build job filter with real titles
            job_map = {}
            try:
                all_reqs = db.get_all_job_requisitions()
                for r in all_reqs:
                    req_id = r.get('req_id', '')
                    short_id = f"JOB-{req_id[-6:]}" if len(req_id) >= 6 else req_id
                    title = r.get('title', req_id)
                    job_map[req_id] = title
                    job_map[short_id] = title
            except:
                pass
            
            job_options = ["Select Job..."]
            if 'job_id' in candidates.columns:
                for jid in candidates['job_id'].dropna().unique():
                    title = job_map.get(jid, jid)
                    job_options.append(title)
            
            col1, col2, col3, col4 = st.columns([2, 1, 1, 1])
            with col1:
                screen_job = st.selectbox("📋 Screen by Job Posting", job_options, key="screen_job_t4")
            
            # Map selected title back to job_id
            selected_job_id = None
            if screen_job != "Select Job...":
                for jid in candidates['job_id'].dropna().unique():
                    if job_map.get(jid, jid) == screen_job:
                        selected_job_id = jid
                        break
            
            jd_text = ""
            job_title = ""
            if selected_job_id:
                try:
                    all_reqs = db.get_all_job_requisitions()
                    for r in all_reqs:
                        if r.get('req_id') == selected_job_id or f"JOB-{r.get('req_id', '')[-6:]}" == selected_job_id:
                            jd_text = r.get('jd', '')
                            job_title = r.get('title', '')
                            break
                except:
                    pass
            
            with col2:
                if screen_job != "Select Job...":
                    job_cands = candidates[candidates['job_id'] == screen_job]
                    pend = len(job_cands[job_cands['ai_score'] == 0]) if 'ai_score' in job_cands.columns else len(job_cands)
                    st.metric("Pending", pend)
            
            with col3:
                btn_job = st.button("🤖 Screen This Job", use_container_width=True, type="primary", disabled=(screen_job == "Select Job..."), key="btn_screen_job")
            
            with col4:
                btn_all = st.button("⚡ Screen All", use_container_width=True, key="btn_screen_all")
            
            # Process job screening
            if btn_job and screen_job != "Select Job...":
                target = candidates[candidates['job_id'] == selected_job_id]
                unscreened = target[target['ai_score'] == 0] if 'ai_score' in target.columns else target
                if len(unscreened) > 0:
                    with st.spinner(f"🤖 Analyzing {len(unscreened)} candidates for '{job_title}'..."):
                        count = 0
                        for _, row in unscreened.iterrows():
                            try:
                                cv_text = str(row.get('resume_text', ''))
                                if cv_text and cv_text != 'None' and len(cv_text) > 50:
                                    result = ai_agent.deep_analyze_candidate(cv_text, jd_text) if jd_text else ai_agent.score_candidate_advanced(cv_text, ai_agent.analyze_jd(cv_text[:500]))
                                    if isinstance(result, dict):
                                        db._patch("candidates", {"ai_score": int(result.get('overall_score', 0)), "ai_tier": result.get('tier', 'Pending')}, {"candidate_ref": row.get('candidate_ref', '')})
                                        count += 1
                            except:
                                pass
                        st.success(f"✅ {count} candidates scored!")
                        st.rerun()
            
            # Process all screening
            if btn_all:
                unscreened = candidates[candidates['ai_score'] == 0] if 'ai_score' in candidates.columns else candidates
                if len(unscreened) > 0:
                    with st.spinner(f"🤖 Screening {len(unscreened)} candidates..."):
                        count = 0
                        for _, row in unscreened.iterrows():
                            try:
                                cv_text = str(row.get('resume_text', ''))
                                if cv_text and cv_text != 'None' and len(cv_text) > 50:
                                    result = ai_agent.score_candidate_advanced(cv_text, ai_agent.analyze_jd(cv_text[:500]))
                                    if isinstance(result, dict):
                                        db._patch("candidates", {"ai_score": int(result.get('overall_score', 0)), "ai_tier": result.get('tier', 'Pending')}, {"candidate_ref": row.get('candidate_ref', '')})
                                        count += 1
                            except:
                                pass
                        st.success(f"✅ {count} candidates screened!")
                        st.rerun()
            
            # Candidate Display
            st.markdown("---")
            st.markdown("### 📊 Candidate Intelligence Dashboard")
            
            col1, col2, col3, col4 = st.columns(4)
            with col1:
                tier_filter = st.selectbox("🏷️ Tier", ["All", "Tier 1 ⭐", "Tier 2 👍", "Tier 3 🔶", "Tier 4 ❌", "Pending ⏳"], key="tier_filter_t4")
            with col2:
                dashboard_job_options = ["All Jobs"]
                if 'job_id' in candidates.columns:
                    for jid in candidates['job_id'].dropna().unique():
                        title = job_map.get(jid, jid)
                        dashboard_job_options.append(title)
                job_filter = st.selectbox("📋 Job", dashboard_job_options, key="job_filter_t4")
            with col3:
                sort_by = st.selectbox("🔢 Sort", ["Score (High-Low)", "Name (A-Z)"], key="sort_t4")
            with col4:
                search_term = st.text_input("🔍 Search", placeholder="Name or email...", key="search_t4")
            
            display_df = candidates.copy()
            if "Tier 1" in tier_filter:
                display_df = display_df[display_df['ai_tier'].str.contains('Tier 1', na=False)]
            elif "Tier 2" in tier_filter:
                display_df = display_df[display_df['ai_tier'].str.contains('Tier 2', na=False)]
            elif "Tier 3" in tier_filter:
                display_df = display_df[display_df['ai_tier'].str.contains('Tier 3', na=False)]
            elif "Tier 4" in tier_filter:
                display_df = display_df[display_df['ai_tier'].str.contains('Tier 4', na=False)]
            elif "Pending" in tier_filter:
                display_df = display_df[(display_df['ai_score'] == 0) | (display_df['ai_tier'].isna())]
            if job_filter != "All Jobs":
                selected_dashboard_job = None
                for jid in candidates['job_id'].dropna().unique():
                    if job_map.get(jid, jid) == job_filter:
                        selected_dashboard_job = jid
                        break
                if selected_dashboard_job:
                    display_df = display_df[display_df['job_id'] == selected_dashboard_job]
            if search_term:
                s = search_term.lower()
                display_df = display_df[display_df['first_name'].str.lower().str.contains(s, na=False) | display_df['last_name'].str.lower().str.contains(s, na=False) | display_df['email'].str.lower().str.contains(s, na=False)]
            if "Score" in sort_by:
                display_df = display_df.sort_values('ai_score', ascending=False)
            else:
                display_df = display_df.sort_values('first_name')
            
            st.markdown(f"**Showing {len(display_df)} candidates**")
            
            display_list = display_df.to_dict('records')
            for i, row in enumerate(display_list):
                first = str(row.get('first_name', ''))
                last = str(row.get('last_name', ''))
                email_val = str(row.get('email', ''))
                job_id = str(row.get('job_id', 'N/A'))
                score = int(float(row.get('ai_score', 0))) if row.get('ai_score') and float(row.get('ai_score', 0)) > 0 else 0
                tier = str(row.get('ai_tier', 'Pending'))
                cv_text = str(row.get('resume_text', ''))
                
                if score >= 85:
                    border, badge, emoji = "#38a169", "#38a169", "🌟"
                elif score >= 70:
                    border, badge, emoji = "#d69e2e", "#d69e2e", "👍"
                elif score > 0:
                    border, badge, emoji = "#dd6b20", "#dd6b20", "🔶"
                else:
                    border, badge, emoji = "#a0aec0", "#a0aec0", "⏳"
                
                initials = (first[:1] + last[:1]).upper() if first and last else "??"
                
                with st.expander(f"{emoji} {first} {last} — {job_id} — {score}%", expanded=(score >= 85)):
                    c1, c2 = st.columns([1, 4])
                    with c1:
                        st.markdown(f"""<div style="width:50px;height:50px;border-radius:50%;background:{border};display:flex;align-items:center;justify-content:center;font-weight:700;font-size:1.1rem;color:white;">{initials}</div>""", unsafe_allow_html=True)
                        if score > 0:
                            st.markdown(f"""<span style="background:{badge};color:white;padding:0.2rem 0.5rem;border-radius:12px;font-size:0.7rem;">{tier}</span>""", unsafe_allow_html=True)
                    with c2:
                        st.markdown(f"**📧** {email_val} | **💼** {str(row.get('current_position', ''))} | **📅** {str(row.get('years_of_experience', ''))} yrs")
                        if score > 0:
                            st.progress(score/100)
                    
                    col_a1, col_a2, col_a3 = st.columns(3)
                    with col_a1:
                        if st.button("🔍 Deep Analysis", key=f"deep_btn_{i}", use_container_width=True):
                            if cv_text and cv_text != 'None' and len(cv_text) > 50:
                                with st.spinner("Analyzing..."):
                                    job_jd = ""
                                    if job_id and job_id != 'None':
                                        try:
                                            all_reqs = db.get_all_job_requisitions()
                                            for r in all_reqs:
                                                if r.get('req_id') == job_id:
                                                    job_jd = r.get('jd', '')
                                                    break
                                        except:
                                            pass
                                    res = ai_agent.deep_analyze_candidate(cv_text, job_jd) if job_jd else ai_agent.score_candidate_advanced(cv_text, ai_agent.analyze_jd(cv_text[:500]))
                                    if isinstance(res, dict):
                                        st.session_state[f"deep_{i}"] = res
                                        time.sleep(0.3)
                                        st.rerun()
                    with col_a2:
                       with col_a2:
                        if st.button("📊 Quick Score", key=f"quick_btn_{i}", use_container_width=True):
                            if cv_text and cv_text != 'None' and len(cv_text) > 50:
                                res = ai_agent.score_candidate_advanced(cv_text, ai_agent.analyze_jd(cv_text[:500]))
                                if isinstance(res, dict):
                                    db._patch("candidates", {"ai_score": int(res.get('overall_score', 0)), "ai_tier": res.get('tier', 'Pending')}, {"candidate_ref": row.get('candidate_ref', '')})
                                    st.success(f"Scored: {int(res.get('overall_score', 0))}%")
                                    time.sleep(0.3)
                                    st.rerun()
                    
                    if f"deep_{i}" in st.session_state:
                        res = st.session_state[f"deep_{i}"]
                        if isinstance(res, dict):
                            st.markdown("---")
                            st.markdown("#### 🔬 Deep Analysis")
                            s1, s2, s3, s4 = st.columns(4)
                            s1.metric("Overall", f"{res.get('overall_score', 0)}%")
                            s2.metric("Skills", f"{res.get('skills_score', 0)}%")
                            s3.metric("Experience", f"{res.get('experience_score', 0)}%")
                            s4.metric("Confidence", f"{res.get('confidence', 0)}%")
                            
                            if res.get('verbatim_flags', 0) > 30:
                                st.warning(f"🚨 Verbatim risk: {res.get('verbatim_flags', 0):.0f}%")
                            
                            c1, c2 = st.columns(2)
                            with c1:
                                st.markdown("**✅ Strengths**")
                                for s in res.get('key_strengths', [])[:3]:
                                    st.markdown(f"- {s}")
                            with c2:
                                st.markdown("**⚠️ Gaps**")
                                for g in res.get('gaps_identified', [])[:3]:
                                    st.markdown(f"- {g}")
                            
                            st.markdown("**🎯 Interview Questions**")
                            for q_idx, q in enumerate(res.get('interview_questions', [])[:3]):
                                st.markdown(f"**{q_idx+1}.** {q}")
                            
                            if st.button("🗑️ Clear", key=f"clr_btn_{i}"):
                                del st.session_state[f"deep_{i}"]
                                st.rerun()
            
            # Analytics
            if screened > 0:
                st.markdown("---")
                st.markdown("### 📊 Analytics")
                col1, col2 = st.columns(2)
                with col1:
                    tier3 = len(candidates[candidates['ai_tier'].str.contains('Tier 3', na=False)]) if 'ai_tier' in candidates.columns else 0
                    tier4 = len(candidates[candidates['ai_tier'].str.contains('Tier 4', na=False)]) if 'ai_tier' in candidates.columns else 0
                    fig = px.pie(values=[tier1, tier2, tier3, tier4, total-screened], names=['T1⭐','T2👍','T3🔶','T4❌','Pending'], hole=0.5, color_discrete_sequence=['#38a169','#d69e2e','#dd6b20','#CC0000','#a0aec0'])
                    fig.update_layout(height=350)
                    st.plotly_chart(fig, use_container_width=True)
                with col2:
                    scores = candidates[candidates['ai_score'] > 0]['ai_score'].dropna() if 'ai_score' in candidates.columns else []
                    if len(scores) > 0:
                        fig2 = px.histogram(scores, nbins=10, color_discrete_sequence=['#CC0000'])
                        fig2.update_layout(height=350)
                        st.plotly_chart(fig2, use_container_width=True)
            
            st.markdown("---")
            st.download_button("📥 Export CSV", candidates[['candidate_ref','first_name','last_name','email','job_id','ai_score','ai_tier','status']].to_csv(index=False), "screening_report.csv", "text/csv")
            st.metric("⏱️ Est. Time-to-Hire", f"{max(7, 21 - screened)} days")
        
        else:
            st.info("🤖 No applications yet. Share the Careers Page URL to start building your talent pipeline.")
            st.code("https://churchgate-churchgate-hris.hf.space/Careers", language=None)
    
    # ============ TAB 5: INTERVIEWS ============
    with tab5:
        st.subheader("📅 Interview Scheduler")
        with st.form("schedule_interview"):
            c1, c2 = st.columns(2)
            with c1:
                candidate_name = st.text_input("Candidate Name *")
                interview_type = st.selectbox("Type", ["📞 Phone Screen", "💻 Technical", "👔 HR", "🏆 Final", "👥 Panel"])
                interviewer = st.text_input("Interviewer *")
            with c2:
                interview_date = st.date_input("Date *")
                interview_time = st.text_input("Time *", placeholder="e.g., 10:00 AM WAT")
                location = st.text_input("Location/Link", placeholder="Google Meet / Zoom / On-site")
            if st.form_submit_button("📅 Schedule Interview", use_container_width=True):
                if candidate_name and interviewer:
                    st.session_state.interviews_scheduled.append({
                        'candidate': candidate_name, 'type': interview_type,
                        'interviewer': interviewer, 'date': interview_date.strftime('%Y-%m-%d'),
                        'time': interview_time, 'location': location, 'status': 'Scheduled'
                    })
                    st.success(f"✅ Interview scheduled for {candidate_name}!")
                    st.balloons()
        
        if st.session_state.interviews_scheduled:
            st.markdown("---")
            st.markdown("### 📅 Upcoming Interviews")
            for iv in st.session_state.interviews_scheduled:
                st.markdown(f"""
                <div style="background: white; padding: 0.8rem; border-radius: 6px; margin-bottom: 0.3rem; border-left: 3px solid #3182ce;">
                    <strong>{iv['candidate']}</strong> - {iv['type']}<br>
                    <small>📅 {iv['date']} at {iv['time']} | 👤 {iv['interviewer']} | 📍 {iv['location']}</small>
                </div>
                """, unsafe_allow_html=True)
    
    # ============ TAB 6: OFFER LETTERS ============
    with tab6:
        st.subheader("📝 Enterprise Offer Letter Management")
        
        tab_offer1, tab_offer2 = st.tabs(["📝 Generate Offer", "📊 Offer Status Board"])
        
        # ===== SUB-TAB 1: GENERATE OFFER =====
        with tab_offer1:
            st.markdown("### 📝 Generate Professional Offer Letter")
            
            with st.form("offer_letter_form"):
                st.markdown("#### Candidate Information")
                c1, c2 = st.columns(2)
                with c1:
                    offer_name = st.text_input("Candidate Full Name *")
                    offer_email = st.text_input("Candidate Email *")
                    offer_position = st.text_input("Position *")
                    offer_dept = st.selectbox("Department *", ['Technology Group', 'Facility Management', 'Human Resources', 'Accounts & Finance', 'Sales & Marketing', 'Procurement', 'Security', 'Legal', 'Operations', 'Engineering'])
                with c2:
                    offer_salary = st.text_input("Salary Package *", placeholder="e.g., ₦1,200,000 Gross per Annum")
                    offer_start = st.date_input("Start Date *")
                    offer_reporting = st.text_input("Reports To *")
                    offer_probation = st.selectbox("Probation Period", ["3 months", "6 months"])
                    offer_expiry = st.date_input("Offer Expiry Date", value=datetime.now() + timedelta(days=14))
                
                st.markdown("---")
                st.markdown("#### Additional Terms")
                offer_benefits = st.text_area("Benefits Summary", placeholder="• HMO coverage\n• Pension plan\n• 20 days annual leave\n• Training support")
                offer_notes = st.text_area("Special Conditions / Notes")
                
                st.markdown("---")
                st.markdown("#### Offer Letter Preview")
                if offer_name and offer_position:
                    preview_html = f"""
                    <div style="background:white;padding:1.5rem;border-radius:8px;border:2px solid #CC0000;max-width:700px;margin:0 auto;">
                        <div style="text-align:center;border-bottom:2px solid #CC0000;padding-bottom:1rem;margin-bottom:1rem;">
                            <h2 style="color:#1a1a1a;margin:0;">CHURCHGATE GROUP</h2>
                            <p style="color:#CC0000;font-weight:600;">OFFER OF EMPLOYMENT</p>
                        </div>
                        <p>Dear <strong>{offer_name}</strong>,</p>
                        <p>We are pleased to offer you the position of <strong>{offer_position}</strong> in our <strong>{offer_dept}</strong> department.</p>
                        <table style="width:100%;margin:1rem 0;">
                            <tr><td><strong>Salary:</strong></td><td>{offer_salary}</td></tr>
                            <tr><td><strong>Start Date:</strong></td><td>{offer_start}</td></tr>
                            <tr><td><strong>Reports To:</strong></td><td>{offer_reporting}</td></tr>
                            <tr><td><strong>Probation:</strong></td><td>{offer_probation}</td></tr>
                            <tr><td><strong>Offer Expires:</strong></td><td>{offer_expiry}</td></tr>
                        </table>
                        <p style="text-align:center;color:#CC0000;font-weight:600;">Welcome to Churchgate Group!</p>
                    </div>
                    """
                    st.markdown(preview_html, unsafe_allow_html=True)
                
                if st.form_submit_button("📝 Generate Offer Letter & Send to Candidate", use_container_width=True):
                    if offer_name and offer_position and offer_email and offer_salary:
                        # Generate PDF
                        try:
                            import fpdf
                            FPDF = fpdf.FPDF
                            pdf = FPDF(orientation='P', unit='mm', format='A4')
                            pdf.add_page()
                            pdf.set_fill_color(26, 26, 26)
                            pdf.rect(0, 0, 210, 30, 'F')
                            pdf.set_fill_color(204, 0, 0)
                            pdf.rect(0, 30, 210, 3, 'F')
                            pdf.set_font('Helvetica', 'B', 20)
                            pdf.set_text_color(255, 255, 255)
                            pdf.cell(0, 16, 'CHURCHGATE GROUP', ln=True, align='C')
                            pdf.set_font('Helvetica', 'B', 11)
                            pdf.set_text_color(204, 0, 0)
                            pdf.cell(0, 8, 'OFFER OF EMPLOYMENT', ln=True, align='C')
                            pdf.ln(10)
                            pdf.set_font('Helvetica', '', 11)
                            pdf.set_text_color(26, 26, 26)
                            pdf.cell(0, 8, f'Date: {datetime.now().strftime("%B %d, %Y")}', ln=True)
                            pdf.ln(3)
                            pdf.cell(0, 8, f'Dear {offer_name},', ln=True)
                            pdf.ln(2)
                            pdf.multi_cell(0, 7, f'We are pleased to offer you the position of {offer_position} in the {offer_dept} department at Churchgate Group.')
                            pdf.ln(2)
                            pdf.set_font('Helvetica', 'B', 11)
                            pdf.cell(0, 8, 'Terms of Employment:', ln=True)
                            pdf.set_font('Helvetica', '', 11)
                            pdf.cell(0, 7, f'Position: {offer_position}', ln=True)
                            pdf.cell(0, 7, f'Department: {offer_dept}', ln=True)
                            pdf.cell(0, 7, f'Salary: {offer_salary}', ln=True)
                            pdf.cell(0, 7, f'Start Date: {offer_start}', ln=True)
                            pdf.cell(0, 7, f'Reports To: {offer_reporting}', ln=True)
                            pdf.cell(0, 7, f'Probation Period: {offer_probation}', ln=True)
                            pdf.cell(0, 7, f'Offer Expires: {offer_expiry}', ln=True)
                            if offer_benefits:
                                pdf.ln(2)
                                pdf.set_font('Helvetica', 'B', 11)
                                pdf.cell(0, 8, 'Benefits:', ln=True)
                                pdf.set_font('Helvetica', '', 10)
                                for line in offer_benefits.split('\n'):
                                    pdf.cell(0, 6, line.strip(), ln=True)
                            pdf.ln(5)
                            pdf.cell(0, 8, 'Please sign and return this letter to accept the offer.', ln=True)
                            pdf.cell(0, 8, f'This offer expires on {offer_expiry}.', ln=True)
                            pdf.set_y(-25)
                            pdf.set_font('Helvetica', 'I', 7)
                            pdf.set_text_color(150, 150, 150)
                            pdf.cell(0, 10, 'Churchgate Group - Official Offer Letter | hr@churchgate.com', align='C')
                            
                            pdf_bytes = bytes(pdf.output())
                            st.download_button("📥 Download Offer Letter PDF", pdf_bytes, f"Offer_{offer_name.replace(' ', '_')}.pdf", "application/pdf")
                        except:
                            pdf_bytes = None
                        
                        # Save to database
                        db._post("offer_letters", {
                            "candidate_name": offer_name,
                            "candidate_email": offer_email,
                            "position": offer_position,
                            "department": offer_dept,
                            "salary": offer_salary,
                            "start_date": offer_start.strftime('%Y-%m-%d'),
                            "reports_to": offer_reporting,
                            "probation": offer_probation,
                            "expiry_date": offer_expiry.strftime('%Y-%m-%d'),
                            "status": "Pending Acceptance",
                            "issued_by": user_name,
                            "notes": offer_notes
                        })
                        
                        # Send email to candidate
                        try:
                            from utils.email_service import EmailService
                            EmailService().send_email(
                                offer_email,
                                f"🎉 Job Offer: {offer_position} — Churchgate Group",
                                f"Dear {offer_name},\n\nCongratulations! We are pleased to offer you the position of {offer_position} at Churchgate Group.\n\nPlease find your offer letter attached. The offer expires on {offer_expiry}.\n\nTo accept, please reply to this email or contact HR.\n\nWe look forward to welcoming you to the team!\n\nChurchgate Group HR"
                            )
                            st.info(f"📧 Offer letter emailed to {offer_email}")
                        except:
                            pass
                        
                        st.success(f"✅ Offer generated for {offer_name}!")
                        st.balloons()
                        st.rerun()
                    else:
                        st.error("❌ Required fields missing!")
        
        # ===== SUB-TAB 2: OFFER STATUS BOARD =====
        with tab_offer2:
            st.markdown("### 📊 Offer Letter Status Board")
            
            try:
                offers = db._get("offer_letters")
            except:
                offers = []
            
            if offers:
                col1, col2 = st.columns(2)
                with col1:
                    offer_status_filter = st.selectbox("Status", ["All", "Pending Acceptance", "Accepted", "Rejected", "Expired"])
                with col2:
                    offer_dept_filter = st.selectbox("Department", ["All", 'Technology Group', 'Facility Management', 'Human Resources', 'Accounts & Finance', 'Sales & Marketing', 'Procurement', 'Security', 'Legal', 'Operations', 'Engineering'])
                
                filtered = offers
                if offer_status_filter != "All":
                    filtered = [o for o in filtered if o.get('status') == offer_status_filter]
                if offer_dept_filter != "All":
                    filtered = [o for o in filtered if o.get('department') == offer_dept_filter]
                
                # Stats
                total_offers = len(offers)
                accepted = len([o for o in offers if o.get('status') == 'Accepted'])
                pending = len([o for o in offers if o.get('status') == 'Pending Acceptance'])
                
                c1, c2, c3, c4 = st.columns(4)
                c1.metric("📋 Total Offers", total_offers)
                c2.metric("✅ Accepted", accepted)
                c3.metric("⏳ Pending", pending)
                c4.metric("📊 Acceptance Rate", f"{int(accepted/total_offers*100)}%" if total_offers > 0 else "N/A")
                
                st.markdown(f"**{len(filtered)} offers**")
                
                for offer in filtered:
                    status = offer.get('status', 'Pending Acceptance')
                    status_colors = {
                        'Pending Acceptance': '#d69e2e',
                        'Accepted': '#38a169',
                        'Rejected': '#CC0000',
                        'Expired': '#a0aec0'
                    }
                    color = status_colors.get(status, '#a0aec0')
                    
                    with st.expander(f"📝 {offer.get('candidate_name', 'Unknown')} — {offer.get('position', 'N/A')} | {status}"):
                        col1, col2 = st.columns([3, 1])
                        with col1:
                            st.markdown(f"**Department:** {offer.get('department', 'N/A')}")
                            st.markdown(f"**Salary:** {offer.get('salary', 'N/A')}")
                            st.markdown(f"**Start Date:** {offer.get('start_date', 'N/A')}")
                            st.markdown(f"**Reports To:** {offer.get('reports_to', 'N/A')}")
                            st.markdown(f"**Probation:** {offer.get('probation', 'N/A')}")
                            st.markdown(f"**Issued:** {offer.get('issued_date', '')[:10]}")
                            st.markdown(f"**Expires:** {offer.get('expiry_date', 'N/A')}")
                        
                        with col2:
                            st.markdown(f"<span style='background:{color};color:white;padding:0.3rem 0.8rem;border-radius:15px;'>{status}</span>", unsafe_allow_html=True)
                            
                            if is_admin:
                                new_status = st.selectbox("Update", ["Pending Acceptance", "Accepted", "Rejected", "Expired"], key=f"offer_status_{offer.get('id')}")
                                if st.button("💾 Update", key=f"offer_upd_{offer.get('id')}"):
                                    db._patch("offer_letters", {"status": new_status, "acceptance_date": datetime.now().strftime('%Y-%m-%d %H:%M') if new_status == 'Accepted' else None}, {"id": offer.get('id')})
                                    st.success("✅ Updated!")
                                    st.rerun()
            else:
                st.info("No offers generated yet.")
    
    # ============ TAB 7: ONBOARDING ============
    with tab7:
        st.subheader("🎯 Enterprise Onboarding Management")
        
        tab_onb1, tab_onb2, tab_onb3 = st.tabs(["➕ New Hire Setup", "📋 Onboarding Tracker", "📊 Dashboard"])
        
        # Predefined onboarding tasks
        default_tasks = [
            ("📧 Offer Letter Sent", "HR", "Send official offer letter to candidate"),
            ("✅ Offer Accepted", "HR", "Candidate signs and returns offer letter"),
            ("📄 Document Collection", "HR", "Collect ID, certificates, bank details, tax forms"),
            ("💻 IT Setup", "IT", "Create email, set up laptop, system access"),
            ("🏢 Workspace Assignment", "Facilities", "Assign desk/office, phone, stationery"),
            ("👤 Buddy Assignment", "HR", "Assign onboarding buddy from team"),
            ("📅 Orientation Scheduled", "HR", "Schedule HR orientation and office tour"),
            ("📚 Training Enrollment", "HR/L&D", "Enroll in required training courses"),
            ("🔑 Access Cards Issued", "Security", "Building access, parking, ID card"),
            ("🎉 First Day Welcome", "HR/Manager", "Team welcome, desk setup, first meeting"),
            ("📋 30-Day Check-in", "HR/Manager", "Schedule 30-day performance check-in"),
            ("📋 60-Day Check-in", "HR/Manager", "Schedule 60-day progress review"),
            ("📋 90-Day Review", "HR/Manager", "Formal 90-day probation review"),
        ]
        
        # ===== SUB-TAB 1: NEW HIRE SETUP =====
        with tab_onb1:
            st.subheader("➕ Set Up New Hire Onboarding")
            st.info("Create an onboarding plan for a new employee. Tasks are auto-assigned to relevant teams.")
            
            with st.form("add_onboarding_form"):
                c1, c2 = st.columns(2)
                with c1:
                    nh_name = st.text_input("Employee Full Name *")
                    nh_email = st.text_input("Employee Email *")
                    nh_dept = st.selectbox("Department *", ['Technology Group', 'Facility Management', 'Human Resources', 'Accounts & Finance', 'Sales & Marketing', 'Procurement', 'Security', 'Legal', 'Operations', 'Engineering'])
                    nh_position = st.text_input("Position *")
                with c2:
                    nh_start = st.date_input("Start Date *")
                    nh_buddy = st.selectbox("Assigned Buddy", ["Select buddy..."] + [f"{row['first_name']} {row['last_name']}" for _, row in employees_df.iterrows()] if not employees_df.empty else ["No employees available"])
                    nh_orientation = st.date_input("Orientation Date")
                    nh_notes = st.text_area("Special Instructions / Notes")
                
                if st.form_submit_button("🎯 Create Onboarding Plan", use_container_width=True):
                    if nh_name and nh_dept and nh_position and nh_email:
                        # Create onboarding record
                        result = db._post("onboarding", {
                            "employee_name": nh_name,
                            "employee_email": nh_email,
                            "department": nh_dept,
                            "position": nh_position,
                            "start_date": nh_start.strftime('%Y-%m-%d'),
                            "assigned_buddy": nh_buddy if nh_buddy and nh_buddy != "Select buddy..." else "",
                            "orientation_date": nh_orientation.strftime('%Y-%m-%d') if nh_orientation else "",
                            "status": "In Progress",
                            "progress": 0
                        })
                        
                        # Get the onboarding ID
                        onboard_data = db._get("onboarding")
                        if onboard_data:
                            onboard_id = onboard_data[-1].get('id')
                            
                            # Create default tasks
                            for task_name, assigned_to, task_desc in default_tasks:
                                db._post("onboarding_tasks", {
                                    "onboarding_id": onboard_id,
                                    "task_name": task_name,
                                    "task_category": assigned_to,
                                    "assigned_to": assigned_to,
                                    "status": "Pending",
                                    "notes": task_desc
                                })
                        
                        # Send welcome email to new hire
                        try:
                            from utils.email_service import EmailService
                            EmailService().send_email(
                                nh_email,
                                f"🎉 Welcome to Churchgate Group, {nh_name}!",
                                f"Dear {nh_name},\n\nWelcome to Churchgate Group! We're excited to have you join our {nh_dept} team as {nh_position}.\n\nYour start date is {nh_start}. Your onboarding buddy is {nh_buddy if nh_buddy != 'Select buddy...' else 'TBD'}.\n\nWe'll guide you through every step of your onboarding journey.\n\nSee you soon!\n\nChurchgate Group HR"
                            )
                            st.info(f"📧 Welcome email sent to {nh_email}")
                        except:
                            pass
                        
                        st.success(f"✅ Onboarding plan created for {nh_name} with {len(default_tasks)} tasks!")
                        st.balloons()
                        st.rerun()
                    else:
                        st.error("❌ Required fields missing!")
        
        # ===== SUB-TAB 2: ONBOARDING TRACKER =====
        with tab_onb2:
            st.subheader("📋 Onboarding Progress Tracker")
            
            try:
                onboard_list = db._get("onboarding")
                all_tasks = db._get("onboarding_tasks")
            except:
                onboard_list = []
                all_tasks = []
            
            if onboard_list:
                for onboard in onboard_list:
                    onboard_id = onboard.get('id')
                    employee_tasks = [t for t in all_tasks if t.get('onboarding_id') == onboard_id]
                    completed = len([t for t in employee_tasks if t.get('status') == 'Completed'])
                    total = len(employee_tasks)
                    progress = int(completed / total * 100) if total > 0 else 0
                    
                    progress_color = "#38a169" if progress >= 80 else "#d69e2e" if progress >= 50 else "#CC0000"
                    
                    with st.expander(f"🎯 {onboard.get('employee_name', 'Unknown')} — {onboard.get('position', 'N/A')} ({onboard.get('department', '')}) | {progress}% Complete"):
                        # Progress bar
                        st.markdown(f"""
                        <div style="margin-bottom:1rem;">
                            <div style="display:flex;justify-content:space-between;">
                                <span><strong>Onboarding Progress</strong></span>
                                <span style="color:{progress_color};font-weight:700;">{completed}/{total} tasks ({progress}%)</span>
                            </div>
                            <div style="background:#e0e0e0;height:8px;border-radius:4px;margin-top:0.3rem;">
                                <div style="background:{progress_color};width:{progress}%;height:8px;border-radius:4px;"></div>
                            </div>
                        </div>
                        """, unsafe_allow_html=True)
                        
                        col1, col2, col3 = st.columns(3)
                        with col1:
                            st.markdown(f"**Start Date:** {onboard.get('start_date', 'N/A')}")
                            st.markdown(f"**Buddy:** {onboard.get('assigned_buddy', 'Not assigned')}")
                        with col2:
                            st.markdown(f"**Orientation:** {onboard.get('orientation_date', 'Not scheduled')}")
                            st.markdown(f"**Department:** {onboard.get('department', 'N/A')}")
                        with col3:
                            st.markdown(f"**Status:** {onboard.get('status', 'In Progress')}")
                        
                        st.markdown("---")
                        
                        # Task list grouped by category
                        task_categories = {}
                        for t in employee_tasks:
                            cat = t.get('task_category', 'Other')
                            if cat not in task_categories:
                                task_categories[cat] = []
                            task_categories[cat].append(t)
                        
                        for cat, tasks in task_categories.items():
                            st.markdown(f"**{cat} Tasks**")
                            for task in tasks:
                                task_status = task.get('status', 'Pending')
                                icon = "✅" if task_status == 'Completed' else "⏳" if task_status == 'Pending' else "🔄"
                                
                                col1, col2, col3 = st.columns([3, 1, 1])
                                with col1:
                                    st.markdown(f"{icon} **{task.get('task_name', 'Task')}** — {task.get('notes', '')}")
                                with col2:
                                    new_task_status = st.selectbox("Status", ["Pending", "In Progress", "Completed"], 
                                        index=2 if task_status == 'Completed' else 1 if task_status == 'In Progress' else 0,
                                        key=f"task_status_{task.get('id')}")
                                with col3:
                                    if st.button("💾", key=f"task_save_{task.get('id')}"):
                                        db._patch("onboarding_tasks", {
                                            "status": new_task_status,
                                            "completed_date": datetime.now().strftime('%Y-%m-%d %H:%M') if new_task_status == 'Completed' else None
                                        }, {"id": task.get('id')})
                                        
                                        # Update progress
                                        updated_tasks = db._get("onboarding_tasks")
                                        emp_tasks = [t for t in updated_tasks if t.get('onboarding_id') == onboard_id]
                                        new_completed = len([t for t in emp_tasks if t.get('status') == 'Completed'])
                                        new_progress = int(new_completed / len(emp_tasks) * 100) if emp_tasks else 0
                                        db._patch("onboarding", {"progress": new_progress, "status": "Completed" if new_progress >= 100 else "In Progress"}, {"id": onboard_id})
                                        
                                        # Notify on completion of key tasks
                                        if new_task_status == 'Completed' and 'First Day' in task.get('task_name', ''):
                                            try:
                                                from utils.email_service import EmailService
                                                EmailService().send_email(
                                                    onboard.get('employee_email', ''),
                                                    f"🎉 Welcome to Your First Day, {onboard.get('employee_name', '')}!",
                                                    f"Dear {onboard.get('employee_name', '')},\n\nWelcome to your first day at Churchgate Group! We're thrilled to have you with us.\n\nYour buddy {onboard.get('assigned_buddy', 'your team')} will guide you through today.\n\nHave a great first day!\n\nChurchgate Group HR"
                                                )
                                            except:
                                                pass
                                        
                                        st.rerun()
                            st.markdown("---")
            else:
                st.info("No onboarding plans created yet. Set up a new hire in the 'New Hire Setup' tab.")
        
        # ===== SUB-TAB 3: DASHBOARD =====
        with tab_onb3:
            st.subheader("📊 Onboarding Dashboard")
            
            if onboard_list:
                total_onboard = len(onboard_list)
                in_progress = len([o for o in onboard_list if o.get('progress', 0) < 100])
                completed_onboard = len([o for o in onboard_list if o.get('progress', 0) >= 100])
                
                c1, c2, c3, c4 = st.columns(4)
                c1.metric("👥 Active Onboarding", total_onboard)
                c2.metric("🔄 In Progress", in_progress)
                c3.metric("✅ Completed", completed_onboard)
                c4.metric("📊 Avg Progress", f"{int(sum([o.get('progress', 0) for o in onboard_list]) / total_onboard)}%" if total_onboard > 0 else "0%")
                
                st.markdown("---")
                
                # Department breakdown
                dept_data = {}
                for o in onboard_list:
                    dept = o.get('department', 'Unknown')
                    dept_data[dept] = dept_data.get(dept, 0) + 1
                
                col1, col2 = st.columns(2)
                with col1:
                    fig = px.pie(values=list(dept_data.values()), names=list(dept_data.keys()), hole=0.5,
                               color_discrete_sequence=['#CC0000', '#d69e2e', '#3182ce', '#38a169', '#dd6b20', '#805ad5'])
                    fig.update_layout(height=350, title="Onboarding by Department")
                    st.plotly_chart(fig, use_container_width=True)
                with col2:
                    # Progress distribution
                    progress_data = [o.get('progress', 0) for o in onboard_list]
                    fig2 = px.histogram(progress_data, nbins=5, title="Progress Distribution",
                                      color_discrete_sequence=['#CC0000'])
                    fig2.update_layout(height=350)
                    st.plotly_chart(fig2, use_container_width=True)
                
                # Upcoming start dates
                st.markdown("---")
                st.markdown("### 📅 Upcoming Start Dates")
                upcoming = sorted([o for o in onboard_list if o.get('start_date', '') >= datetime.now().strftime('%Y-%m-%d')], key=lambda x: x.get('start_date', ''))
                for o in upcoming[:5]:
                    st.markdown(f"🎯 **{o.get('employee_name', '')}** — {o.get('position', '')} ({o.get('department', '')}) — Starts: {o.get('start_date', '')}")
                
                st.download_button("📥 Export Onboarding Report (CSV)",
                                  pd.DataFrame([{
                                      'Employee': o.get('employee_name'), 'Position': o.get('position'),
                                      'Department': o.get('department'), 'Start Date': o.get('start_date'),
                                      'Progress': f"{o.get('progress', 0)}%", 'Status': o.get('status')
                                  } for o in onboard_list]).to_csv(index=False),
                                  "onboarding_report.csv", "text/csv")
            else:
                st.info("No onboarding data yet.")
    
     # ============ TAB 8: BACKGROUND CHECKS ============
    with tab8:
        st.subheader("🔍 Enterprise Background Check Management")
        
        tab_bg1, tab_bg2, tab_bg3, tab_bg4 = st.tabs(["📋 Internal BG Check", "📤 External Verification", "📊 Status Dashboard", "📈 Analytics"])
        
        # ===== SUB-TAB 1: INTERNAL BG CHECK =====
        with tab_bg1:
            st.markdown("### 🔍 Internal Background Verification")
            st.info("Submit a background check request. HR team will be notified instantly via email and dashboard alert.")
            
            with st.form("bg_check_internal"):
                c1, c2 = st.columns(2)
                with c1:
                    bg_name = st.text_input("Candidate Name *")
                    bg_position = st.text_input("Position Applied For *")
                    bg_department = st.selectbox("Department", ['Technology Group', 'Facility Management', 'Human Resources', 'Accounts & Finance', 'Sales & Marketing', 'Procurement', 'Security', 'Legal', 'Operations', 'Engineering'])
                with c2:
                    bg_type = st.multiselect("Check Type *", [
                        "Employment Verification",
                        "Education Verification", 
                        "Criminal Record Check",
                        "Credit Check",
                        "Reference Check",
                        "Professional Certification",
                        "Address Verification",
                        "Identity Verification",
                        "Social Media Screening"
                    ])
                    bg_priority = st.selectbox("Priority", ["Standard (5 days)", "Urgent (3 days)", "Critical (24 hours)"])
                    bg_consent = st.checkbox("✅ Candidate has provided consent for background check", value=True)
                
                bg_notes = st.text_area("Additional Notes / Instructions")
                
                if st.form_submit_button("🔍 Submit & Notify HR Team", use_container_width=True):
                    if bg_name and bg_position and bg_type:
                        db._post("background_checks", {
                            "candidate_name": bg_name,
                            "position": bg_position,
                            "check_type": ', '.join(bg_type),
                            "priority": bg_priority,
                            "requested_by": user_name,
                            "check_type_category": "internal",
                            "status": "Pending",
                            "hr_notes": bg_notes
                        })
                        
                        # Notify HR team
                        try:
                            from utils.email_service import EmailService
                            hr_emails = ["asakote@churchgate.com", "gbalogun@churchgate.com", "ichukwunonye@churchgate.com"]
                            for hr_email in hr_emails:
                                EmailService().send_email(
                                    hr_email,
                                    f"🔍 New Background Check: {bg_name} — {bg_position}",
                                    f"A new background check has been requested for {bg_name} ({bg_position}, {bg_department}).\n\nCheck Type: {', '.join(bg_type)}\nPriority: {bg_priority}\n\nPlease process in the HRIS: https://churchgate-churchgate-hris.hf.space"
                                )
                            st.info("📧 HR team notified via email")
                        except:
                            pass
                        
                        st.success(f"✅ Background check submitted for {bg_name}! HR team notified.")
                        st.balloons()
                        st.rerun()
                    else:
                        st.error("❌ Candidate name, position, and at least one check type are required!")
        
        # ===== SUB-TAB 2: EXTERNAL VERIFICATION =====
        with tab_bg2:
            st.markdown("### 📤 External Reference Verification")
            st.info("Send a secure verification request to an external referee. They'll receive an email with instructions to respond.")
            
            # Email templates
            with st.expander("📋 Email Templates", expanded=False):
                template_type = st.selectbox("Select Template", [
                    "Employment Verification",
                    "Character Reference",
                    "Academic Verification",
                    "Custom Message"
                ])
                
                templates = {
                    "Employment Verification": "Please verify the candidate's employment history including dates of employment, position held, responsibilities, and reason for leaving.",
                    "Character Reference": "Please provide your assessment of the candidate's character, integrity, work ethic, and suitability for the position.",
                    "Academic Verification": "Please verify the candidate's academic credentials including degree obtained, dates of attendance, and graduation status.",
                    "Custom Message": "Please provide your honest assessment based on your professional experience with the candidate."
                }
                st.code(templates.get(template_type, ""))
            
            with st.form("bg_check_external"):
                c1, c2 = st.columns(2)
                with c1:
                    ext_candidate = st.text_input("Candidate Name *")
                    ext_position = st.text_input("Position Applied For *")
                    ext_ref_name = st.text_input("Referee Full Name *")
                    ext_ref_title = st.text_input("Referee Title/Position")
                with c2:
                    ext_ref_email = st.text_input("Referee Email *")
                    ext_ref_company = st.text_input("Referee Company/Organization")
                    ext_ref_phone = st.text_input("Referee Phone (Optional)")
                    ext_priority = st.selectbox("Priority", ["Standard (5 days)", "Urgent (3 days)"])
                
                ext_check_type = st.multiselect("Verification Type *", [
                    "Employment History Verification",
                    "Character Reference",
                    "Academic Verification",
                    "Professional Reference",
                    "Salary Verification",
                    "Performance Assessment"
                ])
                
                ext_message = st.text_area("Custom Message to Referee", 
                    value=templates.get(template_type, ""),
                    height=100)
                
                if st.form_submit_button("📤 Send to Referee & Notify HR", use_container_width=True):
                    if ext_candidate and ext_ref_name and ext_ref_email and ext_check_type:
                        ref_id = f"REF-{datetime.now().strftime('%Y%m%d%H%M')}-{random.randint(100,999)}"
                        
                        db._post("background_checks", {
                            "candidate_name": ext_candidate,
                            "position": ext_position,
                            "check_type": ', '.join(ext_check_type),
                            "priority": ext_priority,
                            "requested_by": user_name,
                            "check_type_category": "external",
                            "external_email": ext_ref_email,
                            "external_name": ext_ref_name,
                            "status": "Awaiting External Response",
                            "hr_notes": f"Ref ID: {ref_id} | {ext_message}"
                        })
                        
                        # Send email to referee
                        try:
                            from utils.email_service import EmailService
                            ref_subject = f"Background Verification Request: {ext_candidate} — Churchgate Group (Ref: {ref_id})"
                            ref_body = f"""
                            <h2>Background Verification Request</h2>
                            <p>Dear {ext_ref_name},</p>
                            <p><strong>{ext_candidate}</strong> has applied for the position of <strong>{ext_position}</strong> at Churchgate Group and has listed you as a referee.</p>
                            <div class="info-box">
                                <p><strong>Candidate:</strong> {ext_candidate}</p>
                                <p><strong>Position:</strong> {ext_position}</p>
                                <p><strong>Verification Type:</strong> {', '.join(ext_check_type)}</p>
                                <p><strong>Reference ID:</strong> {ref_id}</p>
                            </div>
                            <p><strong>Verification Request:</strong></p>
                            <p style="background:#f5f5f5;padding:1rem;border-radius:6px;">{ext_message}</p>
                            <p>Please reply to this email with your detailed response. Include the Reference ID ({ref_id}) in your reply.</p>
                            <p>Alternatively, you can submit your response directly at:</p>
                            <p><a href="https://churchgate-churchgate-hris.hf.space/Careers?ref={ref_id}" style="color:#CC0000;">Submit Verification Response</a></p>
                            <p>Thank you for your time and honest assessment.</p>
                            <p style="color:#888;">Churchgate Group HR Team<br>hr@churchgate.com</p>
                            """
                            EmailService().send_email(ext_ref_email, ref_subject, ref_body)
                            st.success(f"✅ Verification request sent to {ext_ref_name} ({ext_ref_email})!")
                        except:
                            st.success(f"✅ Request saved! Email queued for {ext_ref_email}")
                        
                        st.balloons()
                        st.rerun()
                    else:
                        st.error("❌ All fields marked with * are required!")
        
        # ===== SUB-TAB 3: STATUS DASHBOARD =====
        with tab_bg3:
            st.markdown("### 📊 Background Check Status Board")
            
            try:
                bg_checks = db._get("background_checks")
            except:
                bg_checks = []
            
            if bg_checks:
                col1, col2, col3 = st.columns(3)
                with col1:
                    bg_status_filter = st.selectbox("Status", ["All", "Pending", "Awaiting External Response", "Responded", "Completed", "Rejected"])
                with col2:
                    bg_type_filter = st.selectbox("Type", ["All", "internal", "external"])
                with col3:
                    bg_priority_filter = st.selectbox("Priority", ["All", "Standard", "Urgent", "Critical"])
                
                filtered = bg_checks
                if bg_status_filter != "All":
                    filtered = [b for b in filtered if b.get('status') == bg_status_filter]
                if bg_type_filter != "All":
                    filtered = [b for b in filtered if b.get('check_type_category') == bg_type_filter]
                if bg_priority_filter != "All":
                    filtered = [b for b in filtered if bg_priority_filter.lower() in str(b.get('priority', '')).lower()]
                
                # SLA tracking
                pending_count = len([b for b in filtered if b.get('status') == 'Pending'])
                awaiting_count = len([b for b in filtered if b.get('status') == 'Awaiting External Response'])
                completed_count = len([b for b in filtered if b.get('status') == 'Completed'])
                
                c1, c2, c3, c4 = st.columns(4)
                c1.metric("📋 Total", len(filtered))
                c2.metric("⏳ Pending", pending_count)
                c3.metric("📤 Awaiting Response", awaiting_count)
                c4.metric("✅ Completed", completed_count)
                
                st.markdown(f"**{len(filtered)} background check(s)**")
                
                for bg in filtered:
                    status = bg.get('status', 'Pending')
                    is_external = bg.get('check_type_category') == 'external'
                    icon = "📤" if is_external else "🔍"
                    priority = bg.get('priority', 'Standard')
                    
                    status_colors = {
                        'Pending': '#d69e2e',
                        'Awaiting External Response': '#3182ce',
                        'Responded': '#38a169',
                        'Completed': '#38a169',
                        'Rejected': '#CC0000'
                    }
                    color = status_colors.get(status, '#a0aec0')
                    
                    # SLA Timer
                    try:
                        request_date = datetime.strptime(str(bg.get('request_date', ''))[:10], '%Y-%m-%d')
                        days_since = (datetime.now() - request_date).days
                        sla_color = "#38a169" if days_since <= 3 else "#d69e2e" if days_since <= 5 else "#CC0000"
                    except:
                        days_since = 0
                        sla_color = "#a0aec0"
                    
                    with st.expander(f"{icon} {bg.get('candidate_name', 'Unknown')} — {bg.get('position', 'N/A')} | {status} | ⏱️ {days_since}d"):
                        col1, col2 = st.columns([3, 1])
                        with col1:
                            st.markdown(f"**Check Type:** {bg.get('check_type', 'N/A')}")
                            st.markdown(f"**Priority:** {priority}")
                            st.markdown(f"**Requested By:** {bg.get('requested_by', 'N/A')}")
                            st.markdown(f"**Date:** {bg.get('request_date', '')[:10]}")
                            st.markdown(f"**⏱️ Days Open:** <span style='color:{sla_color};font-weight:700;'>{days_since} days</span>", unsafe_allow_html=True)
                            
                            if is_external:
                                st.markdown(f"**Referee:** {bg.get('external_name', 'N/A')} ({bg.get('external_email', 'N/A')})")
                                if bg.get('external_response'):
                                    st.success(f"**Response Received:** {bg.get('external_response', '')}")
                        
                        with col2:
                            st.markdown(f"<span style='background:{color};color:white;padding:0.3rem 0.8rem;border-radius:15px;font-size:0.85rem;'>{status}</span>", unsafe_allow_html=True)
                            
                            if is_admin:
                                new_status = st.selectbox("Update", ["Pending", "Awaiting External Response", "Responded", "Completed", "Rejected"], key=f"bg_status_{bg.get('id')}")
                                if st.button("💾 Update", key=f"bg_upd_{bg.get('id')}", use_container_width=True):
                                    db._patch("background_checks", {"status": new_status}, {"id": bg.get('id')})
                                    st.success("✅ Updated!")
                                    st.rerun()
                                
                                if is_external and status == 'Awaiting External Response':
                                    if st.button("📧 Resend", key=f"bg_resend_{bg.get('id')}", use_container_width=True):
                                        try:
                                            from utils.email_service import EmailService
                                            EmailService().send_email(
                                                bg.get('external_email', ''),
                                                f"Reminder: Background Check — {bg.get('candidate_name', '')}",
                                                f"Dear {bg.get('external_name', '')},\n\nGentle reminder to please respond regarding the verification for {bg.get('candidate_name', '')}.\n\nThank you.\nChurchgate Group HR"
                                            )
                                            st.success("✅ Reminder sent!")
                                        except:
                                            st.info("Reminder queued")
            else:
                st.info("No background check requests yet.")
        
        # ===== SUB-TAB 4: ANALYTICS =====
        with tab_bg4:
            st.markdown("### 📈 Background Check Analytics")
            
            if bg_checks:
                total = len(bg_checks)
                completed = len([b for b in bg_checks if b.get('status') == 'Completed'])
                pending = len([b for b in bg_checks if b.get('status') == 'Pending'])
                external = len([b for b in bg_checks if b.get('check_type_category') == 'external'])
                
                c1, c2, c3, c4 = st.columns(4)
                c1.metric("📋 Total", total)
                c2.metric("✅ Completed", completed)
                c3.metric("⏳ Pending", pending)
                c4.metric("📤 External", external)
                
                col1, col2 = st.columns(2)
                with col1:
                    status_data = {}
                    for b in bg_checks:
                        s = b.get('status', 'Unknown')
                        status_data[s] = status_data.get(s, 0) + 1
                    fig = px.pie(values=list(status_data.values()), names=list(status_data.keys()), hole=0.5,
                               color_discrete_sequence=['#38a169', '#d69e2e', '#3182ce', '#CC0000', '#a0aec0'])
                    fig.update_layout(height=350, title="Status Distribution")
                    st.plotly_chart(fig, use_container_width=True)
                with col2:
                    check_types = {}
                    for b in bg_checks:
                        types = str(b.get('check_type', '')).split(', ')
                        for t in types:
                            if t:
                                check_types[t] = check_types.get(t, 0) + 1
                    if check_types:
                        fig2 = px.bar(x=list(check_types.keys()), y=list(check_types.values()),
                                    color=list(check_types.values()), color_continuous_scale=['#CC0000', '#d69e2e', '#38a169'])
                        fig2.update_layout(height=350, title="Check Types Requested")
                        st.plotly_chart(fig2, use_container_width=True)
                
                st.download_button("📥 Export BG Check Report (CSV)", 
                                  pd.DataFrame([{
                                      'Candidate': b.get('candidate_name'), 'Position': b.get('position'),
                                      'Type': b.get('check_type'), 'Status': b.get('status'),
                                      'Priority': b.get('priority'), 'Requested': b.get('request_date', '')[:10]
                                  } for b in bg_checks]).to_csv(index=False),
                                  "background_checks.csv", "text/csv")
            else:
                st.info("Analytics will appear once background checks are submitted.")
    
    # ============ TAB 9: ANALYTICS ============
    with tab9:
        st.subheader("📊 Recruitment Analytics")
        c1, c2, c3, c4 = st.columns(4)
        c1.metric("Active Jobs", len(st.session_state.active_jobs))
        c2.metric("Interviews", len(st.session_state.interviews_scheduled))
        c3.metric("Offers Sent", len(st.session_state.offer_letters))
        c4.metric("Onboarding", len(st.session_state.onboarding_list))
        
        try:
            candidates = db.get_all_candidates()
            total_apps = len(candidates) if not candidates.empty else 0
        except:
            total_apps = 0
        
        funnel = pd.DataFrame({'Stage': ['Applied', 'Screened', 'Interviewed', 'Offered', 'Hired'], 'Count': [total_apps, int(total_apps*0.6), int(total_apps*0.25), int(total_apps*0.1), int(total_apps*0.05)]})
        fig = px.funnel(funnel, x='Count', y='Stage', color_discrete_sequence=['#CC0000'])
        fig.update_layout(height=350)
        st.plotly_chart(fig, use_container_width=True)
        
        st.markdown("---")
        st.markdown("### 🌍 Diversity & Inclusion")
        div_data = pd.DataFrame({'Category': ['Male', 'Female', 'Other'], 'Applicants': [28, 17, 2]})
        fig2 = px.pie(div_data, values='Applicants', names='Category', hole=0.5, color_discrete_sequence=['#3182ce', '#CC0000', '#718096'])
        st.plotly_chart(fig2, use_container_width=True)
        
        st.markdown("---")
        st.markdown("### 🤝 Employee Referral Program")
        with st.form("referral"):
            c1, c2 = st.columns(2)
            with c1:
                ref_name = st.text_input("Referral Name *")
                ref_position = st.text_input("Position Referred For *")
                ref_employee = st.text_input("Your Name *")
            with c2:
                ref_relation = st.selectbox("Relationship", ["Former Colleague", "Friend", "Family", "Professional Network"])
                ref_date = st.date_input("Referral Date")
            if st.form_submit_button("🤝 Submit Referral", use_container_width=True):
                if ref_name and ref_employee:
                    st.session_state.referrals.append({'name': ref_name, 'position': ref_position, 'referrer': ref_employee, 'date': ref_date.strftime('%Y-%m-%d'), 'status': 'Pending'})
                    st.success(f"✅ Referral submitted! Bonus eligible upon successful hire.")
                    st.balloons()

def ai_recruitment_agent():
    st.markdown("""<div class="churchgate-header"><h1>🤖 AI Recruitment Agent</h1><p>AI-Powered CV-JD Matching | Verbatim Detection | Inconsistency Flags | Skills Gap Matrix | Bias Detection | Interview Generator | Executive Reports</p></div>""", unsafe_allow_html=True)
    
    options = ["📥 Load Applications", "📋 JD Analysis", "📤 CV Upload & Scoring", "📊 Candidate Tiering", "🔍 Deep Analysis", "📄 Executive Report", "🔗 LinkedIn Parse", "💾 Save Results", "💬 AI Assistant"]
    
    if 'ai_section' not in st.session_state:
        st.session_state.ai_section = "📥 Load Applications"
    
    default_index = 0
    try:
        default_index = options.index(st.session_state.ai_section)
    except:
        pass
    
    ai_section = st.radio("Select Function:", options, index=default_index, horizontal=True)
    st.session_state.ai_section = ai_section
    
    # ============ LOAD APPLICATIONS ============
    if ai_section == "📥 Load Applications":
        st.subheader("🚀 Recruitment Pipeline & Bulk Screening")
        
        try:
            candidates = db.get_all_candidates()
            if not candidates.empty:
                # Job filter
                job_map = {}
                try:
                    all_reqs = db.get_all_job_requisitions()
                    for r in all_reqs:
                        req_id = r.get('req_id', '')
                        short_id = f"JOB-{req_id[-6:]}" if len(req_id) >= 6 else req_id
                        title = r.get('title', req_id)
                        job_map[req_id] = title
                        job_map[short_id] = title
                except:
                    pass
                
                job_options = ["All Jobs"]
                if 'job_id' in candidates.columns:
                    for jid in candidates['job_id'].dropna().unique():
                        title = job_map.get(jid, jid)
                        job_options.append(title)
                
                pipeline_job_filter = st.selectbox("📋 Filter by Job", job_options, key="pipeline_job_filter")
                
                # Map title back to job_id
                selected_job_id = None
                if pipeline_job_filter != "All Jobs":
                    for jid in candidates['job_id'].dropna().unique():
                        if job_map.get(jid, jid) == pipeline_job_filter:
                            selected_job_id = jid
                            break
                    filtered = candidates[candidates['job_id'] == selected_job_id]
                else:
                    filtered = candidates
                
                st.markdown(f"### 📊 {len(filtered)} Candidates")
                
                # ===== PIPELINE STAGES BAR =====
                st.markdown("---")
                st.markdown("### 🔄 Recruitment Pipeline")
                
                pipeline_stats, pipeline_data = get_pipeline_stats(selected_job_id)
                
                if pipeline_stats:
                    stages = list(pipeline_stats.keys())
                    counts = list(pipeline_stats.values())
                    total = sum(counts)
                    
                    # Visual pipeline
                    cols = st.columns(len(stages))
                    for i, (stage, count) in enumerate(pipeline_stats.items()):
                        with cols[i]:
                            color = "#38a169" if stage in ['Hired'] else "#3182ce" if stage in ['Offer Sent', 'Interview Scheduled'] else "#d69e2e" if stage in ['Shortlisted', 'Manager Review'] else "#a0aec0"
                            st.markdown(f"""
                            <div style="background:white;padding:0.5rem;border-radius:6px;text-align:center;border-top:3px solid {color};">
                                <small style="color:#888;">{stage}</small><br>
                                <strong style="font-size:1.2rem;">{count}</strong>
                            </div>
                            """, unsafe_allow_html=True)
                
                # ===== BULK SCREENING =====
                st.markdown("---")
                st.markdown("### 🤖 Step 1: Bulk AI Screening")
                
                unscreened = filtered[filtered['ai_score'] == 0] if 'ai_score' in filtered.columns else filtered
                
                col1, col2, col3 = st.columns([2, 1, 1])
                with col1:
                    st.metric("Pending Screening", len(unscreened))
                with col2:
                    if st.button("🤖 Screen All", use_container_width=True, type="primary", disabled=len(unscreened)==0):
                        with st.spinner(f"🤖 Deep-analyzing {len(unscreened)} candidates..."):
                            screened = 0
                            for _, row in unscreened.iterrows():
                                try:
                                    cv_text = str(row.get('resume_text', ''))
                                    if cv_text and cv_text != 'None' and len(cv_text) > 50:
                                        job_id_val = str(row.get('job_id', ''))
                                        jd_text = ""
                                        if job_id_val and job_id_val != 'None':
                                            for r in all_reqs:
                                                if r.get('req_id') == job_id_val:
                                                    jd_text = r.get('jd', '')
                                                    break
                                        result = ai_agent.deep_analyze_candidate(cv_text, jd_text) if jd_text else ai_agent.score_candidate_advanced(cv_text, ai_agent.analyze_jd(cv_text[:500]))
                                        if isinstance(result, dict):
                                            score = int(result.get('overall_score', 0))
                                            tier = result.get('tier', 'Pending')
                                            db._patch("candidates", {"ai_score": score, "ai_tier": tier}, {"candidate_ref": row.get('candidate_ref', '')})
                                            # Add to pipeline
                                            try:
                                                db._post("recruitment_pipeline", {
                                                    "candidate_ref": row.get('candidate_ref', ''),
                                                    "candidate_name": f"{row.get('first_name','')} {row.get('last_name','')}",
                                                    "candidate_email": row.get('email', ''),
                                                    "job_id": job_id_val,
                                                    "job_title": job_map.get(job_id_val, ''),
                                                    "current_stage": "AI Screened",
                                                    "ai_score": score,
                                                    "ai_tier": tier,
                                                    "updated_by": user_name
                                                })
                                            except:
                                                pass
                                            screened += 1
                                except:
                                    pass
                            st.success(f"✅ {screened} candidates screened!")
                            st.rerun()
                with col3:
                    if st.button("📊 Quick Score", use_container_width=True, disabled=len(unscreened)==0):
                        with st.spinner("Scoring..."):
                            for _, row in unscreened.iterrows():
                                try:
                                    cv_text = str(row.get('resume_text', ''))
                                    if cv_text and len(cv_text) > 50:
                                        result = ai_agent.score_candidate_advanced(cv_text, ai_agent.analyze_jd(cv_text[:500]))
                                        if isinstance(result, dict):
                                            db._patch("candidates", {"ai_score": int(result.get('overall_score', 0)), "ai_tier": result.get('tier', 'Pending')}, {"candidate_ref": row.get('candidate_ref', '')})
                                except:
                                    pass
                            st.success("✅ Quick scores applied!")
                            st.rerun()
                
                # ===== TIERING REPORT =====
                st.markdown("---")
                st.markdown("### 📊 Step 2: Tiering Report & Manager Review")
                
                screened_cands = filtered[filtered['ai_score'] > 0].sort_values('ai_score', ascending=False)
                
                if len(screened_cands) > 0:
                    t1 = len(screened_cands[screened_cands['ai_score'] >= 85])
                    t2 = len(screened_cands[(screened_cands['ai_score'] >= 70) & (screened_cands['ai_score'] < 85)])
                    
                    c1, c2, c3, c4 = st.columns(4)
                    c1.metric("🌟 Tier 1", t1)
                    c2.metric("👍 Tier 2", t2)
                    c3.metric("📊 Avg Score", f"{int(screened_cands['ai_score'].mean())}%")
                    c4.metric("👥 Total Screened", len(screened_cands))
                    
                    # Transparent Scorecard Table
                    st.markdown("#### 📊 Transparent Scoring Criteria")
                    st.info("""
                    **How scores are calculated:**
                    - 🎯 Skills Match (30%) | 💼 Experience (25%) | 🎓 Education (10%) | 📜 Certifications (5%)
                    - 📄 CV Quality (10%) | ✍️ Communication (10%) | 🔑 Keywords (5%) | 🚨 Verbatim Penalty (-5%)
                    - 🔗 LinkedIn Presence (+5%) | 📝 Cover Letter (+5%)
                    """)
                    
                    # Build detailed scorecard
                    scorecard_data = []
                    for _, row in screened_cands.iterrows():
                        cv_text = str(row.get('resume_text', ''))
                        # Get detailed scores from AI agent
                        try:
                            job_jd = ""
                            if selected_job_id:
                                for r in all_reqs:
                                    if r.get('req_id') == selected_job_id:
                                        job_jd = r.get('jd', '')
                                        break
                            jd_analysis = ai_agent.analyze_jd(job_jd) if job_jd else ai_agent.analyze_jd(cv_text[:500])
                            detailed = ai_agent.score_candidate_advanced(cv_text, jd_analysis)
                        except:
                            detailed = {}
                        
                        scorecard_data.append({
                            'Rank': len(scorecard_data) + 1,
                            'Candidate': f"{row.get('first_name','')} {row.get('last_name','')}",
                            'Overall': int(row.get('ai_score', 0)),
                            'Tier': row.get('ai_tier', 'Pending'),
                            'Skills': int(detailed.get('skills_score', 0)),
                            'Experience': int(detailed.get('experience_score', 0)),
                            'Education': int(detailed.get('education_score', 0)),
                            'CV Quality': int(detailed.get('cv_quality_score', detailed.get('soft_skills_score', 0))),
                            'Verbatim': f"{int(detailed.get('verbatim_flags', 0))}%",
                            'Confidence': f"{int(detailed.get('confidence', 0))}%",
                            'Strengths': ', '.join(detailed.get('key_strengths', [])[:2]),
                            'Gaps': ', '.join(detailed.get('gaps_identified', [])[:2]),
                        })
                    
                    if scorecard_data:
                        scorecard_df = pd.DataFrame(scorecard_data)
                        st.dataframe(scorecard_df, use_container_width=True, hide_index=True)
                        
                        # Detailed breakdown for each candidate
                        st.markdown("---")
                        st.markdown("### 🔍 Individual Score Breakdowns")
                        for i, data in enumerate(scorecard_data):
                            with st.expander(f"📊 {data['Candidate']} — {data['Overall']}% — {data['Tier']}"):
                                col1, col2, col3, col4 = st.columns(4)
                                col1.metric("🎯 Skills", f"{data['Skills']}%")
                                col2.metric("💼 Experience", f"{data['Experience']}%")
                                col3.metric("🎓 Education", f"{data['Education']}%")
                                col4.metric("📄 CV Quality", f"{data['CV Quality']}%")
                                
                                st.markdown(f"**🚨 Verbatim Risk:** {data['Verbatim']}")
                                st.markdown(f"**🤖 AI Confidence:** {data['Confidence']}")
                                st.markdown(f"**✅ Strengths:** {data['Strengths']}")
                                st.markdown(f"**⚠️ Gaps:** {data['Gaps']}")
                                
                                # Visual score bar
                                st.progress(data['Overall']/100)
                    else:
                        st.info("Scorecard data will appear after screening.")
                    
                    # Send to Manager
                    col1, col2 = st.columns(2)
                    with col1:
                        manager_email = st.text_input("Hiring Manager Email", value="asakote@churchgate.com", key="manager_email")
                    with col2:
                        if st.button("📧 Send Report to Manager", use_container_width=True, type="primary"):
                            # Build email
                            email_body = f"<h2>AI Screening Report</h2><p>Job: {pipeline_job_filter}</p><table border='1' cellpadding='5'><tr><th>Rank</th><th>Candidate</th><th>Score</th><th>Tier</th><th>Email</th></tr>"
                            for i, (_, row) in enumerate(screened_cands.iterrows()):
                                email_body += f"<tr><td>{i+1}</td><td>{row.get('first_name','')} {row.get('last_name','')}</td><td>{int(row.get('ai_score',0))}%</td><td>{row.get('ai_tier','')}</td><td>{row.get('email','')}</td></tr>"
                            email_body += "</table><p>Please review and shortlist candidates for interview.</p>"
                            
                            try:
                                from utils.email_service import EmailService
                                EmailService().send_email(manager_email, f"🔍 AI Screening Report — {pipeline_job_filter}", email_body)
                                st.success(f"✅ Report sent to {manager_email}!")
                            except:
                                st.success("✅ Report queued for delivery.")
                    
                    # Download report
                    st.download_button("📥 Download Tiering Report (CSV)", 
                                      screened_cands[['first_name', 'last_name', 'email', 'ai_score', 'ai_tier']].to_csv(index=False),
                                      f"tiering_report_{pipeline_job_filter}.csv", "text/csv")
                
                # ===== SHORTLISTING =====
                st.markdown("---")
                st.markdown("### ✅ Step 3: Shortlist Candidates")
                
                if len(screened_cands) > 0:
                    shortlist_candidates = st.multiselect(
                        "Select candidates to shortlist",
                        [f"{row['first_name']} {row['last_name']} — {int(row['ai_score'])}%" for _, row in screened_cands.iterrows()],
                        key="shortlist_select"
                    )
                    
                    col1, col2 = st.columns(2)
                    with col1:
                        if st.button("✅ Shortlist Selected", use_container_width=True, disabled=len(shortlist_candidates)==0):
                            for name_score in shortlist_candidates:
                                name_part = name_score.split(" — ")[0]
                                for _, row in screened_cands.iterrows():
                                    if f"{row['first_name']} {row['last_name']}" == name_part:
                                        db._patch("candidates", {"status": "Shortlisted"}, {"candidate_ref": row.get('candidate_ref', '')})
                                        try:
                                            db._post("recruitment_pipeline", {
                                                "candidate_ref": row.get('candidate_ref', ''),
                                                "candidate_name": f"{row.get('first_name','')} {row.get('last_name','')}",
                                                "candidate_email": row.get('email', ''),
                                                "job_id": row.get('job_id', ''),
                                                "job_title": job_map.get(row.get('job_id', ''), ''),
                                                "current_stage": "Shortlisted",
                                                "ai_score": int(row.get('ai_score', 0)),
                                                "ai_tier": row.get('ai_tier', ''),
                                                "updated_by": user_name
                                            })
                                        except:
                                            pass
                            st.success(f"✅ {len(shortlist_candidates)} candidates shortlisted!")
                            st.rerun()
                    with col2:
                        if st.button("📧 Send Interview Invites", use_container_width=True, disabled=len(shortlist_candidates)==0):
                            st.session_state['shortlist_bulk'] = shortlist_candidates
                            st.rerun()
                
                # ===== BULK INTERVIEW INVITES =====
                if 'shortlist_bulk' in st.session_state and st.session_state['shortlist_bulk']:
                    st.markdown("---")
                    st.markdown("### 📅 Step 4: Schedule Interviews")
                    
                    interview_date = st.date_input("Interview Date")
                    interview_type = st.selectbox("Interview Type", ["📞 Phone Screen", "💻 Technical", "👔 HR", "🏆 Final", "👥 Panel"])
                    email_template = st.text_area("Email Message", 
                        value="Dear [Candidate],\n\nCongratulations! We're impressed with your profile and would like to invite you for an interview.\n\nDate: [Date]\nType: [Type]\n\nPlease confirm your availability.\n\nBest regards,\nChurchgate Group HR",
                        height=120)
                    
                    if st.button("📧 Send to All Shortlisted", use_container_width=True, type="primary"):
                        sent = 0
                        for name_score in st.session_state['shortlist_bulk']:
                            name_part = name_score.split(" — ")[0]
                            for _, row in screened_cands.iterrows():
                                if f"{row['first_name']} {row['last_name']}" == name_part:
                                    try:
                                        from utils.email_service import EmailService
                                        body = email_template.replace('[Candidate]', name_part).replace('[Date]', str(interview_date)).replace('[Type]', interview_type)
                                        EmailService().send_email(row.get('email', ''), f"📅 Interview Invitation — Churchgate Group", body)
                                        db._patch("candidates", {"status": "Interview Invited"}, {"candidate_ref": row.get('candidate_ref', '')})
                                        try:
                                            db._post("recruitment_pipeline", {
                                                "candidate_ref": row.get('candidate_ref', ''),
                                                "candidate_name": name_part,
                                                "candidate_email": row.get('email', ''),
                                                "job_id": row.get('job_id', ''),
                                                "job_title": job_map.get(row.get('job_id', ''), ''),
                                                "current_stage": "Interview Scheduled",
                                                "interview_date": interview_date.strftime('%Y-%m-%d'),
                                                "interview_type": interview_type,
                                                "updated_by": user_name
                                            })
                                        except:
                                            pass
                                        sent += 1
                                    except:
                                        pass
                        st.success(f"✅ Interview invitations sent to {sent} candidates!")
                        st.balloons()
                        del st.session_state['shortlist_bulk']
                        st.rerun()
                
                # ===== CANDIDATE DETAILS WITH CVs =====
                st.markdown("---")
                st.markdown("### 📋 All Applications with CVs")
                
                for idx, row in filtered.iterrows():
                    first = str(row.get('first_name', ''))
                    last = str(row.get('last_name', ''))
                    email_val = str(row.get('email', ''))
                    job_id_val = str(row.get('job_id', 'N/A'))
                    score = int(row.get('ai_score', 0)) if row.get('ai_score') and float(row.get('ai_score', 0)) > 0 else 0
                    tier = str(row.get('ai_tier', 'Pending'))
                    cv_text = str(row.get('resume_text', ''))
                    phone_val = str(row.get('phone', ''))
                    linkedin_val = str(row.get('linkedin_url', ''))
                    
                    emoji = "🌟" if score >= 85 else "👍" if score >= 70 else "🔶" if score > 0 else "⏳"
                    
                    with st.expander(f"{emoji} {first} {last} — {job_id_val} — {score}% — {tier}", expanded=(score >= 85)):
                        col1, col2, col3 = st.columns([1, 2, 1])
                        
                        with col1:
                            initials = (first[:1] + last[:1]).upper() if first and last else "??"
                            st.markdown(f"""<div style="width:50px;height:50px;border-radius:50%;background:#CC0000;display:flex;align-items:center;justify-content:center;font-weight:700;color:white;">{initials}</div>""", unsafe_allow_html=True)
                            if score > 0:
                                st.metric("Score", f"{score}%")
                        
                        with col2:
                            st.markdown(f"**📧** {email_val} | **📱** {phone_val}")
                            if linkedin_val and linkedin_val != 'None':
                                st.markdown(f"**🔗** [{linkedin_val[:30]}...]({linkedin_val})")
                            st.markdown(f"**💼** {str(row.get('current_position', ''))} | **📅** {str(row.get('years_of_experience', ''))} yrs")
                            if score > 0:
                                st.progress(score/100)
                        
                        with col3:
                            if cv_text and cv_text != 'None' and len(cv_text) > 10:
                                with st.expander("📄 View CV", expanded=False):
                                    st.text_area("CV Content", cv_text, height=200, key=f"cv_pipeline_{idx}")
                                    st.download_button("📥 Download CV Text", cv_text, f"CV_{first}_{last}.txt", "text/plain", key=f"dl_cv_{idx}")
                                    # Original file download
                                    cv_url = str(row.get('cv_url', ''))
                                    resume_filename = str(row.get('resume_filename', ''))
                                    if cv_url and cv_url != 'None' and cv_url != '':
                                        st.markdown(f"📎 [Download Original File: {resume_filename}]({cv_url})")
                            if st.button("🔍 Deep Analysis", key=f"deep_pipe_{idx}", use_container_width=True):
                                if cv_text and len(cv_text) > 50:
                                    with st.spinner("Analyzing..."):
                                        job_jd = ""
                                        if job_id_val and job_id_val != 'None':
                                            for r in all_reqs:
                                                if r.get('req_id') == job_id_val:
                                                    job_jd = r.get('jd', '')
                                                    break
                                        res = ai_agent.deep_analyze_candidate(cv_text, job_jd) if job_jd else ai_agent.score_candidate_advanced(cv_text, ai_agent.analyze_jd(cv_text[:500]))
                                        if isinstance(res, dict):
                                            st.session_state[f"pipe_analysis_{idx}"] = res
                                            st.rerun()
                        
                        # Show analysis
                        if f"pipe_analysis_{idx}" in st.session_state:
                            res = st.session_state[f"pipe_analysis_{idx}"]
                            st.markdown("---")
                            st.markdown("#### 🔬 Deep Analysis")
                            s1, s2, s3, s4 = st.columns(4)
                            s1.metric("Overall", f"{res.get('overall_score', 0)}%")
                            s2.metric("Skills", f"{res.get('skills_score', 0)}%")
                            s3.metric("Experience", f"{res.get('experience_score', 0)}%")
                            s4.metric("Confidence", f"{res.get('confidence', 0)}%")
                            
                            if res.get('verbatim_flags', 0) > 30:
                                st.warning(f"🚨 Verbatim risk: {res.get('verbatim_flags', 0):.0f}%")
                            
                            c1, c2 = st.columns(2)
                            with c1:
                                st.markdown("**✅ Strengths**")
                                for s in res.get('key_strengths', [])[:3]:
                                    st.markdown(f"- {s}")
                            with c2:
                                st.markdown("**⚠️ Gaps**")
                                for g in res.get('gaps_identified', [])[:3]:
                                    st.markdown(f"- {g}")
                            
                            if st.button("🗑️ Clear", key=f"clr_pipe_{idx}"):
                                del st.session_state[f"pipe_analysis_{idx}"]
                                st.rerun()
            else:
                st.info("No applications yet. Share the Careers Page URL to start receiving applications.")
        except Exception as e:
            st.warning(f"Loading: {str(e)}")
    
    # ============ AI ASSISTANT ============
    elif ai_section == "💬 AI Assistant":
        st.subheader("💬 AI Recruitment Assistant")
        st.info("Ask me anything about your candidates, jobs, screening results, or hiring best practices.")
        
        if 'ai_chat_history' not in st.session_state:
            st.session_state.ai_chat_history = [
                {"role": "assistant", "content": "👋 Hello! I'm your AI Recruitment Assistant. I can help with analyzing candidates, comparing applications, generating interview questions, and more. What would you like help with?"}
            ]
        
        for msg in st.session_state.ai_chat_history:
            if msg['role'] == 'user':
                st.markdown(f"""<div style="background:#CC0000;color:white;padding:0.8rem;border-radius:10px;margin:0.5rem 0;margin-left:3rem;"><strong>You</strong><p style="margin:0.3rem 0;">{msg['content']}</p></div>""", unsafe_allow_html=True)
            else:
                st.markdown(f"""<div style="background:#f0f0f0;padding:0.8rem;border-radius:10px;margin:0.5rem 0;margin-right:3rem;"><strong>🤖 AI Assistant</strong><p style="margin:0.3rem 0;">{msg['content']}</p></div>""", unsafe_allow_html=True)
        
        qc1, qc2, qc3, qc4 = st.columns(4)
        with qc1:
            if st.button("🏆 Top Candidates", use_container_width=True):
                st.session_state.ai_chat_history.append({"role": "user", "content": "Who are the top candidates and why?"})
                st.rerun()
        with qc2:
            if st.button("📊 Compare All", use_container_width=True):
                st.session_state.ai_chat_history.append({"role": "user", "content": "Compare all screened candidates and recommend who to interview first."})
                st.rerun()
        with qc3:
            if st.button("❓ Questions", use_container_width=True):
                st.session_state.ai_chat_history.append({"role": "user", "content": "Generate targeted interview questions for the top candidates."})
                st.rerun()
        with qc4:
            if st.button("📝 Offer Draft", use_container_width=True):
                st.session_state.ai_chat_history.append({"role": "user", "content": "Draft an offer letter for the top candidate."})
                st.rerun()
        
        with st.form("ai_chat_form", clear_on_submit=True):
            user_message = st.text_input("Ask me anything...", placeholder="e.g., Who should I interview first?", label_visibility="collapsed")
            if st.form_submit_button("📤 Send", use_container_width=True):
                if user_message:
                    st.session_state.ai_chat_history.append({"role": "user", "content": user_message})
                    try:
                        candidates = db.get_all_candidates()
                        screened = candidates[candidates['ai_score'] > 0] if not candidates.empty and 'ai_score' in candidates.columns else []
                        
                        if ai_agent.use_openai:
                            context = f"You are an AI Recruitment Assistant. Pipeline: {len(candidates)} total, {len(screened)} screened.\n"
                            if len(screened) > 0:
                                for _, c in screened.sort_values('ai_score', ascending=False).head(5).iterrows():
                                    context += f"- {c.get('first_name','')} {c.get('last_name','')}: {int(c.get('ai_score',0))}%, {c.get('ai_tier','')}\n"
                            try:
                                response = ai_agent.client.chat.completions.create(
                                    model="gpt-3.5-turbo",
                                    messages=[{"role": "system", "content": context}, *[{"role": m['role'], "content": m['content']} for m in st.session_state.ai_chat_history[-10:]]],
                                    temperature=0.7, max_tokens=800
                                )
                                ai_response = response.choices[0].message.content
                            except:
                                ai_response = get_smart_response(user_message, screened, candidates)
                        else:
                            ai_response = get_smart_response(user_message, screened, candidates)
                        
                        st.session_state.ai_chat_history.append({"role": "assistant", "content": ai_response})
                    except:
                        st.session_state.ai_chat_history.append({"role": "assistant", "content": "I'm having trouble accessing data. Please try again."})
                    st.rerun()
        
        if st.button("🗑️ Clear Chat History"):
            st.session_state.ai_chat_history = [{"role": "assistant", "content": "👋 Hello! How can I help you today?"}]
            st.rerun()
    
    # ============ JD ANALYSIS ============
    elif ai_section == "📋 JD Analysis":
        st.subheader("📋 AI Job Description Analyzer")
        jd_input = st.radio("Input Method:", ["📝 Paste Text", "📄 Upload JD File"], horizontal=True)
        jd_text = ""
        if jd_input == "📝 Paste Text":
            jd_text = st.text_area("Paste Job Description", height=250)
        else:
            jd_file = st.file_uploader("Upload JD", type=['pdf', 'docx', 'txt'], key="jd_file")
            if jd_file:
                jd_text = save_uploaded_file(jd_file)
                st.text_area("Extracted", jd_text[:500] + "...", height=150, disabled=True)
        
        if st.button("🔍 Analyze JD with AI", use_container_width=True, type="primary"):
            if jd_text:
                with st.spinner("🤖 AI analyzing JD..."):
                    time.sleep(1.5)
                    analysis = ai_agent.analyze_jd(jd_text)
                    st.session_state.current_jd = analysis
                    st.success("✅ Analysis Complete!")
                    c1, c2 = st.columns(2)
                    with c1:
                        st.markdown(f"**Title:** {analysis['title']}")
                        st.markdown(f"**Dept:** {analysis['department']}")
                        st.markdown(f"**Experience:** {analysis['experience_level']}")
                    with c2:
                        st.markdown("**Required Skills:**")
                        for skill in analysis['required_skills'][:10]:
                            st.markdown(f"- `{skill['skill'].title()}`")
                    with st.expander("🚨 Bias Detection Report"):
                        bias_words = ['aggressive', 'ninja', 'rockstar', 'young', 'digital native']
                        jd_lower = jd_text.lower()
                        biases = [w for w in bias_words if w in jd_lower]
                        if biases:
                            st.warning(f"⚠️ {len(biases)} potentially biased terms: {', '.join(biases)}")
                        else:
                            st.success("✅ No biased language detected")
    
    # ============ CV UPLOAD ============
    elif ai_section == "📤 CV Upload & Scoring":
        st.subheader("📤 CV Upload & AI Scoring")
        uploaded_files = st.file_uploader("Upload CVs", type=['pdf', 'docx', 'txt'], accept_multiple_files=True)
        if uploaded_files and st.button("🤖 Analyze All CVs", use_container_width=True, type="primary"):
            if st.session_state.get('current_jd'):
                for file in uploaded_files:
                    cv_text = save_uploaded_file(file)
                    if cv_text:
                        score_result = ai_agent.score_candidate_advanced(cv_text, st.session_state.current_jd)
                        st.success(f"✅ {file.name}: {score_result['overall_score']}% — {score_result['tier']}")
            else:
                st.warning("⚠️ Analyze a JD first in JD Analysis tab.")
    
    # ============ TIERING ============
    elif ai_section == "📊 Candidate Tiering":
        st.subheader("📊 Candidate Tiering Dashboard")
        try:
            candidates = db.get_all_candidates()
            if not candidates.empty:
                tier1 = len(candidates[candidates['ai_tier'].str.contains('Tier 1', na=False)])
                tier2 = len(candidates[candidates['ai_tier'].str.contains('Tier 2', na=False)])
                tier3 = len(candidates[candidates['ai_tier'].str.contains('Tier 3', na=False)])
                c1, c2, c3 = st.columns(3)
                c1.metric("🌟 Tier 1 - Strong Fit", tier1)
                c2.metric("👍 Tier 2 - Good Fit", tier2)
                c3.metric("👎 Tier 3 - Not Recommended", tier3)
                st.markdown("---")
                div_data = pd.DataFrame({'Category': ['Male', 'Female', 'Unspecified'], 'Count': [tier1+tier2, tier3, max(0, len(candidates)-tier1-tier2-tier3)]})
                fig = px.pie(div_data, values='Count', names='Category', hole=0.5, color_discrete_sequence=['#3182ce', '#CC0000', '#718096'])
                st.plotly_chart(fig, use_container_width=True)
                st.dataframe(candidates[['first_name', 'last_name', 'email', 'job_id', 'ai_score', 'ai_tier', 'status']], use_container_width=True, hide_index=True)
        except:
            st.info("Tiering data will appear here.")
    
    # ============ DEEP ANALYSIS ============
    elif ai_section == "🔍 Deep Analysis":
        st.subheader("🔍 AI Deep Candidate Analysis")
        if 'analyze_candidate' in st.session_state:
            candidate = st.session_state.analyze_candidate
            st.markdown(f"### Analyzing: {candidate.get('first_name', '')} {candidate.get('last_name', '')}")
            cv_text = str(candidate.get('resume_text', ''))
            
            st.markdown("---")
            st.markdown("### 🎯 Skills Gap Matrix")
            required_skills = ['Leadership', 'Project Management', 'Data Analysis', 'Communication', 'Strategy', 'Team Management', 'Digital Transformation', 'Stakeholder Management', 'Budgeting', 'Innovation']
            skill_scores = {}
            for skill in required_skills:
                if skill.lower() in cv_text.lower():
                    count = cv_text.lower().count(skill.lower())
                    skill_scores[skill] = min(100, 40 + count * 15)
                else:
                    skill_scores[skill] = random.randint(10, 35)
            skills_df = pd.DataFrame({'Skill': list(skill_scores.keys()), 'Score': list(skill_scores.values())})
            fig = px.bar(skills_df, x='Skill', y='Score', color='Score', color_continuous_scale=['#CC0000', '#d69e2e', '#38a169'])
            fig.update_layout(height=350)
            st.plotly_chart(fig, use_container_width=True)
            
            st.markdown("---")
            st.markdown("### 🚨 Inconsistency & Flag Detection")
            flags = []
            if 'manager' in cv_text.lower() and 'managed' not in cv_text.lower():
                flags.append("⚠️ Claims leadership but no evidence of team management")
            import re
            years_match = re.findall(r'(\d+)\s*years', cv_text.lower())
            if years_match and max(int(y) for y in years_match) > 20:
                flags.append("⚠️ Extended experience claims - verify during interview")
            for flag in flags:
                st.warning(flag)
            if not flags:
                st.success("✅ No major inconsistencies detected")
            
            st.markdown("---")
            st.markdown("### 🤖 AI Confidence Score")
            confidence = random.randint(75, 95)
            st.progress(confidence/100)
            st.markdown(f"**AI Confidence: {confidence}%**")
            
            st.markdown("---")
            st.markdown("### 📝 AI-Generated Interview Questions")
            gaps = [s for s, v in skill_scores.items() if v < 50]
            for gap in gaps[:5]:
                st.markdown(f"- *\"Can you describe a situation where you demonstrated {gap.lower()}?\"*")
            
            st.markdown("---")
            st.markdown("### 🤝 Culture Fit Prediction")
            culture_words = ['team', 'collaborat', 'together', 'support', 'mentor', 'grow', 'learn', 'innovate', 'community', 'impact']
            culture_score = sum(1 for w in culture_words if w in cv_text.lower())
            culture_pct = min(100, culture_score * 12)
            st.progress(culture_pct/100)
            st.markdown(f"**Culture Alignment: {culture_pct}%**")
            
            st.markdown("---")
            st.markdown("### 💰 Salary Benchmarking")
            st.info("Market Rate for similar roles: ₦5M - ₦8M annually (Nigerian market, 2026)")
            
            st.markdown("---")
            st.markdown("### 📊 Radar Chart Comparison")
            categories = ['Skills', 'Experience', 'Leadership', 'Culture Fit', 'Communication', 'Technical']
            fig2 = go.Figure()
            fig2.add_trace(go.Scatterpolar(r=[skill_scores.get('Leadership', 50), skill_scores.get('Project Management', 50), random.randint(40,90), culture_pct, random.randint(50,90), random.randint(40,90)], theta=categories, fill='toself', name=candidate.get('first_name', 'Candidate'), line_color='#CC0000'))
            fig2.add_trace(go.Scatterpolar(r=[80, 75, 70, 80, 75, 80], theta=categories, fill='toself', name='Ideal Profile', line_color='#38a169'))
            fig2.update_layout(height=400)
            st.plotly_chart(fig2, use_container_width=True)
        else:
            st.info("Go to '📥 Load Applications' and click '📊 Full Analysis' on a candidate.")
    
    # ============ EXECUTIVE REPORT ============
    elif ai_section == "📄 Executive Report":
        st.subheader("📄 AI Executive Report Generator")
        try:
            candidates = db.get_all_candidates()
            if not candidates.empty:
                if st.button("📊 Generate Executive PDF Report", use_container_width=True, type="primary"):
                    try:
                        import fpdf
                        FPDF = fpdf.FPDF
                        pdf = FPDF(orientation='L', unit='mm', format='A4')
                        pdf.add_page()
                        pdf.set_fill_color(55, 55, 55)
                        pdf.rect(0, 0, 297, 30, 'F')
                        pdf.set_fill_color(204, 0, 0)
                        pdf.rect(0, 30, 297, 3, 'F')
                        pdf.set_font('Helvetica', 'B', 20)
                        pdf.set_text_color(255, 255, 255)
                        pdf.cell(0, 16, 'CHURCHGATE GROUP', ln=True, align='C')
                        pdf.set_font('Helvetica', 'B', 11)
                        pdf.set_text_color(255, 255, 255)
                        pdf.cell(0, 8, 'AI RECRUITMENT EXECUTIVE REPORT', ln=True, align='C')
                        pdf.ln(8)
                        total = len(candidates)
                        tier1 = len(candidates[candidates['ai_tier'].str.contains('Tier 1', na=False)])
                        pdf.set_font('Helvetica', 'B', 10)
                        pdf.set_text_color(26, 26, 26)
                        pdf.cell(0, 8, f'Total: {total} | Tier 1: {tier1} | Date: {datetime.now().strftime("%B %d, %Y")}', ln=True)
                        pdf.ln(5)
                        pdf.set_fill_color(26, 26, 26)
                        pdf.set_text_color(255, 255, 255)
                        pdf.set_font('Helvetica', 'B', 8)
                        pdf.cell(6, 7, ' #', 1, 0, 'C', True)
                        pdf.cell(45, 7, ' NAME', 1, 0, 'L', True)
                        pdf.cell(55, 7, ' POSITION', 1, 0, 'L', True)
                        pdf.cell(18, 7, 'SCORE', 1, 0, 'C', True)
                        pdf.cell(28, 7, 'TIER', 1, 0, 'C', True)
                        pdf.cell(28, 7, 'STATUS', 1, 0, 'C', True)
                        pdf.cell(97, 7, ' RECOMMENDATION', 1, 0, 'L', True)
                        pdf.ln()
                        pdf.set_font('Helvetica', '', 7)
                        for i, (_, row) in enumerate(candidates.iterrows()):
                            score = row.get('ai_score', 0) or 0
                            try:
                                score = int(float(score))
                            except:
                                score = 0
                            if score >= 85: color = (56,161,105)
                            elif score >= 65: color = (214,158,46)
                            else: color = (204,0,0)
                            pdf.set_text_color(26,26,26)
                            pdf.cell(6, 6, str(i+1), 1, 0, 'C')
                            pdf.cell(45, 6, f' {str(row.get("first_name",""))} {str(row.get("last_name",""))}'[:28], 1, 0, 'L')
                            pdf.cell(55, 6, f' {str(row.get("current_position",""))}'[:33], 1, 0, 'L')
                            pdf.set_text_color(*color)
                            pdf.cell(18, 6, f'{int(score)}%', 1, 0, 'C')
                            pdf.set_text_color(26,26,26)
                            pdf.cell(28, 6, str(row.get('ai_tier','Pending'))[:15], 1, 0, 'C')
                            pdf.cell(28, 6, str(row.get('status','New'))[:15], 1, 0, 'C')
                            rec = 'Advance to Interview' if score >= 85 else 'Keep in View' if score >= 65 else 'Not Recommended'
                            pdf.cell(97, 6, f' {rec}', 1, 0, 'L')
                            pdf.ln()
                        pdf.set_y(-15)
                        pdf.set_fill_color(26,26,26)
                        pdf.rect(0, pdf.get_y()-2, 297, 17, 'F')
                        pdf.set_font('Helvetica', 'I', 7)
                        pdf.set_text_color(180,180,180)
                        pdf.cell(0, 5, 'Churchgate Group - AI Recruitment Report - Confidential', align='C')
                        st.download_button("📥 Download Executive Report", bytes(pdf.output()), "ai_recruitment_report.pdf", "application/pdf")
                        st.success("✅ Report generated!")
                    except Exception as e:
                        st.warning(f"PDF Error: {str(e)}")
                st.download_button("📥 Download CSV", candidates.to_csv(index=False), "candidates.csv", "text/csv")
        except:
            st.info("Load applications first.")
    
    # ============ LINKEDIN ============
    elif ai_section == "🔗 LinkedIn Parse":
        st.subheader("🔍 LinkedIn Profile Parser")
        st.info("Paste a LinkedIn URL and optionally the profile text for AI analysis.")
        
        linkedin_url = st.text_input("LinkedIn Profile URL", placeholder="https://www.linkedin.com/in/username/")
        
        st.markdown("---")
        st.markdown("### 📋 Optional: Paste Profile Text for Better Results")
        st.caption("Copy and paste the visible text from the LinkedIn profile page for more accurate parsing.")
        profile_text = st.text_area("Paste LinkedIn Profile Text", height=150, 
            placeholder="Paste the text content from the LinkedIn profile page here...\n\nThis helps the AI extract more accurate information.")
        
        col1, col2 = st.columns(2)
        with col1:
            if st.button("🔍 Parse Profile", use_container_width=True, type="primary"):
                if linkedin_url:
                    with st.spinner("🤖 AI analyzing profile..."):
                        if profile_text:
                            result = linkedin_parser.parse_profile(linkedin_url, profile_text)
                        else:
                            result = linkedin_parser.parse_profile(linkedin_url)
                        
                        if result:
                            st.session_state.parsed_profile = result
                            st.success(f"✅ Profile analyzed! Parsed via: {result.get('parsed_via', 'unknown')}")
                            st.rerun()
                else:
                    st.error("Please enter a LinkedIn URL")
        
        with col2:
            cv_file = st.file_uploader("Or upload CV for enhanced parsing", type=['pdf', 'docx', 'txt'], key="linkedin_cv")
            if cv_file and linkedin_url:
                cv_text = save_uploaded_file(cv_file)
                if st.button("🔍 Parse with CV", use_container_width=True):
                    with st.spinner("🤖 Analyzing with CV..."):
                        result = linkedin_parser.parse_with_cv_text(linkedin_url, cv_text)
                        if result:
                            st.session_state.parsed_profile = result
                            st.success(f"✅ Profile analyzed with CV enhancement!")
                            st.rerun()
        
        if 'parsed_profile' in st.session_state:
            profile = st.session_state.parsed_profile
            st.markdown("---")
            st.markdown("### 📊 Parsed Profile Data")
            
            col1, col2, col3 = st.columns(3)
            with col1:
                st.markdown(f"**👤 Name:** {profile.get('name', 'N/A')}")
                st.markdown(f"**💼 Headline:** {profile.get('headline', 'N/A')}")
                st.markdown(f"**📍 Location:** {profile.get('location', 'N/A')}")
            with col2:
                st.markdown(f"**🏢 Current Company:** {profile.get('current_company', 'N/A')}")
                st.markdown(f"**📋 Position:** {profile.get('current_position', 'N/A')}")
                st.markdown(f"**📅 Experience:** {profile.get('experience_years', 0)} years")
            with col3:
                st.markdown(f"**🎓 Education:** {profile.get('education', 'N/A')}")
                st.markdown(f"**🔑 Skills:** {', '.join(profile.get('skills', [])[:8])}")
                st.markdown(f"**📊 Parsed via:** {profile.get('parsed_via', 'unknown')}")
            
            if profile.get('summary'):
                st.markdown("---")
                st.markdown(f"**📝 Summary:** {profile.get('summary', '')}")
            
            # Use profile button
            if st.button("📋 Use This Profile Data", use_container_width=True, type="primary"):
                st.session_state['prefill_candidate'] = profile
                st.success("✅ Profile data ready! Go to Candidate Portal or Add Candidate to use it.")

def chat_communications():
    st.markdown("""<div class="churchgate-header"><h1>💬 Social Hub</h1><p>Team Chat | Direct Messages | Announcements | Kudos | Polls | Interest Groups | Smart HRIS Bot | Integrations</p></div>""", unsafe_allow_html=True)
    
    # Reset old chat state format
    if isinstance(st.session_state.get('chat_messages'), list):
        st.session_state.chat_messages = {}
    
    user_name = st.session_state.user['name'] if st.session_state.user else 'Staff'
    user_dept = st.session_state.user.get('department', '') if st.session_state.user else ''
    user_email = st.session_state.user.get('email', '') if st.session_state.user else ''
    user_phone = st.session_state.user.get('phone', '') if st.session_state.user else ''
    is_admin = st.session_state.user['role'] in ['Admin', 'HR Director'] if st.session_state.user else False
    
    @st.cache_data(ttl=300)
    def get_employee_list():
        try:
            df = db.get_all_employees()
            if not df.empty:
                return df
        except:
            pass
        return pd.DataFrame()
    
    employees_df = get_employee_list()
    
    if 'chat_messages' not in st.session_state:
        st.session_state.chat_messages = {}
    if 'direct_messages' not in st.session_state:
        st.session_state.direct_messages = {}
    if 'kudos_board' not in st.session_state:
        st.session_state.kudos_board = []
    if 'polls' not in st.session_state:
        st.session_state.polls = []
    if 'announcements_list' not in st.session_state:
        st.session_state.announcements_list = []
    if 'online_users' not in st.session_state:
        st.session_state.online_users = [user_name]
    if 'bot_conversation' not in st.session_state:
        st.session_state.bot_conversation = []
    if 'interest_groups' not in st.session_state:
        st.session_state.interest_groups = {}
    
    if user_name not in st.session_state.online_users:
        st.session_state.online_users.append(user_name)
    
    def send_slack_notification(message, channel="#general"):
        try:
            webhook_url = st.secrets.get("SLACK_WEBHOOK_URL", "")
            if webhook_url:
                import requests
                payload = {"text": message, "channel": channel}
                requests.post(webhook_url, json=payload, timeout=5)
                return True
        except:
            pass
        return False
    
    def send_whatsapp_message(phone, message):
        try:
            account_sid = st.secrets.get("TWILIO_ACCOUNT_SID", "")
            auth_token = st.secrets.get("TWILIO_AUTH_TOKEN", "")
            from_number = st.secrets.get("TWILIO_WHATSAPP_NUMBER", "")
            if account_sid and auth_token and from_number:
                from twilio.rest import Client
                client = Client(account_sid, auth_token)
                client.messages.create(body=message, from_=f"whatsapp:{from_number}", to=f"whatsapp:{phone}")
                return True, "sent"
        except:
            pass
        clean_phone = phone.replace('+', '').replace(' ', '').replace('-', '')
        wa_link = f"https://wa.me/{clean_phone}?text={message.replace(' ', '%20')}"
        return False, wa_link
    
    def create_asana_task(task_name, assignee_email, due_date=None, project_id=None):
        try:
            token = st.secrets.get("ASANA_PERSONAL_TOKEN", "")
            if token and assignee_email:
                import requests
                headers = {"Authorization": f"Bearer {token}", "Content-Type": "application/json"}
                user_resp = requests.get(f"https://app.asana.com/api/1.0/users?email={assignee_email}", headers=headers, timeout=5)
                assignee_id = None
                if user_resp.status_code == 200:
                    users = user_resp.json().get('data', [])
                    if users:
                        assignee_id = users[0]['gid']
                task_data = {"data": {"name": task_name, "notes": f"Auto-created by Churchgate HRIS for {assignee_email}", "assignee": assignee_id, "workspace": st.secrets.get("ASANA_WORKSPACE_ID", "")}}
                if due_date:
                    task_data["data"]["due_on"] = due_date
                resp = requests.post("https://app.asana.com/api/1.0/tasks", headers=headers, json=task_data, timeout=5)
                return resp.status_code == 201
        except:
            pass
        return False
    
    tab1, tab2, tab3, tab4, tab5, tab6, tab7, tab8 = st.tabs([
        "💬 Team Chat", "👥 Direct Messages", "📢 Announcements", 
        "🌟 Kudos Board", "📊 Polls", "🎯 Groups", "🤖 HRIS Bot", "🔌 Integrations"
    ])
    
    with tab1:
        st.subheader("💬 Department & Group Chat")
        online_count = len(set(st.session_state.online_users))
        st.markdown(f"🟢 **{online_count} online now**")
        channels = ["# general", "# water-cooler", "# announcements", "# kudos"]
        if user_dept:
            channels.insert(1, f"# {user_dept.lower().replace(' ', '-')}")
        if not employees_df.empty:
            all_depts = employees_df['department'].dropna().unique()
            for dept in all_depts:
                dept_channel = f"# {dept.lower().replace(' ', '-')}"
                if dept_channel not in channels:
                    channels.append(dept_channel)
        selected_channel = st.selectbox("📢 Channel", channels)
        today = datetime.now()
        if not employees_df.empty:
            birthday_today = []
            for _, emp in employees_df.iterrows():
                dob = emp.get('date_of_birth')
                if dob and str(dob) != 'None' and str(dob) != 'nan':
                    try:
                        dob_date = pd.to_datetime(dob)
                        if dob_date.month == today.month and dob_date.day == today.day:
                            birthday_today.append(f"{emp['first_name']} {emp['last_name']}")
                    except:
                        pass
            if birthday_today:
                for name in birthday_today:
                    st.balloons()
                    st.success(f"🎂 Happy Birthday, {name}! 🎉")
                    if "# general" not in st.session_state.chat_messages:
                        st.session_state.chat_messages["# general"] = []
                    st.session_state.chat_messages["# general"].append({
                        'sender': 'System', 'content': f"🎂 Happy Birthday to {name}! 🎉🎂",
                        'time': datetime.now().strftime('%I:%M %p'), 'type': 'system'
                    })
        if selected_channel not in st.session_state.chat_messages:
            st.session_state.chat_messages[selected_channel] = [
                {"sender": "System", "content": f"Welcome to {selected_channel}! 👋", "time": datetime.now().strftime('%I:%M %p'), "type": "system"}
            ]
        chat_container = st.container()
        with chat_container:
            for msg in st.session_state.chat_messages[selected_channel][-50:]:
                if msg['type'] == 'system':
                    st.markdown(f"<div style='text-align:center;color:#888;font-size:0.8rem;padding:0.3rem;'>{msg['content']} — {msg['time']}</div>", unsafe_allow_html=True)
                elif msg['sender'] == user_name:
                    st.markdown(f"""<div style="background:#CC0000;color:white;padding:0.6rem 1rem;border-radius:10px;margin:0.3rem 0;margin-left:4rem;text-align:right;"><strong>You</strong><p style="margin:0.2rem 0;">{msg['content']}</p><small>{msg['time']}</small></div>""", unsafe_allow_html=True)
                else:
                    initials = generate_initials(msg['sender'])
                    st.markdown(f"""<div style="background:#f0f0f0;padding:0.6rem 1rem;border-radius:10px;margin:0.3rem 0;margin-right:4rem;"><strong>{initials} {msg['sender']}</strong><p style="margin:0.2rem 0;">{msg['content']}</p><small>{msg['time']}</small></div>""", unsafe_allow_html=True)
        with st.form(f"chat_form_{selected_channel.replace('#', '').replace(' ', '_')}", clear_on_submit=True):
            col1, col2 = st.columns([4, 1])
            with col1:
                message = st.text_input("Type a message...", placeholder=f"Message {selected_channel}...", label_visibility="collapsed", key=f"chat_input_{selected_channel.replace('#', '')}")
            with col2:
                send = st.form_submit_button("📤 Send", use_container_width=True)
            if send and message:
                st.session_state.chat_messages[selected_channel].append({
                    'sender': user_name, 'content': message,
                    'time': datetime.now().strftime('%I:%M %p'), 'type': 'user'
                })
                st.rerun()
    
    with tab2:
        st.subheader("👥 Direct Messages")
        
        # Build employee list
        if not employees_df.empty:
            team_list = []
            for _, row in employees_df.iterrows():
                full_name = f"{row['first_name']} {row['last_name']}"
                if full_name != user_name:
                    team_list.append(full_name)
        else:
            team_list = ["Jerome Das", "Sanjeev Purwar", "Ahmed Karim"]
        
        team_list = sorted(team_list)
        
        # Show unread notifications
        try:
            all_my_msgs = db._get("chat_messages")
            if all_my_msgs:
                unread = []
                for m in all_my_msgs:
                    if m.get('receiver_name') == user_name and m.get('is_read') == False:
                        unread.append(m)
                if unread:
                    senders = {}
                    for m in unread:
                        s = m['sender_name']
                        senders[s] = senders.get(s, 0) + 1
                    for sender, count in senders.items():
                        st.info(f"🔴 **{count} unread message{'s' if count > 1 else ''}** from **{sender}** — Select '{sender}' below to read")
        except:
            pass
        
        # Simple dropdown - no fancy auto-select
        dm_with = st.selectbox("💬 Chat with", ["Select colleague..."] + team_list)
        
        if dm_with != "Select colleague...":
            
            # Load conversation from database (fetch all and filter)
            try:
                all_msgs_raw = db._get("chat_messages")
                sent = []
                received = []
                if all_msgs_raw:
                    for m in all_msgs_raw:
                        if m.get('sender_name') == user_name and m.get('receiver_name') == dm_with:
                            sent.append(m)
                        if m.get('sender_name') == dm_with and m.get('receiver_name') == user_name:
                            received.append(m)
                            if m.get('is_read') in [False, 'false', 0, '0']:
                                db._patch("chat_messages", {"is_read": True}, {"id": m['id']})
                
                all_msgs = []
                if sent:
                    for m in sent:
                        all_msgs.append({'sender': m['sender_name'], 'content': m['message'], 'time': str(m.get('sent_at', ''))[:16]})
                if received:
                    for m in received:
                        all_msgs.append({'sender': m['sender_name'], 'content': m['message'], 'time': str(m.get('sent_at', ''))[:16]})
                        # Mark as read
                        if m.get('is_read') in [False, 'false', 0, '0']:
                            db._patch("chat_messages", {"is_read": True}, {"id": m['id']})
                
                all_msgs.sort(key=lambda x: x['time'])
            except:
                all_msgs = []
            
            if all_msgs:
                for msg in all_msgs[-30:]:
                    if msg['sender'] == user_name:
                        st.markdown(f"""<div style="background:#CC0000;color:white;padding:0.6rem 1rem;border-radius:10px;margin:0.3rem 0;margin-left:4rem;"><strong>You</strong><p style="margin:0.2rem 0;">{msg['content']}</p><small>{msg['time']}</small></div>""", unsafe_allow_html=True)
                    else:
                        initials = generate_initials(msg['sender'])
                        st.markdown(f"""<div style="background:#f0f0f0;padding:0.6rem 1rem;border-radius:10px;margin:0.3rem 0;margin-right:4rem;"><strong>{initials} {msg['sender']}</strong><p style="margin:0.2rem 0;">{msg['content']}</p><small>{msg['time']}</small></div>""", unsafe_allow_html=True)
            else:
                st.info(f"No messages with {dm_with} yet. Say hello! 👋")
            
            # Send message
            with st.form(f"dm_form_{dm_with.replace(' ', '_')}", clear_on_submit=True):
                col1, col2 = st.columns([4, 1])
                with col1:
                    dm_message = st.text_input("Type message...", placeholder=f"Message {dm_with}...", label_visibility="collapsed", key=f"dm_input_{dm_with}")
                with col2:
                    if st.form_submit_button("📤", use_container_width=True):
                        if dm_message:
                            db._post("chat_messages", {
                                "sender_name": user_name,
                                "receiver_name": dm_with,
                                "message": dm_message,
                                "sent_at": datetime.now().strftime('%Y-%m-%d %H:%M'),
                                "is_read": False
                            })
                            st.success(f"✅ Sent to {dm_with}!")
                            send_browser_notification("💬 New Message", f"New message from {user_name}")
                            st.rerun()
    
    with tab3:
        st.subheader("📢 Company Announcements")
        if is_admin:
            with st.expander("📝 Create Announcement", expanded=False):
                with st.form("new_announcement"):
                    ann_title = st.text_input("Title *")
                    ann_content = st.text_area("Message *")
                    c1, c2, c3 = st.columns(3)
                    with c1:
                        ann_priority = st.selectbox("Priority", ["Normal", "Important", "Urgent"])
                    with c2:
                        post_to_slack = st.checkbox("Post to Slack", value=True)
                    with c3:
                        post_to_whatsapp = st.checkbox("Send WhatsApp Alert")
                    if st.form_submit_button("📢 Post Announcement", use_container_width=True):
                        if ann_title and ann_content:
                            new_ann = {
                                "title": ann_title, "date": datetime.now().strftime('%Y-%m-%d'),
                                "priority": ann_priority, "content": ann_content,
                                "posted_by": user_name, "time": datetime.now().strftime('%I:%M %p')
                            }
                            st.session_state.announcements_list.append(new_ann)
                            if "# announcements" not in st.session_state.chat_messages:
                                st.session_state.chat_messages["# announcements"] = []
                            st.session_state.chat_messages["# announcements"].append({
                                'sender': 'System', 'content': f"📢 {ann_title}: {ann_content}",
                                'time': datetime.now().strftime('%I:%M %p'), 'type': 'system'
                            })
                            if post_to_slack:
                                slack_msg = f"📢 *{ann_title}*\n{ann_content}\n\nPriority: {ann_priority} | Posted by: {user_name}"
                                if send_slack_notification(slack_msg):
                                    st.info("✅ Posted to Slack")
                            if post_to_whatsapp and not employees_df.empty:
                                wa_count = 0
                                for _, emp in employees_df.head(20).iterrows():
                                    phone = str(emp.get('phone', ''))
                                    if phone and phone != '' and phone != 'nan':
                                        wa_msg = f"📢 Churchgate Group Announcement:\n\n*{ann_title}*\n{ann_content}"
                                        success, _ = send_whatsapp_message(phone, wa_msg)
                                        if success:
                                            wa_count += 1
                                if wa_count > 0:
                                    st.info(f"📱 WhatsApp alerts sent to {wa_count} employees")
                            st.success("✅ Announcement posted!")
                            st.balloons()
                            st.rerun()
        all_announcements = st.session_state.announcements_list.copy()
        if not all_announcements:
            all_announcements = [
                {"title": "Q2 Performance Reviews", "date": "2026-06-01", "priority": "Important", "content": "All departments to submit Q2 performance reviews by June 15, 2026.", "posted_by": "HR", "time": "09:00 AM"},
                {"title": "BMS Implementation Update", "date": "2026-05-28", "priority": "Normal", "content": "Phase 1 of BMS installation complete. Phase 2 starts June 10.", "posted_by": "Technology Group", "time": "02:30 PM"},
                {"title": "Holiday Notice - Democracy Day", "date": "2026-05-25", "priority": "Important", "content": "Office closed on June 12 for Democracy Day. Normal operations resume June 13.", "posted_by": "Admin", "time": "11:00 AM"},
            ]
        for ann in reversed(all_announcements):
            pc = "#CC0000" if ann['priority'] == 'Urgent' else "#d69e2e" if ann['priority'] == 'Important' else "#3182ce"
            st.markdown(f"""<div style="background:white;padding:1rem;border-radius:8px;margin-bottom:0.6rem;border-left:4px solid {pc};"><div style="display:flex;justify-content:space-between;"><strong>{ann['title']}</strong><span style="background:{pc};color:white;padding:0.2rem 0.6rem;border-radius:10px;font-size:0.75rem;">{ann['priority']}</span></div><p style="margin:0.4rem 0;">{ann['content']}</p><small>📅 {ann['date']} | 👤 {ann.get('posted_by', 'System')} | 🕐 {ann.get('time', '')}</small></div>""", unsafe_allow_html=True)
    
    with tab4:
        st.subheader("🌟 Kudos & Appreciation Board")
        st.info("Give public shout-outs to colleagues who've done great work!")
        if not employees_df.empty:
            kudos_list = [f"{row['first_name']} {row['last_name']}" for _, row in employees_df.iterrows() if f"{row['first_name']} {row['last_name']}" != user_name]
            if not kudos_list:
                kudos_list = ["No other employees yet"]
        else:
            kudos_list = ["Jerome Das", "Sanjeev Purwar", "Ahmed Karim"]
        with st.form("give_kudos"):
            c1, c2 = st.columns(2)
            with c1:
                kudos_to = st.selectbox("👤 Give Kudos to", kudos_list)
            with c2:
                kudos_emoji = st.selectbox("🎉 Emoji", ["🌟 Outstanding", "👏 Great Job", "🏆 Top Performer", "💪 Strong Effort", "🎯 On Target", "🔥 On Fire", "💡 Innovation", "🤝 Team Player"])
            kudos_msg = st.text_area("💬 Your appreciation message *", placeholder="What did they do that was exceptional?")
            if st.form_submit_button(f"🌟 Send Kudos!", use_container_width=True):
                if kudos_msg and kudos_to != "No other employees yet":
                    kudos_entry = {
                        'from': user_name, 'to': kudos_to, 'message': kudos_msg,
                        'emoji': kudos_emoji.split()[0], 'time': datetime.now().strftime('%b %d, %I:%M %p'),
                        'from_dept': user_dept
                    }
                    st.session_state.kudos_board.append(kudos_entry)
                    if "# kudos" not in st.session_state.chat_messages:
                        st.session_state.chat_messages["# kudos"] = []
                    st.session_state.chat_messages["# kudos"].append({
                        'sender': 'System', 'content': f"{kudos_emoji.split()[0]} Kudos! {user_name} ({user_dept}) → {kudos_to}: \"{kudos_msg}\"",
                        'time': datetime.now().strftime('%I:%M %p'), 'type': 'system'
                    })
                    send_slack_notification(f"🌟 *Kudos!* {user_name} ({user_dept}) gave kudos to *{kudos_to}*: \"{kudos_msg}\"", "#kudos")
                    st.success(f"🌟 Kudos sent to {kudos_to}!")
                    st.balloons()
                    st.rerun()
        st.markdown("---")
        st.markdown("### 🏆 Recent Kudos")
        if st.session_state.kudos_board:
            for k in reversed(st.session_state.kudos_board[-15:]):
                st.markdown(f"""<div style="background:white;padding:0.8rem;border-radius:8px;margin-bottom:0.4rem;border-left:4px solid #d69e2e;">{k['emoji']} <strong>{k['from']}</strong> <small>({k.get('from_dept', '')})</small> → <strong>{k['to']}</strong><p style="margin:0.2rem 0;">"{k['message']}"</p><small>{k['time']}</small></div>""", unsafe_allow_html=True)
        else:
            st.info("No kudos yet. Be the first to appreciate a colleague! 🌟")
    
    with tab5:
        st.subheader("📊 Quick Polls & Surveys")
        with st.expander("📝 Create New Poll", expanded=not st.session_state.polls):
            with st.form("create_poll"):
                poll_question = st.text_input("Poll Question *")
                c1, c2 = st.columns(2)
                with c1:
                    poll_option1 = st.text_input("Option 1", "Yes 👍")
                    poll_option2 = st.text_input("Option 2", "No 👎")
                with c2:
                    poll_option3 = st.text_input("Option 3 (Optional)")
                    poll_option4 = st.text_input("Option 4 (Optional)")
                if st.form_submit_button("📊 Create Poll", use_container_width=True):
                    if poll_question and poll_option1 and poll_option2:
                        options = [poll_option1, poll_option2]
                        if poll_option3: options.append(poll_option3)
                        if poll_option4: options.append(poll_option4)
                        st.session_state.polls.append({
                            'question': poll_question, 'options': options,
                            'votes': {opt: 0 for opt in options},
                            'voters': [], 'created_by': user_name,
                            'created_at': datetime.now().strftime('%b %d, %I:%M %p'),
                            'active': True
                        })
                        st.success("✅ Poll created!")
                        st.rerun()
        st.markdown("---")
        if st.session_state.polls:
            for i, poll in enumerate(reversed(st.session_state.polls)):
                if poll.get('active', True):
                    with st.expander(f"📊 {poll['question']} — by {poll['created_by']} ({poll['created_at']})", expanded=(i < 3)):
                        total_votes = sum(poll['votes'].values())
                        col1, col2 = st.columns([3, 1])
                        with col1:
                            for opt in poll['options']:
                                votes = poll['votes'][opt]
                                pct = (votes / total_votes * 100) if total_votes > 0 else 0
                                st.markdown(f"**{opt}** — {votes} vote{'s' if votes != 1 else ''} ({pct:.0f}%)")
                                st.progress(pct/100)
                        with col2:
                            st.metric("Total Votes", total_votes)
                        if user_name not in poll['voters']:
                            vote_choice = st.selectbox("Your vote", poll['options'], key=f"vote_{i}")
                            if st.button("🗳️ Vote", key=f"submit_vote_{i}"):
                                poll['votes'][vote_choice] += 1
                                poll['voters'].append(user_name)
                                st.success("✅ Vote recorded!")
                                st.rerun()
                        if user_name == poll['created_by'] or is_admin:
                            if st.button("🔒 Close Poll", key=f"close_poll_{i}"):
                                poll['active'] = False
                                st.rerun()
        else:
            st.info("No polls yet. Create one!")
    
    with tab6:
        st.subheader("🎯 Interest Groups")
        st.info("Join groups and connect with colleagues who share your interests!")
        total_employees = len(employees_df) if not employees_df.empty else 57
        import random
        random.seed(42)
        groups = [
            {"name": "💻 Tech Innovators", "icon": "💻", "description": "AI, IoT, and emerging tech discussions", "members": min(total_employees, random.randint(20, 30))},
            {"name": "⚽ Football Fans", "icon": "⚽", "description": "EPL, La Liga, and Nigerian league discussions", "members": min(total_employees, random.randint(25, 40))},
            {"name": "📚 Book Club", "icon": "📚", "description": "Monthly book recommendations and reviews", "members": min(total_employees, random.randint(10, 20))},
            {"name": "🏋️ Fitness & Wellness", "icon": "🏋️", "description": "Health tips, workout routines, wellness challenges", "members": min(total_employees, random.randint(15, 25))},
            {"name": "🎵 Music Lovers", "icon": "🎵", "description": "Share playlists and discover new music", "members": min(total_employees, random.randint(15, 22))},
            {"name": "📸 Photography", "icon": "📸", "description": "Share photos and photography tips", "members": min(total_employees, random.randint(8, 15))},
            {"name": "🍳 Foodies", "icon": "🍳", "description": "Restaurant reviews, recipes, food events", "members": min(total_employees, random.randint(18, 28))},
            {"name": "🙏 Faith & Fellowship", "icon": "🙏", "description": "Prayer group and spiritual support", "members": min(total_employees, random.randint(20, 35))},
        ]
        for group in groups:
            col1, col2, col3 = st.columns([3, 1, 1])
            with col1:
                st.markdown(f"**{group['name']}** — {group['members']} members<br><small style='color:#666;'>{group['description']}</small>", unsafe_allow_html=True)
            with col2:
                group_key = group['name']
                if group_key not in st.session_state.interest_groups:
                    st.session_state.interest_groups[group_key] = []
                is_member = user_name in st.session_state.interest_groups[group_key]
                if is_member:
                    st.success("✅ Joined")
                else:
                    if st.button(f"Join", key=f"join_{group['name'][:15]}"):
                        st.session_state.interest_groups[group_key].append(user_name)
                        st.success(f"✅ Joined {group['name']}!")
                        st.rerun()
            with col3:
                if is_member:
                    if st.button(f"Leave", key=f"leave_{group['name'][:15]}"):
                        st.session_state.interest_groups[group_key].remove(user_name)
                        st.rerun()
    
    with tab7:
        st.subheader("🤖 Smart HRIS Assistant")
        emp_count = len(employees_df) if not employees_df.empty else 57
        dept_count = len(employees_df['department'].unique()) if not employees_df.empty else 10
        st.info(f"👋 Hi {user_name}! I can help with leave, payroll, KPIs, training, policies, and more. We currently have {emp_count} employees across {dept_count} departments.")
        for msg in st.session_state.bot_conversation:
            if msg['role'] == 'user':
                st.markdown(f"""<div style="background:#CC0000;color:white;padding:0.6rem 1rem;border-radius:10px;margin:0.3rem 0;margin-left:4rem;"><strong>You</strong><p style="margin:0.2rem 0;">{msg['content']}</p><small>{msg['time']}</small></div>""", unsafe_allow_html=True)
            else:
                st.markdown(f"""<div style="background:#f0f0f0;padding:0.6rem 1rem;border-radius:10px;margin:0.3rem 0;margin-right:4rem;"><strong>🤖 HRIS Bot</strong><p style="margin:0.2rem 0;">{msg['content']}</p><small>{msg['time']}</small></div>""", unsafe_allow_html=True)
        policy_kb = {
            'leave': "🏖️ **Leave Policy:** Annual leave varies by grade level. Level 1: 14 days, Level 2-3: 28 days, Level 4-5: 30 days. Apply via HRIS. Maternity: 90 days. Sick leave requires medical certificate.",
            'payroll': "💰 **Payroll:** Processed on 25th monthly. Pay stubs available in My Documents. Salary includes basic, housing, transport, meal, utility, and education allowances.",
            'training': "🎓 **Training:** Available through the LMS with 100+ courses. Check Training & Development tab for enrollment.",
            'kpi': "📊 **KPIs:** Set in Performance & OKRs. Align to 4 strategic pillars. Submit for HOD approval before appraisal.",
            'benefits': "🎁 **Benefits:** HMO coverage, pension (8% employee + 10% employer), annual leave, training support, performance bonuses.",
            'safety': "🦺 **HSE Policy:** Churchgate is committed to workplace safety. Full HSE manual in Knowledge Base. Report incidents immediately.",
            'working hours': "🕐 **Working Hours:** 8:00 AM - 5:30 PM (Abuja), 8:30 AM - 6:00 PM (Lagos). Monday to Friday. Overtime pre-approved only.",
            'disciplinary': "⚖️ **Disciplinary:** Verbal caution → Warning letter → Suspension → Summary dismissal. HR must be copied on all queries.",
        }
        
        st.markdown("**Quick Questions:**")
        quick_questions = [
            ("Leave", "How do I apply for leave?"), 
            ("KPIs", "How do I set my KPIs?"), 
            ("Benefits", "What are my benefits?"), 
            ("Training", "What training is available?"), 
            ("Payroll", "When is payroll processed?"),
            ("HSE", "What is the HSE policy?"),
            ("HR Policy", "What are the working hours?"),
            ("Discipline", "What is the disciplinary procedure?"),
            ("Safety", "What are the fire safety procedures?"),
            ("Maternity", "What is the maternity leave policy?")
        ]
        row1 = st.columns(5)
        row2 = st.columns(5)
        for i, (label, question) in enumerate(quick_questions):
            col = row1[i] if i < 5 else row2[i-5]
            with col:
                if st.button(f"❓ {label}", key=f"qq_{label}", use_container_width=True):
                    st.session_state.bot_conversation.append({'role': 'user', 'content': question, 'time': datetime.now().strftime('%I:%M %p')})
                    # Search policy knowledge base for the answer
                    answer = "I can help with many topics! Try asking about leave, benefits, safety, or HR policies."
                    question_lower = question.lower()
                    for key, val in policy_kb.items():
                        if key in question_lower:
                            answer = val
                            break
                    st.session_state.bot_conversation.append({'role': 'bot', 'content': answer, 'time': datetime.now().strftime('%I:%M %p')})
                    st.rerun()
        with st.form("bot_form", clear_on_submit=True):
            col1, col2 = st.columns([4, 1])
            with col1:
                bot_question = st.text_input("Ask HRIS Bot...", placeholder="e.g., How do I set my KPIs?", label_visibility="collapsed")
            with col2:
                ask = st.form_submit_button("🤖 Ask", use_container_width=True)
            if ask and bot_question:
                # ===== COMBINED KNOWLEDGE BASE =====
                policy_kb = {
                    # HR POLICY MANUAL
                    'working hours': "**Working Hours (Policy Article 2.13):** Normal work hours are 0800hrs to 1730hrs, Monday to Friday at Abuja offices. Lagos offices: 0830hrs to 1800hrs, Monday to Friday. One hour lunch break. Shift work may require continuous coverage. Overtime is compensated when pre-approved.",
                    'overtime': "**Overtime (Policy Article 2.14):** Overtime is time worked beyond normal hours. Rate: Normal rate (100%) per hour after closing hours Mon-Sat, 150% on Sundays and Public Holidays. Only employees within Levels 1 and 2 are eligible. Must be pre-approved and recorded on overtime sheets.",
                    'probation': "**Probation (Policy Article 2.9):** Every new employee undergoes 6 months probation. Two performance evaluations are conducted — at 3 months and 6 months. Confirmation depends on satisfactory performance. Extension possible for 3 more months. Failure after extension leads to termination.",
                    'termination': "**Termination (Policy Article 2.16):** After confirmation, either party may terminate with one month's notice or payment in lieu. Summary dismissal applies for gross misconduct. The Group may refuse resignation during investigation, suspension, or pending disciplinary cases.",
                    'resignation': "**Resignation (Policy Article 2.16):** During probation, resignation can be without notice. After confirmation, one month's notice or salary in lieu is required. The Group reserves the right to refuse resignation if the employee is under investigation, on suspension, or has a pending disciplinary case.",
                    'dress code': "**Dress Code (Policy Article 2.21):** All office staff must dress corporately and decently at all times. Indecent dressing (provocative clothes, exposing body parts, sagging trousers) is prohibited. 1st offense: verbal warning, 2nd: warning letter, 3rd: 3-day suspension without pay. Repeated violations may lead to termination.",
                    'misconduct': "**Acts of Misconduct (Policy Article 2.22):** Includes quarrelling, abusive language, threatening colleagues, fighting, stealing, drunkenness, sleeping on duty, discrimination, smoking in non-designated areas, forgery/fraud, and misuse of office equipment. All proven acts attract disciplinary action.",
                    'disciplinary': "**Disciplinary Procedure (Policy Article 3.3):** When misconduct occurs, a query is issued requiring response within 24 hours. Investigation follows. Depending on gravity, outcomes range from verbal caution, warning letter, penalization, to summary dismissal. HR must be copied on all warnings and queries.",
                    'grievance': "**Grievance Procedure (Policy Article 3.4):** Step 1: Discuss with Unit Head. Step 2: If unresolved within 2 working days, escalate to HOD who must respond within 3 working days. Step 3: If still unresolved, refer to HR department for a hearing committee to resolve.",
                    'suspension': "**Suspension (Policy Article 3.5.4):** Employee may be suspended with half pay or without pay pending investigation (1-6 weeks). Must report daily for 2 hours to designated area and sign compliance. If exonerated, balance salary is paid. If guilty, disciplinary procedure applies.",
                    'summary dismissal': "**Summary Dismissal (Policy Article 3.5.5):** Applies for theft, fraud, wilful disobedience, drunkenness/drugs, divulging confidential information, criminal conviction, prolonged absence, fighting, conflict of interest, and other extreme misconduct. No notice given. All benefits forfeited.",
                    'annual leave': "**Annual Leave (Policy Article 7.1):** Levels 4 & 5: 30 working days. Levels 2 & 3: 28 working days. Level 1H: 21 working days. Level 1: 14 working days. Expatriates: 30 calendar days. Must complete full calendar year to be eligible. Leave must be approved by management. Apply via ERP two weeks prior. Unused leave is forfeited.",
                    'maternity leave': "**Maternity Leave (Policy Article 7.2):** Confirmed female employees entitled to 90 calendar days with full pay, once in 2 years. Must provide medical certificate from Ob-Gyn. Cannot take annual leave and maternity leave in same year. Unconfirmed employees get 90 days without pay. Nursing mothers may close 1 hour early for 3 months after return.",
                    'sick leave': "**Sick Leave (Policy Article 7.5 & 6.5):** Requires medical certificate from Group-approved hospital. Under 6 months service: 1 month full pay. 6 months-3 years: 1 month full, 2nd month half. 3-5 years: 2 months full, 3rd month half. 5-10 years: 2 months full, next 2 months half. 10+ years: 3 months full, next 3 months half. Self-inflicted illness: no pay.",
                    'compassionate leave': "**Compassionate Leave (Policy Article 7.3):** Up to 5 calendar days paid leave per year for death of spouse, parent, or child. Duration depends on circumstances and distance. Death certificate required. Not a right but a privilege at management discretion.",
                    'study leave': "**Study Leave (Policy Article 7.4):** Up to 24 months without pay for Master's Degree in relevant field. Requires minimum 5 years service. Period counts as service if employee returns for at least 2 years. Approval at management discretion.",
                    'pension': "**Pension Scheme (Policy Article 6.1):** Employee contributes 8% of monthly emoluments, Group contributes 10%, remitted to employee's chosen PFA, as required by the Pension Reform Act 2014.",
                    'medical': "**Medical Facilities (Policy Article 6.4):** Available to all permanent employees. Levels 1 & 2: self only. Levels 3-5: self, spouse, and up to 2 children through Group-approved HMO. Full details in separate HMO document.",
                    'salary': "**Salary Policy (Policy Article 5.1 & 5.2):** Annual gross salary includes basic pay, housing, transport, meal subsidy, utility, and education allowances. Paid monthly by last day of each month after deductions (PAYE, Pension, Surcharges). Paid directly to employee's bank account.",
                    'salary advance': "**Salary Advance (Policy Article 8.2):** Cannot exceed 40% of net monthly salary. Not available during probation or in two consecutive months. Requires Executive Director approval. Refunded at next salary payment.",
                    'housing loan': "**Housing Loan/Rent Advance (Policy Article 8.1):** Discretionary. Requires two existing employee guarantors. Repayment within 10 months. Total loan repayments cannot exceed 1/3 of monthly take-home pay. Must authorize salary deduction.",
                    'bonus': "**End of Year Bonus (Policy Article 5.5):** Discretionary payment equal to one month's basic salary. Must have worked from first business day of year to December 31st for full payment. Those with 6+ months service and confirmed by December 31st receive pro-rata.",
                    'promotion': "**Promotion (Policy Article 4.5):** Based on performance evaluation, available positions, demonstrated merit, potential, skills and ability for higher responsibilities. Management may promote without salary increment at its discretion.",
                    'performance evaluation': "**Performance Evaluation (Policy Article 4.1):** Two evaluations per financial year. First in Q1 (or 3 months after employment for new staff). Second in Q1 of preceding year. Average of both forms basis for training, promotion, salary increment decisions. Score below 65% for two consecutive evaluations leads to automatic exit.",
                    'retirement': "**Retirement (Policy Article 2.11):** Compulsory retirement at age 60. Voluntary early retirement allowed. Management may send employee to early retirement for dwindling performance. Employees over 60 may be retained on contract at management discretion.",
                    'redundancy': "**Redundancy (Policy Article 11.8):** Involuntary loss of employment through no fault of employee due to excess manpower or work contraction. Implemented in line with prevailing Nigerian law.",
                    'conflict of interest': "**Conflict of Interest (Policy Article 2.7):** All employees must make periodic Conflict of Interest Declaration. False declaration may result in summary dismissal. Facilitated by HR department.",
                    'confidentiality': "**Confidentiality (Policy Article 2.19):** All employees must maintain confidentiality. No documents removed without authorization. Client information not disclosed without court order or written consent. Breach may lead to summary dismissal.",
                    'sexual harassment': "**Sexual Harassment (Policy Article 11.9):** Unacceptable and against corporate values. Includes unwelcome sexual advances, requests for favors, verbal or physical harassment. Confirmed offenders face 2-week suspension without pay or summary dismissal depending on severity as determined by Disciplinary Committee.",
                    'bullying': "**Intimidation & Bullying (Policy Article 11.11):** Ongoing misuse of power through verbal, threatening, abusive, or manipulative behavior. Will not be tolerated. Offenders face 2-week suspension without pay or summary dismissal.",
                    'religious activities': "**Religious Activities in Office (Policy Article 11.12):** Practicing faith is welcomed but bringing practices to office that make others uncomfortable is prohibited. No solicitation, sharing tracts/fliers, preaching during working hours. 1st offense: verbal warning. 2nd: warning letter. 3rd: 3-day suspension without pay.",
                    'financial impropriety': "**Financial Impropriety (Policy Article 11.13):** Includes fraud, embezzlement, theft, misappropriation, falsifying records, bribery, unauthorized transactions. Offenders face summary dismissal or one-month suspension without pay depending on severity.",
                    'transfer': "**Transfer (Policy Article 2.18 & 11.2):** Employee may be transferred to another subsidiary within the Group. Same-city transfers: no allowance. Different city temporary transfer: hotel/guest house, intra-city transport, meal allowance. Permanent transfer: fixed amount in lieu of accommodation. Employee-initiated transfer: no costs covered.",
                    'training development': "**Training & Development (Policy Article 4.2):** Group provides training based on identified needs from performance evaluations. Employees may suggest courses. Management has sole discretion on training provision. Self-learning is actively supported.",
                    'professional development': "**Professional Development (Policy Article 4.4):** Performing employees may request sponsorship for professional qualifications relevant to their job. If approved, employee signs bond to remain for specified period or repay full amount. Bond length depends on program value and duration.",
                    'long service award': "**Long Service Award (Policy Article 6.2):** Recognition every 5 years of full-time employment. Congratulatory message from Chairman and certificate on anniversary date. 10+ years (in multiples of 5): additional gift and/or certificate presented within the completion year.",
                    'death benefits': "**Death Benefits (Policy Article 10.2):** Group Life Insurance pays 3 times annual emoluments (basic salary, housing, transport allowances) to next of kin. Facilitated through HR department.",
                    'acting allowance': "**Acting Allowance (Policy Article 5.6):** When performing duties in higher grade for not less than 1 month, employee receives monthly acting allowance equal to 20% of monthly basic salary.",
                    'notice period': "**Notice Period (Policy Article 2.16):** After confirmation, employment may be terminated by either side giving one month's notice or payment of equivalent basic salary in lieu of notice. Individual employment contract terms supersede this policy.",
                    'identity card': "**Staff Identification (Policy Article 2.10):** Identity cards issued to all employees and must be displayed at all times. Must be worn when visiting other Churchgate offices or representing the Group. Lost cards: replacement cost borne by employee if due to carelessness. Surrender to HR upon exit.",
                    # HSE MANUAL
                    'safety': "**HSE Policy (HSE Manual Chapter 3):** Churchgate is committed to providing a safe and healthy workplace. The policy covers Preventative Action (risk assessments, PPE, training), Emergency Management (fire alarms, sprinklers, extinguishers, evacuation plans, first-aid kits, fire drills), and Additional Measures (policy updates, incident analysis, expert consultation).",
                    'electrical': "**Electrical Safety (HSE Manual Chapter 4):** Live parts must be de-energized before work. Lockout/Tagout procedures must be followed. Only qualified persons may work on energized equipment. Portable equipment must be visually inspected before use. Damaged items must be removed from service.",
                    'lockout': "**Lockout/Tagout (HSE Manual Chapter 4.4-4.8):** Circuits must be locked out or tagged out before work. Verification of de-energized condition required. Only the person who applied the lock may remove it.",
                    'hot work': "**Hot Work (HSE Manual Chapter 5):** Includes welding, cutting, brazing. Fire hazards must be removed or guarded. Fire extinguishing equipment must be ready. Fire watch required when combustibles within 35 feet.",
                    'fire': "**Fire Safety (HSE Manual Chapter 6):** Emergency plan for fire-fighting and evacuation must be available. Fire extinguishers installed throughout workplace. Fire drills practiced regularly.",
                    'confined space': "**Confined Space Safety (HSE Manual Chapter 7):** Atmosphere must be tested before entry. Oxygen must be 19.5-23.5%. Entry permit required. At least one standby person outside. Rescue plan must be in place.",
                    'ppe': "**PPE Program (HSE Manual Chapter 12):** PPE protects employees from workplace hazards. Facility Manager responsible for ensuring sufficient PPE, training, and periodic re-evaluation. Defective PPE must be immediately replaced.",
                    'heights': "**Safety at Heights (HSE Manual Chapter 11):** Workers must be medically checked. Rope access kit includes working line, safety line, ascender, backup device, helmet, harness. Ladders inspected daily. Use fall arrestor above 6 feet.",
                    'ladder': "**Ladder Safety (HSE Manual Chapter 11.1-11.3):** Daily inspection required. Never stand on top rung. Face ladder when climbing. Extension ladder must extend 3 feet above support point.",
                    'waste': "**Waste Management (HSE Manual Chapter 9):** Segregate waste into Recyclable and Non-Recyclable. Use labelled bins. PPE required. 3R Principle: Recycle, Reduce, Reuse. Dispose through authorized vendor.",
                    'chemical': "**Chemical Management (HSE Manual Chapter 10):** Maintain chemical inventory. MSDS must be available. Hazardous chemicals stored separately. PPE mandatory when handling chemicals.",
                    'air quality': "**Indoor Air Quality (HSE Manual Chapter 14):** HVAC filters changed quarterly. Relative humidity 30-60%. Duct cleaning outside occupied hours. Source control is most effective solution.",
                    'ergonomics': "**Ergonomics (HSE Manual Chapter 13):** Chair: thighs parallel to floor. Monitor: top at or below eye level. Elbows at 90 degrees. Wrists straight. Take regular breaks.",
                    'lift': "**Lift/Elevator Safety (HSE Manual Chapter 15):** Daily checks required. Monthly inspection. Emergency rescue: isolate power, release brake, manually move to nearest floor.",
                    'incident': "**Incident Reporting (HSE Manual Chapter 18):** All injuries must be reported immediately. Near Miss: 48 hours. Minor: 36 hours. Serious: 24 hours. Severe: 18 hours. Investigation determines root cause.",
                    'energy': "**Energy Management (HSE Manual Chapter 17):** Turn off nonessential lighting. Set thermostats at 25°C. Use LED lamps. Regular HVAC maintenance. Switch off equipment when not in use.",
                    'housekeeping': "**Housekeeping Safety (HSE Manual Chapter 8):** Wipe spills immediately. Keep walkways clear. Most common injuries: slips, trips, falls. Use appropriate footwear. Use handrails.",
                    'diesel': "**Diesel Storage (HSE Manual Chapter 16.4-16.6):** No smoking. Authorized persons only. Fire extinguishers required. Two people for filling day tanks. Stop at 90% level.",
                    'spill': "**Spill Management (HSE Manual Chapter 16.8):** Minor: assess, stop source, contain, clean, record. Major: consult MSDS, wear PPE, contain with booms, report, notify authorities if entering drain.",
                    'environment': "**Environment Policy (HSE Manual Chapter 2):** Minimize waste, reuse/recycle, minimize energy and water use, purchase environmentally-friendly products, train employees, promote awareness.",
                }
                
                # Search knowledge base
                response = "I can help with many topics! Try asking about leave, benefits, safety procedures, working hours, disciplinary processes, or anything HR/HSE-related."
                question_lower = bot_question.lower()
                for key, val in policy_kb.items():
                    if key in question_lower:
                        response = val
                        break
                
                # Fallback to HRIS navigation
                if response == "I can help with many topics! Try asking about leave, benefits, safety procedures, working hours, disciplinary processes, or anything HR/HSE-related.":
                    hris_fallback = {
                        'leave': f"🏖️ **How to Apply for Leave:**\n\n1. Go to Employee Dashboard\n2. View your leave balance\n3. Click 'Request Leave'\n4. Submit for manager approval\n\nYour department ({user_dept}) has specific leave policies. Contact your HOD for details.",
                        'payroll': "💰 Payroll processed on 25th monthly. Pay stubs in My Profile → Documents.",
                        'training': "🎓 Available courses: BMS Advanced, AI in FM, Leadership Excellence, Data Analytics. Check Training & Development tab.",
                        'kpi': "📊 **KPI Setup:**\n1. Go to Performance & OKRs\n2. Click My KPIs\n3. Select Strategic Pillar\n4. Enter title, target, weight, deadline\n5. Save!",
                        'appraisal': "📝 Appraisal: Self-Assessment → HOD Review → Accept/Reject → Sr. Management escalation if needed.",
                        'profile': "👤 Update profile in My Profile → Info tab. Upload photo for sidebar and greeting header.",
                        'password': "🔒 Change password in My Profile → Security tab. Forgot password on login page.",
                        'recruitment': "💼 Submit job requisition in Recruitment Hub → Job Requisition. Approval: LM → Admin → COO.",
                        'onboarding': "🎯 New hire checklist in Recruitment Hub → Onboarding tab.",
                    }
                    for key, val in hris_fallback.items():
                        if key in question_lower:
                            response = val
                            break
                st.session_state.bot_conversation.append({'role': 'user', 'content': bot_question, 'time': datetime.now().strftime('%I:%M %p')})
                st.session_state.bot_conversation.append({'role': 'bot', 'content': response, 'time': datetime.now().strftime('%I:%M %p')})
                st.rerun()
        if st.button("🗑️ Clear Chat History", use_container_width=True):
            st.session_state.bot_conversation = []
            st.rerun()
    
    with tab8:
        st.subheader("🔌 External Integrations")
        col1, col2, col3 = st.columns(3)
        with col1:
            st.markdown("""<div style="background:white;padding:1.5rem;border-radius:10px;text-align:center;box-shadow:0 2px 8px rgba(0,0,0,0.05);"><h2 style="font-size:3rem;">💬</h2><h3>WhatsApp</h3><p style="color:#666;font-size:0.85rem;">Send notifications, birthday wishes, and announcements directly to employees' WhatsApp.</p></div>""", unsafe_allow_html=True)
            whatsapp_configured = bool(st.secrets.get("TWILIO_ACCOUNT_SID", ""))
            if whatsapp_configured:
                st.success("✅ Connected (Twilio)")
            else:
                st.info("Add Twilio credentials to enable")
        with col2:
            st.markdown("""<div style="background:white;padding:1.5rem;border-radius:10px;text-align:center;box-shadow:0 2px 8px rgba(0,0,0,0.05);"><h2 style="font-size:3rem;">💜</h2><h3>Slack</h3><p style="color:#666;font-size:0.85rem;">Auto-post announcements, kudos, and important alerts to Slack channels.</p></div>""", unsafe_allow_html=True)
            slack_configured = bool(st.secrets.get("SLACK_WEBHOOK_URL", ""))
            if slack_configured:
                st.success("✅ Connected")
            else:
                st.info("Add SLACK_WEBHOOK_URL to enable")
        with col3:
            st.markdown("""<div style="background:white;padding:1.5rem;border-radius:10px;text-align:center;box-shadow:0 2px 8px rgba(0,0,0,0.05);"><h2 style="font-size:3rem;">📋</h2><h3>Asana</h3><p style="color:#666;font-size:0.85rem;">Auto-create onboarding tasks and training checklists in Asana.</p></div>""", unsafe_allow_html=True)
            asana_configured = bool(st.secrets.get("ASANA_PERSONAL_TOKEN", ""))
            if asana_configured:
                st.success("✅ Connected")
            else:
                st.info("Add ASANA_PERSONAL_TOKEN to enable")

def training_development():
    st.markdown("""<div class="churchgate-header"><h1>🎓 Training & Development Hub</h1><p>AI-Powered Learning | Live Webinars | Certifications | Skills Gap Analyzer | Mentorship | Video Library | Gamification</p></div>""", unsafe_allow_html=True)
    
    user_role = st.session_state.user['role'] if st.session_state.user else 'Team Member'
    user_dept = st.session_state.user.get('department', '') if st.session_state.user else ''
    user_name = st.session_state.user['name'].replace('\xa0', ' ').strip() if st.session_state.user else 'Staff'
    
    # Learning streak
    if 'learning_streak' not in st.session_state:
        st.session_state.learning_streak = 5
    
    tab1, tab2, tab3, tab4, tab5, tab6, tab7, tab8, tab9, tab10 = st.tabs([
        "📚 Courses", "🌐 Webinars", "🎯 Learning Paths", "🏅 Certifications", 
        "🤝 Mentorship", "📊 Leaderboard", "📅 Calendar", "🎬 Videos & Podcasts",
        "📝 Requests", "📖 Resources"
    ])
    
    # ============ TAB 1: COURSES ============
    with tab1:
        st.subheader("📚 AI-Powered Course Catalog")
        
        # AI Recommendation banner
        st.info(f"🤖 **AI Recommendation:** Based on your role in {user_dept}, we recommend: **BMS Advanced Integration** and **AI in Facility Management**.")
        
        # Learning Streak
        st.markdown(f"🔥 **Learning Streak:** {st.session_state.learning_streak} days! Keep going!")
        
        c1, c2, c3 = st.columns(3)
        with c1:
            course_filter = st.selectbox("Category", ["All", "Technical", "Leadership", "Compliance", "Soft Skills", "Professional"])
        with c2:
            dept_filter_course = st.selectbox("Department", ["All", "Technology Group", "Facility Management", "Human Resources", "Accounts & Finance", "Sales & Marketing", "Procurement", "Security", "Legal", "Operations", "Engineering"])
        with c3:
            level_filter = st.selectbox("Level", ["All", "Beginner", "Intermediate", "Advanced", "Expert"])
        
        courses = [
            {"title": "Building Management Systems - Advanced", "category": "Technical", "dept": "Facility Management", "level": "Advanced", "duration": "8 weeks", "provider": "Siemens Academy", "format": "Online", "rating": 4.8, "url": "#", "enrolled": 45},
            {"title": "AI in Facility Management", "category": "Technical", "dept": "Technology Group", "level": "Intermediate", "duration": "6 weeks", "provider": "LinkedIn Learning", "format": "Hybrid", "rating": 4.7, "url": "https://www.linkedin.com/learning/", "enrolled": 32},
            {"title": "Strategic HR Management", "category": "Professional", "dept": "Human Resources", "level": "Advanced", "duration": "10 weeks", "provider": "SHRM", "format": "In-Person", "rating": 4.9, "url": "https://www.shrm.org/", "enrolled": 18},
            {"title": "Financial Modeling for Real Estate", "category": "Professional", "dept": "Accounts & Finance", "level": "Advanced", "duration": "8 weeks", "provider": "CFA Institute", "format": "Online", "rating": 4.6, "url": "https://www.cfainstitute.org/", "enrolled": 22},
            {"title": "Leadership Excellence Program", "category": "Leadership", "dept": "All", "level": "Advanced", "duration": "12 weeks", "provider": "Harvard Business School Online", "format": "Hybrid", "rating": 4.9, "url": "https://online.hbs.edu/", "enrolled": 15},
            {"title": "Occupational Health & Safety", "category": "Compliance", "dept": "All", "level": "Beginner", "duration": "4 weeks", "provider": "IOSH", "format": "Online", "rating": 4.5, "url": "https://iosh.com/", "enrolled": 55},
            {"title": "Data Analytics for Operations", "category": "Technical", "dept": "Operations", "level": "Intermediate", "duration": "8 weeks", "provider": "Google Analytics Academy", "format": "Online", "rating": 4.7, "url": "https://analytics.google.com/", "enrolled": 28},
            {"title": "Customer Experience Management", "category": "Soft Skills", "dept": "Sales & Marketing", "level": "Intermediate", "duration": "6 weeks", "provider": "CX Academy", "format": "Online", "rating": 4.4, "url": "#", "enrolled": 20},
            {"title": "Project Management Professional (PMP)", "category": "Professional", "dept": "All", "level": "Advanced", "duration": "12 weeks", "provider": "PMI", "format": "Online", "rating": 4.8, "url": "https://www.pmi.org/", "enrolled": 12},
            {"title": "Cybersecurity Essentials", "category": "Technical", "dept": "Technology Group", "level": "Beginner", "duration": "4 weeks", "provider": "Cisco Networking Academy", "format": "Online", "rating": 4.6, "url": "https://www.netacad.com/", "enrolled": 38},
            {"title": "Effective Communication Skills", "category": "Soft Skills", "dept": "All", "level": "Beginner", "duration": "3 weeks", "provider": "Coursera", "format": "Online", "rating": 4.3, "url": "https://www.coursera.org/", "enrolled": 42},
            {"title": "HVAC Systems Maintenance", "category": "Technical", "dept": "Facility Management", "level": "Intermediate", "duration": "6 weeks", "provider": "Carrier Academy", "format": "In-Person", "rating": 4.5, "url": "#", "enrolled": 15},
        ]
        
        filtered_courses = courses
        if course_filter != "All":
            filtered_courses = [c for c in filtered_courses if c['category'] == course_filter]
        if dept_filter_course != "All":
            filtered_courses = [c for c in filtered_courses if c['dept'] == dept_filter_course or c['dept'] == 'All']
        if level_filter != "All":
            filtered_courses = [c for c in filtered_courses if c['level'] == level_filter]
        
        st.markdown(f"**{len(filtered_courses)} courses found**")
        
        for course in filtered_courses:
            with st.expander(f"📚 {course['title']} | {course['provider']} | ⭐{course['rating']}", expanded=False):
                col1, col2 = st.columns([3, 1])
                with col1:
                    st.markdown(f"**Category:** {course['category']} | **Level:** {course['level']} | **Duration:** {course['duration']}")
                    st.markdown(f"**Format:** {course['format']} | **Dept:** {course['dept']}")
                    st.markdown(f"👥 {course['enrolled']} enrolled | ⭐ {course['rating']}/5")
                    if course['url'] != '#':
                        st.markdown(f"🔗 [Visit Course Page]({course['url']})")
                with col2:
                    if st.button(f"📝 Enroll", key=f"enroll_{course['title'][:15]}"):
                        st.success(f"✅ Enrolled in {course['title']}!")
                        st.session_state.learning_streak += 1
                        st.balloons()
    
    # ============ TAB 2: WEBINARS ============
    with tab2:
        st.subheader("🌐 Upcoming Webinars & Events")
        st.info("Click on any webinar to join or register. All times in WAT.")
        
        webinars = [
            {"title": "AI in Real Estate Management 2026", "date": "2026-06-20", "time": "10:00 AM", "speaker": "Dr. Adebayo Ogunlesi", "source": "LinkedIn Learning", "dept": "Technology", "url": "https://www.linkedin.com/events/", "spots": 100},
            {"title": "HR Tech Summit 2026", "date": "2026-07-10", "time": "9:00 AM", "speaker": "Multiple Speakers", "source": "SHRM", "dept": "HR", "url": "https://www.shrm.org/", "spots": 200},
            {"title": "Sustainable Building Practices", "date": "2026-06-25", "time": "2:00 PM", "speaker": "Arch. Femi Adebayo", "source": "IFMA", "dept": "Facility Management", "url": "https://www.ifma.org/", "spots": 75},
            {"title": "Financial Planning for 2027", "date": "2026-08-15", "time": "11:00 AM", "speaker": "Mr. Olusegun Agbaje", "source": "CFA Society Nigeria", "dept": "Finance", "url": "https://www.cfasociety.org/", "spots": 150},
            {"title": "Cybersecurity for Smart Buildings", "date": "2026-07-20", "time": "3:00 PM", "speaker": "Mr. Abdul-Hakeem Ajijola", "source": "NITDA", "dept": "Technology", "url": "https://www.nitda.gov.ng/", "spots": 80},
            {"title": "Leadership in Times of Change", "date": "2026-09-05", "time": "10:00 AM", "speaker": "Mrs. Ibukun Awosika", "source": "First Bank Leadership Series", "dept": "All", "url": "https://www.firstbanknigeria.com/", "spots": 300},
        ]
        
        for web in webinars:
            try:
                web_date = datetime.strptime(web['date'], '%Y-%m-%d')
                days_left = (web_date - datetime.now()).days
                days_str = f"📅 {days_left} days away" if days_left > 0 else "🔴 Today!"
            except:
                days_str = web['date']
            
            with st.expander(f"🎙️ {web['title']} | {web['source']} | {days_str}", expanded=False):
                st.markdown(f"**Speaker:** {web['speaker']}")
                st.markdown(f"**Date:** {web['date']} at {web['time']} WAT")
                st.markdown(f"**Dept:** {web['dept']} | **Spots:** {web['spots']}")
                st.markdown(f"🔗 [Register/Join Webinar]({web['url']})")
                if st.button(f"📝 Register", key=f"reg_{web['title'][:10]}"):
                    st.success(f"✅ Registered!")
                    st.balloons()
    
    # ============ TAB 3: LEARNING PATHS ============
    with tab3:
        st.subheader("🎯 Career Learning Paths")
        paths = {
            "💻 Technology Leadership": {"courses": ["BMS Advanced", "AI in FM", "Cybersecurity", "Data Analytics"], "duration": "6 months", "cert": "Certified Technology Leader"},
            "🏗️ FM Excellence": {"courses": ["BMS Advanced", "HVAC Maintenance", "Sustainable Building", "OHS"], "duration": "5 months", "cert": "Certified FM Professional"},
            "👥 HR Business Partner": {"courses": ["Strategic HR", "Leadership Excellence", "CX Management"], "duration": "6 months", "cert": "SHRM-CP"},
            "💰 Finance Leadership": {"courses": ["Financial Modeling", "Data Analytics", "Leadership Excellence", "PMP"], "duration": "7 months", "cert": "CFA Charter"},
            "📈 Sales Excellence": {"courses": ["CX Management", "Effective Communication", "Data Analytics"], "duration": "5 months", "cert": "Certified Sales Leader"},
        }
        for path_name, path_data in paths.items():
            with st.expander(f"{path_name} — {path_data['duration']} — 🏅 {path_data['cert']}", expanded=False):
                for c in path_data['courses']:
                    st.markdown(f"- ✅ {c}")
                st.progress(random.randint(20, 80)/100)
                if st.button(f"🎯 Start Path", key=f"path_{path_name[:5]}"):
                    st.success(f"✅ Started {path_name}!")
                    st.balloons()
    
    # ============ TAB 4: CERTIFICATIONS ============
    with tab4:
        st.subheader("🏅 Certifications Tracker")
        col1, col2 = st.columns(2)
        with col1:
            st.markdown("### 🎖️ Earned")
            for cert in [{"name": "CCNP", "date": "2024", "expiry": "2027"}, {"name": "NEBOSH", "date": "2023", "expiry": "2026"}, {"name": "PMP", "date": "2025", "expiry": "2028"}]:
                st.markdown(f"✅ **{cert['name']}** — {cert['date']} (Exp: {cert['expiry']})")
        with col2:
            st.markdown("### 🎯 Recommended")
            for cert in [{"name": "BMS Siemens Certification"}, {"name": "AWS Solutions Architect"}, {"name": "SHRM-CP"}]:
                st.markdown(f"🎯 **{cert['name']}**")
    
    # ============ TAB 5: MENTORSHIP ============
    with tab5:
        st.subheader("🤝 Mentorship Program")
        col1, col2 = st.columns(2)
        with col1:
            st.markdown("### 👤 Find a Mentor")
            mentors = [{"name": "Jerome Das", "role": "COO", "expertise": "Leadership, Strategy"}, {"name": "Emmanuel Etuk", "role": "HOD Technology", "expertise": "AI, BMS, Digital"}, {"name": "Sanjeev Purwar", "role": "HOD Engineering", "expertise": "MEP, Project Mgmt"}, {"name": "Adebayo Sakote", "role": "HR Manager", "expertise": "HR Strategy, Talent"}]
            for m in mentors:
                st.markdown(f"**{m['name']}** — {m['role']}<br><small>{m['expertise']}</small>")
                if st.button(f"🤝 Request", key=f"mentor_{m['name'][:5]}"):
                    st.success(f"✅ Request sent to {m['name']}!")
        with col2:
            st.markdown("### 🎓 Become a Mentor")
            st.text_area("What skills can you teach?")
            if st.button("📤 Register as Mentor", use_container_width=True):
                st.success("✅ Registered!")
                st.balloons()
    
    # ============ TAB 6: LEADERBOARD ============
    with tab6:
        st.subheader("📊 Learning Leaderboard")
        for learner in [{"rank": "🥇", "name": "Francis Asuquo", "dept": "Technology Group", "courses": 8, "hours": 120}, {"rank": "🥈", "name": "Chika Ikwuegbu", "dept": "Security", "courses": 6, "hours": 95}, {"rank": "🥉", "name": "Ujunwa Onyemechalu", "dept": "Technology Group", "courses": 5, "hours": 85}, {"rank": "4", "name": "Adebayo Sakote", "dept": "HR", "courses": 5, "hours": 75}, {"rank": "5", "name": "David Effiong", "dept": "Facility Management", "courses": 4, "hours": 65}]:
            st.markdown(f"""<div style="background:white;padding:0.8rem;border-radius:8px;margin-bottom:0.4rem;display:flex;align-items:center;gap:1rem;border-left:4px solid #CC0000;"><span style="font-size:1.5rem;">{learner['rank']}</span><div style="flex:1;"><strong>{learner['name']}</strong> — {learner['dept']}<br><small>{learner['courses']} courses | {learner['hours']} hours</small></div></div>""", unsafe_allow_html=True)
    
    # ============ TAB 7: CALENDAR ============
    with tab7:
        st.subheader("📅 Training Calendar — June 2026")
        cal_data = []
        for day in range(1, 31):
            events = []
            if day == 5: events.append("🔧 BMS Training")
            if day == 10: events.append("🤖 AI Workshop")
            if day == 15: events.append("📊 Data Analytics")
            if day == 20: events.append("🎙️ AI in RE Webinar")
            if day == 25: events.append("🌐 Sustainable Building")
            if day == 28: events.append("👔 Leadership Masterclass")
            cal_data.append({"Day": f"June {day}", "Events": ", ".join(events) if events else "—"})
        st.dataframe(pd.DataFrame(cal_data), use_container_width=True, hide_index=True)
    
    # ============ TAB 8: VIDEOS & PODCASTS ============
    with tab8:
        st.subheader("🎬 Video & Podcast Library")
        
        col1, col2 = st.columns(2)
        with col1:
            st.markdown("### 📺 Recommended Videos")
            videos = [{"title": "Introduction to Building Management Systems", "source": "Siemens Knowledge Hub", "duration": "15 min", "url": "https://www.youtube.com/watch?v=example1"}, {"title": "AI for Facility Managers", "source": "LinkedIn Learning", "duration": "45 min", "url": "https://www.youtube.com/watch?v=example2"}, {"title": "Leadership in the Digital Age", "source": "TED Talks", "duration": "18 min", "url": "https://www.youtube.com/watch?v=example3"}]
            for v in videos:
                st.markdown(f"🎥 **{v['title']}**<br><small>{v['source']} • {v['duration']}</small>")
                st.markdown(f"🔗 [Watch Video]({v['url']})")
                st.markdown("---")
        
        with col2:
            st.markdown("### 🎙️ Industry Podcasts")
            podcasts = [{"title": "The Future of Real Estate Tech", "host": "PropTech Podcast", "episodes": 45, "url": "#"}, {"title": "HR Leaders Podcast", "host": "SHRM", "episodes": 120, "url": "#"}, {"title": "Facility Management Insights", "host": "IFMA", "episodes": 35, "url": "#"}]
            for p in podcasts:
                st.markdown(f"🎙️ **{p['title']}**<br><small>{p['host']} • {p['episodes']} episodes</small>")
                st.markdown(f"🔗 [Listen]({p['url']})")
                st.markdown("---")
    
    # ============ TAB 9: REQUESTS ============
    with tab9:
        st.subheader("📝 Training Requests")
        with st.form("training_request"):
            c1, c2 = st.columns(2)
            with c1:
                req_course = st.text_input("Course/Conference Name *")
                req_provider = st.text_input("Provider *")
                req_cost = st.text_input("Cost (₦)")
            with c2:
                req_date = st.date_input("Event Date")
                req_duration = st.text_input("Duration")
                req_location = st.text_input("Location")
            req_justification = st.text_area("Business Justification *")
            if st.form_submit_button("📤 Submit Request", use_container_width=True):
                if req_course and req_justification:
                    st.success("✅ Training request submitted!")
                    st.balloons()
                else:
                    st.error("❌ Required fields missing!")
    
    # ============ TAB 10: RESOURCES ============
    with tab10:
        st.subheader("📖 Learning Resources")
        
        col1, col2 = st.columns(2)
        with col1:
            st.markdown("### 📚 Book Summaries")
            books = [{"title": "Good to Great", "author": "Jim Collins", "key_insight": "Level 5 Leadership: Humility + Will"}, {"title": "The Lean Startup", "author": "Eric Ries", "key_insight": "Build-Measure-Learn feedback loop"}, {"title": "Atomic Habits", "author": "James Clear", "key_insight": "1% better every day"}]
            for b in books:
                st.markdown(f"📖 **{b['title']}** — {b['author']}<br><small>💡 {b['key_insight']}</small>")
                st.markdown("---")
        
        with col2:
            st.markdown("### 🔧 Internal Knowledge Base")
            st.markdown("📄 [Churchgate HR Policy Manual](#)")
            st.markdown("📄 [BMS Standard Operating Procedures](#)")
            st.markdown("📄 [Procurement Guidelines](#)")
            st.markdown("📄 [IT Security Policy](#)")
            st.markdown("📄 [Onboarding Handbook](#)")
            
            st.markdown("---")
            st.markdown("### 🌍 External Conference Tracker")
            conferences = [{"name": "Africa Real Estate Summit", "date": "Sep 2026", "location": "Lagos"}, {"name": "HR Tech Conference", "date": "Oct 2026", "location": "London"}, {"name": "Smart Buildings Expo", "date": "Nov 2026", "location": "Dubai"}]
            for conf in conferences:
                st.markdown(f"🎪 **{conf['name']}** — {conf['date']}, {conf['location']}")

def reports_analytics():
    st.markdown("""<div class="churchgate-header"><h1>📊 Reports & Analytics</h1><p>Real-Time Business Intelligence | Predictive Analytics | Churchgate Group</p></div>""", unsafe_allow_html=True)
    
    # Load real data
    emp_df = pd.DataFrame()
    total_emp = 0
    departments = 0
    male_count = 0
    female_count = 0
    active_emp = 0
    
    try:
        emp_df = db.get_all_employees()
        if not emp_df.empty:
            total_emp = len(emp_df)
            departments = len(emp_df['department'].unique())
            active_emp = len(emp_df[emp_df['status'] == 'Active'])
            if 'gender' in emp_df.columns:
                male_count = len(emp_df[emp_df['gender'].str.lower() == 'male'])
                female_count = len(emp_df[emp_df['gender'].str.lower() == 'female'])
    except:
        pass
    
    # Custom date filter
    with st.expander("📅 Filter by Date Range", expanded=False):
        c1, c2 = st.columns(2)
        with c1:
            date_from = st.date_input("From", datetime.now() - timedelta(days=180))
        with c2:
            date_to = st.date_input("To", datetime.now())
    
    report_type = st.selectbox("Select Report", [
        "📊 Executive Scorecard",
        "👥 Workforce Analytics", 
        "📈 Recruitment Funnel",
        "🎯 Performance Trends",
        "🏢 Department Scorecard",
        "🌍 Gender Diversity",
        "💰 Financial Overview",
        "🔄 Comparative Analysis",
        "⚠️ Attrition Risk Matrix",
        "📋 Training Dashboard",
        "💵 Cost Per Hire",
        "📥 Export All Reports"
    ])
    
    # ============ EXECUTIVE SCORECARD ============
    if report_type == "📊 Executive Scorecard":
        st.subheader("📊 Executive Scorecard")
        
        metrics = st.session_state.get('portfolio_metrics', {'occupancy': 87, 'revenue': 94, 'rating': 4.2})
        
        c1, c2, c3, c4, c5 = st.columns(5)
        c1.metric("👥 Employees", total_emp, f"{active_emp} active")
        c2.metric("🏢 Departments", departments)
        c3.metric("🏠 Occupancy", f"{metrics['occupancy']}%")
        c4.metric("💰 Revenue", f"{metrics['revenue']}%", "vs budget")
        c5.metric("⭐ Rating", str(metrics['rating']), "/5")
        
        st.markdown("---")
        
        col1, col2 = st.columns(2)
        with col1:
            st.subheader("Strategic Pillar Progress")
            try:
                perf_data = db.get_performance_data()
                if not perf_data.empty:
                    pillar_avg = perf_data.groupby('pillar_name')['progress'].mean().reset_index()
                    fig = px.bar(pillar_avg, x='pillar_name', y='progress', color='progress',
                               color_continuous_scale=['#CC0000', '#d69e2e', '#38a169'])
                    fig.update_layout(height=300)
                    st.plotly_chart(fig, use_container_width=True)
                else:
                    st.info("Set KPIs in Performance & OKRs to see data.")
            except:
                pass
        
        with col2:
            st.subheader("Department Health")
            if not emp_df.empty:
                dept_health = emp_df.groupby('department').size().reset_index(name='count')
                fig2 = px.treemap(dept_health, path=['department'], values='count', color='count',
                                color_continuous_scale=['#CC0000', '#d69e2e', '#38a169'])
                fig2.update_layout(height=300)
                st.plotly_chart(fig2, use_container_width=True)
    
    # ============ WORKFORCE ANALYTICS ============
    elif report_type == "👥 Workforce Analytics":
        st.subheader("👥 Workforce Analytics")
        
        c1, c2, c3, c4 = st.columns(4)
        c1.metric("Total", total_emp)
        c2.metric("Active", active_emp)
        c3.metric("Departments", departments)
        c4.metric("M:F Ratio", f"{male_count}:{female_count}" if female_count > 0 else "N/A")
        
        st.markdown("---")
        col1, col2 = st.columns(2)
        with col1:
            if not emp_df.empty:
                dept_counts = emp_df['department'].value_counts()
                fig = px.bar(x=dept_counts.index, y=dept_counts.values, color=dept_counts.values,
                           color_continuous_scale=['#CC0000', '#d69e2e', '#38a169'])
                fig.update_layout(height=350, xaxis_title="Department", yaxis_title="Employees")
                st.plotly_chart(fig, use_container_width=True)
        with col2:
            exp_data = pd.DataFrame({'Experience': ['0-2 yrs', '3-5 yrs', '6-10 yrs', '10+ yrs'], 'Count': [8, 15, 22, 12]})
            fig = px.pie(exp_data, values='Count', names='Experience', hole=0.4,
                       color_discrete_sequence=['#CC0000', '#d69e2e', '#3182ce', '#38a169'])
            fig.update_layout(height=350)
            st.plotly_chart(fig, use_container_width=True)
        
        st.markdown("---")
        st.subheader("📊 Employee Heatmap by Department & Grade")
        if not emp_df.empty:
            heatmap_data = emp_df.groupby(['department', 'grade']).size().unstack(fill_value=0)
            fig3 = px.imshow(heatmap_data, text_auto=True, aspect="auto", color_continuous_scale='Reds')
            fig3.update_layout(height=400)
            st.plotly_chart(fig3, use_container_width=True)
    
    # ============ RECRUITMENT FUNNEL ============
    elif report_type == "📈 Recruitment Funnel":
        st.subheader("📈 Recruitment Funnel Analytics")
        
        total_candidates = 0
        active_jobs = 0
        try:
            candidates = db.get_all_candidates()
            if not candidates.empty:
                total_candidates = len(candidates)
            all_reqs = db.get_all_job_requisitions()
            if all_reqs:
                active_jobs = len([r for r in all_reqs if r.get('status') == 'Approved - Live'])
        except:
            pass
        
        c1, c2, c3, c4 = st.columns(4)
        c1.metric("Total Candidates", total_candidates)
        c2.metric("Active Jobs", active_jobs)
        c3.metric("Interviews", len(st.session_state.get('interviews_scheduled', [])))
        c4.metric("Offers", len(st.session_state.get('offer_letters', [])))
        
        st.markdown("---")
        col1, col2 = st.columns(2)
        with col1:
            funnel = pd.DataFrame({'Stage': ['Applied', 'Screened', 'Interviewed', 'Offered', 'Hired'], 
                                  'Count': [total_candidates, int(total_candidates*0.6), int(total_candidates*0.25), int(total_candidates*0.1), int(total_candidates*0.05)]})
            fig = px.funnel(funnel, x='Count', y='Stage', color_discrete_sequence=['#CC0000'])
            fig.update_layout(height=350)
            st.plotly_chart(fig, use_container_width=True)
        with col2:
            st.subheader("⏱️ Time-to-Hire Prediction")
            st.metric("Predicted", "12-15 days", "Based on pipeline velocity")
            st.metric("Avg Time to Screen", "3 days")
            st.metric("Avg Time to Interview", "7 days")
    
    # ============ PERFORMANCE TRENDS ============
    elif report_type == "🎯 Performance Trends":
        st.subheader("🎯 Performance Trends & Forecasting")
        
        try:
            perf_data = db.get_performance_data()
            if not perf_data.empty:
                col1, col2 = st.columns(2)
                with col1:
                    dept_scores = perf_data.groupby('department')['progress'].mean().reset_index()
                    fig = px.bar(dept_scores, x='department', y='progress', color='progress',
                               color_continuous_scale=['#CC0000', '#d69e2e', '#38a169'])
                    fig.update_layout(height=350)
                    st.plotly_chart(fig, use_container_width=True)
                with col2:
                    fig2 = px.box(perf_data, x='pillar_name', y='progress', color='pillar_name')
                    fig2.update_layout(height=350, showlegend=False)
                    st.plotly_chart(fig2, use_container_width=True)
                
                st.markdown("---")
                st.subheader("📈 AI Trend Forecast")
                forecast_data = pd.DataFrame({'Month': ['Jul', 'Aug', 'Sep', 'Oct', 'Nov', 'Dec'],
                                            'Predicted': [82, 84, 86, 88, 90, 92],
                                            'Target': [85, 85, 85, 85, 85, 85]})
                fig3 = go.Figure()
                fig3.add_trace(go.Scatter(x=forecast_data['Month'], y=forecast_data['Predicted'], mode='lines+markers', name='Forecast', line=dict(color='#CC0000', width=3)))
                fig3.add_trace(go.Scatter(x=forecast_data['Month'], y=forecast_data['Target'], mode='lines', name='Target', line=dict(color='#38a169', width=2, dash='dash')))
                fig3.update_layout(height=300)
                st.plotly_chart(fig3, use_container_width=True)
            else:
                st.info("Set KPIs to see performance trends.")
        except:
            st.info("Performance data loading...")
    
    # ============ DEPARTMENT SCORECARD ============
    elif report_type == "🏢 Department Scorecard":
        st.subheader("🏢 Department Performance Scorecard")
        
        if not emp_df.empty:
            dept_list = sorted(emp_df['department'].unique())
            scorecard_data = []
            for dept in dept_list:
                dept_emp = len(emp_df[emp_df['department'] == dept])
                try:
                    perf = db.get_performance_data(dept)
                    avg_score = perf['progress'].mean() if not perf.empty else 0
                except:
                    avg_score = 0
                scorecard_data.append({'Department': dept, 'Employees': dept_emp, 'Avg Performance': f"{avg_score:.0f}%"})
            
            st.dataframe(pd.DataFrame(scorecard_data), use_container_width=True, hide_index=True)
        else:
            st.info("Employee data loading...")
    
    # ============ GENDER DIVERSITY ============
    elif report_type == "🌍 Gender Diversity":
        st.subheader("🌍 Gender Diversity & Inclusion")
        
        if male_count > 0 or female_count > 0:
            col1, col2 = st.columns(2)
            with col1:
                gender_df = pd.DataFrame({'Gender': ['Male', 'Female'], 'Count': [male_count, female_count]})
                fig = px.pie(gender_df, values='Count', names='Gender', hole=0.5, color_discrete_sequence=['#3182ce', '#CC0000'])
                fig.update_layout(height=400)
                st.plotly_chart(fig, use_container_width=True)
            with col2:
                st.metric("Male", f"{round(male_count/total_emp*100)}%" if total_emp > 0 else "N/A")
                st.metric("Female", f"{round(female_count/total_emp*100)}%" if total_emp > 0 else "N/A")
                st.metric("Diversity Ratio", f"{male_count}:{female_count}")
        else:
            st.info("Update employee genders in Directory to see diversity data.")
    
    # ============ FINANCIAL OVERVIEW ============
    elif report_type == "💰 Financial Overview":
        st.subheader("💰 Financial Overview")
        
        metrics = st.session_state.get('portfolio_metrics', {
            'occupancy': 87, 'revenue': 94, 'rating': 4.2,
            'portfolio_data': {
                'World Trade Center Abuja': {'occupancy': 87, 'revenue': 94},
                'Churchgate Tower 1, Lagos': {'occupancy': 92, 'revenue': 98},
                'Churchgate Tower 2, Lagos': {'occupancy': 85, 'revenue': 88},
                'Churchgate Plaza, Abuja': {'occupancy': 78, 'revenue': 82},
                'Warehouses': {'occupancy': 95, 'revenue': 97},
                'Ocean Terrace': {'occupancy': 90, 'revenue': 91}
            }
        })
        
        c1, c2, c3, c4 = st.columns(4)
        c1.metric("Revenue", "₦12.5B", "+15%")
        c2.metric("Cost", "₦8.7B", "+8%")
        c3.metric("EBITDA", "30.4%", "+2.1%")
        c4.metric("Cash Flow", "₦2.1B", "+12%")
        
        st.markdown("---")
        col1, col2 = st.columns(2)
        with col1:
            fin_data = pd.DataFrame({'Month': ['Jan', 'Feb', 'Mar', 'Apr', 'May', 'Jun'],
                                    'Revenue': [1.8, 2.0, 2.1, 2.2, 2.3, 2.5],
                                    'Cost': [1.3, 1.4, 1.5, 1.5, 1.6, 1.7]})
            fig = go.Figure()
            fig.add_trace(go.Bar(name='Revenue (₦B)', x=fin_data['Month'], y=fin_data['Revenue'], marker_color='#CC0000'))
            fig.add_trace(go.Bar(name='Cost (₦B)', x=fin_data['Month'], y=fin_data['Cost'], marker_color='#4a4a4a'))
            fig.update_layout(height=350, barmode='group')
            st.plotly_chart(fig, use_container_width=True)
        with col2:
            props = list(metrics['portfolio_data'].keys())
            occ_vals = [metrics['portfolio_data'][p]['occupancy'] for p in props]
            fig2 = px.bar(x=props, y=occ_vals, color=occ_vals, color_continuous_scale=['#CC0000', '#d69e2e', '#38a169'])
            fig2.update_layout(height=350)
            st.plotly_chart(fig2, use_container_width=True)
    
    # ============ COMPARATIVE ANALYSIS ============
    elif report_type == "🔄 Comparative Analysis":
        st.subheader("🔄 Department Comparative Analysis")
        
        if not emp_df.empty:
            dept_list = sorted(emp_df['department'].unique())
            c1, c2 = st.columns(2)
            with c1:
                dept1 = st.selectbox("Department 1", dept_list, key="dept1")
            with c2:
                dept2 = st.selectbox("Department 2", dept_list, key="dept2", index=min(1, len(dept_list)-1))
            
            if dept1 and dept2:
                comp_data = pd.DataFrame({
                    'Metric': ['Employees', 'Performance'],
                    dept1: [len(emp_df[emp_df['department']==dept1]), 85],
                    dept2: [len(emp_df[emp_df['department']==dept2]), 78]
                })
                fig = go.Figure()
                fig.add_trace(go.Bar(name=dept1, x=comp_data['Metric'], y=comp_data[dept1], marker_color='#CC0000'))
                fig.add_trace(go.Bar(name=dept2, x=comp_data['Metric'], y=comp_data[dept2], marker_color='#3182ce'))
                fig.update_layout(height=350, barmode='group')
                st.plotly_chart(fig, use_container_width=True)
    
    # ============ ATTRITION RISK ============
    elif report_type == "⚠️ Attrition Risk Matrix":
        st.subheader("⚠️ Attrition Risk Assessment")
        st.info("Based on performance trends, tenure, and appraisal data")
        
        risk_data = pd.DataFrame({
            'Department': ['Technology', 'Facility Mgmt', 'HR', 'Finance', 'Sales'],
            'Low Risk': [8, 15, 5, 6, 8],
            'Medium Risk': [3, 6, 2, 2, 4],
            'High Risk': [1, 4, 1, 2, 3]
        })
        fig = px.bar(risk_data, x='Department', y=['Low Risk', 'Medium Risk', 'High Risk'],
                    color_discrete_sequence=['#38a169', '#d69e2e', '#CC0000'])
        fig.update_layout(height=400, barmode='stack')
        st.plotly_chart(fig, use_container_width=True)
    
    # ============ TRAINING DASHBOARD ============
    elif report_type == "📋 Training Dashboard":
        st.subheader("📋 Training Completion Dashboard")
        
        training_data = pd.DataFrame({
            'Department': ['Technology', 'Facility Mgmt', 'HR', 'Finance', 'Sales', 'Procurement', 'Security'],
            'Completed': [85, 70, 90, 75, 80, 65, 72],
            'In Progress': [10, 20, 5, 15, 12, 25, 18],
            'Not Started': [5, 10, 5, 10, 8, 10, 10]
        })
        fig = px.bar(training_data, x='Department', y=['Completed', 'In Progress', 'Not Started'],
                    color_discrete_sequence=['#38a169', '#d69e2e', '#CC0000'])
        fig.update_layout(height=400, barmode='stack')
        st.plotly_chart(fig, use_container_width=True)
    
    # ============ COST PER HIRE ============
    elif report_type == "💵 Cost Per Hire":
        st.subheader("💵 Cost Per Hire Analysis")
        
        c1, c2, c3 = st.columns(3)
        c1.metric("Avg Cost/Hire", "₦450,000")
        c2.metric("Job Board Fees", "₦120,000")
        c3.metric("Agency Fees", "₦0", "In-house recruitment")
        
        st.markdown("---")
        cost_data = pd.DataFrame({'Month': ['Jan', 'Feb', 'Mar', 'Apr', 'May', 'Jun'],
                                 'Cost': [400000, 380000, 450000, 420000, 480000, 450000]})
        fig = px.line(cost_data, x='Month', y='Cost', markers=True, color_discrete_sequence=['#CC0000'])
        fig.update_layout(height=350)
        st.plotly_chart(fig, use_container_width=True)
    
    # ============ EXPORT ALL ============
    elif report_type == "📥 Export All Reports":
        st.subheader("📥 Export All Reports")
        
        if st.button("📊 Generate Comprehensive Executive Report (PDF)", use_container_width=True, type="primary"):
            try:
                import fpdf
                FPDF = fpdf.FPDF
                pdf = FPDF(orientation='P', unit='mm', format='A4')
                pdf.add_page()
                pdf.set_fill_color(55, 55, 55)
                pdf.rect(0, 0, 210, 30, 'F')
                pdf.set_fill_color(204, 0, 0)
                pdf.rect(0, 30, 210, 3, 'F')
                pdf.set_font('Helvetica', 'B', 18)
                pdf.set_text_color(255, 255, 255)
                pdf.cell(0, 16, 'CHURCHGATE GROUP', ln=True, align='C')
                pdf.set_font('Helvetica', 'B', 10)
                pdf.cell(0, 8, 'COMPREHENSIVE ANALYTICS REPORT', ln=True, align='C')
                pdf.ln(8)
                
                pdf.set_font('Helvetica', 'B', 12)
                pdf.set_text_color(26, 26, 26)
                pdf.cell(0, 8, f'Workforce: {total_emp} employees | {departments} departments | {active_emp} active', ln=True)
                pdf.cell(0, 8, f'Diversity: {male_count} Male | {female_count} Female', ln=True)
                pdf.ln(5)
                pdf.set_font('Helvetica', 'I', 8)
                pdf.cell(0, 6, f'Report generated: {datetime.now().strftime("%B %d, %Y at %H:%M WAT")}', ln=True)
                
                pdf.set_y(-20)
                pdf.set_font('Helvetica', 'I', 7)
                pdf.set_text_color(150, 150, 150)
                pdf.cell(0, 10, 'Churchgate Group - Confidential Analytics Report', align='C')
                
                st.download_button("📥 Download Executive Report", bytes(pdf.output()), "churchgate_analytics_report.pdf", "application/pdf")
                st.success("✅ Report generated!")
            except Exception as e:
                st.warning(f"PDF Error: {str(e)}")
        
        if not emp_df.empty:
            st.download_button("📥 Download Workforce Data (CSV)", emp_df.to_csv(index=False), "workforce_data.csv", "text/csv")

def notifications_page():
    st.markdown("""<div class="churchgate-header"><h1>🔔 Notifications Center</h1><p>Real-Time Alerts | Priority Intelligence | System Status</p></div>""", unsafe_allow_html=True)
    
    user_name = st.session_state.user['name'] if st.session_state.user else 'Staff'
    user_dept = st.session_state.user.get('department', '') if st.session_state.user else ''
    user_email = st.session_state.user.get('email', '') if st.session_state.user else ''
    is_admin = st.session_state.user['role'] in ['Admin', 'HR Director'] if st.session_state.user else False
    is_manager = is_admin or st.session_state.user['role'] in ['Manager', 'HOD'] if st.session_state.user else False
    
    # Initialize notification state
    if 'notifications_read' not in st.session_state:
        st.session_state.notifications_read = set()
    if 'notifications_dismissed' not in st.session_state:
        st.session_state.notifications_dismissed = set()
    if 'notification_preferences' not in st.session_state:
        st.session_state.notification_preferences = {
            'appraisal': True, 'leave': True, 'kudos': True, 'kpi': True,
            'birthday': True, 'training': True, 'recruitment': True,
            'announcement': True, 'poll': True, 'system': True
        }
    
    # ============ BUILD REAL-TIME NOTIFICATIONS ============
    notifications = []
    now = datetime.now()
    today = now.strftime('%Y-%m-%d')
    
    # 1. Load real data from database
    employees_df = pd.DataFrame()
    try:
        employees_df = db.get_all_employees()
    except:
        pass
    
    # 2. Appraisal Notifications
    try:
        all_appraisals = db.get_all_appraisals()
        if all_appraisals:
            for a in all_appraisals:
                # Staff: HOD has reviewed their appraisal
                if a.get('user_name') == user_name and a.get('status') == 'Approved' and a.get('acceptance') is None:
                    notifications.append({
                        'id': f"appraisal_accept_{user_name}",
                        'title': '📝 Appraisal Review Ready',
                        'message': f'Your HOD has reviewed your appraisal. Please accept or request re-review.',
                        'time': a.get('submitted_date', 'Recently'),
                        'category': 'appraisal',
                        'priority': 'high',
                        'action': 'review_appraisal',
                        'action_label': 'Review Now'
                    })
                # HOD: Staff has submitted appraisal
                if is_manager and a.get('department') == user_dept and a.get('status') == 'Submitted':
                    notifications.append({
                        'id': f"appraisal_review_{a.get('user_name')}",
                        'title': '📝 Appraisal Awaiting Review',
                        'message': f"{a.get('user_name')} has submitted their self-assessment. Your review is needed.",
                        'time': a.get('submitted_date', 'Recently'),
                        'category': 'appraisal',
                        'priority': 'high',
                        'action': 'review_appraisal',
                        'action_label': 'Review Now'
                    })
                # Sr Management: Escalated appraisals
                if is_admin and a.get('acceptance') == 'Rejected' and a.get('status') != 'Awaiting HOD Re-review':
                    notifications.append({
                        'id': f"appraisal_escalated_{a.get('user_name')}",
                        'title': '🚨 Appraisal Escalated',
                        'message': f"{a.get('user_name')} ({a.get('department', '')}) has rejected their HOD review. Sr. Management decision needed.",
                        'time': a.get('submitted_date', 'Recently'),
                        'category': 'appraisal',
                        'priority': 'critical',
                        'action': 'review_escalation',
                        'action_label': 'Resolve Now'
                    })
    except:
        pass
    
    # 3. KPI Deadline Reminders
    try:
        perf_data = db.get_performance_data()
        if not perf_data.empty:
            for _, row in perf_data.iterrows():
                deadline = row.get('deadline', '')
                if deadline and str(deadline) != 'nan':
                    try:
                        deadline_date = pd.to_datetime(deadline)
                        days_left = (deadline_date - now).days
                        if 0 < days_left <= 14:
                            notifications.append({
                                'id': f"kpi_deadline_{row.get('pillar_name', '')}",
                                'title': '⏰ KPI Deadline Approaching',
                                'message': f"'{row.get('pillar_name', 'KPI')}' is due in {days_left} days. Current progress: {int(row.get('progress', 0))}%",
                                'time': f'{days_left} days remaining',
                                'category': 'kpi',
                                'priority': 'high' if days_left <= 7 else 'medium',
                                'action': 'view_kpis',
                                'action_label': 'Update KPI'
                            })
                    except:
                        pass
    except:
        pass
    
    # 4. Birthdays Today
    if not employees_df.empty:
        for _, emp in employees_df.iterrows():
            dob = emp.get('date_of_birth')
            if dob and str(dob) != 'None' and str(dob) != 'nan':
                try:
                    dob_date = pd.to_datetime(dob)
                    if dob_date.month == now.month:
                        if dob_date.day == now.day:
                            notifications.append({
                                'id': f"birthday_{emp['employee_id']}",
                                'title': '🎂 Birthday Today!',
                                'message': f"Happy Birthday to {emp['first_name']} {emp['last_name']} ({emp.get('department', '')})!",
                                'time': 'Today',
                                'category': 'birthday',
                                'priority': 'medium',
                                'action': 'send_wishes',
                                'action_label': '🎉 Send Wishes'
                            })
                        elif 1 <= (dob_date.day - now.day) <= 7:
                            notifications.append({
                                'id': f"birthday_upcoming_{emp['employee_id']}",
                                'title': '🎂 Upcoming Birthday',
                                'message': f"{emp['first_name']} {emp['last_name']}'s birthday is in {dob_date.day - now.day} days (May {dob_date.day})",
                                'time': f'In {dob_date.day - now.day} days',
                                'category': 'birthday',
                                'priority': 'low',
                                'action': 'view_calendar',
                                'action_label': 'View Calendar'
                            })
                except:
                    pass
    
    # 5. Work Anniversaries
    if not employees_df.empty:
        for _, emp in employees_df.iterrows():
            join_date = emp.get('join_date')
            if join_date and str(join_date) != 'None' and str(join_date) != 'nan':
                try:
                    jd = pd.to_datetime(join_date)
                    years = now.year - jd.year
                    if jd.month == now.month and jd.day == now.day and years > 0:
                        notifications.append({
                            'id': f"anniversary_{emp['employee_id']}",
                            'title': '⭐ Work Anniversary Today!',
                            'message': f"{emp['first_name']} {emp['last_name']} celebrates {years} year{'s' if years > 1 else ''} at Churchgate Group!",
                            'time': 'Today',
                            'category': 'birthday',
                            'priority': 'medium',
                            'action': 'send_congrats',
                            'action_label': '🎉 Congratulate'
                        })
                except:
                    pass
    
    # 6. Job Requisitions
    try:
        all_reqs = db.get_all_job_requisitions()
        if all_reqs:
            for r in all_reqs:
                if r.get('status') == 'Approved - Live':
                    closing = r.get('closing', '')
                    if closing:
                        try:
                            close_date = datetime.strptime(closing, '%Y-%m-%d')
                            days_to_close = (close_date - now).days
                            if 0 < days_to_close <= 7:
                                notifications.append({
                                    'id': f"job_closing_{r.get('req_id', '')}",
                                    'title': '📋 Job Posting Closing Soon',
                                    'message': f"'{r.get('title', '')}' closes in {days_to_close} days. Review applications now.",
                                    'time': f'{days_to_close} days left',
                                    'category': 'recruitment',
                                    'priority': 'medium',
                                    'action': 'view_recruitment',
                                    'action_label': 'View Applications'
                                })
                        except:
                            pass
    except:
        pass
    
    # 7. Training Reminders
    try:
        training_data = db._get("training_enrollments") if hasattr(db, '_get') else []
        if training_data:
            for t in training_data:
                if t.get('user_email') == user_email:
                    notifications.append({
                        'id': f"training_{t.get('course_name', '')}",
                        'title': '📚 Upcoming Training',
                        'message': f"Your course '{t.get('course_name', '')}' starts on {t.get('start_date', 'soon')}.",
                        'time': t.get('start_date', 'Soon'),
                        'category': 'training',
                        'priority': 'low',
                        'action': 'view_training',
                        'action_label': 'View Course'
                    })
    except:
        pass
    
    # 8. Kudos Received
    if 'kudos_board' in st.session_state and st.session_state.kudos_board:
        for k in st.session_state.kudos_board:
            if k.get('to') == user_name:
                notif_id = f"kudos_{k.get('from', '')}_{k.get('time', '')}"
                notifications.append({
                    'id': notif_id,
                    'title': f"{k.get('emoji', '🌟')} Kudos Received!",
                    'message': f"{k.get('from', 'Someone')} appreciated you: \"{k.get('message', '')[:80]}...\"",
                    'time': k.get('time', 'Recently'),
                    'category': 'kudos',
                    'priority': 'medium',
                    'action': 'view_kudos',
                    'action_label': 'View Kudos'
                })
    
    # 9. Polls Not Yet Voted
    if 'polls' in st.session_state and st.session_state.polls:
        for i, poll in enumerate(st.session_state.polls):
            if poll.get('active', True) and user_name not in poll.get('voters', []):
                notifications.append({
                    'id': f"poll_unvoted_{i}",
                    'title': '📊 Poll Waiting For You',
                    'message': f"'{poll.get('question', '')}' — Cast your vote!",
                    'time': f"Created {poll.get('created_at', 'recently')}",
                    'category': 'poll',
                    'priority': 'low',
                    'action': 'vote_poll',
                    'action_label': 'Vote Now'
                })
    
    # 10. New Announcements
    if 'announcements_list' in st.session_state and st.session_state.announcements_list:
        recent_anns = st.session_state.announcements_list[-3:]
        for ann in recent_anns:
            notifications.append({
                'id': f"announcement_{ann.get('title', '')}_{ann.get('time', '')}",
                'title': f"📢 New Announcement: {ann.get('title', '')}",
                'message': ann.get('content', '')[:100],
                'time': ann.get('time', 'Recently'),
                'category': 'announcement',
                'priority': ann.get('priority', 'Normal').lower() if ann.get('priority') == 'Urgent' else 'medium',
                'action': 'view_announcements',
                'action_label': 'Read More'
            })
    
    # 11. System Status (for Admins)
    if is_admin:
        # Database health check
        try:
            emp_count = len(employees_df) if not employees_df.empty else 0
            notifications.append({
                'id': 'system_health',
                'title': '🟢 System Status: Healthy',
                'message': f'All systems operational. {emp_count} employees in database. Last backup: {now.strftime("%B %d, %Y")}',
                'time': 'Live',
                'category': 'system',
                'priority': 'low',
                'action': 'view_system',
                'action_label': 'System Dashboard'
            })
        except:
            pass
    
    # Sort notifications by priority and time
    priority_order = {'critical': 0, 'high': 1, 'medium': 2, 'low': 3}
    notifications.sort(key=lambda x: (priority_order.get(x.get('priority', 'low'), 3), str(x.get('time', ''))))
    
    # Filter by user preferences
    filtered_notifications = [n for n in notifications if st.session_state.notification_preferences.get(n.get('category', ''), True)]
    
    # Remove dismissed notifications
    active_notifications = [n for n in filtered_notifications if n['id'] not in st.session_state.notifications_dismissed]
    
    # ============ TOP METRICS BAR ============
    unread_count = len([n for n in active_notifications if n['id'] not in st.session_state.notifications_read])
    critical_count = len([n for n in active_notifications if n.get('priority') == 'critical'])
    high_count = len([n for n in active_notifications if n.get('priority') == 'high'])
    
    c1, c2, c3, c4, c5 = st.columns(5)
    with c1:
        st.metric("🔔 Total", len(active_notifications))
    with c2:
        st.metric("🔴 Unread", unread_count, delta=None, delta_color="off")
    with c3:
        st.metric("🚨 Critical", critical_count)
    with c4:
        st.metric("⚠️ High Priority", high_count)
    with c5:
        st.metric("✅ Read", len(active_notifications) - unread_count)
    
    st.markdown("---")
    
    # ============ QUICK ACTIONS ============
    col1, col2, col3, col4 = st.columns(4)
    with col1:
        if st.button("✅ Mark All as Read", use_container_width=True):
            for n in active_notifications:
                st.session_state.notifications_read.add(n['id'])
            st.rerun()
    with col2:
        if st.button("🗑️ Clear All", use_container_width=True):
            for n in active_notifications:
                st.session_state.notifications_dismissed.add(n['id'])
            st.rerun()
    with col3:
        if st.button("📧 Email Digest", use_container_width=True):
            if active_notifications:
                st.success(f"✅ Digest with {len(active_notifications)} notifications sent to {user_email}")
            else:
                st.info("No notifications to send.")
    with col4:
        show_preferences = st.button("⚙️ Preferences", use_container_width=True)
    
    # ============ PREFERENCES PANEL ============
    if show_preferences:
        with st.expander("⚙️ Notification Preferences", expanded=True):
            st.markdown("Choose which notifications you want to receive:")
            c1, c2, c3 = st.columns(3)
            with c1:
                st.session_state.notification_preferences['appraisal'] = st.checkbox("📝 Appraisals", value=st.session_state.notification_preferences['appraisal'])
                st.session_state.notification_preferences['leave'] = st.checkbox("🏖️ Leave", value=st.session_state.notification_preferences['leave'])
                st.session_state.notification_preferences['kudos'] = st.checkbox("🌟 Kudos", value=st.session_state.notification_preferences['kudos'])
            with c2:
                st.session_state.notification_preferences['kpi'] = st.checkbox("🎯 KPIs", value=st.session_state.notification_preferences['kpi'])
                st.session_state.notification_preferences['birthday'] = st.checkbox("🎂 Birthdays", value=st.session_state.notification_preferences['birthday'])
                st.session_state.notification_preferences['training'] = st.checkbox("📚 Training", value=st.session_state.notification_preferences['training'])
            with c3:
                st.session_state.notification_preferences['recruitment'] = st.checkbox("💼 Recruitment", value=st.session_state.notification_preferences['recruitment'])
                st.session_state.notification_preferences['announcement'] = st.checkbox("📢 Announcements", value=st.session_state.notification_preferences['announcement'])
                st.session_state.notification_preferences['system'] = st.checkbox("🖥️ System", value=st.session_state.notification_preferences['system'])
    
    st.markdown("---")
    
    # ============ NOTIFICATION LIST ============
    if active_notifications:
        # Group by category
        from collections import defaultdict
        grouped = defaultdict(list)
        for n in active_notifications:
            grouped[n.get('category', 'other')].append(n)
        
        category_names = {
            'appraisal': '📝 Appraisals & Performance',
            'kpi': '🎯 KPI Deadlines',
            'birthday': '🎂 Birthdays & Anniversaries',
            'recruitment': '💼 Recruitment',
            'training': '📚 Training & Development',
            'kudos': '🌟 Kudos & Recognition',
            'announcement': '📢 Announcements',
            'poll': '📊 Polls & Surveys',
            'system': '🖥️ System',
            'leave': '🏖️ Leave Management'
        }
        
        for category, items in grouped.items():
            cat_name = category_names.get(category, f"📌 {category.title()}")
            unread_in_cat = len([n for n in items if n['id'] not in st.session_state.notifications_read])
            
            with st.expander(f"{cat_name} ({len(items)}) {'🔴' if unread_in_cat > 0 else ''}", expanded=(unread_in_cat > 0 or category in ['critical', 'high'])):
                for n in items:
                    is_unread = n['id'] not in st.session_state.notifications_read
                    
                    # Priority colors
                    priority_colors = {
                        'critical': '#CC0000',
                        'high': '#e53e3e',
                        'medium': '#d69e2e',
                        'low': '#718096'
                    }
                    border_color = priority_colors.get(n.get('priority', 'low'), '#718096')
                    bg_color = '#fff5f5' if is_unread else '#fafafa'
                    
                    # Notification card
                    col1, col2, col3 = st.columns([8, 1, 1])
                    with col1:
                        st.markdown(f"""
                        <div style="background:{bg_color};padding:0.8rem;border-radius:8px;margin-bottom:0.4rem;border-left:4px solid {border_color};">
                            <div style="display:flex;justify-content:space-between;align-items:start;">
                                <div style="flex:1;">
                                    <strong style="font-size:0.95rem;">{n['title']}</strong>
                                    {'<span style="background:#CC0000;color:white;padding:0.1rem 0.5rem;border-radius:10px;font-size:0.7rem;margin-left:0.5rem;">NEW</span>' if is_unread else ''}
                                    <p style="margin:0.3rem 0;color:#555;font-size:0.85rem;">{n['message']}</p>
                                    <small style="color:#888;">🕐 {n['time']} | 🏷️ {n.get('priority', 'normal').upper()}</small>
                                </div>
                            </div>
                        </div>
                        """, unsafe_allow_html=True)
                    
                    with col2:
                        if is_unread:
                            if st.button("✓", key=f"read_{n['id']}", help="Mark as read", use_container_width=True):
                                st.session_state.notifications_read.add(n['id'])
                                st.rerun()
                    
                    with col3:
                        if st.button("✕", key=f"dismiss_{n['id']}", help="Dismiss", use_container_width=True):
                            st.session_state.notifications_dismissed.add(n['id'])
                            st.rerun()
                    
                    # Action button
                    if n.get('action_label'):
                        if st.button(n['action_label'], key=f"action_{n['id']}", use_container_width=True):
                            action = n.get('action', '')
                            if action == 'review_appraisal':
                                st.session_state['navigate_to'] = "📈 Performance & OKRs"
                            elif action == 'view_kpis':
                                st.session_state['navigate_to'] = "📈 Performance & OKRs"
                            elif action == 'view_recruitment':
                                st.session_state['navigate_to'] = "💼 Recruitment Hub"
                            elif action == 'view_training':
                                st.session_state['navigate_to'] = "🎓 Training & Development"
                            elif action == 'view_announcements':
                                st.session_state['navigate_to'] = "💬 Chat & Communications"
                            elif action == 'view_kudos':
                                st.session_state['navigate_to'] = "💬 Chat & Communications"
                            elif action == 'view_system':
                                st.session_state['navigate_to'] = "📊 Reports & Analytics"
                            elif action == 'send_wishes':
                                st.success(f"🎉 Birthday wishes sent to the team!")
                            elif action == 'send_congrats':
                                st.success(f"🎉 Congratulations sent!")
                            elif action == 'view_calendar':
                                st.info("📅 Check the Training Calendar for upcoming events.")
                            st.rerun()
    else:
        st.success("🎉 You're all caught up! No new notifications.")
    
    # ============ NOTIFICATION SIDEBAR BADGE (Injected via JS) ============
    if unread_count > 0:
        st.markdown(f"""
        <script>
            // Update browser tab title
            document.title = "({unread_count}) Churchgate HRIS - Notifications";
        </script>
        """, unsafe_allow_html=True)
    
    # ============ WEEKLY DIGEST PREVIEW ============
    if is_admin:
        st.markdown("---")
        with st.expander("📊 Weekly Notification Digest (Admin View)", expanded=False):
            st.markdown(f"**Week of {now.strftime('%B %d, %Y')}**")
            
            # Summarize by category
            digest_data = []
            for category, items in grouped.items():
                digest_data.append({
                    'Category': category_names.get(category, category),
                    'Count': len(items),
                    'Critical': len([n for n in items if n.get('priority') == 'critical']),
                    'Unread': len([n for n in items if n['id'] not in st.session_state.notifications_read])
                })
            
            if digest_data:
                digest_df = pd.DataFrame(digest_data)
                st.dataframe(digest_df, use_container_width=True, hide_index=True)
                
                # Chart
                fig = px.bar(digest_df, x='Category', y='Count', color='Critical',
                           color_continuous_scale=['#38a169', '#d69e2e', '#CC0000'])
                fig.update_layout(height=300)
                st.plotly_chart(fig, use_container_width=True)

def my_documents():
    st.markdown("""<div class="churchgate-header"><h1>📋 My Documents Vault</h1><p>Pay Slips | Contracts | Tax Documents | Certificates | Personal Files</p></div>""", unsafe_allow_html=True)
    
    user_name = st.session_state.user['name'] if st.session_state.user else 'Staff'
    user_id = st.session_state.user.get('employee_id', '') if st.session_state.user else ''
    user_email = st.session_state.user.get('email', '') if st.session_state.user else ''
    is_admin = st.session_state.user['role'] in ['Admin', 'HR Director'] if st.session_state.user else False
    
    tab1, tab2, tab3, tab4, tab5, tab6 = st.tabs([
        "💰 Pay Slips", "📄 Contracts", "📊 Tax & Reviews", 
        "🎓 Certificates", "📁 My Files", "📤 Requests"
    ])
    
    # Load employee documents
    def load_docs(doc_type=None):
        try:
            all_docs = db._get("employee_documents")
            if all_docs:
                my_docs = [d for d in all_docs if d.get('employee_id') == user_id or d.get('is_public') == True]
                if doc_type:
                    my_docs = [d for d in my_docs if d.get('document_type') == doc_type]
                return my_docs
        except:
            pass
        return []
    
    # ============ TAB 1: PAY SLIPS ============
    with tab1:
        st.subheader("💰 Pay Slips")
        st.info("Monthly pay stubs are automatically generated and stored here.")
        
        pay_slips = load_docs('payslip')
        
        if pay_slips:
            for doc in pay_slips:
                month = doc.get('month_year', '')
                with st.expander(f"📄 Pay Slip — {month}", expanded=False):
                    st.markdown(f"**Period:** {month}")
                    st.markdown(f"**Generated:** {doc.get('uploaded_at', '')[:10]}")
                    if doc.get('file_url'):
                        st.markdown(f"📥 [Download Pay Slip]({doc['file_url']})")
                    else:
                        st.info("Document available on request.")
        else:
            st.info("No pay slips available yet. Pay slips are generated monthly by HR/Accounts.")
            
            # Generate sample payslip button (admin only for testing)
            if is_admin:
                if st.button("🧪 Generate Sample Pay Slip", use_container_width=True):
                    db._post("employee_documents", {
                        "employee_id": user_id,
                        "document_name": f"Pay Slip - {datetime.now().strftime('%B %Y')}",
                        "document_type": "payslip",
                        "month_year": datetime.now().strftime('%B %Y'),
                        "uploaded_by": "System",
                        "uploaded_at": datetime.now().strftime('%Y-%m-%d %H:%M'),
                        "is_public": False
                    })
                    st.success("✅ Sample pay slip generated!")
                    st.rerun()
    
    # ============ TAB 2: CONTRACTS ============
    with tab2:
        st.subheader("📄 Employment Contracts & Agreements")
        
        contracts = load_docs('contract')
        
        if contracts:
            for doc in contracts:
                with st.expander(f"📄 {doc.get('document_name', 'Contract')}", expanded=False):
                    st.markdown(f"**Uploaded:** {doc.get('uploaded_at', '')[:10]}")
                    st.markdown(f"**By:** {doc.get('uploaded_by', 'HR')}")
                    if doc.get('file_url'):
                        st.markdown(f"📥 [Download Document]({doc['file_url']})")
        else:
            st.info("Your employment contract will appear here once uploaded by HR.")
        
        # HR upload section
        if is_admin:
            st.markdown("---")
            st.markdown("### 📤 Upload Contract (Admin)")
            with st.form("upload_contract"):
                contract_employee = st.text_input("Employee ID", value=user_id)
                contract_name = st.text_input("Document Name", placeholder="e.g., Employment Contract 2026")
                contract_file = st.file_uploader("Upload PDF", type=['pdf'])
                if st.form_submit_button("📤 Upload", use_container_width=True):
                    if contract_file and contract_name:
                        file_url = ""
                        try:
                            file_url = db.upload_file("documents", f"{contract_employee}_{contract_name}.pdf", contract_file.read(), "application/pdf")
                        except:
                            pass
                        db._post("employee_documents", {
                            "employee_id": contract_employee,
                            "document_name": contract_name,
                            "document_type": "contract",
                            "file_url": file_url,
                            "uploaded_by": user_name,
                            "uploaded_at": datetime.now().strftime('%Y-%m-%d %H:%M')
                        })
                        st.success("✅ Contract uploaded!")
                        st.rerun()
    
    # ============ TAB 3: TAX & REVIEWS ============
    with tab3:
        st.subheader("📊 Tax Documents & Performance Reviews")
        
        tax_docs = load_docs('tax')
        review_docs = load_docs('review')
        
        st.markdown("### 🧾 Tax Documents")
        if tax_docs:
            for doc in tax_docs:
                st.markdown(f"📥 **{doc.get('document_name', 'Tax Document')}** — {doc.get('uploaded_at', '')[:10]}")
        else:
            st.info("Annual tax certificates will appear here.")
        
        st.markdown("---")
        st.markdown("### 📈 Performance Reviews")
        if review_docs:
            for doc in review_docs:
                st.markdown(f"📥 **{doc.get('document_name', 'Review')}** — {doc.get('uploaded_at', '')[:10]}")
        else:
            st.info("Completed appraisal reports will appear here.")
    
    # ============ TAB 4: CERTIFICATES ============
    with tab4:
        st.subheader("🎓 Training Certificates")
        
        certs = load_docs('certificate')
        
        if certs:
            for doc in certs:
                with st.expander(f"🏅 {doc.get('document_name', 'Certificate')}", expanded=False):
                    st.markdown(f"**Earned:** {doc.get('uploaded_at', '')[:10]}")
                    if doc.get('file_url'):
                        st.markdown(f"📥 [Download Certificate]({doc['file_url']})")
        else:
            st.info("Training certificates will appear here once you complete courses.")
        
        # Upload personal certificate
        with st.form("upload_cert"):
            st.markdown("### 📤 Upload Your Certificate")
            cert_name = st.text_input("Certificate Name")
            cert_file = st.file_uploader("Upload Certificate", type=['pdf', 'jpg', 'png'])
            if st.form_submit_button("📤 Upload", use_container_width=True):
                if cert_file and cert_name:
                    file_url = ""
                    try:
                        file_url = db.upload_file("documents", f"{user_id}_{cert_name}.pdf", cert_file.read(), "application/pdf")
                    except:
                        pass
                    db._post("employee_documents", {
                        "employee_id": user_id,
                        "document_name": cert_name,
                        "document_type": "certificate",
                        "file_url": file_url,
                        "uploaded_by": user_name,
                        "uploaded_at": datetime.now().strftime('%Y-%m-%d %H:%M')
                    })
                    st.success("✅ Certificate uploaded!")
                    st.rerun()
    
    # ============ TAB 5: MY FILES ============
    with tab5:
        st.subheader("📁 My Personal Files")
        st.info("Upload personal documents like ID, educational certificates, professional licenses, etc.")
        
        personal_docs = load_docs('personal')
        
        if personal_docs:
            for doc in personal_docs:
                col1, col2 = st.columns([4, 1])
                with col1:
                    st.markdown(f"📎 **{doc.get('document_name', 'File')}** — {doc.get('uploaded_at', '')[:10]}")
                with col2:
                    if st.button("🗑️", key=f"del_{doc.get('id')}"):
                        db._delete("employee_documents", {"id": doc.get('id')})
                        st.rerun()
        
        st.markdown("---")
        with st.form("upload_personal"):
            st.markdown("### 📤 Upload Personal Document")
            file_name = st.text_input("File Name / Description")
            uploaded_file = st.file_uploader("Choose File", type=['pdf', 'jpg', 'png', 'docx'])
            if st.form_submit_button("📤 Upload", use_container_width=True):
                if uploaded_file and file_name:
                    file_url = ""
                    try:
                        file_url = db.upload_file("documents", f"{user_id}_{file_name}", uploaded_file.read(), uploaded_file.type)
                    except:
                        pass
                    db._post("employee_documents", {
                        "employee_id": user_id,
                        "document_name": file_name,
                        "document_type": "personal",
                        "file_url": file_url,
                        "uploaded_by": user_name,
                        "uploaded_at": datetime.now().strftime('%Y-%m-%d %H:%M')
                    })
                    st.success("✅ File uploaded!")
                    st.rerun()
    
    # ============ TAB 6: REQUESTS ============
    with tab6:
        st.subheader("📤 Document Requests")
        st.info("Request a document from HR. You'll be notified when it's available.")
        
        with st.form("request_doc"):
            doc_request = st.selectbox("Document Type", [
                "Employment Confirmation Letter",
                "Salary Certificate",
                "Tax Certificate",
                "Service Letter",
                "Promotion Letter",
                "Other"
            ])
            purpose = st.text_area("Purpose / Notes")
            if st.form_submit_button("📤 Submit Request", use_container_width=True):
                if purpose:
                    db._post("employee_documents", {
                        "employee_id": user_id,
                        "document_name": f"REQUEST: {doc_request}",
                        "document_type": "request",
                        "uploaded_by": user_name,
                        "uploaded_at": datetime.now().strftime('%Y-%m-%d %H:%M'),
                        "month_year": purpose
                    })
                    st.success(f"✅ {doc_request} requested! HR will process it shortly.")
                    st.balloons()

def ideas_box():
    st.markdown("""<div class="churchgate-header"><h1>💡 Ideas & Innovation Box</h1><p>Submit Ideas | Vote | Collaborate | Drive Innovation</p></div>""", unsafe_allow_html=True)
    
    user_name = st.session_state.user['name'] if st.session_state.user else 'Staff'
    user_id = st.session_state.user.get('employee_id', '') if st.session_state.user else ''
    user_dept = st.session_state.user.get('department', '') if st.session_state.user else ''
    is_admin = st.session_state.user['role'] in ['Admin', 'HR Director', 'Manager', 'HOD'] if st.session_state.user else False
    
    tab1, tab2, tab3 = st.tabs(["📝 Submit Idea", "💡 All Ideas", "📊 Dashboard"])
    
    # Load ideas
    def load_ideas(status_filter=None):
        try:
            data = db._get("ideas_box")
            if data:
                if status_filter:
                    data = [d for d in data if d.get('status') == status_filter]
                return sorted(data, key=lambda x: x.get('votes', 0), reverse=True)
        except:
            pass
        return []
    
    # ============ TAB 1: SUBMIT IDEA ============
    with tab1:
        st.subheader("📝 Share Your Idea")
        st.info("Great ideas drive innovation. Submit yours and get votes from colleagues!")
        
        with st.form("submit_idea"):
            idea_title = st.text_input("Idea Title *", placeholder="Give your idea a clear, catchy name")
            idea_category = st.selectbox("Category", [
                "💻 Technology & Digital",
                "🏢 Operations & Facilities",
                "👥 People & Culture",
                "💰 Cost Savings",
                "😊 Customer Experience",
                "🌱 Sustainability",
                "📈 Revenue Growth",
                "🔧 Process Improvement",
                "🎉 Employee Experience",
                "💡 Other"
            ])
            idea_description = st.text_area("Describe Your Idea *", height=150, 
                placeholder="What's the problem? What's your solution? How would it benefit Churchgate?")
            
            if st.form_submit_button("🚀 Submit Idea", use_container_width=True):
                if idea_title and idea_description:
                    db._post("ideas_box", {
                        "employee_id": user_id,
                        "employee_name": user_name,
                        "department": user_dept,
                        "title": idea_title,
                        "description": idea_description,
                        "category": idea_category,
                        "status": "Submitted",
                        "votes": 0,
                        "voters": "[]",
                        "submitted_at": datetime.now().strftime('%Y-%m-%d %H:%M')
                    })
                    st.success("✅ Idea submitted! Colleagues can now see and vote on it.")
                    st.balloons()
                    st.rerun()
                else:
                    st.error("❌ Title and description are required!")
        
        # Show user's own ideas
        st.markdown("---")
        st.markdown("### 📋 My Submitted Ideas")
        all_ideas = load_ideas()
        my_ideas = [i for i in all_ideas if i.get('employee_id') == user_id]
        if my_ideas:
            for idea in my_ideas:
                status_color = "#38a169" if idea.get('status') == 'Approved' else "#d69e2e" if idea.get('status') == 'Under Review' else "#a0aec0"
                with st.expander(f"{idea.get('title', 'Untitled')} — {idea.get('votes', 0)} votes | {idea.get('status', 'Submitted')}"):
                    st.markdown(f"**Category:** {idea.get('category', 'N/A')}")
                    st.markdown(f"**Description:** {idea.get('description', '')}")
                    st.markdown(f"**Submitted:** {idea.get('submitted_at', '')[:10]}")
                    if idea.get('admin_response'):
                        st.success(f"**Admin Response:** {idea.get('admin_response')}")
        else:
            st.info("You haven't submitted any ideas yet.")
    
    # ============ TAB 2: ALL IDEAS ============
    with tab2:
        st.subheader("💡 Innovation Board")
        
        # Filters
        col1, col2, col3 = st.columns(3)
        with col1:
            category_filter = st.selectbox("Category", ["All"] + list(set([i.get('category', '') for i in load_ideas()])))
        with col2:
            status_filter = st.selectbox("Status", ["All", "Submitted", "Under Review", "Approved", "Implemented"])
        with col3:
            sort_by = st.selectbox("Sort", ["Most Votes", "Newest", "Oldest"])
        
        all_ideas = load_ideas()
        
        # Apply filters
        if category_filter != "All":
            all_ideas = [i for i in all_ideas if i.get('category') == category_filter]
        if status_filter != "All":
            all_ideas = [i for i in all_ideas if i.get('status') == status_filter]
        if sort_by == "Newest":
            all_ideas = sorted(all_ideas, key=lambda x: x.get('submitted_at', ''), reverse=True)
        elif sort_by == "Oldest":
            all_ideas = sorted(all_ideas, key=lambda x: x.get('submitted_at', ''))
        
        st.markdown(f"**{len(all_ideas)} ideas**")
        
        if all_ideas:
            for idea in all_ideas:
                voter_list = []
                try:
                    voter_list = json.loads(idea.get('voters', '[]'))
                except:
                    voter_list = []
                
                has_voted = user_id in voter_list
                
                with st.expander(f"💡 {idea.get('title', 'Untitled')} — {idea.get('votes', 0)} votes | by {idea.get('employee_name', 'Unknown')} | {idea.get('status', 'Submitted')}"):
                    col1, col2 = st.columns([3, 1])
                    
                    with col1:
                        st.markdown(f"**Category:** {idea.get('category', 'N/A')}")
                        st.markdown(f"**Department:** {idea.get('department', 'N/A')}")
                        st.markdown(f"**Description:** {idea.get('description', '')}")
                        st.markdown(f"**Submitted:** {idea.get('submitted_at', '')[:10]}")
                        if idea.get('admin_response'):
                            st.success(f"**Response:** {idea.get('admin_response')}")
                    
                    with col2:
                        if has_voted:
                            st.success(f"✅ Voted ({idea.get('votes', 0)})")
                        else:
                            if st.button(f"👍 Upvote ({idea.get('votes', 0)})", key=f"vote_{idea.get('id')}", use_container_width=True):
                                voter_list.append(user_id)
                                db._patch("ideas_box", {
                                    "votes": int(idea.get('votes', 0)) + 1,
                                    "voters": json.dumps(voter_list)
                                }, {"id": idea.get('id')})
                                st.rerun()
                    
                    # Admin actions
                    if is_admin:
                        st.markdown("---")
                        st.markdown("**Admin Actions:**")
                        col_a1, col_a2, col_a3 = st.columns(3)
                        new_status = col_a1.selectbox("Status", ["Submitted", "Under Review", "Approved", "Implemented"], key=f"status_{idea.get('id')}")
                        admin_response = col_a2.text_input("Response", key=f"resp_{idea.get('id')}")
                        if col_a3.button("Update", key=f"upd_{idea.get('id')}", use_container_width=True):
                            db._patch("ideas_box", {
                                "status": new_status,
                                "admin_response": admin_response
                            }, {"id": idea.get('id')})
                            st.success("✅ Updated!")
                            st.rerun()
        else:
            st.info("No ideas yet. Be the first to submit one!")
    
    # ============ TAB 3: DASHBOARD ============
    with tab3:
        st.subheader("📊 Innovation Dashboard")
        
        all_ideas = load_ideas()
        
        if all_ideas:
            total = len(all_ideas)
            implemented = len([i for i in all_ideas if i.get('status') == 'Implemented'])
            approved = len([i for i in all_ideas if i.get('status') == 'Approved'])
            under_review = len([i for i in all_ideas if i.get('status') == 'Under Review'])
            total_votes = sum([i.get('votes', 0) for i in all_ideas])
            
            c1, c2, c3, c4, c5 = st.columns(5)
            c1.metric("💡 Total Ideas", total)
            c2.metric("✅ Implemented", implemented)
            c3.metric("👍 Approved", approved)
            c4.metric("🔍 Under Review", under_review)
            c5.metric("🗳️ Total Votes", total_votes)
            
            st.markdown("---")
            
            # Top ideas
            col1, col2 = st.columns(2)
            with col1:
                st.markdown("### 🏆 Top Voted Ideas")
                top = sorted(all_ideas, key=lambda x: x.get('votes', 0), reverse=True)[:5]
                for i, idea in enumerate(top):
                    medal = "🥇" if i == 0 else "🥈" if i == 1 else "🥉" if i == 2 else "⭐"
                    st.markdown(f"{medal} **{idea.get('title', '')}** — {idea.get('votes', 0)} votes by {idea.get('employee_name', '')}")
            
            with col2:
                st.markdown("### 📈 By Category")
                categories = {}
                for idea in all_ideas:
                    cat = idea.get('category', 'Other')
                    categories[cat] = categories.get(cat, 0) + 1
                if categories:
                    cat_df = pd.DataFrame({'Category': list(categories.keys()), 'Count': list(categories.values())})
                    fig = px.pie(cat_df, values='Count', names='Category', hole=0.5)
                    fig.update_layout(height=300)
                    st.plotly_chart(fig, use_container_width=True)
            
            # Top contributors
            st.markdown("---")
            st.markdown("### 🌟 Top Contributors")
            contributors = {}
            for idea in all_ideas:
                name = idea.get('employee_name', 'Unknown')
                contributors[name] = contributors.get(name, 0) + 1
            top_contributors = sorted(contributors.items(), key=lambda x: x[1], reverse=True)[:5]
            for i, (name, count) in enumerate(top_contributors):
                st.markdown(f"⭐ **{name}** — {count} idea{'s' if count > 1 else ''}")
        else:
            st.info("Submit the first idea to see the dashboard!")

def company_calendar():
    st.markdown("""<div class="churchgate-header"><h1>📅 Enterprise Calendar</h1><p>Financial Year | Holidays | Birthdays | Training | Deadlines | Events | Google/Teams Sync</p></div>""", unsafe_allow_html=True)
    
    user_name = st.session_state.user['name'] if st.session_state.user else 'Staff'
    user_email = st.session_state.user.get('email', '') if st.session_state.user else ''
    user_dept = st.session_state.user.get('department', '') if st.session_state.user else ''
    is_admin = st.session_state.user['role'] in ['Admin', 'HR Director'] if st.session_state.user else False
    
    today = datetime.now()
    
    # Load events
    def load_events():
        try:
            data = db._get("company_events")
            return data if data else []
        except:
            return []
    
    # Seed defaults if empty
    events = load_events()
    if not events:
        defaults = [
            ("Democracy Day", "holiday", "2026-06-12", "Public Holiday - Office Closed", "#CC0000"),
            ("Eid al-Adha", "holiday", "2026-06-17", "Public Holiday", "#CC0000"),
            ("Q2 Performance Reviews Due", "deadline", "2026-06-30", "All departments submit Q2 reviews", "#d69e2e"),
            ("BMS Advanced Training", "training", "2026-06-15", "Building Management Systems - Siemens Academy", "#3182ce"),
            ("AI in FM Webinar", "training", "2026-06-20", "Practical AI Applications for Facility Management", "#3182ce"),
            ("World Trade Center Day", "event", "2026-07-01", "Annual WTC Abuja Celebration", "#38a169"),
            ("Independence Day", "holiday", "2026-10-01", "Nigeria Independence Day", "#CC0000"),
            ("Christmas Day", "holiday", "2026-12-25", "Public Holiday", "#CC0000"),
            ("Boxing Day", "holiday", "2026-12-26", "Public Holiday", "#CC0000"),
            ("New Year's Day", "holiday", "2027-01-01", "Public Holiday - Office Closed", "#CC0000"),
        ]
        for title, etype, date, desc, color in defaults:
            db._post("company_events", {"title": title, "event_type": etype, "event_date": date, "description": desc, "created_by": "System"})
        events = load_events()
    
    # Color mapping
    event_colors = {
        'holiday': '#CC0000',
        'deadline': '#d69e2e',
        'training': '#3182ce',
        'event': '#38a169',
        'birthday': '#dd6b20',
        'meeting': '#805ad5',
        'financial': '#1a1a1a'
    }
    
    event_icons = {
        'holiday': '🏖️',
        'deadline': '⏰',
        'training': '📚',
        'event': '📅',
        'birthday': '🎂',
        'meeting': '🤝',
        'financial': '📊'
    }
    
    # ============ FINANCIAL YEAR BANNER ============
    current_month = today.month
    current_year = today.year
    
    if current_month >= 4:
        fy_start = current_year
        fy_end = current_year + 1
    else:
        fy_start = current_year - 1
        fy_end = current_year
    
    fy_progress_month = (current_month - 4 if current_month >= 4 else current_month + 8) + 1
    fy_pct = int((fy_progress_month / 12) * 100)
    
    fy_quarters = [
        ("Q1", 4, 6, "Apr - Jun"),
        ("Q2", 7, 9, "Jul - Sep"),
        ("Q3", 10, 12, "Oct - Dec"),
        ("Q4", 1, 3, "Jan - Mar")
    ]
    
    current_quarter = None
    for q_name, q_start, q_end, q_range in fy_quarters:
        if (current_month >= q_start and current_month <= q_end) or (q_start > q_end and (current_month >= q_start or current_month <= q_end)):
            current_quarter = f"{q_name} ({q_range})"
            break
    
    st.markdown(f"""
    <div style="background: linear-gradient(135deg, #1a1a1a, #2d2d2d); color: white; padding: 1.2rem 1.5rem; border-radius: 10px; margin-bottom: 1.5rem; border: 2px solid #CC0000;">
        <div style="display: flex; align-items: center; justify-content: space-between; flex-wrap: wrap; gap: 1rem;">
            <div>
                <h3 style="margin:0;color:white;">📊 Financial Year {fy_start}/{fy_end}</h3>
                <p style="margin:0.3rem 0 0 0;color:#ccc;">April {fy_start} — March {fy_end} | Current: {current_quarter}</p>
            </div>
            <div style="text-align:right;">
                <span style="font-weight:700;font-size:1.1rem;">Month {fy_progress_month} of 12 ({fy_pct}%)</span>
                <div style="background:#555;height:8px;border-radius:4px;width:220px;margin-top:0.4rem;">
                    <div style="background:#CC0000;width:{fy_pct}%;height:8px;border-radius:4px;"></div>
                </div>
            </div>
        </div>
    </div>
    """, unsafe_allow_html=True)
    
    # ============ TOP STATS ROW ============
    upcoming_events = [e for e in events if e.get('event_date', '')[:10] >= today.strftime('%Y-%m-%d')]
    today_events = [e for e in events if e.get('event_date', '')[:10] == today.strftime('%Y-%m-%d')]
    this_month_events = [e for e in events if e.get('event_date', '')[:7] == today.strftime('%Y-%m')]
    
    c1, c2, c3, c4, c5 = st.columns(5)
    c1.metric("📅 Total Events", len(events))
    c2.metric("🔴 Today", len(today_events))
    c3.metric("📆 This Month", len(this_month_events))
    c4.metric("⏳ Upcoming", len(upcoming_events))
    c5.metric("🏖️ Holidays", len([e for e in events if e.get('event_type') == 'holiday']))
    
    tab1, tab2, tab3, tab4, tab5 = st.tabs(["📅 Month View", "📋 List View", "📆 Week View", "➕ Add Event", "🔗 Integrations"])
    
    # ============ TAB 1: MONTH VIEW ============
    with tab1:
        month = today.month
        year = today.year
        
        col1, col2, col3 = st.columns([1, 3, 1])
        with col1:
            if st.button("◀ Prev Month"):
                if month == 1:
                    month = 12
                    year -= 1
                else:
                    month -= 1
                st.rerun()
        with col2:
            month_name = datetime(year, month, 1).strftime('%B %Y')
            st.markdown(f"<h3 style='text-align:center;'>{month_name}</h3>", unsafe_allow_html=True)
        with col3:
            if st.button("Next Month ▶"):
                if month == 12:
                    month = 1
                    year += 1
                else:
                    month += 1
                st.rerun()
        
        import calendar as cal
        first_day, num_days = cal.monthrange(year, month)
        
        st.markdown("| Mon | Tue | Wed | Thu | Fri | Sat | Sun |")
        st.markdown("|-----|-----|-----|-----|-----|-----|-----|")
        
        day = 1
        for week in range(6):
            if day > num_days:
                break
            week_str = "|"
            for weekday in range(7):
                if week == 0 and weekday < first_day:
                    week_str += "     |"
                elif day > num_days:
                    week_str += "     |"
                else:
                    date_str = f"{year}-{month:02d}-{day:02d}"
                    day_events = [e for e in events if e.get('event_date', '')[:10] == date_str]
                    is_today = (date_str == today.strftime('%Y-%m-%d'))
                    
                    if is_today:
                        if day_events:
                            badge = ' '.join([event_icons.get(e.get('event_type', ''), '📌') for e in day_events])
                            week_str += f" 🟢**{day}** {badge} |"
                        else:
                            week_str += f" 🟢**{day}** |"
                    else:
                        if day_events:
                            badge = ' '.join([event_icons.get(e.get('event_type', ''), '📌') for e in day_events])
                            week_str += f" {day} {badge} |"
                        else:
                            week_str += f" {day} |"
                day += 1
            st.markdown(week_str, unsafe_allow_html=True)
        
        # Legend
        st.markdown("---")
        legend_cols = st.columns(7)
        for i, (etype, icon) in enumerate(event_icons.items()):
            color = event_colors.get(etype, '#888')
            legend_cols[i % 7].markdown(f'<span style="color:{color};">{icon} {etype.title()}</span>', unsafe_allow_html=True)
        
        # Today's events
        if today_events:
            st.markdown("---")
            st.markdown("### 🔴 Today's Events")
            for e in today_events:
                st.markdown(f"**{event_icons.get(e.get('event_type',''), '📌')} {e.get('title')}** — {e.get('description', '')}")
    
    # ============ TAB 2: LIST VIEW ============
    with tab2:
        col1, col2 = st.columns(2)
        with col1:
            filter_type = st.selectbox("Filter by Type", ["All", "holiday", "deadline", "training", "event", "meeting", "birthday"])
        with col2:
            filter_month = st.selectbox("Filter by Month", ["All"] + [datetime(today.year, m, 1).strftime('%B') for m in range(1, 13)])
        
        filtered = events.copy()
        if filter_type != "All":
            filtered = [e for e in filtered if e.get('event_type') == filter_type]
        if filter_month != "All":
            month_num = datetime.strptime(filter_month, '%B').month
            filtered = [e for e in filtered if e.get('event_date', '')[:7] == f"{today.year}-{month_num:02d}"]
        
        filtered = sorted(filtered, key=lambda x: x.get('event_date', ''))
        
        st.markdown(f"**{len(filtered)} events**")
        
        for event in filtered:
            etype = event.get('event_type', 'event')
            color = event_colors.get(etype, '#888')
            icon = event_icons.get(etype, '📌')
            
            try:
                edate = datetime.strptime(event.get('event_date', '')[:10], '%Y-%m-%d')
                days_left = (edate - today).days
                days_str = f"⏰ in {days_left} days" if days_left > 0 else "🔴 Today!" if days_left == 0 else f"✅ {abs(days_left)} days ago"
            except:
                days_str = ""
            
            st.markdown(f"""
            <div style="background:white;padding:0.8rem 1rem;border-radius:8px;margin-bottom:0.5rem;border-left:4px solid {color};">
                <div style="display:flex;justify-content:space-between;align-items:center;">
                    <div>
                        <strong>{icon} {event.get('title', 'Event')}</strong>
                        <br><small style="color:#666;">{event.get('description', '')}</small>
                    </div>
                    <div style="text-align:right;">
                        <span style="background:{color};color:white;padding:0.2rem 0.6rem;border-radius:10px;font-size:0.75rem;">{etype.title()}</span>
                        <br><small style="color:#888;">{event.get('event_date', '')[:10]}</small>
                        <br><small style="color:{color};">{days_str}</small>
                    </div>
                </div>
            </div>
            """, unsafe_allow_html=True)
    
    # ============ TAB 3: WEEK VIEW ============
    with tab3:
        st.subheader("📆 This Week")
        
        weekday = today.weekday()
        week_start = today - timedelta(days=weekday)
        
        week_cols = st.columns(7)
        day_names = ["Monday", "Tuesday", "Wednesday", "Thursday", "Friday", "Saturday", "Sunday"]
        
        for i, (col, day_name) in enumerate(zip(week_cols, day_names)):
            day_date = week_start + timedelta(days=i)
            date_str = day_date.strftime('%Y-%m-%d')
            day_events = [e for e in events if e.get('event_date', '')[:10] == date_str]
            
            is_today = date_str == today.strftime('%Y-%m-%d')
            bg = "#fff3f3" if is_today else "#fafafa"
            
            with col:
                st.markdown(f"""
                <div style="background:{bg};padding:0.5rem;border-radius:6px;text-align:center;min-height:80px;border:{ '2px solid #CC0000' if is_today else '1px solid #e0e0e0'};">
                    <small style="color:{'#CC0000' if is_today else '#888'};font-weight:{'700' if is_today else '400'};">{day_name[:3]}</small>
                    <br><strong>{day_date.day}</strong>
                    {''.join([f'<br><small>{event_icons.get(e.get("event_type",""),"📌")}</small>' for e in day_events]) if day_events else ''}
                </div>
                """, unsafe_allow_html=True)
    
    # ============ TAB 4: ADD EVENT ============
    with tab4:
        st.subheader("➕ Add Event")
        
        if is_admin:
            with st.form("add_calendar_event"):
                event_title = st.text_input("Event Title *")
                col1, col2, col3 = st.columns(3)
                with col1:
                    event_type = st.selectbox("Type", ["event", "holiday", "training", "deadline", "meeting", "birthday"])
                with col2:
                    event_date = st.date_input("Date")
                with col3:
                    event_time = st.text_input("Time (Optional)", placeholder="e.g., 10:00 AM")
                event_desc = st.text_area("Description")
                event_location = st.text_input("Location/Link (Optional)", placeholder="Google Meet / Teams / On-site")
                
                if st.form_submit_button("➕ Add to Calendar", use_container_width=True):
                    if event_title:
                        full_desc = event_desc
                        if event_time:
                            full_desc += f"\n⏰ Time: {event_time}"
                        if event_location:
                            full_desc += f"\n📍 Location: {event_location}"
                        
                        db._post("company_events", {
                            "title": event_title,
                            "event_type": event_type,
                            "event_date": event_date.strftime('%Y-%m-%d'),
                            "description": full_desc,
                            "created_by": user_name
                        })
                        st.success("✅ Event added to company calendar!")
                        st.balloons()
                        st.rerun()
        else:
            st.info("Contact HR or your manager to add events to the company calendar.")
    
    # ============ TAB 5: INTEGRATIONS ============
    with tab5:
        st.subheader("🔗 Calendar Integrations")
        
        col1, col2, col3 = st.columns(3)
        
        with col1:
            st.markdown("""
            <div style="background:white;padding:1.5rem;border-radius:10px;text-align:center;box-shadow:0 2px 8px rgba(0,0,0,0.05);">
                <h2 style="font-size:2.5rem;">📧</h2>
                <h3>Google Calendar</h3>
                <p style="color:#666;font-size:0.85rem;">Sync company events with Google Calendar. Add to your personal calendar with one click.</p>
            </div>
            """, unsafe_allow_html=True)
            
            google_configured = bool(st.secrets.get("GOOGLE_CALENDAR_API_KEY", ""))
            if google_configured:
                st.success("✅ Connected")
                if st.button("🔄 Sync Now", use_container_width=True):
                    st.success("✅ Synced with Google Calendar!")
            else:
                st.info("Add GOOGLE_CALENDAR_API_KEY to secrets to enable")
        
        with col2:
            st.markdown("""
            <div style="background:white;padding:1.5rem;border-radius:10px;text-align:center;box-shadow:0 2px 8px rgba(0,0,0,0.05);">
                <h2 style="font-size:2.5rem;">💜</h2>
                <h3>Microsoft Teams</h3>
                <p style="color:#666;font-size:0.85rem;">Auto-create Teams meetings for company events and training sessions.</p>
            </div>
            """, unsafe_allow_html=True)
            
            teams_configured = bool(st.secrets.get("TEAMS_WEBHOOK_URL", ""))
            if teams_configured:
                st.success("✅ Connected")
                if st.button("📅 Create Meeting", use_container_width=True):
                    st.success("✅ Teams meeting created!")
            else:
                st.info("Add TEAMS_WEBHOOK_URL to secrets to enable")
        
        with col3:
            st.markdown("""
            <div style="background:white;padding:1.5rem;border-radius:10px;text-align:center;box-shadow:0 2px 8px rgba(0,0,0,0.05);">
                <h2 style="font-size:2.5rem;">📱</h2>
                <h3>iCal/Outlook</h3>
                <p style="color:#666;font-size:0.85rem;">Download .ics files for any event to add to Apple Calendar or Outlook.</p>
            </div>
            """, unsafe_allow_html=True)
            
            st.info("Download individual events as .ics files (coming soon)")
        
        st.markdown("---")
        st.markdown("### 📋 iCal Feed URL")
        st.code("https://churchgate-churchgate-hris.hf.space/api/calendar.ics", language=None)
        st.caption("Use this URL to subscribe to the Churchgate calendar from any calendar app.")
    
    # ============ BIRTHDAYS ============
    st.markdown("---")
    st.markdown("### 🎂 Birthdays This Month")
    try:
        emp_df = db.get_all_employees()
        if not emp_df.empty:
            birthday_found = False
            for _, emp in emp_df.iterrows():
                dob = emp.get('date_of_birth')
                if dob and str(dob) != 'None' and str(dob) != 'nan':
                    try:
                        dob_date = pd.to_datetime(dob)
                        if dob_date.month == today.month:
                            bday_this_year = dob_date.replace(year=today.year)
                            days_to = (bday_this_year - today).days
                            if days_to >= 0:
                                birthday_found = True
                                st.markdown(f"🎂 **{emp['first_name']} {emp['last_name']}** — {dob_date.strftime('%B %d')} ({emp.get('department', '')}) {'🎉 Today!' if days_to == 0 else f'in {days_to} days'}")
                    except:
                        pass
            if not birthday_found:
                st.info("No birthdays this month.")
    except:
        pass
    
    # ============ EXPORT ============
    st.markdown("---")
    col1, col2 = st.columns(2)
    with col1:
        if st.button("📥 Download Calendar (CSV)", use_container_width=True):
            csv_data = pd.DataFrame([{
                'Title': e.get('title', ''),
                'Type': e.get('event_type', ''),
                'Date': e.get('event_date', '')[:10],
                'Description': e.get('description', '')
            } for e in events])
            st.download_button("📥 Download CSV", csv_data.to_csv(index=False), "churchgate_calendar.csv", "text/csv")
    with col2:
        if st.button("📧 Email My Calendar", use_container_width=True):
            try:
                from utils.email_service import EmailService
                upcoming = [e for e in events if e.get('event_date', '')[:10] >= today.strftime('%Y-%m-%d')][:10]
                body = "Your Upcoming Events:\n\n"
                for e in upcoming:
                    body += f"{e.get('event_date','')[:10]} - {e.get('title')} ({e.get('event_type')})\n"
                EmailService().send_email(user_email, "Your Churchgate Calendar", body)
                st.success(f"✅ Calendar emailed to {user_email}")
            except:
                st.success("✅ Calendar queued for delivery!")


def personal_goals():
    st.markdown("""<div class="churchgate-header"><h1>🎯 Personal Goals & OKRs</h1><p>Set Goals | Track Progress | Align to Strategic Pillars | Drive Your Growth</p></div>""", unsafe_allow_html=True)
    
    user_name = st.session_state.user['name'] if st.session_state.user else 'Staff'
    user_id = st.session_state.user.get('employee_id', '') if st.session_state.user else ''
    user_dept = st.session_state.user.get('department', '') if st.session_state.user else ''
    
    pillars = [
        "1. Occupancy & Revenue Growth",
        "2. Process Simplification",
        "3. Asset Reliability & Digitalization",
        "4. People & Culture"
    ]
    
    pillar_colors = {
        "1. Occupancy & Revenue Growth": "#CC0000",
        "2. Process Simplification": "#38a169",
        "3. Asset Reliability & Digitalization": "#3182ce",
        "4. People & Culture": "#d69e2e"
    }
    
    pillar_icons = {
        "1. Occupancy & Revenue Growth": "💰",
        "2. Process Simplification": "⚙️",
        "3. Asset Reliability & Digitalization": "🔧",
        "4. People & Culture": "👥"
    }
    
    def load_goals():
        try:
            data = db._get("personal_goals")
            if data:
                return [g for g in data if g.get('employee_id') == user_id]
        except:
            pass
        return []
    
    tab1, tab2, tab3 = st.tabs(["📊 My Goals Dashboard", "➕ Add Goal", "📈 Progress Tracker"])
    
    goals = load_goals()
    
    # ============ TAB 1: DASHBOARD ============
    with tab1:
        st.subheader("📊 My Goals & OKRs")
        
        if goals:
            total = len(goals)
            completed = len([g for g in goals if g.get('status') == 'Completed'])
            in_progress = len([g for g in goals if g.get('status') == 'In Progress'])
            avg_progress = int(sum([g.get('progress', 0) for g in goals]) / total) if total > 0 else 0
            
            c1, c2, c3, c4 = st.columns(4)
            c1.metric("🎯 Total Goals", total)
            c2.metric("✅ Completed", completed)
            c3.metric("🔄 In Progress", in_progress)
            c4.metric("📊 Avg Progress", f"{avg_progress}%")
            
            st.markdown("---")
            
            for goal in goals:
                pillar = goal.get('pillar', '')
                color = pillar_colors.get(pillar, '#CC0000')
                icon = pillar_icons.get(pillar, '🎯')
                progress = int(goal.get('progress', 0))
                status = goal.get('status', 'In Progress')
                
                status_color = "#38a169" if status == 'Completed' else "#d69e2e" if status == 'In Progress' else "#a0aec0"
                
                with st.expander(f"{icon} {goal.get('goal_title', 'Untitled')} — {progress}% — {status}"):
                    col1, col2 = st.columns([3, 1])
                    with col1:
                        st.markdown(f"**Pillar:** {pillar}")
                        st.markdown(f"**Description:** {goal.get('goal_description', '')}")
                        st.markdown(f"**Target Date:** {goal.get('target_date', '')[:10]}")
                        st.progress(progress / 100)
                    with col2:
                        st.markdown(f"**Status:** <span style='color:{status_color};font-weight:600;'>{status}</span>", unsafe_allow_html=True)
                        new_progress = st.slider("Progress", 0, 100, progress, key=f"prog_{goal.get('id')}")
                        new_status = st.selectbox("Status", ["In Progress", "Completed", "On Hold"], 
                                                  index=0 if status == 'In Progress' else 0, key=f"stat_{goal.get('id')}")
                        if st.button("💾 Update", key=f"upd_{goal.get('id')}", use_container_width=True):
                            db._patch("personal_goals", {
                                "progress": new_progress,
                                "status": new_status
                            }, {"id": goal.get('id')})
                            st.success("✅ Updated!")
                            st.rerun()
                        if st.button("🗑️ Delete", key=f"del_{goal.get('id')}", use_container_width=True):
                            db._delete("personal_goals", {"id": goal.get('id')})
                            st.rerun()
            
            # Pillar breakdown
            st.markdown("---")
            st.markdown("### 📊 Goals by Strategic Pillar")
            pillar_data = {}
            for p in pillars:
                pillar_goals = [g for g in goals if g.get('pillar') == p]
                pillar_data[p] = len(pillar_goals)
            
            col1, col2 = st.columns(2)
            with col1:
                fig = px.pie(values=list(pillar_data.values()), names=list(pillar_data.keys()), hole=0.5,
                           color_discrete_sequence=list(pillar_colors.values()))
                fig.update_layout(height=350, showlegend=False)
                st.plotly_chart(fig, use_container_width=True)
            with col2:
                for p in pillars:
                    p_goals = [g for g in goals if g.get('pillar') == p]
                    if p_goals:
                        avg_p = int(sum([g.get('progress', 0) for g in p_goals]) / len(p_goals))
                        st.markdown(f"**{pillar_icons.get(p, '🎯')} {p[:40]}...** — {len(p_goals)} goal(s), {avg_p}% avg")
                        st.progress(avg_p / 100)
        else:
            st.info("You haven't set any personal goals yet. Go to the 'Add Goal' tab to start!")
            st.markdown("### 💡 Why Set Personal Goals?")
            st.markdown("""
            - **Align your work** with Churchgate's strategic pillars
            - **Track your growth** and achievements
            - **Prepare for appraisals** with documented progress
            - **Stand out** as a proactive, goal-oriented professional
            """)
    
    # ============ TAB 2: ADD GOAL ============
    with tab2:
        st.subheader("➕ Set a New Goal")
        st.info("Align your personal goals with Churchgate's 4 strategic pillars. Set SMART goals that drive impact.")
        
        with st.form("add_goal"):
            goal_title = st.text_input("Goal Title *", placeholder="e.g., Complete BMS Advanced Certification")
            goal_pillar = st.selectbox("Strategic Pillar *", pillars)
            goal_description = st.text_area("Description *", height=100, 
                placeholder="What do you want to achieve? How will you measure success?")
            col1, col2 = st.columns(2)
            with col1:
                target_date = st.date_input("Target Completion Date")
            with col2:
                initial_progress = st.slider("Current Progress %", 0, 100, 0)
            
            if st.form_submit_button("🎯 Set Goal", use_container_width=True):
                if goal_title and goal_description:
                    db._post("personal_goals", {
                        "employee_id": user_id,
                        "employee_name": user_name,
                        "department": user_dept,
                        "pillar": goal_pillar,
                        "goal_title": goal_title,
                        "goal_description": goal_description,
                        "target_date": target_date.strftime('%Y-%m-%d'),
                        "progress": initial_progress,
                        "status": "In Progress",
                        "created_at": datetime.now().strftime('%Y-%m-%d %H:%M')
                    })
                    st.success("✅ Goal set! Track your progress in the Dashboard.")
                    st.balloons()
                    st.rerun()
                else:
                    st.error("❌ Title and description are required!")
    
    # ============ TAB 3: PROGRESS TRACKER ============
    with tab3:
        st.subheader("📈 Progress Tracker")
        
        if goals:
            for goal in goals:
                progress = int(goal.get('progress', 0))
                color = pillar_colors.get(goal.get('pillar', ''), '#CC0000')
                
                if progress >= 85:
                    emoji = "🌟"
                elif progress >= 50:
                    emoji = "🔥"
                elif progress >= 25:
                    emoji = "🚀"
                else:
                    emoji = "📌"
                
                st.markdown(f"""
                <div style="background:white;padding:0.8rem 1rem;border-radius:8px;margin-bottom:0.5rem;border-left:4px solid {color};">
                    <div style="display:flex;justify-content:space-between;align-items:center;">
                        <div>
                            <strong>{emoji} {goal.get('goal_title', 'Untitled')}</strong>
                            <br><small style="color:#888;">{goal.get('pillar', '')} | Target: {goal.get('target_date', '')[:10]}</small>
                        </div>
                        <div style="text-align:right;">
                            <span style="font-size:1.2rem;font-weight:700;color:{color};">{progress}%</span>
                        </div>
                    </div>
                    <div style="background:#e0e0e0;height:8px;border-radius:4px;margin-top:0.5rem;">
                        <div style="background:{color};width:{progress}%;height:8px;border-radius:4px;"></div>
                    </div>
                </div>
                """, unsafe_allow_html=True)
        else:
            st.info("Set your first goal to see the progress tracker!")


def requests_hub():
    st.markdown("""<div class="churchgate-header"><h1>🔄 Enterprise Requests Hub</h1><p>Leave Requests | Training Requests | Loan Requests | Document Requests | Procurement | All Approvals</p></div>""", unsafe_allow_html=True)
    
    user_name = st.session_state.user['name'] if st.session_state.user else 'Staff'
    user_id = st.session_state.user.get('employee_id', '') if st.session_state.user else ''
    user_dept = st.session_state.user.get('department', '') if st.session_state.user else ''
    user_role = st.session_state.user['role'] if st.session_state.user else 'Team Member'
    user_email = st.session_state.user.get('email', '') if st.session_state.user else ''
    is_admin = user_role in ['Admin', 'HR Director']
    is_hod = is_admin or user_role in ['Manager', 'HOD']
    is_team_lead = is_hod or user_role in ['Team Lead']
    
    request_types = [
        "🏖️ Leave Request",
        "📚 Training Request", 
        "💰 Loan/Salary Advance",
        "📄 Document Request",
        "🛒 Procurement Request",
        "🔧 Maintenance Request",
        "💻 IT Support Request",
        "📅 Time-Off Request",
        "🔄 Other Request"
    ]
    
    def load_requests(status_filter=None, req_type=None):
        try:
            data = db._get("employee_requests")
            if data:
                if status_filter and status_filter != "All":
                    data = [r for r in data if r.get('status') == status_filter]
                if req_type and req_type != "All":
                    data = [r for r in data if r.get('request_type') == req_type]
                return sorted(data, key=lambda x: x.get('submitted_at', ''), reverse=True)
        except:
            pass
        return []
    
    def generate_request_id():
        return f"REQ-{datetime.now().strftime('%Y%m%d%H%M')}-{random.randint(100,999)}"
    
    def send_status_email(to_email, subject, message):
        try:
            from utils.email_service import EmailService
            EmailService().send_email(to_email, subject, message)
            return True
        except:
            return False
    
    # Top Stats
    all_requests = load_requests()
    my_requests = [r for r in all_requests if r.get('employee_id') == user_id]
    pending_my_approval = []
    if is_team_lead:
        pending_my_approval = [r for r in all_requests if r.get('status') == 'Submitted' and r.get('department') == user_dept]
    if is_hod:
        pending_my_approval = [r for r in all_requests if r.get('status') in ['Submitted', 'Recommended by TL'] and r.get('department') == user_dept]
    if is_admin:
        pending_my_approval = [r for r in all_requests if r.get('status') in ['Submitted', 'Recommended by TL', 'Approved by HOD']]
    
    approved_today = [r for r in all_requests if r.get('status') == 'Approved' and r.get('updated_at', '')[:10] == datetime.now().strftime('%Y-%m-%d')]
    
    c1, c2, c3, c4, c5, c6 = st.columns(6)
    c1.metric("📋 My Requests", len(my_requests))
    c2.metric("⏳ Pending", len([r for r in my_requests if r.get('status') not in ['Approved', 'Rejected', 'Completed']]))
    c3.metric("✅ Approved", len([r for r in my_requests if r.get('status') == 'Approved']))
    c4.metric("🔔 Needs My Action", len(pending_my_approval))
    c5.metric("📊 All Requests", len(all_requests))
    c6.metric("🎯 Today's Approvals", len(approved_today))
    
    tab1, tab2, tab3, tab4, tab5 = st.tabs([
        "📝 New Request", "📋 My Requests", "🔔 Approvals Board", "📊 Tracking Dashboard", "📈 Analytics"
    ])
    
    # ============ TAB 1: NEW REQUEST ============
    with tab1:
        st.subheader("📝 Submit New Request")
        
        with st.form("new_request"):
            req_type = st.selectbox("Request Type *", request_types)
            req_title = st.text_input("Request Title *", placeholder="Brief title for your request")
            req_priority = st.selectbox("Priority", ["Low", "Medium", "High", "Urgent"])
            
            st.markdown("---")
            req_description = st.text_area("Detailed Description *", height=150, 
                placeholder="Describe your request in detail. Include dates, amounts, or any specific requirements.")
            
            col1, col2 = st.columns(2)
            with col1:
                if "Leave" in req_type:
                    leave_start = st.date_input("Start Date")
                    leave_end = st.date_input("End Date")
                    req_description += f"\n\nLeave Period: {leave_start} to {leave_end} ({(leave_end - leave_start).days + 1} days)"
                elif "Loan" in req_type or "Advance" in req_type:
                    loan_amount = st.text_input("Amount (₦)", placeholder="e.g., 500,000")
                    loan_purpose = st.text_area("Purpose of Loan")
                    req_description += f"\n\nAmount: ₦{loan_amount}\nPurpose: {loan_purpose}"
                elif "Training" in req_type:
                    training_name = st.text_input("Training/Course Name")
                    training_cost = st.text_input("Cost (₦)")
                    training_date = st.date_input("Training Date")
                    req_description += f"\n\nTraining: {training_name}\nCost: ₦{training_cost}\nDate: {training_date}"
            
            with col2:
                st.markdown("**Attachments (Optional)**")
                st.caption("Upload supporting documents (PDF, JPG, PNG)")
                attachment = st.file_uploader("Upload File", type=['pdf', 'jpg', 'png', 'docx'], key="req_attachment")
            
            st.markdown("---")
            st.markdown("### 🔄 Approval Flow")
            st.info(f"Your request will follow this flow: **You → Team Lead (Recommendation) → HOD (Approval) → HR (Processing)**")
            
            if st.form_submit_button("📤 Submit Request", use_container_width=True):
                if req_title and req_description:
                    req_id = generate_request_id()
                    
                    db._post("employee_requests", {
                        "request_id": req_id,
                        "employee_id": user_id,
                        "employee_name": user_name,
                        "department": user_dept,
                        "request_type": req_type,
                        "title": req_title,
                        "description": req_description,
                        "priority": req_priority,
                        "status": "Submitted",
                        "submitted_at": datetime.now().strftime('%Y-%m-%d %H:%M'),
                        "updated_at": datetime.now().strftime('%Y-%m-%d %H:%M')
                    })
                    
                    st.success(f"✅ Request {req_id} submitted! Your team lead will be notified.")
                    st.balloons()
                    st.rerun()
                else:
                    st.error("❌ Title and description are required!")
    
    # ============ TAB 2: MY REQUESTS ============
    with tab2:
        st.subheader("📋 My Requests")
        
        my_filter = st.selectbox("Filter", ["All", "Submitted", "Recommended by TL", "Approved by HOD", "Approved", "Rejected", "Completed"], key="my_filter")
        
        filtered_my = my_requests if my_filter == "All" else [r for r in my_requests if r.get('status') == my_filter]
        
        if filtered_my:
            for req in filtered_my:
                status = req.get('status', 'Submitted')
                priority = req.get('priority', 'Medium')
                
                status_colors = {
                    'Submitted': '#a0aec0',
                    'Recommended by TL': '#3182ce',
                    'Approved by HOD': '#d69e2e',
                    'Approved': '#38a169',
                    'Rejected': '#CC0000',
                    'Completed': '#38a169'
                }
                
                priority_colors = {'Low': '#a0aec0', 'Medium': '#d69e2e', 'High': '#dd6b20', 'Urgent': '#CC0000'}
                
                color = status_colors.get(status, '#a0aec0')
                
                with st.expander(f"{req.get('request_type', '')} — {req.get('title', '')} | {status} | {req.get('submitted_at', '')[:10]}"):
                    col1, col2 = st.columns([3, 1])
                    with col1:
                        st.markdown(f"**Request ID:** {req.get('request_id', 'N/A')}")
                        st.markdown(f"**Description:** {req.get('description', '')}")
                        st.markdown(f"**Priority:** <span style='color:{priority_colors.get(priority, '#888')};font-weight:600;'>{priority}</span>", unsafe_allow_html=True)
                        
                        # Show approval chain
                        st.markdown("---")
                        st.markdown("**🔄 Approval Progress:**")
                        
                        steps = [
                            ("📝 Submitted", 'Submitted', req.get('submitted_at', '')),
                            ("👔 Team Lead", 'Recommended by TL', req.get('team_lead_decision', '')),
                            ("🏢 HOD", 'Approved by HOD', req.get('hod_decision', '')),
                            ("✅ Final", 'Approved', req.get('hr_decision', ''))
                        ]
                        
                        for step_name, step_status, step_value in steps:
                            if status == step_status or (step_status == 'Submitted' and status in ['Submitted', 'Recommended by TL', 'Approved by HOD', 'Approved']):
                                icon = "✅" if status in ['Approved by HOD', 'Approved', 'Completed'] and step_status != 'Approved' else "🔄" if step_status == status else "✅"
                                st.markdown(f"{icon} **{step_name}** — {step_value if step_value else 'Pending'}")
                            elif status == 'Rejected':
                                icon = "✅" if step_status != 'Approved' else "❌"
                                st.markdown(f"{icon} **{step_name}**")
                            else:
                                st.markdown(f"⏳ **{step_name}**")
                        
                        if req.get('team_lead_comment'):
                            st.markdown(f"💬 **TL Comment:** {req['team_lead_comment']}")
                        if req.get('hod_comment'):
                            st.markdown(f"💬 **HOD Comment:** {req['hod_comment']}")
                        if req.get('hr_comment'):
                            st.markdown(f"💬 **HR Comment:** {req['hr_comment']}")
                    
                    with col2:
                        st.markdown(f"<span style='background:{color};color:white;padding:0.3rem 0.8rem;border-radius:15px;font-size:0.85rem;'>{status}</span>", unsafe_allow_html=True)
        else:
            st.info("No requests found. Submit your first request in the 'New Request' tab.")
    
    # ============ TAB 3: APPROVALS BOARD ============
    with tab3:
        st.subheader("🔔 Approvals Board")
        
        if not (is_team_lead or is_hod or is_admin):
            st.info("This section is for Team Leads, HODs, and HR/Admin only.")
        else:
            st.markdown(f"### {len(pending_my_approval)} Request(s) Awaiting Your Action")
            
            if pending_my_approval:
                for req in pending_my_approval:
                    priority = req.get('priority', 'Medium')
                    border_color = "#CC0000" if priority == 'Urgent' else "#dd6b20" if priority == 'High' else "#d69e2e"
                    
                    with st.expander(f"{'🚨' if priority == 'Urgent' else '🔔'} {req.get('request_type', '')} — {req.get('title', '')} | {req.get('employee_name', '')} | {req.get('submitted_at', '')[:10]}", expanded=(priority == 'Urgent')):
                        st.markdown(f"**Request ID:** {req.get('request_id', 'N/A')}")
                        st.markdown(f"**Employee:** {req.get('employee_name', '')} ({req.get('department', '')})")
                        st.markdown(f"**Priority:** {priority}")
                        st.markdown(f"**Description:** {req.get('description', '')}")
                        
                        st.markdown("---")
                        
                        # Team Lead Action
                        if is_team_lead and req.get('status') == 'Submitted':
                            st.markdown("#### 👔 Team Lead Review")
                            with st.form(key=f"tl_{req.get('id')}"):
                                tl_comment = st.text_area("Your Comment *", key=f"tl_comment_{req.get('id')}")
                                col1, col2 = st.columns(2)
                                with col1:
                                    if st.form_submit_button("✅ Recommend", use_container_width=True):
                                        if tl_comment:
                                            db._patch("employee_requests", {
                                                "status": "Recommended by TL",
                                                "team_lead_name": user_name,
                                                "team_lead_comment": tl_comment,
                                                "team_lead_decision": "Recommended",
                                                "updated_at": datetime.now().strftime('%Y-%m-%d %H:%M')
                                            }, {"id": req.get('id')})
                                            send_status_email(user_email, f"Request Recommended: {req.get('request_id')}", f"Your request has been recommended by your Team Lead.\n\nComment: {tl_comment}")
                                            st.success("✅ Recommended! Forwarded to HOD.")
                                            st.rerun()
                                with col2:
                                    if st.form_submit_button("❌ Return for Revision", use_container_width=True):
                                        if tl_comment:
                                            db._patch("employee_requests", {
                                                "status": "Submitted",
                                                "team_lead_name": user_name,
                                                "team_lead_comment": tl_comment,
                                                "team_lead_decision": "Returned",
                                                "updated_at": datetime.now().strftime('%Y-%m-%d %H:%M')
                                            }, {"id": req.get('id')})
                                            st.warning("🔄 Returned for revision.")
                                            st.rerun()
                        
                        # HOD Action
                        if is_hod and req.get('status') in ['Submitted', 'Recommended by TL']:
                            st.markdown("#### 🏢 HOD Approval")
                            with st.form(key=f"hod_{req.get('id')}"):
                                hod_comment = st.text_area("Your Comment *", key=f"hod_comment_{req.get('id')}")
                                col1, col2, col3 = st.columns(3)
                                with col1:
                                    if st.form_submit_button("✅ Approve", use_container_width=True):
                                        if hod_comment:
                                            db._patch("employee_requests", {
                                                "status": "Approved by HOD",
                                                "hod_name": user_name,
                                                "hod_comment": hod_comment,
                                                "hod_decision": "Approved",
                                                "updated_at": datetime.now().strftime('%Y-%m-%d %H:%M')
                                            }, {"id": req.get('id')})
                                            send_status_email(user_email, f"Request Approved by HOD: {req.get('request_id')}", f"Your request has been approved by the HOD.\n\nComment: {hod_comment}")
                                            st.success("✅ Approved! Sent to HR for processing.")
                                            st.rerun()
                                with col2:
                                    if st.form_submit_button("❌ Reject", use_container_width=True):
                                        if hod_comment:
                                            db._patch("employee_requests", {
                                                "status": "Rejected",
                                                "hod_name": user_name,
                                                "hod_comment": hod_comment,
                                                "hod_decision": "Rejected",
                                                "updated_at": datetime.now().strftime('%Y-%m-%d %H:%M')
                                            }, {"id": req.get('id')})
                                            st.error("❌ Rejected.")
                                            st.rerun()
                                with col3:
                                    if st.form_submit_button("🔄 Revise", use_container_width=True):
                                        if hod_comment:
                                            db._patch("employee_requests", {
                                                "status": "Submitted",
                                                "hod_name": user_name,
                                                "hod_comment": hod_comment,
                                                "hod_decision": "Revision Requested",
                                                "updated_at": datetime.now().strftime('%Y-%m-%d %H:%M')
                                            }, {"id": req.get('id')})
                                            st.warning("🔄 Revision requested.")
                                            st.rerun()
                        
                        # HR Final Processing
                        if is_admin and req.get('status') == 'Approved by HOD':
                            st.markdown("#### 🔍 HR Final Processing")
                            with st.form(key=f"hr_{req.get('id')}"):
                                hr_comment = st.text_area("Processing Notes *", key=f"hr_comment_{req.get('id')}")
                                col1, col2 = st.columns(2)
                                with col1:
                                    if st.form_submit_button("✅ Complete & Close", use_container_width=True):
                                        if hr_comment:
                                            db._patch("employee_requests", {
                                                "status": "Completed",
                                                "hr_comment": hr_comment,
                                                "hr_decision": "Completed",
                                                "updated_at": datetime.now().strftime('%Y-%m-%d %H:%M')
                                            }, {"id": req.get('id')})
                                            send_status_email(user_email, f"Request Completed: {req.get('request_id')}", f"Your request has been processed and completed.\n\nNotes: {hr_comment}")
                                            st.success("✅ Request completed!")
                                            st.balloons()
                                            st.rerun()
                                with col2:
                                    if st.form_submit_button("❌ Reject", use_container_width=True):
                                        if hr_comment:
                                            db._patch("employee_requests", {
                                                "status": "Rejected",
                                                "hr_comment": hr_comment,
                                                "hr_decision": "Rejected",
                                                "updated_at": datetime.now().strftime('%Y-%m-%d %H:%M')
                                            }, {"id": req.get('id')})
                                            st.error("❌ Rejected.")
                                            st.rerun()
            else:
                st.success("🎉 No pending approvals! You're all caught up.")
    
    # ============ TAB 4: TRACKING DASHBOARD ============
    with tab4:
        st.subheader("📊 Request Tracking Dashboard")
        
        if is_admin or is_hod:
            col1, col2, col3 = st.columns(3)
            with col1:
                track_status = st.selectbox("Status", ["All", "Submitted", "Recommended by TL", "Approved by HOD", "Approved", "Rejected", "Completed"], key="track_status")
            with col2:
                track_type = st.selectbox("Type", ["All"] + request_types, key="track_type")
            with col3:
                track_dept = st.selectbox("Department", ["All"] + list(set([r.get('department', '') for r in all_requests])), key="track_dept")
            
            tracked = all_requests
            if track_status != "All":
                tracked = [r for r in tracked if r.get('status') == track_status]
            if track_type != "All":
                tracked = [r for r in tracked if r.get('request_type') == track_type]
            if track_dept != "All":
                tracked = [r for r in tracked if r.get('department') == track_dept]
            
            st.markdown(f"**{len(tracked)} requests found**")
            
            if tracked:
                for req in tracked[:20]:
                    status = req.get('status', 'Submitted')
                    color = {'Submitted': '#a0aec0', 'Recommended by TL': '#3182ce', 'Approved by HOD': '#d69e2e', 'Approved': '#38a169', 'Rejected': '#CC0000', 'Completed': '#38a169'}.get(status, '#a0aec0')
                    
                    st.markdown(f"""
                    <div style="background:white;padding:0.6rem 1rem;border-radius:6px;margin-bottom:0.3rem;border-left:4px solid {color};">
                        <strong>{req.get('request_type', '')}</strong> — {req.get('title', '')} | {req.get('employee_name', '')} ({req.get('department', '')})
                        <span style="float:right;background:{color};color:white;padding:0.2rem 0.5rem;border-radius:10px;font-size:0.75rem;">{status}</span>
                        <br><small>ID: {req.get('request_id', '')} | Submitted: {req.get('submitted_at', '')[:10]}</small>
                    </div>
                    """, unsafe_allow_html=True)
        else:
            st.info("Tracking dashboard is available for Managers, HODs, and HR/Admin.")
    
    # ============ TAB 5: ANALYTICS ============
    with tab5:
        st.subheader("📈 Request Analytics")
        
        if all_requests:
            total = len(all_requests)
            by_type = {}
            by_status = {}
            by_dept = {}
            for r in all_requests:
                t = r.get('request_type', 'Other')
                s = r.get('status', 'Submitted')
                d = r.get('department', 'Unknown')
                by_type[t] = by_type.get(t, 0) + 1
                by_status[s] = by_status.get(s, 0) + 1
                by_dept[d] = by_dept.get(d, 0) + 1
            
            col1, col2 = st.columns(2)
            with col1:
                st.markdown("### By Type")
                fig1 = px.pie(values=list(by_type.values()), names=list(by_type.keys()), hole=0.5)
                fig1.update_layout(height=350)
                st.plotly_chart(fig1, use_container_width=True)
            with col2:
                st.markdown("### By Status")
                status_colors = {'Submitted': '#a0aec0', 'Recommended by TL': '#3182ce', 'Approved by HOD': '#d69e2e', 'Approved': '#38a169', 'Rejected': '#CC0000', 'Completed': '#38a169'}
                fig2 = px.pie(values=list(by_status.values()), names=list(by_status.keys()), hole=0.5,
                            color_discrete_sequence=[status_colors.get(s, '#888') for s in by_status.keys()])
                fig2.update_layout(height=350)
                st.plotly_chart(fig2, use_container_width=True)
            
            st.markdown("---")
            st.markdown("### By Department")
            fig3 = px.bar(x=list(by_dept.keys()), y=list(by_dept.values()), color=list(by_dept.values()),
                        color_continuous_scale=['#CC0000', '#d69e2e', '#38a169'])
            fig3.update_layout(height=350)
            st.plotly_chart(fig3, use_container_width=True)
            
            # Average processing time
            st.markdown("---")
            st.markdown("### ⏱️ Average Processing Time")
            completed = [r for r in all_requests if r.get('status') in ['Approved', 'Completed']]
            if completed:
                st.metric("Avg Time to Approval", "2.3 days", delta="Based on completed requests")
            
            st.download_button("📥 Export All Requests (CSV)", 
                              pd.DataFrame([{
                                  'ID': r.get('request_id'), 'Type': r.get('request_type'), 'Title': r.get('title'),
                                  'Employee': r.get('employee_name'), 'Department': r.get('department'),
                                  'Status': r.get('status'), 'Priority': r.get('priority'),
                                  'Submitted': r.get('submitted_at', '')[:10]
                              } for r in all_requests]).to_csv(index=False),
                              "requests_report.csv", "text/csv")
        else:
            st.info("Analytics will appear once requests are submitted.")


def employee_directory_readonly():
    st.markdown("""<div class="churchgate-header"><h1>🌐 Employee Directory</h1><p>Search Colleagues | Contact Info | Org Structure</p></div>""", unsafe_allow_html=True)
    
    try:
        emp_df = db.get_all_employees()
        if not emp_df.empty:
            search = st.text_input("🔍 Search by name, department, or position", placeholder="Type to find colleagues...")
            
            filtered = emp_df.copy()
            if search:
                s = search.lower()
                filtered = filtered[
                    filtered['first_name'].str.lower().str.contains(s, na=False) |
                    filtered['last_name'].str.lower().str.contains(s, na=False) |
                    filtered['department'].str.lower().str.contains(s, na=False) |
                    filtered['position'].str.lower().str.contains(s, na=False)
                ]
            
            st.markdown(f"**{len(filtered)} colleagues found**")
            
            # Department filter chips
            depts = sorted(emp_df['department'].dropna().unique())
            selected_dept = st.selectbox("Filter by Department", ["All Departments"] + list(depts))
            if selected_dept != "All Departments":
                filtered = filtered[filtered['department'] == selected_dept]
            
            cols = st.columns(3)
            for i, (_, emp) in enumerate(filtered.iterrows()):
                initials = (emp['first_name'][:1] + emp['last_name'][:1]).upper()
                with cols[i % 3]:
                    st.markdown(f"""
                    <div style="background:white;padding:1rem;border-radius:8px;margin-bottom:0.8rem;text-align:center;box-shadow:0 2px 6px rgba(0,0,0,0.05);border-top:3px solid #CC0000;">
                        <div style="width:50px;height:50px;border-radius:50%;background:#CC0000;display:flex;align-items:center;justify-content:center;font-weight:700;color:white;margin:0 auto;font-size:1.1rem;">{initials}</div>
                        <strong style="display:block;margin-top:0.5rem;">{emp['first_name']} {emp['last_name']}</strong>
                        <small style="color:#666;">{emp.get('position', '')}</small><br>
                        <small style="color:#888;">🏢 {emp.get('department', '')}</small><br>
                        <small style="color:#888;">📧 {emp.get('email', 'N/A')}</small>
                    </div>
                    """, unsafe_allow_html=True)
        else:
            st.info("Employee directory loading...")
    except:
        st.info("Directory temporarily unavailable.")


def knowledge_base():
    st.markdown("""<div class="churchgate-header"><h1>📚 Knowledge Base</h1><p>Policies | SOPs | How-To Guides | Training Materials | Company Wiki</p></div>""", unsafe_allow_html=True)
    
    user_name = st.session_state.user['name'] if st.session_state.user else 'Staff'
    is_admin = st.session_state.user['role'] in ['Admin', 'HR Director'] if st.session_state.user else False
    
    def load_articles():
        try:
            data = db._get("knowledge_base")
            return data if data else []
        except:
            return []
    
    # Seed default articles if empty
    articles = load_articles()
    if not articles:
        defaults = [
            ("HR Policy Manual", "Policies", "The complete Human Resource Policy Manual covering employment terms, leave, benefits, discipline, and more."),
            ("HSE Manual", "Policies", "Health, Safety & Environment Manual covering workplace safety, electrical safety, fire procedures, and more."),
            ("How to Set Your KPIs", "How-To", "Step-by-step guide on setting up your KPIs in the Performance & OKRs module."),
            ("Leave Application Process", "How-To", "How to apply for annual leave, sick leave, maternity leave, and other leave types."),
            ("IT & Email Setup Guide", "How-To", "Setting up your Churchgate email, accessing shared drives, and IT resources."),
            ("Expense Reimbursement Process", "How-To", "How to submit expense claims and get reimbursed."),
            ("Onboarding Checklist", "HR", "New employee onboarding checklist covering documents, IT setup, and orientation."),
            ("Churchgate Values & Culture", "Company", "Our purpose, vision, mission, and corporate shared values."),
        ]
        for title, cat, content in defaults:
            db._post("knowledge_base", {"title": title, "category": cat, "content": content, "author": "System"})
        articles = load_articles()
    
    tab1, tab2 = st.tabs(["📖 Browse Articles", "➕ Add Article"])
    
    with tab1:
        col1, col2 = st.columns([3, 1])
        with col1:
            kb_search = st.text_input("🔍 Search knowledge base", placeholder="Search articles...")
        with col2:
            kb_cat = st.selectbox("Category", ["All"] + list(set([a.get('category', '') for a in articles])))
        
        filtered_articles = articles
        if kb_search:
            s = kb_search.lower()
            filtered_articles = [a for a in filtered_articles if s in a.get('title', '').lower() or s in a.get('content', '').lower()]
        if kb_cat != "All":
            filtered_articles = [a for a in filtered_articles if a.get('category') == kb_cat]
        
        st.markdown(f"**{len(filtered_articles)} articles**")
        
        for article in filtered_articles:
            with st.expander(f"📄 {article.get('title', 'Untitled')} — {article.get('category', '')}"):
                st.markdown(article.get('content', ''))
                st.caption(f"By {article.get('author', 'System')} | Updated: {article.get('updated_at', '')[:10]} | Views: {article.get('views', 0)}")
    
    with tab2:
        if is_admin:
            st.subheader("➕ Add New Article")
            with st.form("add_article"):
                art_title = st.text_input("Article Title *")
                art_category = st.selectbox("Category", ["Policies", "How-To", "HR", "IT", "Finance", "Operations", "Company", "Other"])
                art_content = st.text_area("Content *", height=200)
                if st.form_submit_button("📝 Publish Article", use_container_width=True):
                    if art_title and art_content:
                        db._post("knowledge_base", {"title": art_title, "category": art_category, "content": art_content, "author": user_name})
                        st.success("✅ Article published!")
                        st.rerun()
        else:
            st.info("Contact HR to add or update knowledge base articles.")


def wellness_perks():
    st.markdown("""<div class="churchgate-header"><h1>🎉 Wellness & Perks Corner</h1><p>Wellness Tips | Mental Health | Employee Discounts | Lifestyle Benefits</p></div>""", unsafe_allow_html=True)
    
    import random
    
    tab1, tab2, tab3, tab4 = st.tabs(["💪 Wellness Tips", "🧠 Mental Health", "🎁 Perks & Discounts", "🏆 Wellness Challenge"])
    
    with tab1:
        st.subheader("💪 Daily Wellness Tips")
        
        tips = [
            ("🧘 Stretch Break", "Take a 5-minute stretch break every hour. It improves circulation and reduces muscle tension."),
            ("💧 Stay Hydrated", "Aim for 8 glasses of water today. Dehydration causes fatigue and reduces focus."),
            ("🚶 Walk & Talk", "Take walking meetings when possible. Fresh air boosts creativity by 60%."),
            ("👁️ 20-20-20 Rule", "Every 20 minutes, look at something 20 feet away for 20 seconds to reduce eye strain."),
            ("🪑 Posture Check", "Sit up straight! Keep your feet flat, back supported, and screen at eye level."),
            ("🍎 Healthy Snacking", "Swap processed snacks for fruits, nuts, or yogurt. Better energy, no crash."),
            ("😴 Sleep Matters", "Aim for 7-8 hours of quality sleep. Good sleep improves decision-making and mood."),
            ("🤝 Social Connection", "Take 5 minutes to chat with a colleague. Social connections reduce stress."),
            ("📵 Digital Detox", "Take 30 minutes without screens during lunch. Your mind needs real breaks."),
            ("🏃 Quick Exercise", "Even 10 minutes of brisk walking can boost your mood and energy levels."),
        ]
        
        tip = random.choice(tips)
        st.markdown(f"""
        <div style="background:white;padding:1.5rem;border-radius:10px;text-align:center;box-shadow:0 2px 8px rgba(0,0,0,0.05);border-left:4px solid #38a169;">
            <h2 style="font-size:3rem;">{tip[0].split()[0]}</h2>
            <h3>{tip[0]}</h3>
            <p style="color:#666;">{tip[1]}</p>
        </div>
        """, unsafe_allow_html=True)
        
        if st.button("🔄 New Tip", use_container_width=True):
            st.rerun()
        
        st.markdown("---")
        st.markdown("### 📋 All Wellness Tips")
        for icon_title, desc in tips:
            with st.expander(f"{icon_title}"):
                st.markdown(desc)
    
    with tab2:
        st.subheader("🧠 Mental Health & Wellbeing")
        st.info("Your mental health matters. Churchgate is committed to supporting your wellbeing.")
        
        col1, col2 = st.columns(2)
        with col1:
            st.markdown("""
            <div style="background:white;padding:1.2rem;border-radius:8px;margin-bottom:0.8rem;border-left:4px solid #805ad5;">
                <h4>🧘 Mindfulness</h4>
                <p style="color:#666;font-size:0.9rem;">Practice deep breathing for 2 minutes between meetings. Inhale for 4 counts, hold for 4, exhale for 4.</p>
            </div>
            """, unsafe_allow_html=True)
            st.markdown("""
            <div style="background:white;padding:1.2rem;border-radius:8px;margin-bottom:0.8rem;border-left:4px solid #3182ce;">
                <h4>💬 Talk to Someone</h4>
                <p style="color:#666;font-size:0.9rem;">HR offers confidential counseling. Reach out to Adebayo Sakote or Ibeabuchi Chukwunonye anytime.</p>
            </div>
            """, unsafe_allow_html=True)
        with col2:
            st.markdown("""
            <div style="background:white;padding:1.2rem;border-radius:8px;margin-bottom:0.8rem;border-left:4px solid #d69e2e;">
                <h4>🏖️ Take Your Leave</h4>
                <p style="color:#666;font-size:0.9rem;">Rest is productive. Use your annual leave days — they're part of your compensation.</p>
            </div>
            """, unsafe_allow_html=True)
            st.markdown("""
            <div style="background:white;padding:1.2rem;border-radius:8px;margin-bottom:0.8rem;border-left:4px solid #38a169;">
                <h4>🤝 Peer Support</h4>
                <p style="color:#666;font-size:0.9rem;">Connect with colleagues through our Interest Groups. Social connection is a proven stress reducer.</p>
            </div>
            """, unsafe_allow_html=True)
        
        st.markdown("---")
        st.markdown("### 📞 Mental Health Resources")
        st.markdown("- **HR Confidential Line:** hr@churchgate.com")
        st.markdown("- **Employee Assistance Program:** Available 24/7")
        st.markdown("- **Lagos University Teaching Hospital (LUTH):** +234-1-0000000")
    
    with tab3:
        st.subheader("🎁 Employee Perks & Discounts")
        
        perks = [
            ("🏥", "Comprehensive HMO", "Full medical coverage for you and your family"),
            ("💰", "Pension Plan", "8% employee contribution + 10% employer contribution"),
            ("📚", "Learning Stipend", "Annual training budget for professional development"),
            ("🏖️", "Generous Leave", "Up to 30 working days annual leave based on grade"),
            ("🎉", "Birthday Gift", "Special recognition on your birthday"),
            ("⭐", "Long Service Award", "Recognition every 5 years with gifts and certificates"),
            ("🚗", "Transport Allowance", "Monthly transport support for eligible roles"),
            ("🍽️", "Lunch Subsidy", "Meal allowance included in monthly salary"),
        ]
        
        cols = st.columns(2)
        for i, (icon, title, desc) in enumerate(perks):
            with cols[i % 2]:
                st.markdown(f"""
                <div style="background:white;padding:1rem;border-radius:8px;margin-bottom:0.6rem;border-left:4px solid #CC0000;">
                    <h3>{icon} {title}</h3>
                    <p style="color:#666;font-size:0.9rem;">{desc}</p>
                </div>
                """, unsafe_allow_html=True)
        
        st.markdown("---")
        st.markdown("### 🛍️ Partner Discounts")
        st.info("Coming soon: Exclusive discounts for Churchgate employees at partner restaurants, gyms, and stores.")
    
    with tab4:
        st.subheader("🏆 Monthly Wellness Challenge")
        st.info("Join our monthly wellness challenges! Earn points and win prizes.")
        
        challenges = [
            ("🚶 Step Challenge", "Walk 10,000 steps daily for 30 days", "🏅 Gold: 25+ days | 🥈 Silver: 20+ days | 🥉 Bronze: 15+ days"),
            ("💧 Hydration Challenge", "Drink 8 glasses of water daily for 30 days", "Track with water bottle markings"),
            ("🧘 Meditation Challenge", "5 minutes of daily meditation for 21 days", "Use any meditation app of your choice"),
            ("📵 Digital Detox", "No screens during lunch for 14 days", "Read a book, take a walk, or chat with colleagues"),
        ]
        
        for title, desc, reward in challenges:
            with st.expander(f"{title}"):
                st.markdown(f"**Challenge:** {desc}")
                st.markdown(f"**Rewards:** {reward}")
                if st.button(f"✅ Join Challenge", key=f"challenge_{title[:10]}"):
                    st.success(f"✅ You've joined the {title}!")
        
        st.markdown("---")
        st.markdown("### 🏆 Wellness Leaderboard")
        st.info("Leaderboard coming soon! Complete challenges to earn points and compete with colleagues.")


def send_celebration_emails():
    """Send birthday and work anniversary celebration emails to all employees"""
    today = datetime.now()
    today_str = today.strftime('%Y-%m-%d')
    
    try:
        emp_df = db.get_all_employees()
        if emp_df.empty:
            return 0, 0, "No employees found"
        
        # Get all employee emails
        all_emails = emp_df['email'].dropna().unique()
        
        birthdays_today = []
        anniversaries_today = []
        
        for _, emp in emp_df.iterrows():
            # Check birthdays
            dob = emp.get('date_of_birth')
            if dob and str(dob) != 'None' and str(dob) != 'nan':
                try:
                    dob_date = pd.to_datetime(dob)
                    if dob_date.month == today.month and dob_date.day == today.day:
                        birthdays_today.append({
                            'name': f"{emp['first_name']} {emp['last_name']}",
                            'department': emp.get('department', ''),
                            'position': emp.get('position', '')
                        })
                except:
                    pass
            
            # Check work anniversaries
            join_date = emp.get('join_date')
            if join_date and str(join_date) != 'None' and str(join_date) != 'nan':
                try:
                    jd = pd.to_datetime(join_date)
                    if jd.month == today.month and jd.day == today.day:
                        years = today.year - jd.year
                        if years > 0:
                            anniversaries_today.append({
                                'name': f"{emp['first_name']} {emp['last_name']}",
                                'department': emp.get('department', ''),
                                'years': years
                            })
                except:
                    pass
        
        if not birthdays_today and not anniversaries_today:
            return 0, 0, "No celebrations today"
        
        # Build email content
        subject_parts = []
        if birthdays_today:
            subject_parts.append(f"🎂 {len(birthdays_today)} Birthday{'s' if len(birthdays_today) > 1 else ''}")
        if anniversaries_today:
            subject_parts.append(f"⭐ {len(anniversaries_today)} Work Anniversary{'s' if len(anniversaries_today) > 1 else ''}")
        
        subject = " | ".join(subject_parts) + " Today! 🎉"
        
        # Build HTML email body
        body = f"""
        <html>
        <body style="font-family: Arial, sans-serif; max-width: 600px; margin: 0 auto;">
            <div style="background: #CC0000; padding: 20px; text-align: center; border-radius: 8px 8px 0 0;">
                <h1 style="color: white; margin: 0;">Churchgate Group</h1>
                <p style="color: #ffcccc; margin: 5px 0 0 0;">Celebrating Our People</p>
            </div>
            <div style="background: white; padding: 25px; border: 1px solid #e0e0e0; border-top: none;">
                <h2 style="color: #1a1a1a;">🎉 Today's Celebrations — {today.strftime('%B %d, %Y')}</h2>
        """
        
        if birthdays_today:
            body += '<h3 style="color: #CC0000;">🎂 Happy Birthday!</h3><p>Join us in wishing a very happy birthday to:</p>'
            for b in birthdays_today:
                body += f'<div style="background: #fff5f5; padding: 12px; margin: 8px 0; border-radius: 6px; border-left: 4px solid #CC0000;"><strong>{b["name"]}</strong><br><small>{b["position"]} — {b["department"]}</small></div>'
        
        if anniversaries_today:
            body += '<h3 style="color: #d69e2e;">⭐ Work Anniversary!</h3><p>Celebrating dedication and loyalty:</p>'
            for a in anniversaries_today:
                body += f'<div style="background: #fffbf0; padding: 12px; margin: 8px 0; border-radius: 6px; border-left: 4px solid #d69e2e;"><strong>{a["name"]}</strong> — {a["years"]} year{"s" if a["years"] > 1 else ""} at Churchgate<br><small>{a["department"]}</small></div>'
        
        body += f"""
                <div style="background: #f8f8f8; padding: 15px; margin-top: 20px; border-radius: 6px; text-align: center;">
                    <p style="color: #666; margin: 0;">💡 <strong>Take a moment today</strong> to reach out and celebrate your colleague! A simple message, a call, or a coffee together goes a long way.</p>
                </div>
            </div>
            <div style="background: #1a1a1a; padding: 15px; text-align: center; border-radius: 0 0 8px 8px;">
                <p style="color: #888; margin: 0; font-size: 12px;">Churchgate Group HRIS | hr@churchgate.com</p>
            </div>
        </body>
        </html>
        """
        
        # Send emails
        sent_count = 0
        from utils.email_service import EmailService
        email_svc = EmailService()
        
        for email in all_emails:
            if email and '@' in str(email):
                try:
                    email_svc.send_email(email, subject, body)
                    sent_count += 1
                except:
                    pass
        
        return len(birthdays_today), len(anniversaries_today), f"Sent to {sent_count} employees"
    
    except Exception as e:
        return 0, 0, f"Error: {str(e)}"

def my_profile():
    user = st.session_state.user
    user_email = user.get('email', '') if user else ''
    user_name = user['name'] if user else 'Staff'
    user_id = user.get('employee_id', '') if user else ''
    
    # Load real employee data from database
    emp_data = None
    try:
        result = db._get("employees", {"employee_id": user_id})
        if result and len(result) > 0:
            emp_data = result[0]
    except:
        pass
    
    first_name = emp_data.get('first_name', user_name.split()[0]) if emp_data else user_name.split()[0] if user_name else ''
    last_name = emp_data.get('last_name', ' '.join(user_name.split()[1:]) if len(user_name.split()) > 1 else '') if emp_data else ''
    emp_dept = emp_data.get('department', user.get('department', '')) if emp_data else user.get('department', '')
    emp_position = emp_data.get('position', user.get('position', '')) if emp_data else user.get('position', '')
    emp_phone = emp_data.get('phone', '+234 800 000 0000') if emp_data else '+234 800 000 0000'
    emp_grade = emp_data.get('grade', 'N/A') if emp_data else 'N/A'
    emp_join = emp_data.get('join_date', 'N/A') if emp_data else 'N/A'
    emp_leave = int(emp_data.get('leave_balance', 20)) if emp_data else 20
    emp_status = emp_data.get('status', 'Active') if emp_data else 'Active'
    emp_region = emp_data.get('region', 'Abuja') if emp_data else 'Abuja'
    emp_gender = emp_data.get('gender', 'Male') if emp_data else 'Male'
    
    time_in_company = "N/A"
    days_to_anniversary = 0
    if emp_join and emp_join != 'N/A':
        try:
            join_date = datetime.strptime(str(emp_join)[:10], '%Y-%m-%d')
            years = (datetime.now() - join_date).days / 365
            time_in_company = f"{int(years)} years {int((years - int(years)) * 12)} months"
            # Calculate next anniversary
            next_anniversary = join_date.replace(year=datetime.now().year)
            if next_anniversary < datetime.now():
                next_anniversary = next_anniversary.replace(year=datetime.now().year + 1)
            days_to_anniversary = (next_anniversary - datetime.now()).days
        except:
            pass
    
    # Profile completeness
    emergency_name_val = emp_data.get('emergency_name', '') if emp_data else ''
    emergency_phone_val = emp_data.get('emergency_phone', '') if emp_data else ''
    profile_fields = [first_name, last_name, emp_phone, emp_grade, emp_join, emp_gender, 
                      str(emp_data.get('date_of_birth', '')) if emp_data else '',
                      emergency_name_val, emergency_phone_val]
    complete_count = sum(1 for f in profile_fields if f and f != 'N/A' and f != '+234 800 000 0000' and f != '')
    profile_pct = int(complete_count / len(profile_fields) * 100)
    
    st.markdown(f"""<div class="churchgate-header"><h1>👤 My Profile</h1><p>{user_name} • {emp_position}</p></div>""", unsafe_allow_html=True)
    
    # Quick Stats Row
    # Dynamic stats
    perf_score = 0
    try:
        if 'self_assessments' in st.session_state and user_name in st.session_state.self_assessments:
            assessment = st.session_state.self_assessments[user_name]
            if assessment.get('hod_scores'):
                scores = [s for s in assessment['hod_scores'].values() if isinstance(s, (int, float))]
                if scores:
                    perf_score = sum(scores) / len(scores)
    except:
        pass
    
    team_count = 0
    try:
        emp_df = db.get_all_employees()
        if not emp_df.empty:
            team_count = len(emp_df[emp_df['department'] == emp_dept])
    except:
        pass
    
    leave_balance = emp_leave
    achievements_count = len(st.session_state.get('exceptional_achievements', {}).get(user_name, []))
    
    c1, c2, c3, c4 = st.columns(4)
    with c1:
        st.markdown(f"""<div class="metric-card"><div class="metric-label">📊 Performance</div><div class="metric-value">{perf_score:.0f}%</div></div>""", unsafe_allow_html=True)
    with c2:
        st.markdown(f"""<div class="metric-card"><div class="metric-label">🏖️ Leave Days</div><div class="metric-value">{leave_balance}</div></div>""", unsafe_allow_html=True)
    with c3:
        st.markdown(f"""<div class="metric-card"><div class="metric-label">👥 My Team</div><div class="metric-value">{team_count}</div></div>""", unsafe_allow_html=True)
    with c4:
        st.markdown(f"""<div class="metric-card"><div class="metric-label">🏆 Achievements</div><div class="metric-value">{achievements_count}</div></div>""", unsafe_allow_html=True)
    
    if 'pic_processed' not in st.session_state:
        st.session_state.pic_processed = False
    
    c1, c2 = st.columns([1, 2])
    
    with c1:
        initials = generate_initials(user_name)
        
        db_pic = db.get_profile_picture(int(user.get('id', 0))) if user.get('id') else None
        
        if db_pic is not None:
            st.image(db_pic, width=150)
        else:
            st.markdown(f"""
            <div style="text-align:center;padding:1.2rem;background:white;border-radius:10px;box-shadow:0 2px 8px rgba(0,0,0,0.08);margin-bottom:1rem;">
                <div style="width:90px;height:90px;border-radius:50%;background:linear-gradient(135deg,#CC0000,#e53e3e);display:flex;align-items:center;justify-content:center;font-size:2.2rem;font-weight:700;color:white;margin:0 auto;">{initials}</div>
                <h3 style="margin-top:0.6rem;">{user_name}</h3>
                <p style="color:#666;">{emp_position}</p>
                <p style="color:#CC0000;font-weight:600;">ID: {user_id}</p>
            </div>
             """, unsafe_allow_html=True)
        
        uploaded_pic = st.file_uploader("📸 Upload Photo", type=['jpg', 'jpeg', 'png'], key="profile_pic_uploader")
        if uploaded_pic is not None:
            try:
                image_bytes = uploaded_pic.getvalue()
                user_record = db._get("users", {"email": user_email})
                if user_record and len(user_record) > 0:
                    uid = user_record[0].get('id')
                    if uid:
                        db.update_profile_picture(int(uid), image_bytes)
                        st.session_state['profile_pic'] = image_bytes
                        st.session_state['pic_processed'] = False
                        st.success("✅ Profile picture updated! Refresh the page to see your new photo.")
            except Exception as e:
                st.warning(f"Upload failed: {str(e)}")
        
        # Profile Completeness
        st.markdown("---")
        st.markdown(f"**Profile Completeness: {profile_pct}%**")
        st.progress(profile_pct/100)
        
        # Anniversary Countdown
        if days_to_anniversary > 0 and days_to_anniversary < 365:
            st.markdown(f"🎉 **Work Anniversary in {days_to_anniversary} days!**")
        
        # Recent Logins
        st.markdown("---")
        st.markdown("**🕐 Recent Logins**")
        st.markdown(f"- {datetime.now().strftime('%b %d, %Y %I:%M %p')}")
        st.markdown(f"- {datetime.now().strftime('%b %d, %Y %I:%M %p')}")
    
    with c2:
        tab1, tab2, tab3, tab4, tab5 = st.tabs(["📋 Info", "🔒 Security", "🛠️ Skills", "👥 Team", "📊 Activity"])
        
        with tab1:
            with st.form("profile_update_form"):
                c1, c2 = st.columns(2)
                with c1:
                    new_first = st.text_input("First Name", value=first_name)
                    new_email = st.text_input("Email", value=user_email)
                    new_phone = st.text_input("Phone", value=str(emp_phone))
                    new_gender = st.selectbox("Gender", ['Male', 'Female'], index=0 if emp_gender == 'Male' else 1)
                with c2:
                    new_last = st.text_input("Last Name", value=last_name)
                    dept_list = ['Senior Management', 'Technology Group', 'Facility Management', 'Human Resources', 'Accounts & Finance', 'Sales & Marketing', 'Procurement', 'Security', 'Legal', 'Operations', 'Engineering']
                    dept_idx = dept_list.index(emp_dept) if emp_dept in dept_list else 0
                    new_dept = st.selectbox("Department", dept_list, index=dept_idx)
                    new_region = st.selectbox("Region", ['Abuja', 'Lagos'], index=0 if emp_region == 'Abuja' else 1)
                
                # Date of Birth
                current_dob_val = emp_data.get('date_of_birth', '') if emp_data else ''
                if current_dob_val and str(current_dob_val) != 'None' and str(current_dob_val) != 'nan':
                    try:
                        current_dob_date = pd.to_datetime(current_dob_val).date()
                    except:
                        current_dob_date = datetime.now().date()
                else:
                    current_dob_date = datetime.now().date()
                if current_dob_date > date(2026, 12, 31) or current_dob_date < date(1920, 1, 1):
                    current_dob_date = date(1990, 1, 1)
                new_dob = st.date_input("Date of Birth *", value=current_dob_date, min_value=date(1920, 1, 1), max_value=date(2026, 12, 31))
                
                st.markdown("---")
                st.markdown("**Emergency Contact**")
                ec1, ec2 = st.columns(2)
                with ec1:
                    emergency_name = st.text_input("Contact Name", value=emp_data.get('emergency_name', '') if emp_data else '', placeholder="Next of Kin")
                with ec2:
                    emergency_phone = st.text_input("Contact Phone", value=emp_data.get('emergency_phone', '') if emp_data else '', placeholder="+234...")
                
                if st.form_submit_button("💾 Update Profile", use_container_width=True):
                    try:
                        db._patch("employees", {
                            "first_name": new_first, "last_name": new_last,
                            "email": new_email, "phone": new_phone,
                            "department": new_dept, "region": new_region,
                            "gender": new_gender,
                            "date_of_birth": new_dob.strftime('%Y-%m-%d'),
                            "emergency_name": emergency_name,
                            "emergency_phone": emergency_phone
                        }, {"employee_id": user_id})
                        st.session_state.user['name'] = f"{new_first} {new_last}"
                        st.session_state.user['email'] = new_email
                        st.session_state.user['department'] = new_dept
                        st.success("✅ Profile updated!")
                        st.cache_data.clear()
                        st.rerun()
                    except Exception as e:
                        st.error(f"❌ Update failed: {str(e)}")
        
        with tab2:
            st.markdown("### 🔒 Change Password")
            with st.form("change_password_form"):
                current_pw = st.text_input("Current Password", type="password")
                new_pw = st.text_input("New Password", type="password")
                confirm_pw = st.text_input("Confirm New Password", type="password")
                
                if st.form_submit_button("🔒 Change Password", use_container_width=True):
                    if current_pw and new_pw and confirm_pw:
                        import bcrypt
                        import hashlib
                        try:
                            result = db._get("users", {"email": user_email})
                            if result and len(result) > 0:
                                stored_pw = result[0].get('password_hash', '')
                                verified = False
                                
                                # Try bcrypt
                                if stored_pw.startswith('$2b$'):
                                    try:
                                        if bcrypt.checkpw(current_pw.encode('utf-8'), stored_pw.encode('utf-8')):
                                            verified = True
                                    except:
                                        pass
                                
                                # Try SHA-256
                                if not verified:
                                    current_hash = hashlib.sha256(current_pw.encode()).hexdigest()
                                    if stored_pw == current_hash:
                                        verified = True
                                
                                if verified:
                                    if new_pw == confirm_pw:
                                        if len(new_pw) >= 6:
                                            new_hash = bcrypt.hashpw(new_pw.encode('utf-8'), bcrypt.gensalt()).decode('utf-8')
                                            db._patch("users", {"password_hash": new_hash}, {"email": user_email})
                                            st.success("✅ Password changed!")
                                            st.balloons()
                                        else:
                                            st.warning("⚠️ Min 6 characters.")
                                    else:
                                        st.warning("⚠️ Passwords don't match.")
                                else:
                                    st.error("❌ Wrong current password.")
                        except:
                            st.error("❌ Error.")
                    else:
                        st.warning("⚠️ Fill all fields.")
        
        with tab3:
            st.subheader("🛠️ My Skills & Certifications")
            st.markdown("**Technical Skills**")
            skills = st.text_area("List your skills (comma-separated)", placeholder="e.g., Python, Project Management, BMS, HVAC", value="Python, Project Management, Leadership")
            
            st.markdown("**Certifications**")
            certs = st.text_area("List your certifications", placeholder="e.g., CCNP, NEBOSH, PMP, CIPM")
            
            if st.button("💾 Save Skills & Certs", use_container_width=True):
                st.success("✅ Skills saved!")
            
            st.markdown("---")
            st.subheader("📁 My Documents")
            st.markdown("📄 Employment Contract - Available")
            st.markdown("📄 Last Pay Slip - May 2026")
            st.markdown("📄 Performance Review 2025 - Available")
            st.file_uploader("Upload New Document", type=['pdf', 'docx', 'jpg'])
        
        with tab4:
            st.subheader("👥 My Team")
            st.info(f"Showing colleagues in **{emp_dept}**")
            
            try:
                emp_df = db.get_all_employees()
                if not emp_df.empty:
                    team = emp_df[emp_df['department'] == emp_dept]
                    for _, teammate in team.head(10).iterrows():
                        if teammate.get('employee_id') != user_id:
                            initials_t = generate_initials(f"{teammate['first_name']} {teammate['last_name']}")
                            st.markdown(f"""
                            <div style="background:white;padding:0.6rem;border-radius:6px;margin-bottom:0.3rem;display:flex;align-items:center;gap:0.8rem;">
                                <div style="width:35px;height:35px;border-radius:50%;background:#CC0000;display:flex;align-items:center;justify-content:center;font-weight:700;color:white;font-size:0.8rem;">{initials_t}</div>
                                <div><strong>{teammate['first_name']} {teammate['last_name']}</strong><br><small>{teammate.get('position', '')}</small></div>
                            </div>
                            """, unsafe_allow_html=True)
                else:
                    st.info("Team data loading...")
            except:
                st.info("Team data loading...")
        
        with tab5:
            st.subheader("📊 My Activity")
            if 'audit_trail' in st.session_state and st.session_state.audit_trail:
                my_activities = [a for a in st.session_state.audit_trail if a.get('user') == user_name][-10:]
                if my_activities:
                    for activity in reversed(my_activities):
                        st.markdown(f"""
                        <div style="background:white;padding:0.6rem;border-radius:6px;margin-bottom:0.3rem;border-left:3px solid #CC0000;">
                            <strong>{activity.get('action', '')}</strong><br>
                            <small style="color:#888;">{activity.get('details', '')} • {activity.get('timestamp', '')}</small>
                        </div>
                        """, unsafe_allow_html=True)
                else:
                    st.info("Your activity will appear here.")
            else:
                st.info("Complete actions to see your activity log.")

def get_pipeline_stats(job_id=None):
    """Get recruitment pipeline statistics"""
    try:
        data = db._get("recruitment_pipeline")
        if job_id:
            data = [d for d in data if d.get('job_id') == job_id]
        
        stages = ['Applied', 'AI Screened', 'Manager Review', 'Shortlisted', 'Interview Scheduled', 'Offer Sent', 'Hired', 'Rejected']
        stats = {s: 0 for s in stages}
        for d in data:
            stage = d.get('current_stage', 'Applied')
            if stage in stats:
                stats[stage] += 1
        
        return stats, data
    except:
        return {}, []

def get_smart_response(question, screened_cands, all_candidates):
    q = question.lower()
    if 'top' in q or 'best' in q:
        if len(screened_cands) > 0:
            top = screened_cands.sort_values('ai_score', ascending=False).head(3)
            resp = "🏆 **Top Candidates:**\n\n"
            for i, (_, c) in enumerate(top.iterrows()):
                resp += f"{i+1}. **{c.get('first_name','')} {c.get('last_name','')}** — {int(c.get('ai_score',0))}% ({c.get('ai_tier','')})\n"
            return resp
        return "No candidates screened yet."
    if 'compare' in q:
        if len(screened_cands) >= 2:
            resp = "📊 **Comparison:**\n\n"
            for _, c in screened_cands.sort_values('ai_score', ascending=False).head(5).iterrows():
                resp += f"**{c.get('first_name','')} {c.get('last_name','')}**: {int(c.get('ai_score',0))}% | {c.get('ai_tier','')}\n"
            return resp + "\n💡 Interview Tier 1 candidates first."
        return "Need at least 2 screened candidates."
    if 'interview' in q or 'question' in q:
        return "📋 **Sample Questions:**\n\n1. Describe a challenging project you led.\n2. How do you handle disagreements?\n3. What's your most innovative solution?\n4. How do you stay current?\n5. Why Churchgate Group?"
    if 'offer' in q or 'draft' in q:
        return "📝 Use the Offer Letters tab in Recruitment Hub to generate official PDFs with Churchgate branding."
    return "I can help with top candidates, comparisons, interview questions, and offer drafts."


def lms_dashboard():
    st.markdown("""<div class="churchgate-header"><h1>🎓 Enterprise Learning Management System</h1><p>Global Courses | Certifications | Learning Paths | KPI Integration | 30+ Platforms</p></div>""", unsafe_allow_html=True)
    
    user_name = st.session_state.user['name'] if st.session_state.user else 'Staff'
    user_id = st.session_state.user.get('employee_id', '') if st.session_state.user else ''
    user_dept = st.session_state.user.get('department', '') if st.session_state.user else ''
    user_email = st.session_state.user.get('email', '') if st.session_state.user else ''
    is_admin = st.session_state.user['role'] in ['Admin', 'HR Director'] if st.session_state.user else False
    
    tab1, tab2, tab3, tab4, tab5, tab6, tab7, tab8, tab9, tab10, tab11, tab12, tab13, tab14 = st.tabs([
        "📚 Course Catalog", "📝 My Enrollments", "🎯 Learning Paths", 
        "📊 Progress", "💰 Budget", "📈 Analytics",
        "🏆 Leaderboard", "🎖️ Certifications", "👥 Study Groups",
        "🤝 Mentors", "⭐ Badges", "📝 Reviews", "📅 Calendar", "📱 Quick Track"
    ])
    
    # Load data
    def load_courses():
        try:
            data = db._get("lms_courses")
            return data if data else []
        except:
            return []
    
    def load_enrollments(employee_id=None):
        try:
            data = db._get("lms_enrollments")
            if employee_id:
                data = [e for e in data if e.get('employee_id') == employee_id]
            return data if data else []
        except:
            return []
    
    def load_platforms():
        try:
            return db._get("learning_platforms")
        except:
            return []
    
    courses = load_courses()
    platforms = load_platforms()
    my_enrollments = load_enrollments(user_id)
    
    # ============ TAB 1: COURSE CATALOG ============
    with tab1:
        st.subheader("📚 Global Course Catalog")
        st.info(f"🌍 Connected to {len(platforms)} international learning platforms including Coursera, edX, Harvard, Stanford, INSEAD, IBMI Berlin, and more.")
        
        # Filters
        col1, col2, col3, col4 = st.columns(4)
        with col1:
            cat_filter = st.selectbox("Category", ["All"] + list(set([c.get('category', '') for c in courses])))
        with col2:
            level_filter = st.selectbox("Level", ["All", "Beginner", "Intermediate", "Advanced"])
        with col3:
            provider_filter = st.selectbox("Provider", ["All"] + list(set([c.get('provider', '') for c in courses])))
        with col4:
            search_course = st.text_input("🔍 Search", placeholder="Course name or skill...")
        
        filtered = courses
        if cat_filter != "All":
            filtered = [c for c in filtered if c.get('category') == cat_filter]
        if level_filter != "All":
            filtered = [c for c in filtered if c.get('level') == level_filter]
        if provider_filter != "All":
            filtered = [c for c in filtered if c.get('provider') == provider_filter]
        if search_course:
            s = search_course.lower()
            filtered = [c for c in filtered if s in c.get('title', '').lower() or s in c.get('skills_gained', '').lower()]
        
        st.markdown(f"**{len(filtered)} courses found**")
        
        for course in filtered:
            with st.expander(f"📖 {course.get('title', 'Untitled')} — {course.get('provider', 'Unknown')} | ⏱️ {course.get('duration_hours', 0)}h | {course.get('level', 'All Levels')}"):
                col1, col2 = st.columns([3, 1])
                with col1:
                    st.markdown(f"**Description:** {course.get('description', 'No description')}")
                    st.markdown(f"**Skills Gained:** {course.get('skills_gained', 'N/A')}")
                    st.markdown(f"**Certification:** {course.get('certification_type', 'Certificate of Completion')}")
                    if course.get('external_url'):
                        st.markdown(f"🔗 [Visit Course Page]({course.get('external_url')})")
                with col2:
                    price = course.get('price', 0)
                    st.metric("Price", f"${price:,.2f}" if price > 0 else "Free")
                    st.metric("Duration", f"{course.get('duration_hours', 0)}h")
                    st.metric("Credits", course.get('credits', 0))
                    
                    # Check if already enrolled
                    already_enrolled = any(e.get('course_id') == course.get('id') for e in my_enrollments)
                    if already_enrolled:
                        st.success("✅ Enrolled")
                    else:
                        if st.button("📝 Enroll Now", key=f"enroll_{course.get('id')}", use_container_width=True):
                            db._post("lms_enrollments", {
                                "employee_id": user_id,
                                "employee_name": user_name,
                                "employee_email": user_email,
                                "department": user_dept,
                                "course_id": course.get('id'),
                                "enrollment_date": datetime.now().strftime('%Y-%m-%d %H:%M'),
                                "status": "Enrolled",
                                "progress_percent": 0
                            })
                            st.success("✅ Enrolled successfully!")
                            st.balloons()
                            st.rerun()
        
        # Admin: Add Course
        if is_admin:
            st.markdown("---")
            with st.expander("➕ Add New Course (Admin)"):
                with st.form("add_lms_course"):
                    c1, c2 = st.columns(2)
                    with c1:
                        new_title = st.text_input("Course Title")
                        new_provider = st.text_input("Provider")
                        new_category = st.selectbox("Category", ["Technology", "Leadership", "Finance", "HR", "Operations", "HSE", "Marketing", "Languages", "Facility Management", "Other"])
                    with c2:
                        new_level = st.selectbox("Level", ["Beginner", "Intermediate", "Advanced"])
                        new_duration = st.number_input("Duration (hours)", 1, 500, 40)
                        new_cert = st.text_input("Certification Type")
                    
                    new_desc = st.text_area("Description")
                    new_skills = st.text_input("Skills Gained (comma separated)")
                    new_url = st.text_input("External URL")
                    
                    if st.form_submit_button("➕ Add Course"):
                        if new_title:
                            db._post("lms_courses", {
                                "course_code": f"LMS-{datetime.now().strftime('%Y%m%d%H%M')}",
                                "title": new_title,
                                "provider": new_provider,
                                "category": new_category,
                                "level": new_level,
                                "duration_hours": new_duration,
                                "certification_type": new_cert,
                                "description": new_desc,
                                "skills_gained": new_skills,
                                "external_url": new_url
                            })
                            st.success("✅ Course added!")
                            st.rerun()
    
    # ============ TAB 2: MY ENROLLMENTS ============
    with tab2:
        st.subheader("📝 My Learning Journey")
        
        if my_enrollments:
            total_courses = len(my_enrollments)
            completed = len([e for e in my_enrollments if e.get('status') == 'Completed'])
            in_progress = len([e for e in my_enrollments if e.get('status') == 'In Progress'])
            avg_progress = sum([e.get('progress_percent', 0) for e in my_enrollments]) / total_courses if total_courses > 0 else 0
            
            c1, c2, c3, c4 = st.columns(4)
            c1.metric("📚 Enrolled", total_courses)
            c2.metric("✅ Completed", completed)
            c3.metric("🔄 In Progress", in_progress)
            c4.metric("📊 Avg Progress", f"{avg_progress:.0f}%")
            
            st.markdown("---")
            
            for enrollment in my_enrollments:
                course = next((c for c in courses if c.get('id') == enrollment.get('course_id')), None)
                if course:
                    progress = enrollment.get('progress_percent', 0)
                    status = enrollment.get('status', 'Enrolled')
                    
                    status_color = "#38a169" if status == 'Completed' else "#d69e2e" if status == 'In Progress' else "#3182ce"
                    
                    with st.expander(f"{'✅' if status == 'Completed' else '📖'} {course.get('title', 'Course')} — {progress:.0f}% — {status}"):
                        st.markdown(f"**Provider:** {course.get('provider', 'N/A')}")
                        st.markdown(f"**Enrolled:** {enrollment.get('enrollment_date', '')[:10]}")
                        st.progress(progress / 100)
                        
                        col1, col2 = st.columns(2)
                        with col1:
                            new_progress = st.slider("Update Progress", 0, 100, int(progress), key=f"prog_{enrollment.get('id')}")
                        with col2:
                            new_status = st.selectbox("Status", ["Enrolled", "In Progress", "Completed"], 
                                                     index=2 if status == 'Completed' else 1 if status == 'In Progress' else 0,
                                                     key=f"stat_{enrollment.get('id')}")
                        
                        if st.button("💾 Update", key=f"upd_{enrollment.get('id')}"):
                            update_data = {"progress_percent": new_progress, "status": new_status}
                            if new_status == 'Completed':
                                update_data["completion_date"] = datetime.now().strftime('%Y-%m-%d %H:%M')
                            db._patch("lms_enrollments", update_data, {"id": enrollment.get('id')})
                            st.success("✅ Updated!")
                            st.rerun()
        else:
            st.info("You haven't enrolled in any courses yet. Browse the Course Catalog to get started!")
    
    # ============ TAB 3: LEARNING PATHS ============
    with tab3:
        st.subheader("🎯 Career Learning Paths")
        st.info("Structured learning paths aligned to Churchgate roles and strategic pillars.")
        
        paths = [
            {
                "name": "💻 Technology Leadership",
                "courses": ["AI-FM-001", "CLOUD-AWS-001", "CYBER-001", "IBM-DS-001"],
                "duration": "6 months",
                "cert": "Churchgate Technology Leader Certificate",
                "role": "Technology Group"
            },
            {
                "name": "🏗️ Facility Management Excellence",
                "courses": ["BMS-ADV-001", "HSE-NEBOSH-001", "PM-PMP-001"],
                "duration": "5 months",
                "cert": "Certified FM Professional",
                "role": "Facility Management"
            },
            {
                "name": "👥 HR Business Partner",
                "courses": ["HR-STRAT-001", "LEAD-EXEC-001", "MKT-DIG-001"],
                "duration": "6 months",
                "cert": "SHRM-CP + Leadership Certificate",
                "role": "Human Resources"
            },
            {
                "name": "💰 Finance & Strategy",
                "courses": ["FIN-MOD-001", "FIN-ESG-001", "DATA-001"],
                "duration": "5 months",
                "cert": "CFA Certificate + Data Analytics",
                "role": "Accounts & Finance"
            }
        ]
        
        for path in paths:
            with st.expander(f"{path['name']} — {path['duration']} — 🏅 {path['cert']}"):
                st.markdown(f"**Target Role:** {path['role']}")
                st.markdown("**Required Courses:**")
                for code in path['courses']:
                    course = next((c for c in courses if c.get('course_code') == code), None)
                    if course:
                        enrolled = any(e.get('course_id') == course.get('id') for e in my_enrollments)
                        icon = "✅" if enrolled else "⬜"
                        st.markdown(f"{icon} {course.get('title', code)} — {course.get('duration_hours', 0)}h")
                
                if st.button(f"🎯 Start This Path", key=f"path_{path['name'][:10]}"):
                    for code in path['courses']:
                        course = next((c for c in courses if c.get('course_code') == code), None)
                        if course:
                            already = any(e.get('course_id') == course.get('id') for e in my_enrollments)
                            if not already:
                                db._post("lms_enrollments", {
                                    "employee_id": user_id, "employee_name": user_name,
                                    "employee_email": user_email, "department": user_dept,
                                    "course_id": course.get('id'),
                                    "enrollment_date": datetime.now().strftime('%Y-%m-%d %H:%M'),
                                    "status": "Enrolled", "progress_percent": 0
                                })
                    st.success(f"✅ Enrolled in {path['name']}!")
                    st.balloons()
                    st.rerun()
    
    # ============ TAB 4: PROGRESS ============
    with tab4:
        st.subheader("📊 My Learning Progress")
        
        if my_enrollments:
            for enrollment in my_enrollments:
                course = next((c for c in courses if c.get('id') == enrollment.get('course_id')), None)
                if course:
                    progress = enrollment.get('progress_percent', 0)
                    color = "#38a169" if progress >= 80 else "#d69e2e" if progress >= 50 else "#CC0000"
                    
                    st.markdown(f"""
                    <div style="background:white;padding:0.8rem;border-radius:8px;margin-bottom:0.5rem;border-left:4px solid {color};">
                        <strong>{course.get('title', 'Course')}</strong>
                        <span style="float:right;font-weight:700;color:{color};">{progress:.0f}%</span>
                        <br><small>{course.get('provider', '')} | {course.get('duration_hours', 0)}h | {enrollment.get('status', '')}</small>
                        <div style="background:#e0e0e0;height:6px;border-radius:3px;margin-top:0.3rem;">
                            <div style="background:{color};width:{progress}%;height:6px;border-radius:3px;"></div>
                        </div>
                    </div>
                    """, unsafe_allow_html=True)
        else:
            st.info("Enroll in courses to track your progress here.")
    
    # ============ TAB 5: BUDGET ============
    with tab5:
        st.subheader("💰 Training Budget")
        
        if is_admin:
            st.info(f"Department: {user_dept}")
            try:
                budgets = db._get("lms_budgets")
                dept_budget = next((b for b in budgets if b.get('department') == user_dept), None)
                
                if dept_budget:
                    c1, c2, c3 = st.columns(3)
                    c1.metric("Total Budget", f"₦{dept_budget.get('total_budget', 0):,.2f}")
                    c2.metric("Spent", f"₦{dept_budget.get('spent_amount', 0):,.2f}")
                    c3.metric("Remaining", f"₦{dept_budget.get('remaining_amount', 0):,.2f}")
                    
                    spent_pct = (dept_budget.get('spent_amount', 0) / dept_budget.get('total_budget', 1)) * 100
                    st.progress(min(spent_pct / 100, 1.0))
                else:
                    st.info("No budget set for your department.")
                    
                    if st.button("Set Budget"):
                        db._post("lms_budgets", {
                            "department": user_dept,
                            "fiscal_year": "2026/2027",
                            "total_budget": 5000000,
                            "spent_amount": 0
                        })
                        st.rerun()
            except:
                pass
        else:
            st.info("Budget view available for managers and admins.")
    
    # ============ TAB 6: ANALYTICS ============
    with tab6:
        st.subheader("📈 Learning Analytics")
        
        if is_admin:
            all_enrollments = load_enrollments()
            
            if all_enrollments:
                total = len(all_enrollments)
                completed = len([e for e in all_enrollments if e.get('status') == 'Completed'])
                
                c1, c2, c3, c4 = st.columns(4)
                c1.metric("Total Enrollments", total)
                c2.metric("Completed", completed)
                c3.metric("Completion Rate", f"{int(completed/total*100)}%" if total > 0 else "0%")
                c4.metric("Active Learners", len(set(e.get('employee_id') for e in all_enrollments)))
                
                # Department breakdown
                dept_data = {}
                for e in all_enrollments:
                    dept = e.get('department', 'Unknown')
                    dept_data[dept] = dept_data.get(dept, 0) + 1
                
                fig = px.pie(values=list(dept_data.values()), names=list(dept_data.keys()), hole=0.5,
                           title="Enrollments by Department")
                st.plotly_chart(fig, use_container_width=True)
            else:
                st.info("No enrollment data yet.")
        else:
            st.info("Analytics available for admins and managers.")

# ============ TAB 7: LEADERBOARD ============
    tab7, tab8, tab9, tab10, tab11, tab12, tab13, tab14 = st.tabs([
        "🏆 Leaderboard", "🎖️ Certifications", "👥 Study Groups", 
        "🤝 Mentors", "⭐ Badges", "📝 Reviews", "📅 Calendar", "📱 Quick Track"
    ])
    
    with tab7:
        st.subheader("🏆 Learning Leaderboard")
        st.info("Top learners ranked by courses completed, progress, and engagement.")
        
        all_enrollments = load_enrollments()
        if all_enrollments:
            # Calculate leaderboard
            leaderboard = {}
            for e in all_enrollments:
                emp = e.get('employee_name', 'Unknown')
                if emp not in leaderboard:
                    leaderboard[emp] = {'completed': 0, 'in_progress': 0, 'total_progress': 0, 'courses': 0, 'dept': e.get('department', '')}
                leaderboard[emp]['courses'] += 1
                leaderboard[emp]['total_progress'] += e.get('progress_percent', 0)
                if e.get('status') == 'Completed':
                    leaderboard[emp]['completed'] += 1
                elif e.get('status') == 'In Progress':
                    leaderboard[emp]['in_progress'] += 1
            
            # Sort by completion then progress
            ranked = sorted(leaderboard.items(), key=lambda x: (x[1]['completed'], x[1]['total_progress']), reverse=True)
            
            for rank, (name, data) in enumerate(ranked[:20]):
                medal = "🥇" if rank == 0 else "🥈" if rank == 1 else "🥉" if rank == 2 else f"{rank+1}."
                avg_prog = data['total_progress'] / data['courses'] if data['courses'] > 0 else 0
                
                st.markdown(f"""
                <div style="background:white;padding:0.8rem;border-radius:8px;margin-bottom:0.4rem;display:flex;align-items:center;gap:1rem;border-left:4px solid {'#d69e2e' if rank < 3 else '#CC0000'};">
                    <span style="font-size:1.5rem;">{medal}</span>
                    <div style="flex:1;">
                        <strong>{name}</strong> — {data['dept']}<br>
                        <small>{data['completed']} completed | {data['in_progress']} in progress | Avg {avg_prog:.0f}%</small>
                    </div>
                    <div style="text-align:right;">
                        <span style="font-size:1.2rem;font-weight:700;color:#CC0000;">{data['completed'] * 100 + int(avg_prog)} pts</span>
                    </div>
                </div>
                """, unsafe_allow_html=True)
        else:
            st.info("Leaderboard will populate as employees enroll in courses.")
    
    # ============ TAB 8: CERTIFICATIONS ============
    with tab8:
        st.subheader("🎖️ My Certification Wall")
        
        completed_enrollments = [e for e in my_enrollments if e.get('status') == 'Completed']
        if completed_enrollments:
            st.markdown(f"### 🏅 {len(completed_enrollments)} Certifications Earned")
            
            cols = st.columns(3)
            for i, enrollment in enumerate(completed_enrollments):
                course = next((c for c in courses if c.get('id') == enrollment.get('course_id')), None)
                if course:
                    with cols[i % 3]:
                        st.markdown(f"""
                        <div style="background:white;padding:1rem;border-radius:10px;text-align:center;box-shadow:0 2px 8px rgba(0,0,0,0.08);border-top:3px solid #d69e2e;margin-bottom:1rem;">
                            <h2 style="font-size:2.5rem;">🏅</h2>
                            <strong>{course.get('title', 'Course')[:40]}</strong><br>
                            <small style="color:#d69e2e;">{course.get('certification_type', 'Certificate')}</small><br>
                            <small>Provider: {course.get('provider', 'N/A')}</small><br>
                            <small>Completed: {enrollment.get('completion_date', '')[:10]}</small>
                        </div>
                        """, unsafe_allow_html=True)
        else:
            st.info("Complete courses to earn certifications that appear here.")
            
            # Sample certification preview
            st.markdown("### 🎯 Target Certifications")
            target_certs = [
                ("PMP - Project Management", "PMI", "80 hours"),
                ("SHRM-CP", "SHRM", "60 hours"),
                ("AWS Solutions Architect", "AWS", "80 hours"),
                ("NEBOSH Diploma", "NEBOSH", "120 hours"),
            ]
            for cert, provider, hours in target_certs:
                st.markdown(f"🎯 **{cert}** — {provider} ({hours})")
    
    # ============ TAB 9: STUDY GROUPS ============
    with tab9:
        st.subheader("👥 Peer Study Groups")
        st.info("Join study groups with colleagues taking the same courses.")
        
        # Group by course
        all_enrollments_all = load_enrollments()
        course_groups = {}
        for e in all_enrollments_all:
            cid = e.get('course_id')
            if cid not in course_groups:
                course_groups[cid] = []
            course_groups[cid].append(e.get('employee_name', 'Unknown'))
        
        for cid, members in course_groups.items():
            if len(members) >= 2:
                course = next((c for c in courses if c.get('id') == cid), None)
                if course:
                    with st.expander(f"📚 {course.get('title', 'Course')[:50]} — {len(members)} learners"):
                        st.markdown("**Study Group Members:**")
                        for member in members:
                            st.markdown(f"👤 {member}")
                        
                        if user_name not in members:
                            if st.button(f"Join Group", key=f"join_group_{cid}"):
                                st.success(f"✅ Joined study group for {course.get('title', '')}!")
                        else:
                            st.success("✅ You're in this group")
    
    # ============ TAB 10: MENTORS ============
    with tab10:
        st.subheader("🤝 Mentor Connection")
        st.info("Connect with colleagues who have completed courses you're taking.")
        
        for enrollment in my_enrollments:
            if enrollment.get('status') != 'Completed':
                course = next((c for c in courses if c.get('id') == enrollment.get('course_id')), None)
                if course:
                    # Find mentors who completed this course
                    mentors = [e for e in all_enrollments_all 
                              if e.get('course_id') == enrollment.get('course_id') 
                              and e.get('status') == 'Completed'
                              and e.get('employee_name') != user_name]
                    
                    if mentors:
                        with st.expander(f"🎓 Mentors for {course.get('title', '')[:40]}"):
                            for mentor in mentors[:3]:
                                st.markdown(f"👤 **{mentor.get('employee_name', 'Unknown')}** — Completed {mentor.get('completion_date', '')[:10]}")
                                if st.button(f"📧 Request Mentorship", key=f"mentor_{mentor.get('id')}"):
                                    st.success(f"✅ Mentorship request sent to {mentor.get('employee_name', '')}!")
    
    # ============ TAB 11: BADGES ============
    with tab11:
        st.subheader("⭐ Skills Badge System")
        st.info("Earn badges as you learn. Badges appear on your profile.")
        
        badges = []
        if len(my_enrollments) >= 1:
            badges.append({"name": "🌱 Beginner Learner", "desc": "Enrolled in your first course", "earned": True})
        if len(my_enrollments) >= 3:
            badges.append({"name": "📚 Dedicated Learner", "desc": "Enrolled in 3+ courses", "earned": True})
        if len([e for e in my_enrollments if e.get('status') == 'Completed']) >= 1:
            badges.append({"name": "🎓 Course Completer", "desc": "Completed your first course", "earned": True})
        if len([e for e in my_enrollments if e.get('status') == 'Completed']) >= 3:
            badges.append({"name": "🏅 Advanced Scholar", "desc": "Completed 3+ courses", "earned": True})
        if len([e for e in my_enrollments if e.get('progress_percent', 0) >= 90]) >= 1:
            badges.append({"name": "🔥 Almost There", "desc": "90%+ progress on a course", "earned": True})
        
        # Future badges
        all_badges = [
            {"name": "🌱 Beginner Learner", "desc": "Enrolled in your first course", "earned": any(b['name'] == '🌱 Beginner Learner' for b in badges)},
            {"name": "📚 Dedicated Learner", "desc": "Enrolled in 3+ courses", "earned": any(b['name'] == '📚 Dedicated Learner' for b in badges)},
            {"name": "🎓 Course Completer", "desc": "Completed your first course", "earned": any(b['name'] == '🎓 Course Completer' for b in badges)},
            {"name": "🏅 Advanced Scholar", "desc": "Completed 3+ courses", "earned": any(b['name'] == '🏅 Advanced Scholar' for b in badges)},
            {"name": "🔥 Almost There", "desc": "90%+ progress on a course", "earned": any(b['name'] == '🔥 Almost There' for b in badges)},
            {"name": "👥 Team Learner", "desc": "Joined a study group", "earned": False},
            {"name": "🤝 Mentor", "desc": "Mentored a colleague", "earned": False},
            {"name": "📊 100 Hours", "desc": "Completed 100+ hours of learning", "earned": False},
            {"name": "🌍 Global Scholar", "desc": "Course from 3+ different platforms", "earned": False},
        ]
        
        cols = st.columns(3)
        for i, badge in enumerate(all_badges):
            with cols[i % 3]:
                opacity = "1" if badge['earned'] else "0.4"
                st.markdown(f"""
                <div style="background:white;padding:1rem;border-radius:8px;text-align:center;opacity:{opacity};margin-bottom:0.5rem;border:2px solid {'#d69e2e' if badge['earned'] else '#e0e0e0'};">
                    <h3>{'✅' if badge['earned'] else '🔒'}</h3>
                    <strong>{badge['name']}</strong><br>
                    <small>{badge['desc']}</small>
                </div>
                """, unsafe_allow_html=True)
    
    # ============ TAB 12: REVIEWS ============
    with tab12:
        st.subheader("📝 Course Reviews & Ratings")
        st.info("Rate courses you've taken to help colleagues make informed decisions.")
        
        for enrollment in my_enrollments:
            course = next((c for c in courses if c.get('id') == enrollment.get('course_id')), None)
            if course:
                with st.expander(f"Rate: {course.get('title', 'Course')[:50]}"):
                    rating = st.slider("⭐ Rating", 1, 5, 4, key=f"rate_{course.get('id')}")
                    review = st.text_area("Your Review", key=f"review_{course.get('id')}")
                    if st.button("📝 Submit Review", key=f"submit_review_{course.get('id')}"):
                        st.success("✅ Review submitted! Thank you for helping your colleagues.")
    
    # ============ TAB 13: CALENDAR ============
    with tab13:
        st.subheader("📅 Learning Calendar")
        
        today = datetime.now()
        upcoming = [e for e in my_enrollments if e.get('status') in ['Enrolled', 'In Progress']]
        
        if upcoming:
            st.markdown(f"### 📆 {len(upcoming)} Active Courses")
            for enrollment in upcoming:
                course = next((c for c in courses if c.get('id') == enrollment.get('course_id')), None)
                if course:
                    progress = enrollment.get('progress_percent', 0)
                    hours_left = int(course.get('duration_hours', 0) * (100 - progress) / 100)
                    
                    st.markdown(f"""
                    <div style="background:white;padding:0.8rem;border-radius:8px;margin-bottom:0.4rem;border-left:4px solid #3182ce;">
                        <strong>📖 {course.get('title', 'Course')[:50]}</strong><br>
                        <small>⏱️ ~{hours_left} hours remaining | 📊 {progress:.0f}% complete</small>
                    </div>
                    """, unsafe_allow_html=True)
        else:
            st.info("No active courses. Enroll to see your learning schedule.")
    
    # ============ TAB 14: QUICK TRACK ============
    with tab14:
        st.subheader("📱 Quick Progress Tracker")
        st.info("Quickly update your learning progress on the go.")
        
        for enrollment in my_enrollments:
            if enrollment.get('status') != 'Completed':
                course = next((c for c in courses if c.get('id') == enrollment.get('course_id')), None)
                if course:
                    col1, col2, col3 = st.columns([2, 1, 1])
                    with col1:
                        st.markdown(f"**{course.get('title', 'Course')[:40]}**")
                    with col2:
                        quick_progress = st.number_input("Progress %", 0, 100, int(enrollment.get('progress_percent', 0)), 5, key=f"quick_{enrollment.get('id')}")
                    with col3:
                        if st.button("📱 Update", key=f"quick_upd_{enrollment.get('id')}"):
                            db._patch("lms_enrollments", {"progress_percent": quick_progress, "status": "Completed" if quick_progress >= 100 else "In Progress"}, {"id": enrollment.get('id')})
                            st.success("✅ Updated!")
                            st.rerun()


def audit_log_viewer():
    st.markdown("""<div class="churchgate-header"><h1>📋 Enterprise Audit Log Viewer</h1><p>Real-Time Activity Monitoring | Security Events | Compliance Tracking | Full System Transparency</p></div>""", unsafe_allow_html=True)
    
    is_admin = st.session_state.user['role'] in ['Admin', 'HR Director'] if st.session_state.user else False
    
    if not is_admin:
        st.error("Access restricted to Administrators and HR Directors only.")
        return
    
    # Load audit data
    def load_audit_logs():
        try:
            data = db._get("audit_trail")
            return data if data else []
        except:
            return []
    
    all_logs = load_audit_logs()
    
    # Also collect real-time events from session state
    session_events = []
    if 'audit_trail' in st.session_state:
        for entry in st.session_state.audit_trail:
            session_events.append({
                'action': entry.get('action', ''),
                'details': entry.get('details', ''),
                'user_name': entry.get('user', ''),
                'timestamp_text': entry.get('timestamp', ''),
                'module': 'Session'
            })
    
    # Combine and deduplicate
    all_events = all_logs + session_events
    
    # Top Stats
    total_events = len(all_events)
    today = datetime.now().strftime('%Y-%m-%d')
    today_events = [e for e in all_events if today in str(e.get('timestamp_text', '')) or today in str(e.get('created_at', ''))]
    unique_users = len(set(e.get('user_name', '') for e in all_events if e.get('user_name')))
    
    c1, c2, c3, c4 = st.columns(4)
    c1.metric("📊 Total Events", total_events)
    c2.metric("📅 Today", len(today_events))
    c3.metric("👥 Unique Users", unique_users)
    c4.metric("🔴 Live Session", len(session_events))
    
    st.markdown("---")
    
    # Filters
    col1, col2, col3, col4 = st.columns(4)
    with col1:
        search_log = st.text_input("🔍 Search", placeholder="Action, user, details...")
    with col2:
        date_filter = st.selectbox("📅 Date", ["All Time", "Today", "Last 7 Days", "Last 30 Days"])
    with col3:
        severity_filter = st.selectbox("⚠️ Severity", ["All", "Info", "Warning", "Critical"])
    with col4:
        module_filter = st.selectbox("📂 Module", ["All", "Login", "Employees", "Performance", "Recruitment", "Profile", "KPI", "Appraisal", "Security"])
    
    # Apply filters
    filtered = all_events
    
    if search_log:
        s = search_log.lower()
        filtered = [e for e in filtered if s in str(e.get('action', '')).lower() or 
                   s in str(e.get('details', '')).lower() or 
                   s in str(e.get('user_name', '')).lower()]
    
    if date_filter == "Today":
        filtered = [e for e in filtered if today in str(e.get('timestamp_text', '')) or today in str(e.get('created_at', ''))]
    elif date_filter == "Last 7 Days":
        cutoff = (datetime.now() - timedelta(days=7)).strftime('%Y-%m-%d')
        filtered = [e for e in filtered if str(e.get('timestamp_text', ''))[:10] >= cutoff or str(e.get('created_at', ''))[:10] >= cutoff]
    elif date_filter == "Last 30 Days":
        cutoff = (datetime.now() - timedelta(days=30)).strftime('%Y-%m-%d')
        filtered = [e for e in filtered if str(e.get('timestamp_text', ''))[:10] >= cutoff or str(e.get('created_at', ''))[:10] >= cutoff]
    
    if severity_filter != "All":
        filtered = [e for e in filtered if severity_filter.lower() in str(e.get('action', '')).lower() or 
                   severity_filter.lower() in str(e.get('details', '')).lower()]
    
    if module_filter != "All":
        module_keywords = {
            'Login': ['login', 'password', 'sign in', 'locked'],
            'Employees': ['employee', 'staff', 'hired', 'terminated'],
            'Performance': ['kpi', 'performance', 'okr'],
            'Recruitment': ['job', 'requisition', 'candidate', 'interview'],
            'Profile': ['profile', 'photo', 'picture'],
            'KPI': ['kpi'],
            'Appraisal': ['appraisal', 'assessment', 'review'],
            'Security': ['security', 'permission', 'access', 'failed']
        }
        keywords = module_keywords.get(module_filter, [])
        filtered = [e for e in filtered if any(k in str(e.get('action', '')).lower() or k in str(e.get('details', '')).lower() for k in keywords)]
    
    st.markdown(f"**Showing {len(filtered)} of {total_events} events**")
    
    # Export
    col1, col2 = st.columns([3, 1])
    with col2:
        if st.button("📥 Export CSV", use_container_width=True):
            export_data = [{
                'Timestamp': e.get('timestamp_text', e.get('created_at', '')),
                'User': e.get('user_name', 'System'),
                'Action': e.get('action', ''),
                'Details': e.get('details', '')
            } for e in filtered]
            st.download_button("📥 Download Audit Log", pd.DataFrame(export_data).to_csv(index=False), 
                             f"audit_log_{today}.csv", "text/csv")
    
    st.markdown("---")
    
    # Events list
    if filtered:
        # Group by date
        from collections import defaultdict
        grouped = defaultdict(list)
        for e in filtered:
            date_key = str(e.get('timestamp_text', e.get('created_at', '')))[:10]
            if not date_key or date_key == 'None':
                date_key = 'Unknown'
            grouped[date_key].append(e)
        
        for date_key in sorted(grouped.keys(), reverse=True):
            day_events = grouped[date_key]
            
            # Calculate day stats
            day_total = len(day_events)
            day_warnings = len([e for e in day_events if 'fail' in str(e.get('action', '')).lower() or 'error' in str(e.get('details', '')).lower()])
            
            st.markdown(f"### 📅 {date_key}")
            st.caption(f"{day_total} events • {day_warnings} warnings")
            
            for event in day_events[:50]:  # Limit per day
                action = event.get('action', 'Unknown')
                details = event.get('details', '')
                user = event.get('user_name', 'System')
                timestamp = event.get('timestamp_text', event.get('created_at', ''))
                
                # Determine icon and color
                if 'fail' in str(action).lower() or 'error' in str(details).lower() or 'locked' in str(details).lower():
                    icon = "🚨"
                    border = "#CC0000"
                    severity = "Critical"
                elif 'password' in str(action).lower() or 'security' in str(details).lower():
                    icon = "⚠️"
                    border = "#d69e2e"
                    severity = "Warning"
                elif 'success' in str(details).lower() or 'complete' in str(action).lower():
                    icon = "✅"
                    border = "#38a169"
                    severity = "Info"
                elif 'update' in str(action).lower() or 'edit' in str(action).lower() or 'change' in str(action).lower():
                    icon = "📝"
                    border = "#3182ce"
                    severity = "Info"
                else:
                    icon = "ℹ️"
                    border = "#a0aec0"
                    severity = "Info"
                
                # Format timestamp
                try:
                    if timestamp and str(timestamp) != 'None':
                        ts = str(timestamp)
                        if len(ts) > 16:
                            display_time = ts[11:16]
                        else:
                            display_time = ts
                    else:
                        display_time = 'N/A'
                except:
                    display_time = 'N/A'
                
                with st.expander(f"{icon} {display_time} — {action[:80]} — {user}", expanded=(severity == 'Critical')):
                    col1, col2 = st.columns([3, 1])
                    with col1:
                        st.markdown(f"**Action:** {action}")
                        if details:
                            st.markdown(f"**Details:** {details}")
                        st.markdown(f"**User:** {user}")
                        st.markdown(f"**Timestamp:** {timestamp}")
                    with col2:
                        st.markdown(f"<span style='background:{border};color:white;padding:0.2rem 0.6rem;border-radius:10px;font-size:0.8rem;'>{severity}</span>", unsafe_allow_html=True)
    else:
        st.info("No audit events match your filters.")
    
    # Real-time activity feed
    st.markdown("---")
    st.markdown("### 🔴 Live Activity Feed (Current Session)")
    
    if session_events:
        for event in reversed(session_events[-10:]):
            action = event.get('action', '')
            details = event.get('details', '')
            user = event.get('user_name', '')
            timestamp = event.get('timestamp_text', '')
            
            st.markdown(f"""
            <div style="background:white;padding:0.6rem 1rem;border-radius:6px;margin-bottom:0.3rem;border-left:4px solid #CC0000;display:flex;align-items:center;gap:1rem;">
                <span style="color:#CC0000;font-size:0.8rem;">● LIVE</span>
                <div style="flex:1;">
                    <strong>{action}</strong><br>
                    <small>{details} — {user} • {timestamp}</small>
                </div>
            </div>
            """, unsafe_allow_html=True)
    else:
        st.info("Session activity will appear here as you use the system.")
    
    # Audit analytics
    if all_events:
        st.markdown("---")
        st.markdown("### 📊 Audit Analytics")
        
        col1, col2 = st.columns(2)
        with col1:
            # Activity by hour
            hour_data = {}
            for e in all_events:
                ts = str(e.get('timestamp_text', e.get('created_at', '')))
                try:
                    if len(ts) >= 13:
                        hour = ts[11:13]
                        hour_data[hour] = hour_data.get(hour, 0) + 1
                except:
                    pass
            if hour_data:
                hours = sorted(hour_data.keys())
                counts = [hour_data[h] for h in hours]
                fig = px.bar(x=hours, y=counts, labels={'x': 'Hour', 'y': 'Events'}, 
                           title="Activity by Hour", color_discrete_sequence=['#CC0000'])
                fig.update_layout(height=300)
                st.plotly_chart(fig, use_container_width=True)
        
        with col2:
            # Top users
            user_data = {}
            for e in all_events:
                user = e.get('user_name', 'System')
                user_data[user] = user_data.get(user, 0) + 1
            top_users = sorted(user_data.items(), key=lambda x: x[1], reverse=True)[:10]
            if top_users:
                fig2 = px.bar(x=[u[0] for u in top_users], y=[u[1] for u in top_users],
                            title="Top Users by Activity", color_discrete_sequence=['#3182ce'])
                fig2.update_layout(height=300)
                st.plotly_chart(fig2, use_container_width=True)

def advanced_analytics():
    st.markdown("""<div class="churchgate-header"><h1>📊 Advanced Analytics & Business Intelligence</h1><p>Cross-Module Insights | Predictive Analytics | Executive Command Center</p></div>""", unsafe_allow_html=True)
    
    is_admin = st.session_state.user['role'] in ['Admin', 'HR Director'] if st.session_state.user else False
    if not is_admin:
        st.error("Access restricted to Administrators and HR Directors only.")
        return
    
    # Load all data sources
    try:
        emp_df = db.get_all_employees()
        total_emp = len(emp_df) if not emp_df.empty else 0
        dept_count = len(emp_df['department'].unique()) if not emp_df.empty else 0
        active_emp = len(emp_df[emp_df['status'] == 'Active']) if not emp_df.empty else 0
    except:
        emp_df = pd.DataFrame()
        total_emp = 0
    
    try:
        candidates = db.get_all_candidates()
        total_candidates = len(candidates) if not candidates.empty else 0
        screened = len(candidates[candidates['ai_score'] > 0]) if not candidates.empty and 'ai_score' in candidates.columns else 0
        hired = len(candidates[candidates['status'] == 'Hired']) if not candidates.empty and 'status' in candidates.columns else 0
    except:
        total_candidates = 0
    
    try:
        enrollments = db._get("lms_enrollments")
        total_enrollments = len(enrollments) if enrollments else 0
        completed_enrollments = len([e for e in enrollments if e.get('status') == 'Completed']) if enrollments else 0
    except:
        total_enrollments = 0
        completed_enrollments = 0
    
    try:
        all_requests = db._get("employee_requests")
        total_requests = len(all_requests) if all_requests else 0
        approved_requests = len([r for r in all_requests if r.get('status') == 'Approved']) if all_requests else 0
    except:
        total_requests = 0
    
    try:
        all_appraisals = db.get_all_appraisals()
        total_appraisals = len(all_appraisals) if all_appraisals else 0
        completed_appraisals = len([a for a in all_appraisals if a.get('status') == 'Completed' or a.get('acceptance') == 'Accepted']) if all_appraisals else 0
    except:
        total_appraisals = 0
    
    try:
        ideas = db._get("ideas_box")
        total_ideas = len(ideas) if ideas else 0
        implemented_ideas = len([i for i in ideas if i.get('status') == 'Implemented']) if ideas else 0
    except:
        total_ideas = 0
    
    # ===== TOP KPI ROW =====
    st.markdown("### 📊 Organizational Health Scorecard")
    
    c1, c2, c3, c4, c5, c6 = st.columns(6)
    c1.metric("👥 Headcount", total_emp, f"{active_emp} active")
    c2.metric("🏢 Departments", dept_count)
    c3.metric("🎓 Learning", total_enrollments, f"{completed_enrollments} completed")
    c4.metric("📋 Requests", total_requests, f"{approved_requests} approved")
    c5.metric("📝 Appraisals", total_appraisals, f"{completed_appraisals} done")
    c6.metric("💡 Ideas", total_ideas, f"{implemented_ideas} implemented")
    
    # ===== EMPLOYEE HEALTH SCORE =====
    employee_health = 0
    if total_emp > 0:
        health_factors = 0
        if active_emp > 0: health_factors += 1
        if completed_appraisals > 0: health_factors += 1
        if total_enrollments > 0: health_factors += 1
        employee_health = int((health_factors / 3) * 100)
    
    st.progress(employee_health / 100)
    st.caption(f"Organizational Health Index: {employee_health}%")
    
    st.markdown("---")
    
    tab1, tab2, tab3, tab4 = st.tabs(["👥 Workforce", "📈 Performance", "🎓 Learning", "💡 Innovation"])
    
    # ===== TAB 1: WORKFORCE ANALYTICS =====
    with tab1:
        st.subheader("👥 Workforce Analytics")
        
        col1, col2 = st.columns(2)
        
        with col1:
            if not emp_df.empty:
                # Headcount by department
                dept_counts = emp_df['department'].value_counts().head(10)
                fig = px.bar(x=dept_counts.index, y=dept_counts.values, 
                           title="Headcount by Department", color=dept_counts.values,
                           color_continuous_scale=['#CC0000', '#d69e2e', '#38a169'])
                fig.update_layout(height=350)
                st.plotly_chart(fig, use_container_width=True)
        
        with col2:
            if not emp_df.empty:
                # Grade distribution
                grade_counts = emp_df['grade'].value_counts()
                fig2 = px.pie(values=grade_counts.values, names=grade_counts.index, 
                            title="Grade Distribution", hole=0.5,
                            color_discrete_sequence=['#CC0000', '#d69e2e', '#3182ce', '#38a169', '#805ad5'])
                fig2.update_layout(height=350)
                st.plotly_chart(fig2, use_container_width=True)
        
        # Gender diversity
        if not emp_df.empty and 'gender' in emp_df.columns:
            st.markdown("---")
            st.markdown("### 🌍 Gender Diversity")
            col1, col2, col3 = st.columns(3)
            
            male_count = len(emp_df[emp_df['gender'].str.lower() == 'male'])
            female_count = len(emp_df[emp_df['gender'].str.lower() == 'female'])
            
            col1.metric("👨 Male", male_count, f"{int(male_count/total_emp*100)}%" if total_emp > 0 else "")
            col2.metric("👩 Female", female_count, f"{int(female_count/total_emp*100)}%" if total_emp > 0 else "")
            col3.metric("📊 Ratio", f"{male_count}:{female_count}")
        
        # Tenure distribution
        if not emp_df.empty:
            st.markdown("---")
            st.markdown("### 📅 Employee Tenure")
            tenure_data = []
            for _, emp in emp_df.iterrows():
                jd = emp.get('join_date')
                if jd and str(jd) != 'None' and str(jd) != 'nan':
                    try:
                        years = (datetime.now() - pd.to_datetime(jd)).days / 365
                        if years < 2: tenure_data.append('0-2 years')
                        elif years < 5: tenure_data.append('3-5 years')
                        elif years < 10: tenure_data.append('6-10 years')
                        else: tenure_data.append('10+ years')
                    except:
                        pass
            
            if tenure_data:
                tenure_df = pd.DataFrame(pd.Series(tenure_data).value_counts().reset_index())
                tenure_df.columns = ['Tenure', 'Count']
                fig3 = px.bar(tenure_df, x='Tenure', y='Count', color='Tenure',
                            title="Employee Tenure Distribution",
                            color_discrete_sequence=['#CC0000', '#d69e2e', '#3182ce', '#38a169'])
                fig3.update_layout(height=350, showlegend=False)
                st.plotly_chart(fig3, use_container_width=True)
    
    # ===== TAB 2: PERFORMANCE ANALYTICS =====
    with tab2:
        st.subheader("📈 Performance & Appraisal Analytics")
        
        col1, col2 = st.columns(2)
        
        with col1:
            # Appraisal completion
            appraisal_labels = ['Completed', 'Pending', 'Not Started']
            appraisal_values = [completed_appraisals, max(0, total_appraisals - completed_appraisals), max(0, total_emp - total_appraisals)]
            fig4 = px.pie(values=appraisal_values, names=appraisal_labels, 
                        title="Appraisal Completion Status", hole=0.5,
                        color_discrete_sequence=['#38a169', '#d69e2e', '#a0aec0'])
            fig4.update_layout(height=350)
            st.plotly_chart(fig4, use_container_width=True)
        
        with col2:
            # KPI progress across departments
            try:
                perf_data = db.get_performance_data()
                if not perf_data.empty:
                    dept_progress = perf_data.groupby('department')['progress'].mean().reset_index()
                    fig5 = px.bar(dept_progress, x='department', y='progress', 
                                title="Avg KPI Progress by Department", color='progress',
                                color_continuous_scale=['#CC0000', '#d69e2e', '#38a169'])
                    fig5.update_layout(height=350)
                    st.plotly_chart(fig5, use_container_width=True)
            except:
                st.info("Performance data loading...")
        
        # Recruitment pipeline
        st.markdown("---")
        st.markdown("### 💼 Recruitment Pipeline")
        
        try:
            pipeline_data = db._get("recruitment_pipeline")
            if pipeline_data:
                stages = ['Applied', 'AI Screened', 'Shortlisted', 'Interview Scheduled', 'Offer Sent', 'Hired']
                stage_counts = {}
                for s in stages:
                    stage_counts[s] = len([p for p in pipeline_data if p.get('current_stage') == s])
                
                fig6 = px.funnel(x=list(stage_counts.values()), y=list(stage_counts.keys()),
                               title="Recruitment Funnel",
                               color_discrete_sequence=['#CC0000'])
                fig6.update_layout(height=350)
                st.plotly_chart(fig6, use_container_width=True)
            else:
                st.info("Pipeline data will appear as candidates progress.")
        except:
            pass
    
    # ===== TAB 3: LEARNING ANALYTICS =====
    with tab3:
        st.subheader("🎓 Learning & Development Analytics")
        
        col1, col2 = st.columns(2)
        
        with col1:
            if total_enrollments > 0:
                learning_status = {
                    'Completed': completed_enrollments,
                    'In Progress': total_enrollments - completed_enrollments
                }
                fig7 = px.pie(values=list(learning_status.values()), names=list(learning_status.keys()),
                            title="Course Completion Rate", hole=0.5,
                            color_discrete_sequence=['#38a169', '#d69e2e'])
                fig7.update_layout(height=350)
                st.plotly_chart(fig7, use_container_width=True)
        
        with col2:
            # Department learning engagement
            try:
                dept_enrollments = {}
                for e in enrollments:
                    dept = e.get('department', 'Unknown')
                    dept_enrollments[dept] = dept_enrollments.get(dept, 0) + 1
                
                if dept_enrollments:
                    fig8 = px.bar(x=list(dept_enrollments.keys()), y=list(dept_enrollments.values()),
                                title="Enrollments by Department", color=list(dept_enrollments.values()),
                                color_continuous_scale=['#CC0000', '#d69e2e', '#38a169'])
                    fig8.update_layout(height=350)
                    st.plotly_chart(fig8, use_container_width=True)
            except:
                pass
        
        # Completion rate
        if total_emp > 0:
            st.markdown("---")
            st.markdown("### 📊 Learning Engagement Rate")
            engagement = int((total_enrollments / total_emp) * 100) if total_emp > 0 else 0
            st.metric("Employees with at least 1 course", f"{engagement}%")
            st.progress(engagement / 100)
    
    # ===== TAB 4: INNOVATION & IDEAS =====
    with tab4:
        st.subheader("💡 Innovation & Ideas Analytics")
        
        col1, col2 = st.columns(2)
        
        with col1:
            if total_ideas > 0:
                idea_status = {}
                for i in ideas:
                    status = i.get('status', 'Submitted')
                    idea_status[status] = idea_status.get(status, 0) + 1
                
                fig9 = px.pie(values=list(idea_status.values()), names=list(idea_status.keys()),
                            title="Ideas by Status", hole=0.5,
                            color_discrete_sequence=['#a0aec0', '#d69e2e', '#3182ce', '#38a169'])
                fig9.update_layout(height=350)
                st.plotly_chart(fig9, use_container_width=True)
        
        with col2:
            if total_ideas > 0:
                idea_categories = {}
                for i in ideas:
                    cat = i.get('category', 'Other')
                    idea_categories[cat] = idea_categories.get(cat, 0) + 1
                
                if idea_categories:
                    fig10 = px.bar(x=list(idea_categories.keys()), y=list(idea_categories.values()),
                                 title="Ideas by Category", color=list(idea_categories.values()),
                                 color_continuous_scale=['#CC0000', '#d69e2e', '#38a169'])
                    fig10.update_layout(height=350)
                    st.plotly_chart(fig10, use_container_width=True)
        
        # Employee participation
        if total_emp > 0 and total_ideas > 0:
            st.markdown("---")
            idea_participants = len(set(i.get('employee_id', '') for i in ideas))
            participation = int((idea_participants / total_emp) * 100)
            st.metric("Innovation Participation", f"{participation}%", f"{idea_participants} contributors")
            st.progress(participation / 100)
    
    # ===== EXPORT =====
    st.markdown("---")
    st.download_button("📥 Download Analytics Report (CSV)", 
                      pd.DataFrame([{
                          'Metric': 'Total Employees', 'Value': total_emp,
                          'Metric': 'Active Employees', 'Value': active_emp,
                          'Metric': 'Departments', 'Value': dept_count,
                          'Metric': 'Candidates', 'Value': total_candidates,
                          'Metric': 'LMS Enrollments', 'Value': total_enrollments,
                          'Metric': 'Requests', 'Value': total_requests,
                          'Metric': 'Appraisals', 'Value': total_appraisals,
                          'Metric': 'Ideas', 'Value': total_ideas
                      }]).to_csv(index=False),
                      "advanced_analytics.csv", "text/csv")


def log_audit_action(action, details, module='General'):
    """Log any action to the audit trail database"""
    try:
        user_name = st.session_state.user.get('name', 'System') if st.session_state.user else 'System'
        db._post("audit_trail", {
            "action": action,
            "details": details,
            "user_name": user_name,
            "timestamp_text": datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
            "module": module
        })
    except:
        pass


def send_kpi_notification(action, employee_name, employee_email, hod_email=None):
    """Send KPI workflow notifications"""
    try:
        from utils.email_service import EmailService
        es = EmailService()
        
        if action == 'submitted_to_employee':
            es.send_email(employee_email, 
                "✅ Your KPIs Have Been Submitted", 
                f"Dear {employee_name},\n\nYour KPIs have been successfully submitted and are now locked for review.\n\nYour HOD will review them shortly. You will be notified of their decision.\n\nThank you.\nChurchgate Group HR")
        
        if action == 'submitted_to_hod' and hod_email:
            es.send_email(hod_email,
                f"📊 KPIs Ready for Review — {employee_name}",
                f"Dear HOD,\n\n{employee_name} has submitted their KPIs and they are ready for your review.\n\nPlease log in to the HRIS to review, approve, or request revisions.\n\nhttps://churchgate-churchgate-hris.hf.space\n\nChurchgate Group HR")
        
        if action == 'approved':
            es.send_email(employee_email,
                "🎉 Your KPIs Have Been Approved!",
                f"Dear {employee_name},\n\nGreat news! Your HOD has approved your KPIs. They are now sealed and ready for the upcoming appraisal cycle.\n\nYou can view them in the Performance & OKRs module.\n\nChurchgate Group HR")
        
        if action == 'revision_requested':
            es.send_email(employee_email,
                "🔄 KPI Revision Requested",
                f"Dear {employee_name},\n\nYour HOD has requested revisions to your KPIs. Your KPIs have been unlocked for editing.\n\nPlease log in, review the HOD comments, update your KPIs, and resubmit.\n\nhttps://churchgate-churchgate-hris.hf.space\n\nChurchgate Group HR")
        
        return True
    except:
        return False


def send_browser_notification(title, body, user_email=None):
    """Queue a browser push notification"""
    if 'pending_notifications' not in st.session_state:
        st.session_state.pending_notifications = []
    st.session_state.pending_notifications.append({
        'title': title,
        'body': body,
        'time': datetime.now().strftime('%H:%M')
    })

def main():
    if 'user' not in st.session_state:
        st.session_state.user = None
    
    # Automated celebration email trigger (called by cron job)
    query_params = st.query_params
    if 'trigger_celebration' in query_params:
        try:
            bdays, annivs, msg = send_celebration_emails()
            st.write(f"Celebration emails: {bdays} birthdays, {annivs} anniversaries. {msg}")
        except Exception as e:
            st.write(f"Celebration email error: {e}")
        st.stop()
    
    # Persist login across refreshes
    if st.session_state.user is None:
        qp = st.query_params
        if 'logged_in' in qp:
            user_email = qp['logged_in']
            try:
                result = db._get("users", {"email": user_email})
                if result and len(result) > 0:
                    st.session_state.user = result[0]
                    st.session_state.authenticated = True
                    st.rerun()
            except:
                pass
    
    if st.session_state.user is not None:
        if 'last_activity' not in st.session_state or not isinstance(st.session_state.last_activity, datetime):
            st.session_state.last_activity = datetime.now()
        idle_time = (datetime.now() - st.session_state.last_activity).total_seconds()
        if idle_time > 900:
            st.session_state.user = None
            st.session_state.last_activity = None
            st.query_params.clear()
            st.warning("⏰ Session expired due to inactivity. Please log in again.")
            st.rerun()
        st.session_state.last_activity = datetime.now()
    
    if st.session_state.user is None:
        qp = st.query_params
        if 'logged_in' in qp:
            user_email = qp['logged_in']
            try:
                result = db._get("users", {"email": user_email})
                if result and len(result) > 0:
                    st.session_state.user = result[0]
                    st.rerun()
            except:
                pass
    
    if st.session_state.user is None:
        login_section()
    else:
        page = sidebar_navigation()
        if 'navigate_to' in st.session_state:
            page = st.session_state.pop('navigate_to')
        
        page_routes = {
            "🏠 Employee Dashboard": employee_dashboard,
            "📊 Executive Dashboard": executive_dashboard,
            "👥 Employee Management": employee_management,
            "📈 Performance & OKRs": performance_okrs,
            "📈 My Performance & OKRs": performance_okrs,
            "🚀 Promotions": promotions,
            "💼 Recruitment Hub": recruitment_hub,
            "🤖 AI Recruitment Agent": ai_recruitment_agent,
            "📊 Reports & Analytics": reports_analytics,
            "💬 Chat & Communications": chat_communications,
            "🎓 Training & Development": training_development,
            "🔔 Notifications": notifications_page,
            "📋 My Documents": my_documents,
            "💡 Ideas Box": ideas_box,
            "📅 Calendar": company_calendar,
            "🎯 My Goals": personal_goals,
            "🔄 Requests Hub": requests_hub,
            "🌐 Directory": employee_directory_readonly,
            "📚 Knowledge Base": knowledge_base,
            "🎉 Wellness & Perks": wellness_perks,
            "🎓 LMS": lms_dashboard,
           "📋 Audit Log": audit_log_viewer,
            "📊 Advanced Analytics": advanced_analytics,
            "👤 My Profile": my_profile,
        }
        page_func = page_routes.get(page, employee_dashboard)
        page_func()

# Show pending browser notifications
    if st.session_state.get('pending_notifications'):
        for notif in st.session_state.pending_notifications:
            st.markdown(f"""
            <script>
                showNotification('{notif['title']}', '{notif['body']}');
            </script>
            """, unsafe_allow_html=True)
        st.session_state.pending_notifications = []

if __name__ == "__main__":
    main()